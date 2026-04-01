import tkinter as tk
from tkinter import ttk, font, colorchooser, filedialog, messagebox
import json
import os
from datetime import datetime
from PIL import Image, ImageTk
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class WritingApp:
    def __init__(self, root):
        self.root = root
        self.root.title("专业码字软件 - 网文创作助手")
        self.root.geometry("1400x800")
        
        # 数据存储
        self.chapters = {}  # 存储章节内容
        self.current_chapter = "第一章"
        self.notes = ""
        self.extra_windows = []  # 存储额外的章节窗口
        self.split_mode = False  # 分屏模式状态
        self.split_text_areas = []  # 存储分屏的文本框
        
        # 配置文件路径
        self.config_file = "writing_config.json"
        self.load_config()
        
        # 初始化章节
        if not self.chapters:
            self.chapters[self.current_chapter] = ""
        
        # 创建主界面
        self.create_menu()
        self.create_main_layout()
        
        # 绑定事件
        self.text_area.bind('<KeyRelease>', self.update_stats)
        self.text_area.bind('<FocusIn>', self.update_stats)
        
        # 自动保存
        self.auto_save()
        
        # 设置默认排版
        self.apply_default_format()
        
    def create_menu(self):
        """创建菜单栏"""
        menubar = tk.Menu(self.root)
        self.root.config(menu=menubar)
        
        # 文件菜单
        file_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="文件", menu=file_menu)
        file_menu.add_command(label="新建项目", command=self.new_project)
        file_menu.add_command(label="打开项目", command=self.open_project)
        file_menu.add_command(label="保存项目", command=self.save_project)
        file_menu.add_separator()
        file_menu.add_command(label="一键导出", command=self.export_with_selection)
        file_menu.add_separator()
        file_menu.add_command(label="更改存储路径", command=self.change_storage_path)
        file_menu.add_separator()
        file_menu.add_command(label="退出", command=self.root.quit)
        
        # 编辑菜单
        edit_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="编辑", menu=edit_menu)
        edit_menu.add_command(label="撤销", command=self.undo)
        edit_menu.add_command(label="重做", command=self.redo)
        edit_menu.add_separator()
        edit_menu.add_command(label="剪切", command=self.cut)
        edit_menu.add_command(label="复制", command=self.copy)
        edit_menu.add_command(label="粘贴", command=self.paste)
        edit_menu.add_separator()
        edit_menu.add_command(label="一键排版", command=self.one_key_format)
        
        # 视图菜单
        view_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="视图", menu=view_menu)
        view_menu.add_command(label="分屏对比模式", command=self.toggle_split_mode)
        view_menu.add_command(label="退出分屏模式", command=self.exit_split_mode)
        
    def create_main_layout(self):
        """创建主布局"""
        # 主容器
        self.main_container = tk.Frame(self.root)
        self.main_container.pack(fill=tk.BOTH, expand=True)
        
        # 创建侧边栏（自动隐藏）
        self.create_auto_hide_sidebar()
        
        # 创建编辑区域
        self.create_edit_area()
        
    def create_auto_hide_sidebar(self):
        """创建自动隐藏侧边栏"""
        # 侧边栏容器
        self.sidebar_container = tk.Frame(self.main_container, width=280)
        self.sidebar_container.pack(side=tk.LEFT, fill=tk.Y)
        self.sidebar_container.pack_propagate(False)
        
        # 侧边栏内容
        self.sidebar = tk.Frame(self.sidebar_container, bg='#f5f5f5')
        self.sidebar.pack(fill=tk.BOTH, expand=True)
        
        # 触发区域（鼠标移入移出）
        self.sidebar.bind('<Enter>', self.show_sidebar)
        self.sidebar_container.bind('<Leave>', self.hide_sidebar)
        
        # 默认隐藏侧边栏
        self.sidebar_container.config(width=0)
        
        # 创建侧边栏内容
        self.create_sidebar_content()
        
    def show_sidebar(self, event=None):
        """显示侧边栏"""
        self.sidebar_container.config(width=280)
        
    def hide_sidebar(self, event=None):
        """隐藏侧边栏"""
        # 检查鼠标是否在侧边栏内
        x, y = self.root.winfo_pointerxy()
        widget = self.root.winfo_containing(x, y)
        if widget and (widget == self.sidebar or self.is_child_of(widget, self.sidebar)):
            return
        self.sidebar_container.config(width=0)
        
    def is_child_of(self, child, parent):
        """检查widget是否是parent的子级"""
        while child:
            if child == parent:
                return True
            child = child.master
        return False
        
    def create_sidebar_content(self):
        """创建侧边栏内容"""
        # 标题
        title_label = tk.Label(self.sidebar, text="控制面板", bg='#f5f5f5', 
                                font=('微软雅黑', 14, 'bold'), fg='#333')
        title_label.pack(pady=10)
        
        # 章节切换（带滚动条）
        chapter_frame = tk.LabelFrame(self.sidebar, text="章节管理", bg='#f5f5f5', 
                                       padx=5, pady=5, font=('微软雅黑', 10))
        chapter_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        # 创建滚动条框架
        chapter_scroll_frame = tk.Frame(chapter_frame)
        chapter_scroll_frame.pack(fill=tk.BOTH, expand=True)
        
        chapter_canvas = tk.Canvas(chapter_scroll_frame, bg='#f5f5f5', highlightthickness=0)
        chapter_scrollbar = tk.Scrollbar(chapter_scroll_frame, orient="vertical", command=chapter_canvas.yview)
        self.chapter_list_frame = tk.Frame(chapter_canvas, bg='#f5f5f5')
        
        chapter_canvas.configure(yscrollcommand=chapter_scrollbar.set)
        
        chapter_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        chapter_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        canvas_window = chapter_canvas.create_window((0, 0), window=self.chapter_list_frame, anchor="nw")
        
        self.chapter_buttons = {}
        self.update_chapter_list()
        
        self.chapter_list_frame.bind("<Configure>", lambda e: chapter_canvas.configure(scrollregion=chapter_canvas.bbox("all")))
        chapter_canvas.bind('<Configure>', lambda e: chapter_canvas.itemconfig(canvas_window, width=e.width))
        
        # 新建章节按钮
        tk.Button(chapter_frame, text="+ 新建章节", command=self.new_chapter,
                  bg='#4CAF50', fg='white', font=('微软雅黑', 9)).pack(pady=5, fill=tk.X)
        
        # 外观设置
        appearance_frame = tk.LabelFrame(self.sidebar, text="外观设置", bg='#f5f5f5', padx=5, pady=5)
        appearance_frame.pack(fill=tk.X, padx=10, pady=5)
        
        # 字体选择
        tk.Label(appearance_frame, text="字体:", bg='#f5f5f5', font=('微软雅黑', 9)).pack(anchor=tk.W)
        self.font_family = tk.StringVar(value="微软雅黑")
        font_menu = ttk.Combobox(appearance_frame, textvariable=self.font_family, 
                                   values=list(font.families())[:50], width=20)
        font_menu.pack(pady=2)
        font_menu.bind('<<ComboboxSelected>>', self.change_font)
        
        # 字体大小
        tk.Label(appearance_frame, text="字体大小:", bg='#f5f5f5', font=('微软雅黑', 9)).pack(anchor=tk.W)
        self.font_size = tk.IntVar(value=14)
        size_scale = tk.Scale(appearance_frame, from_=8, to=72, orient=tk.HORIZONTAL,
                               variable=self.font_size, command=self.change_font_size)
        size_scale.pack(fill=tk.X)
        
        # 背景颜色
        tk.Button(appearance_frame, text="选择背景颜色", command=self.change_bg_color,
                  bg='#FF9800', fg='white', font=('微软雅黑', 9)).pack(pady=5, fill=tk.X)
        
        # 背景图片
        tk.Button(appearance_frame, text="选择背景图片", command=self.change_bg_image,
                  bg='#9C27B0', fg='white', font=('微软雅黑', 9)).pack(pady=5, fill=tk.X)
        
        # 亮度调节
        tk.Label(appearance_frame, text="亮度调节:", bg='#f5f5f5', font=('微软雅黑', 9)).pack(anchor=tk.W)
        self.brightness = tk.Scale(appearance_frame, from_=0, to=100, orient=tk.HORIZONTAL,
                                    command=self.adjust_brightness)
        self.brightness.set(50)
        self.brightness.pack(fill=tk.X)
        
        # 备忘
        notes_frame = tk.LabelFrame(self.sidebar, text="备忘", bg='#f5f5f5', padx=5, pady=5)
        notes_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        self.notes_text = tk.Text(notes_frame, height=8, width=25, bg='white', 
                                   font=('微软雅黑', 9), wrap=tk.WORD)
        self.notes_text.pack(fill=tk.BOTH, expand=True)
        self.notes_text.insert(1.0, self.notes)
        
        tk.Button(notes_frame, text="独立窗口", command=self.open_notes_window,
                  bg='#00BCD4', fg='white', font=('微软雅黑', 9)).pack(pady=5, fill=tk.X)
        
    def update_chapter_list(self):
        """更新章节列表"""
        # 清除现有按钮
        for widget in self.chapter_list_frame.winfo_children():
            widget.destroy()
        self.chapter_buttons.clear()
        
        # 创建章节按钮
        for chapter in self.chapters.keys():
            btn_frame = tk.Frame(self.chapter_list_frame, bg='#f5f5f5')
            btn_frame.pack(fill=tk.X, pady=2)
            
            btn = tk.Button(btn_frame, text=chapter, command=lambda c=chapter: self.switch_to_chapter(c),
                           bg='#e0e0e0' if chapter == self.current_chapter else 'white',
                           fg='#333', font=('微软雅黑', 9), anchor='w', padx=10)
            btn.pack(side=tk.LEFT, fill=tk.X, expand=True)
            
            del_btn = tk.Button(btn_frame, text="×", command=lambda c=chapter: self.delete_chapter(c),
                               bg='#ff4444', fg='white', font=('微软雅黑', 9, 'bold'), width=3)
            del_btn.pack(side=tk.RIGHT)
            
            self.chapter_buttons[chapter] = btn
            
    def switch_to_chapter(self, chapter):
        """切换到指定章节"""
        # 保存当前章节
        self.chapters[self.current_chapter] = self.text_area.get(1.0, tk.END)
        # 切换到新章节
        self.current_chapter = chapter
        self.load_chapter_content()
        self.update_chapter_list()
        
    def delete_chapter(self, chapter):
        """删除章节"""
        if len(self.chapters) <= 1:
            messagebox.showwarning("警告", "至少需要保留一个章节！")
            return
            
        if messagebox.askyesno("确认", f"确定要删除章节 '{chapter}' 吗？"):
            del self.chapters[chapter]
            if self.current_chapter == chapter:
                self.current_chapter = list(self.chapters.keys())[0]
            self.update_chapter_list()
            self.load_chapter_content()
            
    def create_edit_area(self):
        """创建编辑区域"""
        edit_container = tk.Frame(self.main_container)
        edit_container.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)
        
        # 工具栏
        toolbar = tk.Frame(edit_container, bg='#e0e0e0', height=40)
        toolbar.pack(fill=tk.X)
        
        # 统计信息
        self.stats_label = tk.Label(toolbar, text="字数: 0 | 字符数: 0 | 行数: 1", 
                                    bg='#e0e0e0', font=('微软雅黑', 10))
        self.stats_label.pack(side=tk.LEFT, padx=10)
        
        # 当前章节标签
        self.chapter_label = tk.Label(toolbar, text=f"当前: {self.current_chapter}", 
                                      bg='#e0e0e0', font=('微软雅黑', 10, 'bold'))
        self.chapter_label.pack(side=tk.LEFT, padx=20)
        
        # 排版按钮
        tk.Button(toolbar, text="一键排版", command=self.one_key_format,
                 bg='#2196F3', fg='white', font=('微软雅黑', 9)).pack(side=tk.RIGHT, padx=5, pady=2)
        
        # 文本编辑区域容器
        self.text_container = tk.Frame(edit_container)
        self.text_container.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # 创建主文本框
        self.create_text_widgets()
        
    def create_text_widgets(self):
        """创建文本框组件"""
        # 清除现有组件
        for widget in self.text_container.winfo_children():
            widget.destroy()
        
        if self.split_mode:
            # 分屏模式：创建两个文本框
            self.text_area1 = self.create_text_widget(self.text_container, side=tk.LEFT)
            self.text_area2 = self.create_text_widget(self.text_container, side=tk.RIGHT)
            self.text_area = self.text_area1  # 主文本框设为第一个
            self.split_text_areas = [self.text_area1, self.text_area2]
            
            # 加载内容
            self.text_area1.insert(1.0, self.chapters.get(self.current_chapter, ""))
            chapters_list = list(self.chapters.keys())
            second_chapter = chapters_list[1] if len(chapters_list) > 1 else chapters_list[0]
            self.text_area2.insert(1.0, self.chapters.get(second_chapter, ""))
        else:
            # 普通模式：创建一个文本框
            self.text_area = self.create_text_widget(self.text_container)
            self.load_chapter_content()
            
    def create_text_widget(self, parent, side=None):
        """创建文本框组件"""
        frame = tk.Frame(parent)
        if side:
            frame.pack(side=side, fill=tk.BOTH, expand=True, padx=5)
        else:
            frame.pack(fill=tk.BOTH, expand=True)
            
        scrollbar = tk.Scrollbar(frame)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        text_area = tk.Text(frame, wrap=tk.WORD, font=('微软雅黑', 14),
                            yscrollcommand=scrollbar.set, undo=True, maxundo=100,
                            padx=15, pady=15, spacing1=5, spacing2=2, spacing3=5)
        text_area.pack(fill=tk.BOTH, expand=True)
        scrollbar.config(command=text_area.yview)
        
        text_area.bind('<KeyRelease>', self.update_stats)
        
        return text_area
        
    def toggle_split_mode(self):
        """切换分屏模式"""
        if not self.split_mode:
            # 保存当前内容
            self.chapters[self.current_chapter] = self.text_area.get(1.0, tk.END)
            self.split_mode = True
            self.create_text_widgets()
            messagebox.showinfo("分屏模式", "已进入分屏对比模式，右侧为第一个其他章节")
        else:
            messagebox.showinfo("提示", "已在分屏模式中")
            
    def exit_split_mode(self):
        """退出分屏模式"""
        if self.split_mode:
            # 保存分屏内容
            if hasattr(self, 'text_area1'):
                self.chapters[self.current_chapter] = self.text_area1.get(1.0, tk.END)
            self.split_mode = False
            self.create_text_widgets()
            self.load_chapter_content()
            
    def apply_default_format(self):
        """应用默认排版（网文风格）"""
        self.text_area.config(
            wrap=tk.WORD,
            padx=20,
            pady=20,
            spacing1=5,  # 段落上间距
            spacing2=2,  # 行间距
            spacing3=5,  # 段落下间距
            font=('微软雅黑', 14)
        )
        
    def one_key_format(self):
        """一键排版功能"""
        content = self.text_area.get(1.0, tk.END)
        
        # 基本排版规则
        lines = content.split('\n')
        formatted_lines = []
        
        for line in lines:
            line = line.strip()
            if line:
                # 段落开头自动空两格
                if not line.startswith((' ', '\t')):
                    line = '    ' + line
                formatted_lines.append(line)
            else:
                formatted_lines.append('')
        
        formatted_content = '\n'.join(formatted_lines)
        
        self.text_area.delete(1.0, tk.END)
        self.text_area.insert(1.0, formatted_content)
        
        # 设置段落格式
        self.text_area.tag_configure('paragraph', spacing1=10, spacing2=5, spacing3=10)
        
        messagebox.showinfo("排版完成", "已应用标准网文排版格式")
        
    def export_with_selection(self):
        """带章节选择的一键导出功能"""
        # 创建选择窗口
        export_window = tk.Toplevel(self.root)
        export_window.title("导出章节选择")
        export_window.geometry("500x400")
        
        tk.Label(export_window, text="请选择要导出的章节：", font=('微软雅黑', 12, 'bold')).pack(pady=10)
        
        # 章节选择列表
        frame = tk.Frame(export_window)
        frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        
        canvas = tk.Canvas(frame)
        scrollbar = tk.Scrollbar(frame, orient="vertical", command=canvas.yview)
        scrollable_frame = tk.Frame(canvas)
        
        scrollable_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # 创建复选框
        check_vars = {}
        for chapter in self.chapters.keys():
            var = tk.BooleanVar(value=True)
            check_vars[chapter] = var
            tk.Checkbutton(scrollable_frame, text=chapter, variable=var, 
                          font=('微软雅黑', 10)).pack(anchor=tk.W, pady=2)
            
        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # 导出格式选择
        format_frame = tk.Frame(export_window)
        format_frame.pack(fill=tk.X, padx=20, pady=10)
        
        tk.Label(format_frame, text="导出格式：", font=('微软雅黑', 10)).pack(side=tk.LEFT)
        export_format = tk.StringVar(value="txt")
        tk.Radiobutton(format_frame, text="TXT", variable=export_format, value="txt").pack(side=tk.LEFT, padx=10)
        tk.Radiobutton(format_frame, text="Word", variable=export_format, value="docx").pack(side=tk.LEFT, padx=10)
        
        def do_export():
            selected_chapters = [ch for ch, var in check_vars.items() if var.get()]
            if not selected_chapters:
                messagebox.showwarning("警告", "请至少选择一个章节")
                return
                
            format_type = export_format.get()
            self.perform_export(selected_chapters, format_type)
            export_window.destroy()
            
        tk.Button(export_window, text="确定导出", command=do_export,
                 bg='#4CAF50', fg='white', font=('微软雅黑', 10)).pack(pady=20)
                 
    def perform_export(self, chapters, format_type):
        """执行导出"""
        filename = filedialog.asksaveasfilename(
            defaultextension=f".{format_type}",
            filetypes=[(f"{format_type.upper()} files", f"*.{format_type}"), ("All files", "*.*")]
        )
        
        if filename:
            if format_type == "txt":
                with open(filename, 'w', encoding='utf-8') as f:
                    for chapter in chapters:
                        f.write(f"\n{'='*50}\n")
                        f.write(f"{chapter}\n")
                        f.write(f"{'='*50}\n\n")
                        f.write(self.chapters.get(chapter, ""))
                        f.write("\n\n")
            else:  # docx
                doc = docx.Document()
                for chapter in chapters:
                    doc.add_heading(chapter, level=1)
                    paragraphs = self.chapters.get(chapter, "").split('\n')
                    for para in paragraphs:
                        if para.strip():
                            p = doc.add_paragraph(para)
                            p.paragraph_format.first_line_indent = Pt(28)
                            p.paragraph_format.line_spacing = 1.5
                    doc.add_page_break()
                doc.save(filename)
                
            messagebox.showinfo("导出成功", f"已成功导出{len(chapters)}个章节到：\n{filename}")
            
    def change_storage_path(self):
        """更改存储路径"""
        new_path = filedialog.askdirectory(title="选择新的存储路径")
        if new_path:
            self.storage_path = new_path
            self.save_config()
            messagebox.showinfo("成功", f"存储路径已更改为：\n{new_path}")
            
    def load_config(self):
        """加载配置"""
        if os.path.exists(self.config_file):
            try:
                with open(self.config_file, 'r', encoding='utf-8') as f:
                    config = json.load(f)
                    self.storage_path = config.get('storage_path', os.getcwd())
            except:
                self.storage_path = os.getcwd()
        else:
            self.storage_path = os.getcwd()
            
    def save_config(self):
        """保存配置"""
        config = {
            'storage_path': self.storage_path
        }
        with open(self.config_file, 'w', encoding='utf-8') as f:
            json.dump(config, f)
            
    def change_bg_image(self):
        """更改背景图片"""
        filename = filedialog.askopenfilename(filetypes=[("Image files", "*.png *.jpg *.jpeg *.gif *.bmp")])
        if filename:
            try:
                image = Image.open(filename)
                # 调整图片大小以适应文本框
                image = image.resize((800, 600), Image.Resampling.LANCZOS)
                self.bg_image = ImageTk.PhotoImage(image)
                self.text_area.config(bg='white')  # 先设置背景色
                # 使用tag来设置背景图片比较复杂，这里简化处理
                self.text_area.image_create("1.0", image=self.bg_image)
                messagebox.showinfo("成功", "背景图片已设置")
            except Exception as e:
                messagebox.showerror("错误", f"无法加载图片：{str(e)}")
                
    def open_notes_window(self):
        """打开独立备忘窗口（自动保存）"""
        notes_window = tk.Toplevel(self.root)
        notes_window.title("独立备忘窗口")
        notes_window.geometry("500x600")
        
        text_area = tk.Text(notes_window, wrap=tk.WORD, font=('微软雅黑', 11), padx=10, pady=10)
        text_area.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        text_area.insert(1.0, self.notes_text.get(1.0, tk.END))
        
        def on_close():
            # 关闭时自动保存并更新
            content = text_area.get(1.0, tk.END)
            self.notes_text.delete(1.0, tk.END)
            self.notes_text.insert(1.0, content)
            self.notes = content
            notes_window.destroy()
            
        notes_window.protocol("WM_DELETE_WINDOW", on_close)
        
    def update_stats(self, event=None):
        """更新统计信息"""
        if hasattr(self, 'text_area'):
            content = self.text_area.get(1.0, tk.END)
            # 字数统计
            chinese_chars = sum(1 for char in content if '\u4e00' <= char <= '\u9fff')
            words = len(content.split())
            total_chars = len(content.replace('\n', '').replace(' ', ''))
            lines = int(self.text_area.index('end-1c').split('.')[0])
            
            self.stats_label.config(text=f"字数: {chinese_chars + words} | 字符数: {total_chars} | 行数: {lines}")
            
            # 自动保存当前章节
            if not self.split_mode:
                self.chapters[self.current_chapter] = self.text_area.get(1.0, tk.END)
            else:
                if hasattr(self, 'text_area1'):
                    self.chapters[self.current_chapter] = self.text_area1.get(1.0, tk.END)
                    
    def load_chapter_content(self):
        """加载章节内容"""
        if not self.split_mode and hasattr(self, 'text_area'):
            self.text_area.delete(1.0, tk.END)
            self.text_area.insert(1.0, self.chapters.get(self.current_chapter, ""))
            self.update_stats()
            self.chapter_label.config(text=f"当前: {self.current_chapter}")
            
    def new_chapter(self):
        """新建章节"""
        chapter_name = f"第{len(self.chapters) + 1}章"
        self.chapters[chapter_name] = ""
        self.update_chapter_list()
        self.switch_to_chapter(chapter_name)
        
    def change_font(self, event=None):
        """更改字体"""
        if hasattr(self, 'text_area'):
            current_font = font.Font(font=self.text_area['font'])
            self.text_area.config(font=(self.font_family.get(), current_font.cget('size')))
            
    def change_font_size(self, value):
        """更改字体大小"""
        if hasattr(self, 'text_area'):
            self.text_area.config(font=(self.font_family.get(), self.font_size.get()))
            
    def change_bg_color(self):
        """更改背景颜色"""
        color = colorchooser.askcolor()[1]
        if color and hasattr(self, 'text_area'):
            self.text_area.config(bg=color)
            
    def adjust_brightness(self, value):
        """调节亮度"""
        if hasattr(self, 'text_area'):
            brightness = int(value) / 100
            gray_value = int(255 * (1 - brightness * 0.5))
            color = f'#{gray_value:02x}{gray_value:02x}{gray_value:02x}'
            self.text_area.config(bg=color)
            
    def auto_save(self):
        """自动保存"""
        if not self.split_mode and hasattr(self, 'text_area'):
            self.chapters[self.current_chapter] = self.text_area.get(1.0, tk.END)
        elif self.split_mode and hasattr(self, 'text_area1'):
            self.chapters[self.current_chapter] = self.text_area1.get(1.0, tk.END)
        self.root.after(30000, self.auto_save)
        
    def new_project(self):
        """新建项目"""
        if messagebox.askyesno("确认", "新建项目会清空当前内容，确定吗？"):
            self.chapters = {"第一章": ""}
            self.current_chapter = "第一章"
            self.update_chapter_list()
            self.load_chapter_content()
            self.notes_text.delete(1.0, tk.END)
            
    def save_project(self):
        """保存项目"""
        filename = filedialog.asksaveasfilename(defaultextension=".json", 
                                                filetypes=[("JSON files", "*.json"), ("All files", "*.*")])
        if filename:
            if not self.split_mode:
                self.chapters[self.current_chapter] = self.text_area.get(1.0, tk.END)
            else:
                self.chapters[self.current_chapter] = self.text_area1.get(1.0, tk.END)
                
            data = {
                "chapters": self.chapters,
                "notes": self.notes_text.get(1.0, tk.END),
                "settings": {
                    "font_family": self.font_family.get(),
                    "font_size": self.font_size.get()
                }
            }
            
            with open(filename, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
                
            messagebox.showinfo("成功", "项目已保存！")
            
    def open_project(self):
        """打开项目"""
        filename = filedialog.askopenfilename(filetypes=[("JSON files", "*.json"), ("All files", "*.*")])
        if filename:
            try:
                with open(filename, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    
                self.chapters = data.get("chapters", {"第一章": ""})
                self.current_chapter = list(self.chapters.keys())[0]
                self.update_chapter_list()
                self.load_chapter_content()
                
                notes = data.get("notes", "")
                self.notes_text.delete(1.0, tk.END)
                self.notes_text.insert(1.0, notes)
                
                settings = data.get("settings", {})
                if settings:
                    self.font_family.set(settings.get("font_family", "微软雅黑"))
                    self.font_size.set(settings.get("font_size", 14))
                    self.change_font()
                    self.change_font_size(self.font_size.get())
                    
                messagebox.showinfo("成功", "项目已打开！")
            except Exception as e:
                messagebox.showerror("错误", f"打开失败：{str(e)}")
                
    def undo(self):
        if hasattr(self, 'text_area'):
            self.text_area.edit_undo()
            
    def redo(self):
        if hasattr(self, 'text_area'):
            self.text_area.edit_redo()
            
    def cut(self):
        if hasattr(self, 'text_area'):
            self.text_area.event_generate("<<Cut>>")
            
    def copy(self):
        if hasattr(self, 'text_area'):
            self.text_area.event_generate("<<Copy>>")
            
    def paste(self):
        if hasattr(self, 'text_area'):
            self.text_area.event_generate("<<Paste>>")

if __name__ == "__main__":
    root = tk.Tk()
    app = WritingApp(root)
    root.mainloop()