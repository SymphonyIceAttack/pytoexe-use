"""
Курсовая работа: Течение вязкой жидкости по плоскому каналу
Метод Контрольных Объёмов + решение СЛАУ (Гаусс, Гаусс-Зейдель, TDMA)
GUI-версия + генерация Word-отчёта
"""

import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import numpy as np
import matplotlib
matplotlib.use("TkAgg")
from matplotlib.figure import Figure
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg, NavigationToolbar2Tk

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import time

# ═════════════════════════════════════════════════════════════════════════════
#                              СОЛВЕРЫ
# ═════════════════════════════════════════════════════════════════════════════

def gauss_elimination(A, b):
    n = len(b)
    Ab = np.column_stack([A.astype(float), b.astype(float)])
    for k in range(n):
        p = int(np.argmax(np.abs(Ab[k:, k]))) + k
        Ab[[k, p]] = Ab[[p, k]]
        for i in range(k + 1, n):
            if Ab[k, k] == 0.0: continue
            m = Ab[i, k] / Ab[k, k]
            Ab[i, k:] -= m * Ab[k, k:]
    x = np.zeros(n)
    for i in range(n - 1, -1, -1):
        x[i] = (Ab[i, n] - np.dot(Ab[i, i+1:n], x[i+1:])) / Ab[i, i]
    return x


def gauss_seidel(A, b, tol=1e-12, max_iter=100_000):
    n = len(b)
    x = np.zeros(n)
    history = []
    for _ in range(max_iter):
        x_old = x.copy()
        for i in range(n):
            sigma = b[i] - sum(A[i, j] * x[j] for j in range(n) if j != i)
            x[i] = sigma / A[i, i]
        delta = np.max(np.abs(x - x_old))
        history.append(delta)
        if delta < tol: break
    return x, len(history), history


def tdma(A, b):
    n = len(b)
    lo = np.array([A[i, i-1] for i in range(1, n)], dtype=float)
    diag = np.array([A[i, i] for i in range(n)], dtype=float)
    up = np.array([A[i, i+1] for i in range(n-1)], dtype=float)
    d = b.astype(float).copy()
    for i in range(1, n):
        w = lo[i-1] / diag[i-1]
        diag[i] -= w * up[i-1]
        d[i] -= w * d[i-1]
    x = np.zeros(n)
    x[n-1] = d[n-1] / diag[n-1]
    for i in range(n-2, -1, -1):
        x[i] = (d[i] - up[i] * x[i+1]) / diag[i]
    return x


def solve(Re, dPdx, U0, Ua, N):
    a = 1.0
    dy = a / (N - 1)
    y = np.linspace(0.0, a, N)
    Gamma = 2.0 / Re
    aE = aW = Gamma / dy
    aP = aE + aW
    b_src = -dPdx * dy

    A = np.zeros((N, N))
    rhs = np.zeros(N)
    A[0, 0] = 1.0; rhs[0] = U0
    A[N-1, N-1] = 1.0; rhs[N-1] = Ua
    for i in range(1, N-1):
        A[i, i-1] = -aW
        A[i, i] = aP
        A[i, i+1] = -aE
        rhs[i] = b_src

    C2 = U0
    C1 = Ua - U0 - (Re / 4.0) * dPdx
    v_exact = (Re / 4.0) * dPdx * y**2 + C1 * y + C2

    # Максимум параболы
    A_coeff = (Re / 4.0) * dPdx
    if A_coeff != 0:
        y_max = -C1 / (2 * A_coeff)
        y_max = max(0, min(1, y_max))
        v_max = A_coeff * y_max**2 + C1 * y_max + C2
    else:
        v_max = max(v_exact)
        y_max = y[np.argmax(v_exact)]

    t0 = time.perf_counter()
    v_g = gauss_elimination(A.copy(), rhs.copy())
    t_g = time.perf_counter() - t0

    t0 = time.perf_counter()
    v_gs, iters, history = gauss_seidel(A.copy(), rhs.copy())
    t_gs = time.perf_counter() - t0

    t0 = time.perf_counter()
    v_t = tdma(A.copy(), rhs.copy())
    t_t = time.perf_counter() - t0

    coeffs = dict(Gamma=Gamma, aE=aE, aW=aW, aP=aP, b=b_src, dy=dy, C1=C1, C2=C2, v_max=v_max, y_max=y_max)
    times = dict(gauss=t_g, gs=t_gs, tdma=t_t)
    return y, v_g, v_gs, v_t, v_exact, iters, history, coeffs, times


# ═════════════════════════════════════════════════════════════════════════════
#                              GUI
# ═════════════════════════════════════════════════════════════════════════════

class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Курсовая работа — Течение вязкой жидкости (Метод КО)")
        self.geometry("1200x750")
        self.resizable(True, True)
        self._build_ui()

    def _build_ui(self):
        self.configure(bg="#f5f6fa")

        # Верхняя панель с параметрами
        top = tk.Frame(self, bg="#2c3e50", pady=8)
        top.pack(fill="x")

        params = [
            ("Re", "16"),
            ("dP/dx", "-4"),
            ("U(0)", "0"),
            ("U(a)", "3"),
            ("N (узлов)", "11"),
        ]
        self.entries = {}
        for label, default in params:
            frm = tk.Frame(top, bg="#2c3e50", padx=8)
            frm.pack(side="left")
            tk.Label(frm, text=label, bg="#2c3e50", fg="white",
                     font=("Segoe UI", 10, "bold")).pack()
            e = tk.Entry(frm, width=8, font=("Segoe UI", 12),
                         justify="center", relief="flat",
                         bg="#ecf0f1", highlightthickness=1,
                         highlightbackground="#bdc3c7")
            e.insert(0, default)
            e.pack()
            self.entries[label] = e

        # Кнопки
        btn_frame = tk.Frame(top, bg="#2c3e50")
        btn_frame.pack(side="left", padx=15)

        tk.Button(btn_frame, text="  Рассчитать  ", command=self._run,
                  bg="#27ae60", fg="white", font=("Segoe UI", 11, "bold"),
                  relief="flat", padx=12, pady=4).pack(side="left", padx=5)

        tk.Button(btn_frame, text="  Сгенерировать Word-отчёт  ", command=self._generate_report,
                  bg="#2980b9", fg="white", font=("Segoe UI", 11, "bold"),
                  relief="flat", padx=12, pady=4).pack(side="left", padx=5)

        self.status = tk.StringVar(value="Введите параметры и нажмите «Рассчитать»")
        tk.Label(top, textvariable=self.status, bg="#2c3e50", fg="#bdc3c7",
                 font=("Segoe UI", 9)).pack(side="left", padx=15)

        # Основная область
        main = tk.Frame(self, bg="#f5f6fa")
        main.pack(fill="both", expand=True, padx=8, pady=8)

        # Левая часть — графики
        left = tk.Frame(main, bg="white", relief="flat",
                        highlightbackground="#ddd", highlightthickness=1)
        left.pack(side="left", fill="both", expand=True)

        self.fig = Figure(figsize=(7, 7), dpi=100)
        self.ax1 = self.fig.add_subplot(2, 1, 1)
        self.ax2 = self.fig.add_subplot(2, 1, 2)
        self.fig.tight_layout(pad=3)

        self.canvas = FigureCanvasTkAgg(self.fig, master=left)
        self.canvas.get_tk_widget().pack(fill="both", expand=True)

        toolbar_frame = tk.Frame(left)
        toolbar_frame.pack(fill="x")
        NavigationToolbar2Tk(self.canvas, toolbar_frame)

        # Правая часть — коэффициенты + таблица
        right = tk.Frame(main, bg="#f5f6fa", width=420)
        right.pack(side="right", fill="both", padx=(8, 0))
        right.pack_propagate(False)

        # Коэффициенты
        coeff_frame = tk.LabelFrame(right, text=" Коэффициенты и результаты ",
                                    font=("Segoe UI", 10, "bold"),
                                    bg="#f5f6fa", fg="#2c3e50", padx=8, pady=6)
        coeff_frame.pack(fill="x", pady=(0, 6))

        self.coeff_text = tk.Text(coeff_frame, height=12, width=48,
                                  font=("Consolas", 9), bg="#f8f9fa",
                                  relief="flat", borderwidth=0)
        self.coeff_text.pack()

        # Таблица
        table_frame = tk.LabelFrame(right, text=" Таблица значений скорости ",
                                    font=("Segoe UI", 10, "bold"),
                                    bg="#f5f6fa", fg="#2c3e50")
        table_frame.pack(fill="both", expand=True)

        cols = ("i", "y", "Гаусс", "Г.-Зейдель", "TDMA")
        self.tree = ttk.Treeview(table_frame, columns=cols, show="headings", height=18)
        widths = (30, 60, 100, 100, 100)
        for col, w in zip(cols, widths):
            self.tree.heading(col, text=col)
            self.tree.column(col, width=w, anchor="center", stretch=False)

        vsb = ttk.Scrollbar(table_frame, orient="vertical", command=self.tree.yview)
        self.tree.configure(yscrollcommand=vsb.set)
        vsb.pack(side="right", fill="y")
        self.tree.pack(fill="both", expand=True)

        self.tree.tag_configure("odd", background="#f0f4f8")
        self.tree.tag_configure("even", background="#ffffff")

    def _get_params(self):
        try:
            Re = float(self.entries["Re"].get())
            dPdx = float(self.entries["dP/dx"].get())
            U0 = float(self.entries["U(0)"].get())
            Ua = float(self.entries["U(a)"].get())
            N = int(self.entries["N (узлов)"].get())
        except ValueError:
            messagebox.showerror("Ошибка", "Все поля должны быть числами!\nN — целое число.")
            return None
        if Re <= 0 or N < 3:
            messagebox.showerror("Ошибка", "Re > 0 и N ≥ 3")
            return None
        return Re, dPdx, U0, Ua, N

    def _run(self):
        params = self._get_params()
        if params is None:
            return
        Re, dPdx, U0, Ua, N = params

        try:
            y, v_g, v_gs, v_t, v_ex, iters, history, cf, times = solve(Re, dPdx, U0, Ua, N)
        except Exception as e:
            messagebox.showerror("Ошибка расчёта", str(e))
            return

        self._update_results(y, v_g, v_gs, v_t, v_ex, iters, history, cf, times, Re, dPdx, N)
        self.status.set(f"Готово! Гаусс-Зейдель: {iters} итераций")

    def _update_results(self, y, v_g, v_gs, v_t, v_ex, iters, history, cf, times, Re, dPdx, N):
        # Коэффициенты
        self.coeff_text.delete("1.0", "end")
        text = (
            f"Γ   = 2/Re          = {cf['Gamma']:.6f}\n"
            f"dy  = 1/(N-1)       = {cf['dy']:.6f}\n"
            f"a_E = Γ/dy          = {cf['aE']:.6f}\n"
            f"a_W = Γ/dy          = {cf['aW']:.6f}\n"
            f"a_P = a_E + a_W     = {cf['aP']:.6f}\n"
            f"b   = -dP/dx·dy     = {cf['b']:.6f}\n\n"
            f"C1  = {cf['C1']:.6f}\n"
            f"C2  = {cf['C2']:.6f}\n\n"
            f"Макс. скорость v_max = {cf['v_max']:.6f}\n"
            f"при y = {cf['y_max']:.4f}\n\n"
            f"── Время решения ──\n"
            f"Гаусс:         {times['gauss']*1000:.3f} мс\n"
            f"Гаусс-Зейдель: {times['gs']*1000:.3f} мс\n"
            f"TDMA:          {times['tdma']*1000:.3f} мс"
        )
        self.coeff_text.insert("1.0", text)

        # Таблица
        for row in self.tree.get_children():
            self.tree.delete(row)
        for i in range(len(y)):
            tag = "odd" if i % 2 else "even"
            self.tree.insert("", "end", tag=tag, values=(
                i, f"{y[i]:.4f}",
                f"{v_g[i]:.8f}", f"{v_gs[i]:.8f}", f"{v_t[i]:.8f}"
            ))

        # Графики
        self.ax1.clear()
        self.ax1.plot(v_g, y, "b-o", lw=1.6, ms=5, label="Гаусс")
        self.ax1.plot(v_gs, y, "g--s", lw=1.6, ms=5, label="Гаусс-Зейдель")
        self.ax1.plot(v_t, y, "r-.^", lw=1.6, ms=5, label="TDMA")
        self.ax1.plot(v_ex, y, "k-", lw=2.5, label="Аналитика")
        self.ax1.set_xlabel("v (скорость)")
        self.ax1.set_ylabel("y (координата)")
        self.ax1.set_title(f"Профиль скорости  |  Re={Re}, dP/dx={dPdx}, N={N}")
        self.ax1.legend(fontsize=8)
        self.ax1.grid(True, linestyle="--", alpha=0.4)

        self.ax2.clear()
        self.ax2.semilogy(range(1, len(history)+1), history, "g-", lw=1.8)
        self.ax2.set_xlabel("Итерация")
        self.ax2.set_ylabel("max |Δx|")
        self.ax2.set_title(f"Сходимость Гаусса-Зейделя ({iters} итераций)")
        self.ax2.grid(True, linestyle="--", alpha=0.4)

        self.fig.tight_layout(pad=2.5)
        self.canvas.draw()

    def _generate_report(self):
        params = self._get_params()
        if params is None:
            return
        Re, dPdx, U0, Ua, N = params

        try:
            y, v_g, v_gs, v_t, v_ex, iters, history, cf, times = solve(Re, dPdx, U0, Ua, N)
        except Exception as e:
            messagebox.showerror("Ошибка", str(e))
            return

        # Создаём Word-отчёт (упрощённая версия)
        doc = Document()
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2)

        # Заголовок
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("ОТЧЁТ ПО КУРСОВОЙ РАБОТЕ")
        run.bold = True
        run.font.size = Pt(14)

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(f"Течение вязкой жидкости | Re={Re}, dP/dx={dPdx}, N={N}")
        run2.font.size = Pt(12)

        doc.add_paragraph()

        # Параметры
        doc.add_heading("Параметры варианта", level=2)
        tbl = doc.add_table(rows=6, cols=2)
        tbl.style = "Table Grid"
        data = [("Параметр", "Значение"),
                ("Re", str(Re)), ("dP/dx", str(dPdx)),
                ("U(0)", str(U0)), ("U(a)", str(Ua)), ("N", str(N))]
        for i, (k, v) in enumerate(data):
            tbl.rows[i].cells[0].text = k
            tbl.rows[i].cells[1].text = v
            if i == 0:
                for cell in tbl.rows[i].cells:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True

        doc.add_paragraph()
        doc.add_heading("Результаты", level=2)
        doc.add_paragraph(f"Максимальная скорость: v_max = {cf['v_max']:.6f} при y = {cf['y_max']:.4f}")
        doc.add_paragraph(f"Гаусс-Зейдель сошёлся за {iters} итераций")

        # Сохранение
        path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")],
            title="Сохранить отчёт как..."
        )
        if path:
            doc.save(path)
            messagebox.showinfo("Готово", f"Отчёт сохранён:\n{path}")


if __name__ == "__main__":
    app = App()
    app.mainloop()