import tkinter as tk
from tkinter import ttk, messagebox, scrolledtext, filedialog
from datetime import datetime
import json
import os
import csv
import subprocess
import platform
import hashlib

APP_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_FILE = os.path.join(APP_DIR, "journal.json")
PERSONNEL_FILE = os.path.join(APP_DIR, "personnel.json")
HISTORY_FILE = os.path.join(APP_DIR, "field_history.json")
USERS_FILE = os.path.join(APP_DIR, "users.json")
TEMPLATES_FILE = os.path.join(APP_DIR, "templates.json")

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def hash_password(pwd):
    return hashlib.sha256(pwd.encode("utf-8")).hexdigest()


class LoginWindow:
    def __init__(self, root):
        self.root = root
        self.root.title("Вход в систему")
        self.root.geometry("400x320")
        self.root.resizable(False, False)
        self.root.configure(bg="#f3f4f6")
        self.ensure_admin_exists()
        self.build_ui()
        self.center_window()

    def ensure_admin_exists(self):
        users = self.load_users()
        if not users:
            users.append({
                "username": "admin",
                "password": hash_password("adm"),
                "role": "admin",
                "name": "Администратор"
            })
            self.save_users(users)

    def load_users(self):
        if os.path.exists(USERS_FILE):
            try:
                with open(USERS_FILE, "r", encoding="utf-8") as f:
                    return json.load(f)
            except:
                return []
        return []

    def save_users(self, data):
        with open(USERS_FILE, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

    def center_window(self):
        self.root.update_idletasks()
        w, h = self.root.winfo_width(), self.root.winfo_height()
        x = (self.root.winfo_screenwidth() // 2) - (w // 2)
        y = (self.root.winfo_screenheight() // 2) - (h // 2)
        self.root.geometry(f"{w}x{h}+{x}+{y}")

    def build_ui(self):
        frame = tk.Frame(self.root, bg="#f3f4f6")
        frame.pack(expand=True, fill="both", padx=40, pady=30)

        tk.Label(frame, text="🔐 Вход в систему", font=("Segoe UI", 18, "bold"),
                 bg="#f3f4f6", fg="#111827").pack(pady=(0, 20))

        tk.Label(frame, text="Логин:", font=("Segoe UI", 11, "bold"),
                 bg="#f3f4f6", fg="#374151").pack(anchor="w", pady=(0, 4))

        # Combobox со списком пользователей
        users = self.load_users()
        logins = [u.get("username", "") for u in users if u.get("username")]
        self.login_combo = ttk.Combobox(frame, font=("Segoe UI", 12), values=logins, height=8)
        self.login_combo.pack(fill="x", pady=(0, 12))
        if logins:
            self.login_combo.set("Выберите логин...")

        tk.Label(frame, text="Пароль:", font=("Segoe UI", 11, "bold"),
                 bg="#f3f4f6", fg="#374151").pack(anchor="w", pady=(0, 4))
        self.pass_entry = tk.Entry(frame, font=("Segoe UI", 12), relief="solid", bd=1, show="•")
        self.pass_entry.pack(fill="x", pady=(0, 20))

        btn = tk.Button(frame, text="Войти", font=("Segoe UI", 12, "bold"),
                        bg="#2563eb", fg="white", bd=0, padx=20, pady=10,
                        cursor="hand2", command=self.do_login)
        btn.pack(fill="x")

        self.status = tk.Label(frame, text="", font=("Segoe UI", 10),
                               bg="#f3f4f6", fg="#ef4444")
        self.status.pack(pady=(10, 0))

        self.pass_entry.bind("<Return>", lambda e: self.do_login())
        self.login_combo.bind("<Return>", lambda e: self.pass_entry.focus())

    def do_login(self):
        login = self.login_combo.get().strip()
        pwd = self.pass_entry.get().strip()

        if not login or login == "Выберите логин...":
            self.status.config(text="❌ Выберите логин из списка")
            return

        users = self.load_users()

        for u in users:
            if u["username"] == login and u["password"] == hash_password(pwd):
                self.root.destroy()
                root = tk.Tk()
                app = RaportApp(root, user=u)
                root.mainloop()
                return

        self.status.config(text="❌ Неверный логин или пароль")


class UsersManager:
    def __init__(self, parent, app):
        self.app = app
        self.win = tk.Toplevel(parent)
        self.win.title("👤 Управление пользователями")
        self.win.geometry("550x500")
        self.win.configure(bg="#f3f4f6")
        self.win.transient(parent)
        self.win.grab_set()
        self.build_ui()

    def load_users(self):
        if os.path.exists(USERS_FILE):
            try:
                with open(USERS_FILE, "r", encoding="utf-8") as f:
                    return json.load(f)
            except:
                return []
        return []

    def save_users(self, data):
        with open(USERS_FILE, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

    def build_ui(self):
        tk.Label(self.win, text="👤 Управление пользователями",
                 font=("Segoe UI", 16, "bold"), bg="#f3f4f6", fg="#111827").pack(pady=(15, 5))

        table_frame = tk.Frame(self.win, bg="#f3f4f6")
        table_frame.pack(fill="both", expand=True, padx=20, pady=10)

        columns = ("№", "Логин", "Имя", "Роль")
        self.tree = ttk.Treeview(table_frame, columns=columns, show="headings", height=10)
        for col in columns:
            self.tree.heading(col, text=col)
        self.tree.column("№", width=40, anchor="center")
        self.tree.column("Логин", width=120)
        self.tree.column("Имя", width=200)
        self.tree.column("Роль", width=100, anchor="center")

        sb = ttk.Scrollbar(table_frame, orient="vertical", command=self.tree.yview)
        self.tree.configure(yscrollcommand=sb.set)
        self.tree.pack(side="left", fill="both", expand=True)
        sb.pack(side="right", fill="y")

        self.refresh_tree()

        add_frame = tk.LabelFrame(self.win, text=" Добавить пользователя ",
                                  font=("Segoe UI", 11, "bold"), bg="#f3f4f6", fg="#374151",
                                  bd=1, relief="solid")
        add_frame.pack(fill="x", padx=20, pady=(0, 10))

        inner = tk.Frame(add_frame, bg="#f3f4f6")
        inner.pack(fill="x", padx=15, pady=12)

        tk.Label(inner, text="Логин:", font=("Segoe UI", 10, "bold"),
                 bg="#f3f4f6", fg="#374151").grid(row=0, column=0, sticky="w", pady=(0, 5))
        self.en_login = tk.Entry(inner, font=("Segoe UI", 11), width=25)
        self.en_login.grid(row=0, column=1, sticky="w", padx=(8, 0), pady=(0, 5))

        tk.Label(inner, text="Имя:", font=("Segoe UI", 10, "bold"),
                 bg="#f3f4f6", fg="#374151").grid(row=1, column=0, sticky="w", pady=(0, 5))
        self.en_name = tk.Entry(inner, font=("Segoe UI", 11), width=25)
        self.en_name.grid(row=1, column=1, sticky="w", padx=(8, 0), pady=(0, 5))

        tk.Label(inner, text="Пароль:", font=("Segoe UI", 10, "bold"),
                 bg="#f3f4f6", fg="#374151").grid(row=2, column=0, sticky="w", pady=(0, 5))
        self.en_pwd = tk.Entry(inner, font=("Segoe UI", 11), width=25, show="•")
        self.en_pwd.grid(row=2, column=1, sticky="w", padx=(8, 0), pady=(0, 5))

        tk.Label(inner, text="Роль:", font=("Segoe UI", 10, "bold"),
                 bg="#f3f4f6", fg="#374151").grid(row=3, column=0, sticky="w")
        self.en_role = ttk.Combobox(inner, values=["user", "admin"], font=("Segoe UI", 11), width=22, state="readonly")
        self.en_role.grid(row=3, column=1, sticky="w", padx=(8, 0))
        self.en_role.set("user")

        tk.Button(add_frame, text="➕ Добавить пользователя", font=("Segoe UI", 11, "bold"),
                  bg="#10b981", fg="white", bd=0, padx=20, pady=8, cursor="hand2",
                  command=self.add_user).pack(pady=(0, 12))

        btn_frame = tk.Frame(self.win, bg="#f3f4f6")
        btn_frame.pack(fill="x", padx=20, pady=(5, 15))

        tk.Button(btn_frame, text="🗑 Удалить", font=("Segoe UI", 10, "bold"),
                  bg="#ef4444", fg="white", bd=0, padx=20, pady=8, cursor="hand2",
                  command=self.delete_user).pack(side="left", padx=(0, 10))

        tk.Button(btn_frame, text="Закрыть", font=("Segoe UI", 10, "bold"),
                  bg="#6b7280", fg="white", bd=0, padx=25, pady=8,
                  command=self.win.destroy).pack(side="right")

    def refresh_tree(self):
        for item in self.tree.get_children():
            self.tree.delete(item)
        for i, u in enumerate(self.load_users(), 1):
            self.tree.insert("", "end", values=(
                i, u.get("username", ""), u.get("name", ""), u.get("role", "user")
            ))

    def add_user(self):
        login = self.en_login.get().strip()
        name = self.en_name.get().strip()
        pwd = self.en_pwd.get().strip()
        role = self.en_role.get()

        if not login or not pwd:
            messagebox.showwarning("Ошибка", "Введите логин и пароль.")
            return

        users = self.load_users()
        if any(u["username"] == login for u in users):
            messagebox.showwarning("Ошибка", f"Пользователь '{login}' уже существует.")
            return

        users.append({
            "username": login,
            "password": hash_password(pwd),
            "name": name or login,
            "role": role
        })
        self.save_users(users)
        self.refresh_tree()
        self.en_login.delete(0, tk.END)
        self.en_name.delete(0, tk.END)
        self.en_pwd.delete(0, tk.END)
        messagebox.showinfo("Готово", f"Пользователь '{login}' добавлен.")

    def delete_user(self):
        sel = self.tree.selection()
        if not sel:
            messagebox.showwarning("Выбор", "Выберите пользователя для удаления.")
            return
        idx = self.tree.index(sel[0])
        users = self.load_users()
        u = users[idx]
        if u.get("username") == "admin":
            messagebox.showwarning("Ошибка", "Администратора удалить нельзя.")
            return
        if messagebox.askyesno("Подтверждение", f"Удалить пользователя '{u.get('username')}'?"):
            del users[idx]
            self.save_users(users)
            self.refresh_tree()


class TemplateEditor:
    def __init__(self, parent, app):
        self.app = app
        self.win = tk.Toplevel(parent)
        self.win.title("📝 Шаблоны полей")
        self.win.geometry("650x520")
        self.win.configure(bg="#f3f4f6")
        self.win.transient(parent)
        self.win.grab_set()
        self.build_ui()

    def load_templates(self):
        if os.path.exists(TEMPLATES_FILE):
            try:
                with open(TEMPLATES_FILE, "r", encoding="utf-8") as f:
                    return json.load(f)
            except:
                pass
        return {
            "situation": "Происшествий не было. Объект находится под охраной. Средства связи и наблюдения исправны.",
            "items": "Ключи от поста №1, №2; журнал учёта; рация Motorola"
        }

    def save_templates(self, data):
        with open(TEMPLATES_FILE, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

    def build_ui(self):
        tk.Label(self.win, text="📝 Шаблоны полей по умолчанию",
                 font=("Segoe UI", 16, "bold"), bg="#f3f4f6", fg="#111827").pack(pady=(15, 5))
        tk.Label(self.win, text="Эти тексты будут подставляться в поля при создании рапорта",
                 font=("Segoe UI", 10), bg="#f3f4f6", fg="#6b7280").pack(pady=(0, 10))

        tmpl = self.load_templates()

        tk.Label(self.win, text="Шаблон 'Состояние дел на посту / объекте':",
                 font=("Segoe UI", 10, "bold"), bg="#f3f4f6", fg="#374151").pack(anchor="w", padx=20, pady=(10, 4))
        self.txt_situation = scrolledtext.ScrolledText(self.win, wrap=tk.WORD, font=("Segoe UI", 11),
                                                        bg="white", fg="#1f2937", height=6,
                                                        padx=10, pady=10, borderwidth=1, relief="solid")
        self.txt_situation.pack(fill="x", padx=20, pady=(0, 10))
        self.txt_situation.insert("1.0", tmpl.get("situation", ""))

        tk.Label(self.win, text="Шаблон 'Переданные материалы / ключи / документы':",
                 font=("Segoe UI", 10, "bold"), bg="#f3f4f6", fg="#374151").pack(anchor="w", padx=20, pady=(10, 4))
        self.txt_items = scrolledtext.ScrolledText(self.win, wrap=tk.WORD, font=("Segoe UI", 11),
                                                    bg="white", fg="#1f2937", height=6,
                                                    padx=10, pady=10, borderwidth=1, relief="solid")
        self.txt_items.pack(fill="x", padx=20, pady=(0, 10))
        self.txt_items.insert("1.0", tmpl.get("items", ""))

        btn_frame = tk.Frame(self.win, bg="#f3f4f6")
        btn_frame.pack(fill="x", padx=20, pady=(10, 15))

        tk.Button(btn_frame, text="💾 Сохранить шаблоны", font=("Segoe UI", 11, "bold"),
                  bg="#2563eb", fg="white", bd=0, padx=25, pady=8, cursor="hand2",
                  command=self.save).pack(side="left", padx=(0, 10))

        tk.Button(btn_frame, text="Закрыть", font=("Segoe UI", 11, "bold"),
                  bg="#6b7280", fg="white", bd=0, padx=25, pady=8,
                  command=self.win.destroy).pack(side="right")

    def save(self):
        data = {
            "situation": self.txt_situation.get("1.0", "end-1c").strip(),
            "items": self.txt_items.get("1.0", "end-1c").strip()
        }
        self.save_templates(data)
        self.app.update_templates()
        messagebox.showinfo("Готово", "Шаблоны сохранены.")


class RaportApp:
    def __init__(self, root, user):
        self.root = root
        self.user = user
        self.is_admin = user.get("role") == "admin"

        self.root.title(f"Рапорт сдачи дежурства — {user.get('name', user.get('username', ''))}")
        self.root.state("zoomed")
        self.root.minsize(900, 700)

        self.bg_color = "#f3f4f6"
        self.accent = "#2563eb"
        self.text_bg = "#ffffff"
        self.root.configure(bg=self.bg_color)

        self.personnel = self.load_personnel()
        self.field_history = self.load_field_history()

        self.build_ui()
        self.set_defaults()

    def load_personnel(self):
        if os.path.exists(PERSONNEL_FILE):
            try:
                with open(PERSONNEL_FILE, "r", encoding="utf-8") as f:
                    return json.load(f)
            except:
                return []
        return []

    def save_personnel(self):
        with open(PERSONNEL_FILE, "w", encoding="utf-8") as f:
            json.dump(self.personnel, f, ensure_ascii=False, indent=2)

    def load_field_history(self):
        if os.path.exists(HISTORY_FILE):
            try:
                with open(HISTORY_FILE, "r", encoding="utf-8") as f:
                    return json.load(f)
            except:
                return {}
        return {}

    def save_field_history(self):
        with open(HISTORY_FILE, "w", encoding="utf-8") as f:
            json.dump(self.field_history, f, ensure_ascii=False, indent=2)

    def add_to_history(self, field_name, value):
        if not value or len(value) < 2:
            return
        if field_name not in self.field_history:
            self.field_history[field_name] = []
        hist = self.field_history[field_name]
        if value in hist:
            hist.remove(value)
        hist.insert(0, value)
        self.field_history[field_name] = hist[:30]
        self.save_field_history()

    def get_combo_values(self, field_name):
        hist = self.field_history.get(field_name, [])
        if field_name in ("to_whom_post", "to_whom_rank", "to_whom_name", "from_whom", "accepted_by"):
            persons = [f"{p.get('rank','')} {p.get('name','')}".strip() for p in self.personnel if p.get('name')]
            seen = set()
            result = []
            for v in persons + hist:
                if v and v not in seen:
                    seen.add(v)
                    result.append(v)
            return result
        return hist

    def load_templates(self):
        if os.path.exists(TEMPLATES_FILE):
            try:
                with open(TEMPLATES_FILE, "r", encoding="utf-8") as f:
                    return json.load(f)
            except:
                pass
        return {
            "situation": "Происшествий не было. Объект находится под охраной. Средства связи и наблюдения исправны.",
            "items": "Ключи от поста №1, №2; журнал учёта; рация Motorola"
        }

    def update_templates(self):
        tmpl = self.load_templates()
        if "situation" in self.fields and isinstance(self.fields["situation"], tk.Text):
            self.fields["situation"].delete("1.0", tk.END)
            self.fields["situation"].insert("1.0", tmpl.get("situation", ""))
        if "items" in self.fields and isinstance(self.fields["items"], tk.Text):
            self.fields["items"].delete("1.0", tk.END)
            self.fields["items"].insert("1.0", tmpl.get("items", ""))

    def build_ui(self):
        tmpl = self.load_templates()

        header = tk.Frame(self.root, bg=self.bg_color)
        header.pack(fill="x", padx=20, pady=(10, 6))

        role_label = " (Администратор)" if self.is_admin else ""
        tk.Label(header, text=f"📋 Рапорт сдачи дежурства{role_label}",
                 font=("Segoe UI", 18, "bold"), bg=self.bg_color, fg="#111827").pack()
        tk.Label(header, text=f"Пользователь: {self.user.get('name', self.user.get('username', ''))}",
                 font=("Segoe UI", 10), bg=self.bg_color, fg="#6b7280").pack()

        top_btns = tk.Frame(self.root, bg=self.bg_color)
        top_btns.pack(fill="x", padx=20, pady=(0, 5))

        tk.Button(top_btns, text="📖 Журнал сдачи дежурств", font=("Segoe UI", 10, "bold"),
                  bg="#4f46e5", fg="white", bd=0, padx=12, pady=5, cursor="hand2",
                  command=self.open_journal).pack(side="left", padx=(0, 6))

        tk.Button(top_btns, text="👥 Список дежурных", font=("Segoe UI", 10, "bold"),
                  bg="#0891b2", fg="white", bd=0, padx=12, pady=5, cursor="hand2",
                  command=self.open_personnel).pack(side="left", padx=(0, 6))

        tk.Button(top_btns, text="📤 Экспорт журнала", font=("Segoe UI", 10),
                  bg="#7c3aed", fg="white", bd=0, padx=12, pady=5, cursor="hand2",
                  command=self.export_journal).pack(side="left", padx=(0, 6))

        if self.is_admin:
            tk.Button(top_btns, text="👤 Пользователи", font=("Segoe UI", 10, "bold"),
                      bg="#dc2626", fg="white", bd=0, padx=12, pady=5, cursor="hand2",
                      command=self.open_users).pack(side="left", padx=(0, 6))

            tk.Button(top_btns, text="📝 Шаблоны полей", font=("Segoe UI", 10, "bold"),
                      bg="#ea580c", fg="white", bd=0, padx=12, pady=5, cursor="hand2",
                      command=self.open_templates).pack(side="left")

        tk.Button(top_btns, text="🚪 Выход", font=("Segoe UI", 10),
                  bg="#6b7280", fg="white", bd=0, padx=12, pady=5, cursor="hand2",
                  command=self.logout).pack(side="right")

        main_frame = tk.Frame(self.root, bg=self.bg_color)
        main_frame.pack(fill="both", expand=True, padx=20, pady=5)
        main_frame.columnconfigure(0, weight=2, minsize=420)
        main_frame.columnconfigure(1, weight=3, minsize=500)
        main_frame.rowconfigure(0, weight=1)

        left_frame = tk.Frame(main_frame, bg=self.bg_color)
        left_frame.grid(row=0, column=0, sticky="nsew", padx=(0, 10))
        left_frame.rowconfigure(0, weight=1)
        left_frame.columnconfigure(0, weight=1)

        canvas = tk.Canvas(left_frame, bg=self.bg_color, highlightthickness=0)
        scrollbar = ttk.Scrollbar(left_frame, orient="vertical", command=canvas.yview)
        form_frame = tk.Frame(canvas, bg=self.bg_color)

        def on_canvas_configure(event):
            canvas.itemconfig(canvas_window, width=event.width)
        canvas.bind("<Configure>", on_canvas_configure)

        form_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas_window = canvas.create_window((0, 0), window=form_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        canvas.grid(row=0, column=0, sticky="nsew")
        scrollbar.grid(row=0, column=1, sticky="ns")

        def on_mousewheel(event):
            canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        canvas.bind_all("<MouseWheel>", on_mousewheel)

        self.fields = {}
        row = 0

        self.add_label(form_frame, "Кому (должность, звание, ФИО):", row)
        row += 1
        self.add_label(form_frame, "  Должность (1-я строка):", row)
        self.fields['to_whom_post'] = self.add_combo(form_frame, "to_whom_post",
            "Начальнику караула", row+1)
        row += 2
        self.add_label(form_frame, "  Звание (2-я строка):", row)
        self.fields['to_whom_rank'] = self.add_combo(form_frame, "to_whom_rank",
            "капитану", row+1)
        row += 2
        self.add_label(form_frame, "  ФИО (3-я строка):", row)
        self.fields['to_whom_name'] = self.add_combo(form_frame, "to_whom_name",
            "Иванову И.И.", row+1)
        row += 2

        self.add_label(form_frame, "Кто сдаёт (должность, ФИО):", row)
        self.fields['from_whom'] = self.add_combo(form_frame, "from_whom",
            "Сержант Петров П.П.", row+1)
        row += 2

        date_time_frame = tk.Frame(form_frame, bg=self.bg_color)
        date_time_frame.grid(row=row, column=0, sticky="ew", pady=(4,0))
        date_time_frame.columnconfigure(0, weight=1)
        date_time_frame.columnconfigure(1, weight=1)
        self.add_label_inline(date_time_frame, "Дата начала:", 0, 0)
        self.fields['date_from'] = self.add_date_entry(date_time_frame, 0, 1)
        self.add_label_inline(date_time_frame, "Дата окончания:", 1, 0)
        self.fields['date_to'] = self.add_date_entry(date_time_frame, 1, 1)
        row += 1

        time_frame = tk.Frame(form_frame, bg=self.bg_color)
        time_frame.grid(row=row, column=0, sticky="ew", pady=(4,0))
        time_frame.columnconfigure(0, weight=1)
        time_frame.columnconfigure(1, weight=1)
        self.add_label_inline(time_frame, "Время сдачи:", 0, 0)
        self.fields['time'] = self.add_time_entry(time_frame, 0, 1)
        row += 1

        self.add_label(form_frame, "Состояние дел на посту / объекте:", row)
        self.fields['situation'] = self.add_text(form_frame,
            tmpl.get("situation", ""), row+1)
        row += 2

        self.add_label(form_frame, "Переданные материалы / ключи / документы:", row)
        self.fields['items'] = self.add_text(form_frame,
            tmpl.get("items", ""), row+1)
        row += 2

        self.add_label(form_frame, "Особые замечания (если есть):", row)
        self.fields['notes'] = self.add_text(form_frame, "Замечаний нет.", row+1)
        row += 2

        self.add_label(form_frame, "Кто принял (должность, ФИО):", row)
        self.fields['accepted_by'] = self.add_combo(form_frame, "accepted_by",
            "Старший сержант Сидоров С.С.", row+1)
        row += 2

        self.add_label(form_frame, "Номер рапорта (при наличии):", row)
        self.fields['raport_num'] = self.add_combo(form_frame, "raport_num",
            "№ 15", row+1, person_field=False)
        row += 2

        gen_btn = tk.Button(form_frame, text="✍️ Сформировать рапорт",
                            font=("Segoe UI", 12, "bold"),
                            bg=self.accent, fg="white", bd=0, padx=20, pady=10,
                            cursor="hand2", command=self.generate_raport)
        gen_btn.grid(row=row, column=0, sticky="ew", pady=(8, 0))

        right_frame = tk.Frame(main_frame, bg=self.bg_color)
        right_frame.grid(row=0, column=1, sticky="nsew", padx=(10, 0))
        right_frame.rowconfigure(1, weight=1)
        right_frame.columnconfigure(0, weight=1)

        btn_bar = tk.Frame(right_frame, bg=self.bg_color)
        btn_bar.pack(fill="x", pady=(0, 6))

        tk.Label(btn_bar, text="Готовый рапорт:",
                 font=("Segoe UI", 12, "bold"), bg=self.bg_color, fg="#374151").pack(side="left")

        btn_copy = tk.Button(btn_bar, text="📋 Копировать",
                             font=("Segoe UI", 9, "bold"),
                             bg="#10b981", fg="white", bd=0, padx=12, pady=5,
                             cursor="hand2", command=self.copy_to_clipboard)
        btn_copy.pack(side="right", padx=(6, 0))

        btn_word = tk.Button(btn_bar, text="📝 Открыть в Word",
                             font=("Segoe UI", 9, "bold"),
                             bg="#2b579a", fg="white", bd=0, padx=12, pady=5,
                             cursor="hand2", command=self.open_in_word)
        btn_word.pack(side="right", padx=(6, 0))

        btn_save = tk.Button(btn_bar, text="💾 Сохранить .docx",
                             font=("Segoe UI", 9, "bold"),
                             bg="#7c3aed", fg="white", bd=0, padx=12, pady=5,
                             cursor="hand2", command=self.save_as_docx)
        btn_save.pack(side="right")

        self.result_text = scrolledtext.ScrolledText(
            right_frame, wrap=tk.WORD, font=("Georgia", 11),
            bg=self.text_bg, fg="#1f2937", padx=15, pady=15,
            borderwidth=1, relief="solid"
        )
        self.result_text.pack(fill="both", expand=True)

        self.status_var = tk.StringVar()
        self.status_var.set("")
        self.status_label = tk.Label(self.root, textvariable=self.status_var,
                                      font=("Segoe UI", 9), bg=self.bg_color, fg="#10b981")
        self.status_label.pack(pady=(0, 6))

        if not DOCX_AVAILABLE:
            hint = tk.Label(self.root,
                text="⚠️ Для работы с Word установите: pip install python-docx",
                font=("Segoe UI", 9), bg=self.bg_color, fg="#f59e0b")
            hint.pack(pady=(0, 5))

    def add_label(self, parent, text, row):
        lbl = tk.Label(parent, text=text, font=("Segoe UI", 10, "bold"),
                       bg=self.bg_color, fg="#374151", anchor="w")
        lbl.grid(row=row, column=0, sticky="w", pady=(8, 2))

    def add_label_inline(self, parent, text, col, row):
        lbl = tk.Label(parent, text=text, font=("Segoe UI", 10, "bold"),
                       bg=self.bg_color, fg="#374151", anchor="w")
        lbl.grid(row=row, column=col, sticky="w", padx=(0, 5))

    def add_combo(self, parent, field_name, placeholder, row, person_field=True):
        values = self.get_combo_values(field_name)
        combo = ttk.Combobox(parent, font=("Segoe UI", 11), values=values, height=10)
        combo.grid(row=row, column=0, sticky="ew", pady=(0, 4))
        combo.set(placeholder)
        style = ttk.Style()
        style.configure("TCombobox", fieldbackground="white", background="white")
        combo.bind("<Button-1>", lambda e: self.on_combo_click(combo, field_name, person_field))
        combo.bind("<FocusIn>", lambda e: self.on_combo_focus(combo, field_name, person_field))
        combo.bind("<KeyRelease>", lambda e: self.on_combo_type(combo, field_name, person_field))
        combo._field_name = field_name
        combo._placeholder = placeholder
        combo._person_field = person_field
        return combo

    def on_combo_click(self, combo, field_name, person_field):
        values = self.get_combo_values(field_name)
        combo["values"] = values
        if combo.get() == combo._placeholder:
            combo.set("")

    def on_combo_focus(self, combo, field_name, person_field):
        values = self.get_combo_values(field_name)
        combo["values"] = values

    def on_combo_type(self, combo, field_name, person_field):
        text = combo.get().lower()
        all_values = self.get_combo_values(field_name)
        filtered = [v for v in all_values if text in v.lower()]
        combo["values"] = filtered

    def add_date_entry(self, parent, col, row):
        entry = tk.Entry(parent, font=("Segoe UI", 11), bg="white", fg="#1f2937",
                         relief="solid", bd=1, width=15,
                         highlightcolor=self.accent, highlightbackground="#d1d5db")
        entry.grid(row=row, column=col, sticky="w", padx=(0, 10))
        return entry

    def add_time_entry(self, parent, col, row):
        entry = tk.Entry(parent, font=("Segoe UI", 11), bg="white", fg="#1f2937",
                         relief="solid", bd=1, width=10,
                         highlightcolor=self.accent, highlightbackground="#d1d5db")
        entry.grid(row=row, column=col, sticky="w")
        return entry

    def add_text(self, parent, placeholder, row):
        text = tk.Text(parent, font=("Segoe UI", 11), bg="white", fg="#1f2937",
                       relief="solid", bd=1, height=2, wrap=tk.WORD,
                       highlightcolor=self.accent, highlightbackground="#d1d5db")
        text.grid(row=row, column=0, sticky="ew", pady=(0, 4))
        text.insert("1.0", placeholder)
        text.bind("<FocusIn>", lambda e, w=text, p=placeholder: self.on_text_focus_in(w, p))
        text.bind("<FocusOut>", lambda e, w=text, p=placeholder: self.on_text_focus_out(w, p))
        return text

    def on_text_focus_in(self, widget, placeholder):
        if widget.get("1.0", "end-1c") == placeholder:
            widget.delete("1.0", tk.END)
            widget.config(fg="#1f2937")

    def on_text_focus_out(self, widget, placeholder):
        if widget.get("1.0", "end-1c").strip() == "":
            widget.insert("1.0", placeholder)
            widget.config(fg="#9ca3af")

    def set_defaults(self):
        now = datetime.now()
        self.fields['date_from'].delete(0, tk.END)
        self.fields['date_from'].insert(0, now.strftime("%d.%m.%Y"))
        self.fields['date_to'].delete(0, tk.END)
        self.fields['date_to'].insert(0, now.strftime("%d.%m.%Y"))
        self.fields['time'].delete(0, tk.END)
        self.fields['time'].insert(0, now.strftime("%H:%M"))

    def get_value(self, field, placeholder=""):
        if isinstance(field, ttk.Combobox):
            val = field.get().strip()
            if val == placeholder:
                return ""
            return val
        elif isinstance(field, tk.Entry):
            val = field.get().strip()
            if val == placeholder:
                return ""
            return val
        else:
            val = field.get("1.0", "end-1c").strip()
            if val == placeholder:
                return ""
            return val

    def logout(self):
        if messagebox.askyesno("Выход", "Выйти из системы?"):
            self.root.destroy()
            root = tk.Tk()
            LoginWindow(root)
            root.mainloop()

    def open_users(self):
        UsersManager(self.root, self)

    def open_templates(self):
        TemplateEditor(self.root, self)

    def open_personnel(self):
        win = tk.Toplevel(self.root)
        win.title("👥 Список дежурных")
        win.geometry("700x720")
        win.minsize(600, 600)
        win.configure(bg=self.bg_color)
        win.transient(self.root)
        win.grab_set()

        tk.Label(win, text="👥 Список дежурных",
                 font=("Segoe UI", 16, "bold"), bg=self.bg_color, fg="#111827").pack(pady=(15, 5))
        tk.Label(win, text="Добавьте людей — они появятся в выпадающих списках полей",
                 font=("Segoe UI", 10), bg=self.bg_color, fg="#6b7280").pack(pady=(0, 10))

        table_frame = tk.Frame(win, bg=self.bg_color)
        table_frame.pack(fill="both", expand=True, padx=20, pady=5)

        columns = ("№", "Фамилия И.О.", "Звание", "Должность")
        tree = ttk.Treeview(table_frame, columns=columns, show="headings", height=10)
        tree.heading("№", text="№")
        tree.heading("Фамилия И.О.", text="Фамилия И.О.")
        tree.heading("Звание", text="Звание")
        tree.heading("Должность", text="Должность")
        tree.column("№", width=40, anchor="center")
        tree.column("Фамилия И.О.", width=200)
        tree.column("Звание", width=120)
        tree.column("Должность", width=200)

        scrollbar = ttk.Scrollbar(table_frame, orient="vertical", command=tree.yview)
        tree.configure(yscrollcommand=scrollbar.set)

        def refresh_tree():
            for item in tree.get_children():
                tree.delete(item)
            for i, p in enumerate(self.personnel, 1):
                tree.insert("", "end", values=(
                    i, p.get("name", ""), p.get("rank", ""), p.get("post", "")
                ))

        refresh_tree()
        tree.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        tk.Frame(win, bg="#d1d5db", height=1).pack(fill="x", padx=20, pady=(10, 10))

        add_frame = tk.LabelFrame(win, text=" Добавить нового дежурного ",
                                   font=("Segoe UI", 11, "bold"),
                                   bg=self.bg_color, fg="#374151", bd=1, relief="solid")
        add_frame.pack(fill="x", padx=20, pady=(0, 10))

        inner = tk.Frame(add_frame, bg=self.bg_color)
        inner.pack(fill="x", padx=15, pady=12)

        tk.Label(inner, text="Фамилия И.О.:", font=("Segoe UI", 10, "bold"),
                 bg=self.bg_color, fg="#374151").grid(row=0, column=0, sticky="w", pady=(0, 5))
        name_entry = tk.Entry(inner, font=("Segoe UI", 11), width=30)
        name_entry.grid(row=0, column=1, sticky="ew", padx=(8, 0), pady=(0, 5))

        tk.Label(inner, text="Звание:", font=("Segoe UI", 10, "bold"),
                 bg=self.bg_color, fg="#374151").grid(row=1, column=0, sticky="w", pady=(0, 5))
        rank_entry = tk.Entry(inner, font=("Segoe UI", 11), width=30)
        rank_entry.grid(row=1, column=1, sticky="ew", padx=(8, 0), pady=(0, 5))

        tk.Label(inner, text="Должность:", font=("Segoe UI", 10, "bold"),
                 bg=self.bg_color, fg="#374151").grid(row=2, column=0, sticky="w", pady=(0, 5))
        post_entry = tk.Entry(inner, font=("Segoe UI", 11), width=30)
        post_entry.grid(row=2, column=1, sticky="ew", padx=(8, 0), pady=(0, 5))

        inner.columnconfigure(1, weight=1)

        add_btn = tk.Button(add_frame, text="➕ Добавить в список", font=("Segoe UI", 11, "bold"),
                            bg="#10b981", fg="white", bd=0, padx=25, pady=8, cursor="hand2")
        add_btn.pack(pady=(0, 12))

        def add_person():
            name = name_entry.get().strip()
            rank = rank_entry.get().strip()
            post = post_entry.get().strip()
            if not name:
                messagebox.showwarning("Ошибка", "Введите фамилию и инициалы.")
                return
            self.personnel.append({"name": name, "rank": rank, "post": post})
            self.save_personnel()
            refresh_tree()
            name_entry.delete(0, tk.END)
            rank_entry.delete(0, tk.END)
            post_entry.delete(0, tk.END)
            self.refresh_all_combos()
            messagebox.showinfo("Готово", f"{name} добавлен в список.")

        add_btn.config(command=add_person)

        btn_frame = tk.Frame(win, bg=self.bg_color)
        btn_frame.pack(fill="x", padx=20, pady=(5, 15))

        def delete_person():
            sel = tree.selection()
            if not sel:
                messagebox.showwarning("Выбор", "Выберите запись для удаления.")
                return
            idx = tree.index(sel[0])
            name = self.personnel[idx].get("name", "")
            if messagebox.askyesno("Подтверждение", f"Удалить {name} из списка?"):
                del self.personnel[idx]
                self.save_personnel()
                refresh_tree()
                self.refresh_all_combos()

        def edit_person():
            sel = tree.selection()
            if not sel:
                messagebox.showwarning("Выбор", "Выберите запись для редактирования.")
                return
            idx = tree.index(sel[0])
            p = self.personnel[idx]

            edit_win = tk.Toplevel(win)
            edit_win.title("Редактирование")
            edit_win.geometry("450x250")
            edit_win.configure(bg=self.bg_color)
            edit_win.transient(win)
            edit_win.grab_set()

            tk.Label(edit_win, text="Фамилия И.О.:", font=("Segoe UI", 10, "bold"),
                     bg=self.bg_color).grid(row=0, column=0, sticky="w", padx=20, pady=(20, 5))
            en_name = tk.Entry(edit_win, font=("Segoe UI", 11), width=35)
            en_name.grid(row=0, column=1, sticky="w", pady=(20, 5))
            en_name.insert(0, p.get("name", ""))

            tk.Label(edit_win, text="Звание:", font=("Segoe UI", 10, "bold"),
                     bg=self.bg_color).grid(row=1, column=0, sticky="w", padx=20, pady=5)
            en_rank = tk.Entry(edit_win, font=("Segoe UI", 11), width=35)
            en_rank.grid(row=1, column=1, sticky="w", pady=5)
            en_rank.insert(0, p.get("rank", ""))

            tk.Label(edit_win, text="Должность:", font=("Segoe UI", 10, "bold"),
                     bg=self.bg_color).grid(row=2, column=0, sticky="w", padx=20, pady=5)
            en_post = tk.Entry(edit_win, font=("Segoe UI", 11), width=35)
            en_post.grid(row=2, column=1, sticky="w", pady=5)
            en_post.insert(0, p.get("post", ""))

            def save_edit():
                p["name"] = en_name.get().strip()
                p["rank"] = en_rank.get().strip()
                p["post"] = en_post.get().strip()
                if not p["name"]:
                    messagebox.showwarning("Ошибка", "Фамилия не может быть пустой.")
                    return
                self.save_personnel()
                refresh_tree()
                self.refresh_all_combos()
                edit_win.destroy()

            tk.Button(edit_win, text="💾 Сохранить изменения", font=("Segoe UI", 10, "bold"),
                      bg=self.accent, fg="white", bd=0, padx=25, pady=8,
                      command=save_edit).grid(row=3, column=0, columnspan=2, pady=20)

        tk.Button(btn_frame, text="✏️ Редактировать", font=("Segoe UI", 10, "bold"),
                  bg="#f59e0b", fg="white", bd=0, padx=15, pady=6, cursor="hand2",
                  command=edit_person).pack(side="left", padx=(0, 6))

        tk.Button(btn_frame, text="🗑 Удалить", font=("Segoe UI", 10, "bold"),
                  bg="#ef4444", fg="white", bd=0, padx=15, pady=6, cursor="hand2",
                  command=delete_person).pack(side="left", padx=(0, 6))

        tk.Button(btn_frame, text="Закрыть", font=("Segoe UI", 10, "bold"),
                  bg="#6b7280", fg="white", bd=0, padx=15, pady=6,
                  command=win.destroy).pack(side="right")

    def refresh_all_combos(self):
        for field_name, widget in self.fields.items():
            if isinstance(widget, ttk.Combobox):
                widget["values"] = self.get_combo_values(field_name)

    def parse_rank_name(self, full_text):
        if not full_text or full_text.strip() == "":
            return "", ""
        full_text = full_text.strip()
        ranks = [
            "старший прапорщик", "старший сержант", "младший сержант", "младший лейтенант",
            "старший лейтенант", "генерал-полковник", "генерал-лейтенант", "генерал-майор",
            "генерал армии", "ефрейтор", "рядовой", "сержант", "старшина", "прапорщик",
            "лейтенант", "капитан", "майор", "подполковник", "полковник"
        ]
        lower = full_text.lower()
        for rank in ranks:
            if lower.startswith(rank):
                name = full_text[len(rank):].strip()
                return full_text[:len(rank)], name
        parts = full_text.split(None, 1)
        if len(parts) == 2:
            return parts[0], parts[1]
        return full_text, ""

    def generate_raport(self):
        to_whom_post = self.get_value(self.fields['to_whom_post'], "Начальнику караула")
        to_whom_rank = self.get_value(self.fields['to_whom_rank'], "капитану")
        to_whom_name = self.get_value(self.fields['to_whom_name'], "Иванову И.И.")
        from_whom = self.get_value(self.fields['from_whom'], "Сержант Петров П.П.")
        date_from = self.get_value(self.fields['date_from'])
        date_to = self.get_value(self.fields['date_to'])
        time_val = self.get_value(self.fields['time'])
        situation = self.get_value(self.fields['situation'],
            "Происшествий не было. Объект находится под охраной. Средства связи и наблюдения исправны.")
        items = self.get_value(self.fields['items'],
            "Ключи от поста №1, №2; журнал учёта; рация Motorola")
        notes = self.get_value(self.fields['notes'], "Замечаний нет.")
        accepted_by = self.get_value(self.fields['accepted_by'], "Старший сержант Сидоров С.С.")
        raport_num = self.get_value(self.fields['raport_num'], "№ 15")

        if not (to_whom_post or to_whom_name) or not from_whom or not date_from or not date_to or not time_val:
            messagebox.showwarning("Незаполненные поля",
                "Пожалуйста, заполните обязательные поля: Кому, Кто сдаёт, Даты и Время.")
            return

        self.add_to_history("to_whom_post", to_whom_post)
        self.add_to_history("to_whom_rank", to_whom_rank)
        self.add_to_history("to_whom_name", to_whom_name)
        self.add_to_history("from_whom", from_whom)
        self.add_to_history("accepted_by", accepted_by)
        self.add_to_history("raport_num", raport_num)

        def fmt_date(dval):
            try:
                if "." in dval:
                    parts = dval.split(".")
                    day, month, year = int(parts[0]), int(parts[1]), int(parts[2])
                elif "-" in dval:
                    parts = dval.split("-")
                    if len(parts[0]) == 4:
                        year, month, day = int(parts[0]), int(parts[1]), int(parts[2])
                    else:
                        day, month, year = int(parts[0]), int(parts[1]), int(parts[2])
                else:
                    return dval
                months = ['января','февраля','марта','апреля','мая','июня',
                          'июля','августа','сентября','октября','ноября','декабря']
                if 1 <= month <= 12 and year > 0:
                    return f"{day} {months[month-1]} {year} г."
                return dval
            except:
                return dval

        fmt_from = fmt_date(date_from)
        fmt_to = fmt_date(date_to)
        if date_from == date_to:
            formatted_date = fmt_from
        else:
            formatted_date = f"с {fmt_from} по {fmt_to}"

        lines = []
        if raport_num:
            lines.append(f"Рапорт {raport_num}")
            lines.append("")
        if to_whom_post:
            lines.append(to_whom_post)
        rank_name = " ".join(filter(None, [to_whom_rank, to_whom_name]))
        if rank_name:
            lines.append(rank_name)
        lines.append("")
        lines.append(f"Докладываю, что {formatted_date} в {time_val} мною, {from_whom}, сдано дежурство.")

        if situation:
            lines.append("")
            lines.append("Состояние дел:")
            lines.append(situation)

        if items:
            lines.append("")
            lines.append("Передано:")
            lines.append(items)

        if notes and notes.lower() not in ("замечаний нет.", "замечаний нет"):
            lines.append("")
            lines.append("Особые замечания:")
            lines.append(notes)

        rank_from, name_from = self.parse_rank_name(from_whom)
        rank_acc, name_acc = self.parse_rank_name(accepted_by)

        lines.append("")
        lines.append(f"Дежурство сдал:   {rank_from} ___________ {name_from}")
        lines.append(f"Дежурство принял: {rank_acc or '___________'} ___________ {name_acc or '___________'}")
        lines.append("")
        lines.append("«___» __________ 20__ г.")

        result = "\\n".join(lines)

        self.result_text.delete("1.0", tk.END)
        self.result_text.insert("1.0", result)

        journal = self.load_journal()
        record = {
            "id": len(journal) + 1,
            "timestamp": datetime.now().isoformat(),
            "date_from": date_from,
            "date_to": date_to,
            "time": time_val,
            "to_whom_post": to_whom_post,
            "to_whom_rank": to_whom_rank,
            "to_whom_name": to_whom_name,
            "from_whom": from_whom,
            "accepted_by": accepted_by,
            "raport_num": raport_num,
            "situation": situation,
            "items": items,
            "notes": notes,
            "raport_text": result,
            "user": self.user.get("username", "")
        }
        journal.append(record)
        self.save_journal(journal)

        self.status_var.set(f"✅ Рапорт сформирован и записан в журнал (запись №{record['id']})")
        self.root.after(4000, lambda: self.status_var.set(""))
        self.refresh_all_combos()

    def copy_to_clipboard(self):
        text = self.result_text.get("1.0", "end-1c")
        if not text.strip():
            messagebox.showinfo("Информация", "Сначала сформируйте рапорт.")
            return
        self.root.clipboard_clear()
        self.root.clipboard_append(text)
        self.status_var.set("📋 Текст скопирован в буфер обмена")
        self.root.after(2000, lambda: self.status_var.set(""))

    def _check_docx_available(self):
        if not DOCX_AVAILABLE:
            messagebox.showerror(
                "Модуль не установлен",
                "Для работы с Word необходимо установить модуль python-docx.\\n\\n"
                "Выполните в командной строке:\\n"
                "pip install python-docx"
            )
            return False
        return True

    def _create_docx(self, text, filepath):
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

        sections = doc.sections[0]
        sections.top_margin = Inches(20/25.4)
        sections.bottom_margin = Inches(20/25.4)
        sections.left_margin = Inches(30/25.4)
        sections.right_margin = Inches(10/25.4)

        lines = text.split("\\n")
        for line in lines:
            if not line.strip():
                doc.add_paragraph()
                continue
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

            if line.strip().startswith("Рапорт"):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.bold = True
                run.font.size = Pt(16)
            elif line.strip() and not line.startswith("Докладываю") and not line.startswith("Состояние") and not line.startswith("Передано") and not line.startswith("Особые") and not line.startswith("Дежурство") and not line.startswith("«___»"):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            p.paragraph_format.line_spacing = Pt(18)
            p.paragraph_format.space_after = Pt(0)

        doc.save(filepath)

    def save_as_docx(self):
        text = self.result_text.get("1.0", "end-1c")
        if not text.strip():
            messagebox.showinfo("Информация", "Сначала сформируйте рапорт.")
            return
        if not self._check_docx_available():
            return

        date_val = self.get_value(self.fields['date_from'])
        raport_num = self.get_value(self.fields['raport_num'], "№ 15")
        default_name = f"Рапорт_{date_val.replace('.', '-')}_{raport_num.replace(' ', '')}.docx"

        filepath = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word документ", "*.docx"), ("Все файлы", "*.*")],
            title="Сохранить рапорт как...",
            initialfile=default_name
        )
        if not filepath:
            return
        try:
            self._create_docx(text, filepath)
            self.status_var.set(f"💾 Сохранено: {os.path.basename(filepath)}")
            self.root.after(3000, lambda: self.status_var.set(""))
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось сохранить файл:\\n{e}")

    def open_in_word(self):
        text = self.result_text.get("1.0", "end-1c")
        if not text.strip():
            messagebox.showinfo("Информация", "Сначала сформируйте рапорт.")
            return
        if not self._check_docx_available():
            return

        temp_dir = os.path.join(APP_DIR, "temp_docs")
        os.makedirs(temp_dir, exist_ok=True)

        date_val = self.get_value(self.fields['date_from'])
        raport_num = self.get_value(self.fields['raport_num'], "№ 15")
        filename = f"Рапорт_{date_val.replace('.', '-')}_{raport_num.replace(' ', '')}_{datetime.now().strftime('%H%M%S')}.docx"
        filepath = os.path.join(temp_dir, filename)

        try:
            self._create_docx(text, filepath)
            system = platform.system()
            if system == "Windows":
                os.startfile(filepath)
            elif system == "Darwin":
                subprocess.call(["open", filepath])
            else:
                subprocess.call(["xdg-open", filepath])
            self.status_var.set("📝 Рапорт открыт в Word")
            self.root.after(3000, lambda: self.status_var.set(""))
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось открыть Word:\\n{e}")

    def load_journal(self):
        if os.path.exists(DATA_FILE):
            try:
                with open(DATA_FILE, "r", encoding="utf-8") as f:
                    return json.load(f)
            except:
                return []
        return []

    def save_journal(self, data):
        with open(DATA_FILE, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

    def open_journal(self):
        journal = self.load_journal()
        if not journal:
            messagebox.showinfo("Журнал", "Журнал пока пуст. Сформируйте первый рапорт.")
            return

        win = tk.Toplevel(self.root)
        win.title("📖 Журнал сдачи дежурств")
        win.geometry("950x650")
        win.minsize(800, 500)
        win.configure(bg=self.bg_color)

        header = tk.Frame(win, bg=self.bg_color)
        header.pack(fill="x", padx=20, pady=(15, 5))
        tk.Label(header, text="📖 Журнал сдачи дежурств",
                 font=("Segoe UI", 16, "bold"), bg=self.bg_color, fg="#111827").pack()
        tk.Label(header, text=f"Всего записей: {len(journal)}",
                 font=("Segoe UI", 10), bg=self.bg_color, fg="#6b7280").pack()

        table_frame = tk.Frame(win, bg=self.bg_color)
        table_frame.pack(fill="both", expand=True, padx=20, pady=10)

        columns = ("№", "Период", "Время", "Сдал", "Принял", "Кому", "Номер рапорта")
        tree = ttk.Treeview(table_frame, columns=columns, show="headings", height=10)
        for col in columns:
            tree.heading(col, text=col)
        tree.column("№", width=40, anchor="center")
        tree.column("Период", width=110, anchor="center")
        tree.column("Время", width=70, anchor="center")
        tree.column("Сдал", width=170)
        tree.column("Принял", width=170)
        tree.column("Кому", width=220)
        tree.column("Номер рапорта", width=100, anchor="center")

        scrollbar = ttk.Scrollbar(table_frame, orient="vertical", command=tree.yview)
        tree.configure(yscrollcommand=scrollbar.set)

        for rec in journal:
            tree.insert("", "end", values=(
                rec.get("id", ""),
                f"{rec.get('date_from','')} — {rec.get('date_to','')}",
                rec.get("time", ""),
                rec.get("from_whom", ""),
                rec.get("accepted_by", ""),
                " ".join(filter(None, [rec.get("to_whom_post", ""), rec.get("to_whom_rank", ""), rec.get("to_whom_name", "")])),
                rec.get("raport_num", "—")
            ))

        tree.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        btn_frame = tk.Frame(win, bg=self.bg_color)
        btn_frame.pack(fill="x", padx=20, pady=(5, 15))

        def view_record():
            sel = tree.selection()
            if not sel:
                messagebox.showwarning("Выбор", "Выберите запись для просмотра.")
                return
            item = tree.item(sel[0])
            rec_id = item["values"][0]
            rec = next((r for r in journal if r.get("id") == rec_id), None)
            if rec:
                view_win = tk.Toplevel(win)
                view_win.title(f"Рапорт №{rec_id}")
                view_win.geometry("600x500")
                view_win.configure(bg=self.bg_color)

                tk.Label(view_win, text=f"Рапорт записи №{rec_id}",
                         font=("Segoe UI", 14, "bold"), bg=self.bg_color).pack(pady=(15, 5))

                txt = scrolledtext.ScrolledText(view_win, wrap=tk.WORD, font=("Georgia", 11),
                                                bg="white", fg="#1f2937", padx=15, pady=15,
                                                borderwidth=1, relief="solid")
                txt.pack(fill="both", expand=True, padx=20, pady=5)
                txt.insert("1.0", rec.get("raport_text", ""))
                txt.config(state="disabled")

                tk.Button(view_win, text="Закрыть", command=view_win.destroy,
                          font=("Segoe UI", 10), bg="#6b7280", fg="white", bd=0,
                          padx=20, pady=6).pack(pady=(0, 15))

        def delete_record():
            sel = tree.selection()
            if not sel:
                messagebox.showwarning("Выбор", "Выберите запись для удаления.")
                return
            item = tree.item(sel[0])
            rec_id = item["values"][0]
            if messagebox.askyesno("Подтверждение", f"Удалить запись №{rec_id}?"):
                journal_new = [r for r in journal if r.get("id") != rec_id]
                for i, r in enumerate(journal_new, 1):
                    r["id"] = i
                self.save_journal(journal_new)
                tree.delete(sel[0])
                messagebox.showinfo("Готово", "Запись удалена.")
                win.destroy()
                self.open_journal()

        def clear_all():
            if messagebox.askyesno("Подтверждение", "Очистить ВЕСЬ журнал? Это действие нельзя отменить."):
                self.save_journal([])
                messagebox.showinfo("Готово", "Журнал очищен.")
                win.destroy()

        tk.Button(btn_frame, text="👁 Просмотреть рапорт", font=("Segoe UI", 10, "bold"),
                  bg=self.accent, fg="white", bd=0, padx=20, pady=8, cursor="hand2",
                  command=view_record).pack(side="left", padx=(0, 10))

        tk.Button(btn_frame, text="🗑 Удалить запись", font=("Segoe UI", 10, "bold"),
                  bg="#ef4444", fg="white", bd=0, padx=20, pady=8, cursor="hand2",
                  command=delete_record).pack(side="left", padx=(0, 10))

        tk.Button(btn_frame, text="🧹 Очистить весь журнал", font=("Segoe UI", 10, "bold"),
                  bg="#f59e0b", fg="white", bd=0, padx=20, pady=8, cursor="hand2",
                  command=clear_all).pack(side="left")

        tk.Button(btn_frame, text="Закрыть", font=("Segoe UI", 10, "bold"),
                  bg="#6b7280", fg="white", bd=0, padx=25, pady=8,
                  command=win.destroy).pack(side="right")

    def export_journal(self):
        journal = self.load_journal()
        if not journal:
            messagebox.showinfo("Экспорт", "Журнал пуст — нечего экспортировать.")
            return

        filepath = filedialog.asksaveasfilename(
            defaultextension=".csv",
            filetypes=[("CSV файлы", "*.csv"), ("Все файлы", "*.*")],
            title="Сохранить журнал как..."
        )
        if not filepath:
            return

        try:
            with open(filepath, "w", newline="", encoding="utf-8-sig") as f:
                writer = csv.writer(f, delimiter=";")
                writer.writerow(["№", "Дата_начала", "Дата_окончания", "Время", "Сдал", "Принял", "Кому",
                                 "Номер рапорта", "Состояние дел", "Передано", "Замечания"])
                for rec in journal:
                    writer.writerow([
                        rec.get("id", ""),
                        rec.get("date_from", ""),
                        rec.get("date_to", ""),
                        rec.get("time", ""),
                        rec.get("from_whom", ""),
                        rec.get("accepted_by", ""),
                        " ".join(filter(None, [rec.get("to_whom_post", ""), rec.get("to_whom_rank", ""), rec.get("to_whom_name", "")])),
                        rec.get("raport_num", ""),
                        rec.get("situation", "").replace("\\n", " "),
                        rec.get("items", "").replace("\\n", " "),
                        rec.get("notes", "").replace("\\n", " ")
                    ])
            messagebox.showinfo("Готово", f"Журнал экспортирован:\\n{filepath}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось экспортировать:\\n{e}")


def main():
    root = tk.Tk()
    LoginWindow(root)
    root.mainloop()


if __name__ == "__main__":
    main()
