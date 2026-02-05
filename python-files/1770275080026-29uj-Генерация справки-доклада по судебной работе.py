#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from __future__ import annotations
import sys
import re
import string
from pathlib import Path
from typing import Dict, List, Tuple

import pandas as pd
from PyQt5 import QtCore, QtWidgets, QtGui
from PyQt5.QtCore import Qt, QSize
from PyQt5.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QPushButton,
    QLineEdit, QTextEdit, QFileDialog, QGroupBox, QSpinBox,
    QCheckBox, QListWidget, QStackedWidget, QMessageBox, QFrame,
    QSizePolicy, QComboBox, QScrollArea, QAbstractItemView, QListView
)
from PyQt5.QtGui import QFont, QIcon, QColor
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DEFAULT_UNITS = [
    "упр. 3 мсд", "252 мсп", "752 мсп", "245 мсп", "237 тп",
    "337 оисб", "159 оптадн", "84 орб", "99 сап"
]

STYLESHEET = """
    QWidget {
        font-family: "Segoe UI", "Roboto", "Helvetica", sans-serif;
        font-size: 10pt;
        background-color: #F2F4F8; /* Более приятный светло-серый фон */
        color: #2D2D2D;
    }
    
    QGroupBox {
        background-color: #FFFFFF;
        border: 1px solid #E1E4E8;
        border-radius: 10px;
        margin-top: 1.2em;
        padding-top: 20px;
        font-weight: 600;
        font-size: 11pt;
    }
    
    QGroupBox::title {
        subcontrol-origin: margin;
        subcontrol-position: top left;
        left: 15px;
        padding: 0 5px;
        color: #0066CC;
    }
    
    QPushButton {
        background-color: #007AFF;
        color: white;
        border: none;
        padding: 10px 24px;
        border-radius: 8px;
        font-weight: 600;
        font-size: 11pt;
        min-height: 24px;
        outline: none;
    }
    
    QPushButton:hover {
        background-color: #0062CC;
    }
    
    QPushButton:pressed {
        background-color: #004999;
        padding-top: 11px; /* Эффект нажатия */
    }
    
    QPushButton:disabled {
        background-color: #D1D5DB;
        color: #9CA3AF;
        border: 1px solid #D1D5DB;
    }
    
    QPushButton#secondary {
        background-color: #FFFFFF;
        color: #007AFF;
        border: 1px solid #D1D5DB;
    }
    
    QPushButton#secondary:hover {
        background-color: #F3F4F6;
        border-color: #007AFF;
    }

    QPushButton#danger {
        background-color: #EF4444;
        color: white;
    }
    QPushButton#danger:hover {
        background-color: #DC2626;
    }
    
    QPushButton#success {
        background-color: #10B981;
        color: white;
        font-size: 12pt;
        padding: 14px;
        border-radius: 8px;
    }
    QPushButton#success:hover {
        background-color: #059669;
    }
    
    QLineEdit, QSpinBox {
        border: 1px solid #D1D5DB;
        border-radius: 8px;
        padding: 8px 12px;
        background-color: #FFFFFF;
        selection-background-color: #007AFF;
        min-height: 20px;
    }
    
    QLineEdit:focus, QSpinBox:focus {
        border: 2px solid #007AFF;
        padding: 7px 11px; 
    }
    
    QComboBox {
        border: 1px solid #D1D5DB;
        border-radius: 8px;
        padding: 8px 12px;
        background-color: #FFFFFF;
        selection-background-color: #007AFF;
        min-height: 24px;
    }

    QComboBox:focus {
        border: 2px solid #007AFF;
        padding: 7px 11px;
    }
    
    QComboBox::drop-down {
        subcontrol-origin: padding;
        subcontrol-position: top right;
        width: 30px;
        border-left-width: 0px;
        border-top-right-radius: 8px;
        border-bottom-right-radius: 8px;
    }
    
    QComboBox::down-arrow {
        width: 0; 
        height: 0; 
        border-left: 5px solid transparent;
        border-right: 5px solid transparent;
        border-top: 6px solid #6B7280;
        margin-right: 10px;
    }
    
    QComboBox QAbstractItemView {
        border: 1px solid #D1D5DB;
        background-color: #FFFFFF;
        selection-background-color: #E6F2FF;
        selection-color: #007AFF;
        outline: none;
        padding: 5px;
        border-radius: 8px;
    }
    
    /* --- Lists & Text --- */
    QTextEdit, QListWidget {
        border: 1px solid #D1D5DB;
        border-radius: 8px;
        background-color: #FFFFFF;
        padding: 10px;
        font-family: "Consolas", "Monaco", monospace;
    }

    QFrame#Card {
        background-color: #FFFFFF;
        border-radius: 12px;
        border: 1px solid #E5E7EB;
    }
"""

def column_letter_to_index(letter: str) -> int:
    letter = letter.upper().strip()
    num = 0
    for char in letter:
        if char in string.ascii_uppercase:
            num = num * 26 + (ord(char) - ord('A')) + 1
    return num - 1

def safe_str(value) -> str:
    if value is None:
        return ""
    try:
        if pd.isna(value):
            return ""
    except:
        pass
    return str(value).strip()

def normalize_text(text: str) -> str:
    text = safe_str(text)
    text = text.replace("ё", "е").lower()
    text = re.sub(r"[\n\r\t]+", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text

def plural_ru(number: int, forms: Tuple[str, str, str]) -> str:
    number = abs(int(number))
    if 11 <= number % 100 <= 19:
        return forms[2]
    if number % 10 == 1:
        return forms[0]
    if 2 <= number % 10 <= 4:
        return forms[1]
    return forms[2]


class ExcelAnalyzer:
    def __init__(self, file_path: Path):
        self.file_path = file_path
        self.df = None
        self.sheet_names = []
        
    def load_sheets(self) -> List[str]:
        try:
            excel_file = pd.ExcelFile(self.file_path, engine='openpyxl')
            self.sheet_names = excel_file.sheet_names
            return self.sheet_names
        except Exception as e:
            raise ValueError(f"Ошибка чтения файла: {e}")
    
    def find_header_row(self, sheet_name: str) -> Tuple[int, Dict[str, int]]:
        df = pd.read_excel(self.file_path, sheet_name=sheet_name, 
                          header=None, nrows=20, engine='openpyxl')
        
        for row_idx in range(len(df)):
            row_values = [normalize_text(v) for v in df.iloc[row_idx]]
            unit_col = None
            status_col = None
            
            for col_idx, val in enumerate(row_values):
                if any(k in val for k in ["в/ч", "вч", "часть", "подразделение"]):
                    unit_col = col_idx
                if any(k in val for k in ["новая неделя"]):
                    status_col = col_idx
            
            if unit_col is not None and status_col is not None:
                return row_idx + 1, {"unit": unit_col, "status": status_col}
        
        raise ValueError("Не удалось автоматически найти заголовки.")
    
    def load_data(self, sheet_name: str, header_row: int, 
                  unit_col: int, status_col: int) -> pd.DataFrame:
        pd_header = (header_row - 1) if header_row > 0 else None
        try:
            df = pd.read_excel(
                self.file_path,
                sheet_name=sheet_name,
                header=pd_header,
                usecols=[unit_col, status_col],
                engine='openpyxl'
            )
            if df.shape[1] < 2:
                full_df = pd.read_excel(self.file_path, sheet_name=sheet_name, header=pd_header, engine='openpyxl')
                df = full_df.iloc[:, [unit_col, status_col]]

            self.df = df
            return df
        except Exception as e:
            raise ValueError(f"Ошибка загрузки: {e}")
    
    def preview_data(self, rows: int = 10) -> str:
        if self.df is None:
            return "Нет данных"
        
        lines = []
        for i in range(min(rows, len(self.df))):
            u = safe_str(self.df.iloc[i, 0])
            s = safe_str(self.df.iloc[i, 1])
            if u or s:
                lines.append(f"[{i+1}] В/Ч: {u} | Статус: {s}")
        return "\n".join(lines) if lines else "Таблица пуста или данные не распознаны"


class DataProcessor:
    def __init__(self, df: pd.DataFrame, known_units: List[str]):
        self.df = df
        self.known_units = known_units
        self.filed_counts = {}
        self.decision_counts = {}
        self.display_names = {}
        
    def process(self):
        for unit in self.known_units:
            key = normalize_text(unit)
            self.display_names[key] = unit
        
        for i in range(len(self.df)):
            unit_raw = safe_str(self.df.iloc[i, 0])
            status_raw = safe_str(self.df.iloc[i, 1])
            
            if not unit_raw:
                continue
            
            unit_key = normalize_text(unit_raw)
            if unit_key not in self.display_names:
                self.display_names[unit_key] = unit_raw
            
            status_norm = normalize_text(status_raw)
            if "суд" in status_norm:
                self.filed_counts[unit_key] = self.filed_counts.get(unit_key, 0) + 1
            if "бво" in status_norm:
                self.decision_counts[unit_key] = self.decision_counts.get(unit_key, 0) + 1
    
    def get_sorted_units(self) -> List[str]:
        result = []
        seen = set()
        for unit in self.known_units:
            key = normalize_text(unit)
            if key in self.display_names and key not in seen:
                result.append(key)
                seen.add(key)
        
        others = sorted(
            [k for k in self.display_names.keys() if k not in seen],
            key=lambda k: self.display_names[k]
        )
        result.extend(others)
        return result
    
    def get_summary(self) -> str:
        tf = sum(self.filed_counts.values())
        td = sum(self.decision_counts.values())
        lines = [f"ИТОГО: Подано: {tf} | Решений: {td}\n" + "-"*30]
        
        for uk in self.get_sorted_units():
            f = self.filed_counts.get(uk, 0)
            d = self.decision_counts.get(uk, 0)
            if f > 0 or d > 0:
                lines.append(f"{self.display_names[uk]}: Исков {f}, Решений {d}")
        return "\n".join(lines)


class DocxGenerator:
    def __init__(self, processor: DataProcessor, date_range: str):
        self.processor = processor
        self.date_range = date_range
        self.doc = Document()
        
    def set_run_font(self, run, font_name: str, size_pt: int, bold: bool = None, italic: bool = False):
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        if bold is not None: run.bold = bold
        if italic: run.italic = italic
        rpr = run._element.get_or_add_rPr()
        rpr.rFonts.set(qn("w:eastAsia"), font_name)
    
    def add_paragraph_with_indent(self, text: str, font_name: str, font_size: int, 
                                   first_line_indent: float = 0, bold_parts: bool = False):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.08
        if first_line_indent > 0:
            p.paragraph_format.first_line_indent = Cm(first_line_indent)
        
        if bold_parts:
            parts = re.split(r'(\*\*[^*]+\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    self.set_run_font(p.add_run(part[2:-2]), font_name, font_size, bold=True)
                elif part:
                    self.set_run_font(p.add_run(part), font_name, font_size, bold=False)
        else:
            self.set_run_font(p.add_run(text), font_name, font_size)
    
    def set_cell(self, cell, text: str, font="Times New Roman", size=12, bold=False, italic=False):
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        self.set_run_font(p.add_run(text), font, size, bold, italic)
    
    def generate(self, output_path: Path):
        section = self.doc.sections[0]
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1)
        
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.set_run_font(p.add_run("СПРАВКА – ДОКЛАД ПО ПРОДЕЛАННОЙ СУДЕБНОЙ РАБОТЕ ЗА НЕДЕЛЮ "), "Times New Roman", 16)
        self.set_run_font(p.add_run(self.date_range), "Times New Roman", 16, bold=True)
        
        self.add_paragraph_with_indent("", "Times New Roman", 14)
        
        total_filed = sum(self.processor.filed_counts.values())
        filed_w = plural_ru(total_filed, ('иск', 'иска', 'исков'))
        
        text_filed = f"За неделю подано **{total_filed} {filed_w}** в суд"
        if total_filed > 0:
            text_filed += ", из них:"
            self.add_paragraph_with_indent(text_filed, "Times New Roman", 14, 1.25, True)
            for uk in self.processor.get_sorted_units():
                cnt = self.processor.filed_counts.get(uk, 0)
                if cnt > 0:
                    w = plural_ru(cnt, ('иск', 'иска', 'исков'))
                    self.add_paragraph_with_indent(f"{self.processor.display_names[uk]} – {cnt} {w}.", "Times New Roman", 14)
        else:
            text_filed += "."
            self.add_paragraph_with_indent(text_filed, "Times New Roman", 14, 1.25, True)

        self.add_paragraph_with_indent("", "Times New Roman", 14)

        total_dec = sum(self.processor.decision_counts.values())
        text_dec = f"За неделю получено решений судов – **{total_dec}**"
        if total_dec > 0:
            text_dec += ", из них:"
            self.add_paragraph_with_indent(text_dec, "Times New Roman", 14, 1.25, True)
            dec_list = [(uk, self.processor.decision_counts[uk]) for uk in self.processor.get_sorted_units() if self.processor.decision_counts.get(uk,0) > 0]
            for i, (uk, cnt) in enumerate(dec_list):
                w = plural_ru(cnt, ('решение', 'решения', 'решений'))
                punct = "." if i == len(dec_list)-1 else ";"
                self.add_paragraph_with_indent(f"{self.processor.display_names[uk]} – {cnt} {w}{punct}", "Times New Roman", 14)
        else:
            text_dec += "."
            self.add_paragraph_with_indent(text_dec, "Times New Roman", 14, 1.25, True)
            
        self.add_paragraph_with_indent("", "Times New Roman", 14)

        table = self.doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        
        headers = ["№\nп/п", "Подразделения", "За сутки\nподано", "За неделю\nподано", "За неделю\nрешений", "Прим."]
        widths = [1.2, 4.2, 2.7, 2.7, 2.7, 2.7]
        
        for i, h in enumerate(headers):
            self.set_cell(table.rows[0].cells[i], h)
            table.rows[0].cells[i].width = Cm(widths[i])
            
        row_num = 1
        for uk in self.processor.get_sorted_units():
            f = self.processor.filed_counts.get(uk, 0)
            d = self.processor.decision_counts.get(uk, 0)
            
            row = table.add_row()
            cells = row.cells
            self.set_cell(cells[0], f"{row_num}.")
            self.set_cell(cells[1], self.processor.display_names[uk], "Calibri", 12, True, True)
            self.set_cell(cells[2], "0")
            self.set_cell(cells[3], str(f))
            self.set_cell(cells[4], str(d))
            self.set_cell(cells[5], "")
            row_num += 1
            
        row = table.add_row()
        self.set_cell(row.cells[0], f"{row_num}.")
        self.set_cell(row.cells[1], "ИТОГО ЗА ДИВИЗИЮ", "Calibri", 12, True, True)
        self.set_cell(row.cells[2], "0")
        self.set_cell(row.cells[3], str(total_filed))
        self.set_cell(row.cells[4], str(total_dec))
        self.set_cell(row.cells[5], "")

        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(output_path)


class StepIndicator(QWidget):
    def __init__(self, steps: List[str], parent=None):
        super().__init__(parent)
        self.steps = steps
        self.labels = []
        self.init_ui()
    
    def init_ui(self):
        layout = QHBoxLayout()
        layout.setContentsMargins(20, 10, 20, 10)
        layout.setSpacing(5)
        
        for i, text in enumerate(self.steps):
            lbl = QLabel(f"{i+1}. {text}")
            lbl.setAlignment(Qt.AlignCenter)
            lbl.setProperty("active", False)
            lbl.setStyleSheet("""
                QLabel {
                    background-color: #E5E5EA;
                    color: #8E8E93;
                    border-radius: 12px;
                    padding: 6px 12px;
                    font-weight: 500;
                }
                QLabel[active="true"] {
                    background-color: #007AFF;
                    color: white;
                    font-weight: bold;
                }
            """)
            self.labels.append(lbl)
            layout.addWidget(lbl)
            
            if i < len(self.steps) - 1:
                line = QFrame()
                line.setFrameShape(QFrame.HLine)
                line.setStyleSheet("background-color: #C7C7CC;")
                line.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
                layout.addWidget(line)
        
        self.setLayout(layout)
    
    def set_step(self, index):
        for i, lbl in enumerate(self.labels):
            is_active = (i == index)
            lbl.setProperty("active", is_active)
            lbl.style().unpolish(lbl)
            lbl.style().polish(lbl)


class MainApp(QWidget):
    def __init__(self):
        super().__init__()
        self.analyzer = None
        self.processor = None
        self.current_step = 0
        self.init_ui()
        
    def init_ui(self):
        self.setWindowTitle("Генератор судебных отчётов v2.0")
        self.setMinimumSize(950, 650)
        self.setStyleSheet(STYLESHEET)
        
        main_layout = QVBoxLayout(self)
        main_layout.setSpacing(0)
        main_layout.setContentsMargins(0, 0, 0, 0)
        
        header = QWidget()
        header.setStyleSheet("background-color: white; border-bottom: 1px solid #D1D1D6;")
        hl = QVBoxLayout(header)
        hl.setContentsMargins(30, 20, 30, 20)
        title = QLabel("Генератор отчётов")
        title.setStyleSheet("font-size: 22pt; font-weight: 300; color: #1C1C1E;")
        subtitle = QLabel("Автоматизация обработки судебных реестров")
        subtitle.setStyleSheet("color: #8E8E93; font-size: 11pt;")
        hl.addWidget(title)
        hl.addWidget(subtitle)
        main_layout.addWidget(header)
        
        self.step_indicator = StepIndicator(["Файл", "Колонки", "Фильтр", "Финиш"])
        self.step_indicator.setStyleSheet("background-color: #F9F9F9; border-bottom: 1px solid #D1D1D6;")
        main_layout.addWidget(self.step_indicator)
        
        self.stack = QStackedWidget()
        self.stack.addWidget(self.ui_step1())
        self.stack.addWidget(self.ui_step2())
        self.stack.addWidget(self.ui_step3())
        self.stack.addWidget(self.ui_step4())
        
        container = QWidget()
        cl = QVBoxLayout(container)
        cl.setContentsMargins(30, 20, 30, 10)
        cl.addWidget(self.stack)
        main_layout.addWidget(container, 1)
        
        footer = QWidget()
        footer.setStyleSheet("background-color: white; border-top: 1px solid #D1D1D6;")
        fl = QHBoxLayout(footer)
        fl.setContentsMargins(30, 15, 30, 15)
        
        self.btn_back = QPushButton("Назад")
        self.btn_back.setObjectName("secondary")
        self.btn_back.setMinimumWidth(120)
        self.btn_back.clicked.connect(self.go_back)
        self.btn_back.setEnabled(False)
        
        self.btn_next = QPushButton("Далее")
        self.btn_next.setObjectName("secondary")
        self.btn_next.setMinimumWidth(120)
        self.btn_next.clicked.connect(self.go_next)
        
        fl.addWidget(self.btn_back)
        fl.addStretch()
        fl.addWidget(self.btn_next)
        
        main_layout.addWidget(footer)
        
        self.update_nav_state()
    
    def ui_step1(self):
        w = QWidget()
        l = QVBoxLayout(w)
        
        card = QFrame()
        card.setObjectName("Card")
        cl = QVBoxLayout(card)
        cl.setSpacing(15)
        cl.setContentsMargins(20, 20, 20, 20)
        
        lbl = QLabel("Выберите Excel файл (.xlsx) для анализа")
        lbl.setStyleSheet("font-size: 12pt; font-weight: 500;")
        cl.addWidget(lbl)
        
        file_row = QHBoxLayout()
        self.inp_path = QLineEdit()
        self.inp_path.setPlaceholderText("Путь к файлу...")
        self.inp_path.setReadOnly(True)
        btn_browse = QPushButton("Обзор")
        btn_browse.clicked.connect(self.action_browse)
        file_row.addWidget(self.inp_path)
        file_row.addWidget(btn_browse)
        cl.addLayout(file_row)
        
        sheet_row = QHBoxLayout()
        sheet_row.addWidget(QLabel("Лист с данными:"))
        self.combo_sheet = QComboBox()
        self.combo_sheet.setView(QListView())
        self.combo_sheet.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
        sheet_row.addWidget(self.combo_sheet)
        cl.addLayout(sheet_row)
        
        cl.addStretch()
        l.addWidget(card)
        l.addStretch()
        return w

    def ui_step2(self):
        w = QWidget()
        l = QVBoxLayout(w)
        
        card = QFrame()
        card.setObjectName("Card")
        cl = QVBoxLayout(card)
        cl.setSpacing(15)
        cl.setContentsMargins(20, 20, 20, 20)
        
        cl.addWidget(QLabel("Проверьте, правильно ли определены координаты:"))
        
        grid = QtWidgets.QGridLayout()
        grid.setHorizontalSpacing(20)
        grid.setVerticalSpacing(15)
        
        grid.addWidget(QLabel("Колонка 'В/Ч':"), 0, 0)
        self.inp_col_unit = QLineEdit()
        self.inp_col_unit.setAlignment(Qt.AlignCenter)
        self.inp_col_unit.setPlaceholderText("A")
        grid.addWidget(self.inp_col_unit, 0, 1)
        
        grid.addWidget(QLabel("Колонка 'НОВАЯ НЕДЕЛЯ':"), 1, 0)
        self.inp_col_status = QLineEdit()
        self.inp_col_status.setAlignment(Qt.AlignCenter)
        self.inp_col_status.setPlaceholderText("B")
        grid.addWidget(self.inp_col_status, 1, 1)
        
        grid.addWidget(QLabel("Номер строки заголовка:"), 2, 0)
        self.inp_row_head = QSpinBox()
        self.inp_row_head.setRange(1, 100)
        grid.addWidget(self.inp_row_head, 2, 1)
        
        cl.addLayout(grid)
        
        btn_check = QPushButton("Проверить и показать пример")
        btn_check.setObjectName("secondary")
        btn_check.clicked.connect(self.action_refresh_preview)
        cl.addWidget(btn_check)
        
        self.txt_preview = QTextEdit()
        self.txt_preview.setPlaceholderText("Здесь отобразится пример данных...")
        self.txt_preview.setReadOnly(True)
        cl.addWidget(self.txt_preview)
        
        l.addWidget(card)
        return w

    def ui_step3(self):
        w = QWidget()
        l = QVBoxLayout(w)
        
        top_layout = QHBoxLayout()
        
        left_layout = QVBoxLayout()
        left_layout.addWidget(QLabel("Порядок подразделений в отчёте:"))
        self.list_units = QListWidget()
        self.list_units.setDragDropMode(QAbstractItemView.InternalMove)
        self.list_units.addItems(DEFAULT_UNITS)
        left_layout.addWidget(self.list_units)
        top_layout.addLayout(left_layout, 2)
        
        btns_layout = QVBoxLayout()
        btns_layout.setContentsMargins(10, 25, 0, 0)
        btns_layout.setSpacing(10)
        
        btn_add = QPushButton("+ Добавить")
        btn_add.setObjectName("secondary")
        btn_add.clicked.connect(self.action_add_unit)
        
        btn_del = QPushButton("- Удалить")
        btn_del.setObjectName("danger")
        btn_del.clicked.connect(self.action_del_unit)
        
        btns_layout.addWidget(btn_add)
        btns_layout.addWidget(btn_del)
        btns_layout.addStretch()
        top_layout.addLayout(btns_layout, 1)
        
        l.addLayout(top_layout)
        
        self.check_others = QCheckBox("Включать в отчёт остальные подразделения (найденные в файле, но не в списке)")
        self.check_others.setChecked(True)
        self.check_others.setStyleSheet("margin-top: 10px; font-size: 11pt;")
        l.addWidget(self.check_others)
        
        return w

    def ui_step4(self):
        w = QWidget()
        l = QVBoxLayout(w)
        
        card = QFrame()
        card.setObjectName("Card")
        cl = QVBoxLayout(card)
        cl.setContentsMargins(20, 20, 20, 20)
        cl.setSpacing(15)
        
        cl.addWidget(QLabel("Параметры документа:"))
        
        form = QtWidgets.QFormLayout()
        self.inp_date = QLineEdit()
        self.inp_date.setPlaceholderText("Пример: 03.02-09.02.2025")
        form.addRow("Период (в шапку):", self.inp_date)
        
        path_layout = QHBoxLayout()
        self.inp_out_path = QLineEdit()
        btn_out = QPushButton("...")
        btn_out.setObjectName("secondary")
        btn_out.setMinimumWidth(40)
        btn_out.clicked.connect(self.action_choose_out)
        path_layout.addWidget(self.inp_out_path)
        path_layout.addWidget(btn_out)
        form.addRow("Куда сохранить:", path_layout)
        cl.addLayout(form)
        
        self.txt_summary = QTextEdit()
        self.txt_summary.setReadOnly(True)
        self.txt_summary.setMaximumHeight(150)
        cl.addWidget(QLabel("Итоговые цифры:"))
        cl.addWidget(self.txt_summary)
        
        btn_gen = QPushButton("СФОРМИРОВАТЬ ОТЧЁТ")
        btn_gen.setObjectName("success")
        btn_gen.setCursor(Qt.PointingHandCursor)
        btn_gen.clicked.connect(self.action_generate)
        cl.addWidget(btn_gen)
        
        l.addWidget(card)
        l.addStretch()
        return w

    def action_browse(self):
        path, _ = QFileDialog.getOpenFileName(self, "Excel файл", "", "Excel Files (*.xlsx)")
        if path:
            self.inp_path.setText(path)
            try:
                self.analyzer = ExcelAnalyzer(Path(path))
                sheets = self.analyzer.load_sheets()
                self.combo_sheet.clear()
                self.combo_sheet.addItems(sheets)
                out = Path(path).parent / "Судебный_отчёт.docx"
                self.inp_out_path.setText(str(out))
            except Exception as e:
                QMessageBox.critical(self, "Ошибка", str(e))

    def action_refresh_preview(self):
        if not self.analyzer: return
        try:
            sheet = self.combo_sheet.currentText()
            u_idx = column_letter_to_index(self.inp_col_unit.text())
            s_idx = column_letter_to_index(self.inp_col_status.text())
            
            if u_idx < 0 or s_idx < 0:
                self.txt_preview.setText("Ошибка: Некорректные буквы колонок")
                return
                
            header = self.inp_row_head.value()
            self.analyzer.load_data(sheet, header, u_idx, s_idx)
            text = self.analyzer.preview_data(8)
            self.txt_preview.setText(text)
        except Exception as e:
            self.txt_preview.setText(f"Ошибка чтения:\n{e}")

    def action_add_unit(self):
        text, ok = QtWidgets.QInputDialog.getText(self, "Новое подразделение", "Название:")
        if ok and text:
            self.list_units.addItem(text)
            
    def action_del_unit(self):
        row = self.list_units.currentRow()
        if row >= 0:
            self.list_units.takeItem(row)
            
    def action_choose_out(self):
        path, _ = QFileDialog.getSaveFileName(self, "Сохранить", self.inp_out_path.text(), "Word (*.docx)")
        if path: self.inp_out_path.setText(path)

    def action_calc_summary(self):
        if not self.analyzer or self.analyzer.df is None: return
        try:
            units = [self.list_units.item(i).text() for i in range(self.list_units.count())]
            self.processor = DataProcessor(self.analyzer.df, units)
            self.processor.process()
            self.txt_summary.setText(self.processor.get_summary())
        except Exception as e:
            self.txt_summary.setText(f"Ошибка обработки: {e}")

    def action_generate(self):
        if not self.processor: self.action_calc_summary()
        if not self.inp_date.text():
            QMessageBox.warning(self, "Внимание", "Введите период дат для шапки документа!")
            self.inp_date.setFocus()
            return
            
        try:
            gen = DocxGenerator(self.processor, self.inp_date.text())
            gen.generate(Path(self.inp_out_path.text()))
            QMessageBox.information(self, "Успех", "Файл успешно создан!")
            if sys.platform == 'win32':
                import os
                os.startfile(str(Path(self.inp_out_path.text()).parent))
        except Exception as e:
            QMessageBox.critical(self, "Ошибка записи", str(e))

    def go_next(self):
        idx = self.stack.currentIndex()
        
        if idx == 0:
            if not self.inp_path.text():
                QMessageBox.warning(self, "Стоп", "Сначала выберите файл!")
                return
            try:
                if not self.inp_col_unit.text():
                    sheet = self.combo_sheet.currentText()
                    r, cols = self.analyzer.find_header_row(sheet)
                    
                    def to_char(i): 
                        return chr(ord('A') + i) if 0 <= i < 26 else "?"
                    
                    self.inp_col_unit.setText(to_char(cols['unit']))
                    self.inp_col_status.setText(to_char(cols['status']))
                    self.inp_row_head.setValue(r)
            except:
                pass

        if idx == 1:
            if not self.analyzer.df is not None:
                self.action_refresh_preview()
                if self.analyzer.df is None:
                    QMessageBox.warning(self, "Стоп", "Данные не загружены. Нажмите 'Проверить'!")
                    return

        if idx == 2:
            self.action_calc_summary()

        if idx < 3:
            self.stack.setCurrentIndex(idx + 1)
            self.update_nav_state()

    def go_back(self):
        idx = self.stack.currentIndex()
        if idx > 0:
            self.stack.setCurrentIndex(idx - 1)
            self.update_nav_state()

    def update_nav_state(self):
        idx = self.stack.currentIndex()
        self.step_indicator.set_step(idx)
        
        self.btn_back.setEnabled(idx > 0)
        
        if idx == 3:
            self.btn_next.hide() 
        else:
            self.btn_next.show()
            self.btn_next.setText("Далее")
            self.btn_next.setEnabled(True)

def main():
    app = QtWidgets.QApplication(sys.argv)
    app.setStyle("Fusion")
    
    font = QFont("Segoe UI", 10)
    app.setFont(font)
    
    w = MainApp()
    w.show()
    sys.exit(app.exec_())

if __name__ == "__main__":
    main()