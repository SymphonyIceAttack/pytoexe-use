import os
import sys
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import pypandoc
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import logging
import shutil
from pathlib import Path
import tempfile
import threading
import re
 
 
class MarkdownConverterApp:
    def __init__(self, master):
        self.master = master
        self.master.title("Markdown转Word工具 v2.6")
        self.master.geometry("900x700")
        self.master.minsize(800, 600)
 
        # 初始化变量
        self.current_font = "微软雅黑"
        self.template_path = None
        self.is_converting = False
 
        self.setup_logging()
        self.setup_template()
        self.setup_ui()
 
        # 设置正确的拖拽事件绑定
        self.setup_drag_drop()
 
    def setup_logging(self):
        """配置日志系统"""
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s',
            handlers=[
                logging.FileHandler('converter.log', encoding='utf-8'),
                logging.StreamHandler(sys.stdout)
            ]
        )
        self.logger = logging.getLogger(__name__)
 
    def setup_template(self):
        """创建或验证Word模板"""
        try:
            template_path = Path.home() / ".markdown_converter_template.docx"
            if not template_path.exists():
                self._create_default_template(template_path)
            self.template_path = str(template_path)
            self.logger.info(f"模板路径: {self.template_path}")
        except Exception as e:
            self.logger.error(f"模板设置失败: {str(e)}")
            self.template_path = self._create_temp_template()
 
    def _create_default_template(self, template_path):
        """创建包含中文排版和首行缩进的Word模板，支持1-6级标题"""
        try:
            doc = Document()
 
            # 设置正文样式 - 添加首行缩进两个字符
            style = doc.styles['Normal']
            style.font.name = '微软雅黑'
            style.font.size = Pt(10.5)
            # 设置首行缩进为0.3英寸（约两个字符）
            style.paragraph_format.first_line_indent = Inches(0.3)
            style.paragraph_format.line_spacing = 1.5
            style.paragraph_format.space_after = Pt(6)
 
            if hasattr(style.font, '_element'):
                style.font._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
 
            # 配置标题样式（1-6级）
            for level in range(1, 7):  # 扩展为1-6级标题
                heading_style_name = f'Heading {level}'
                try:
                    heading_style = doc.styles[heading_style_name]
                    heading_style.font.name = '微软雅黑'
                    # 设置不同级别的字体大小
                    base_sizes = [16, 14, 12, 11, 10.5, 10]
                    heading_style.font.size = Pt(base_sizes[level - 1] if level <= len(base_sizes) else 10)
                    heading_style.font.bold = True
 
                    # 设置标题间距
                    heading_style.paragraph_format.space_before = Pt(12 if level == 1 else 8)
                    heading_style.paragraph_format.space_after = Pt(6)
 
                    # 标题不需要首行缩进
                    heading_style.paragraph_format.first_line_indent = Pt(0)
 
                    if hasattr(heading_style.font, '_element'):
                        heading_style.font._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                except KeyError:
                    # 如果样式不存在，跳过
                    continue
 
            # 保存模板
            template_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(template_path)
            self.logger.info("默认模板创建成功，支持1-6级标题")
            return str(template_path)
 
        except Exception as e:
            self.logger.error(f"创建默认模板失败: {str(e)}")
            return self._create_temp_template()
 
    def _create_temp_template(self):
        """创建临时模板作为备选"""
        try:
            temp_dir = tempfile.gettempdir()
            template_path = Path(temp_dir) / "markdown_temp_template.docx"
            doc = Document()
            doc.save(template_path)
            self.logger.info(f"临时模板创建于: {template_path}")
            return str(template_path)
        except Exception as e:
            self.logger.error(f"创建临时模板失败: {str(e)}")
            return None
 
    def setup_drag_drop(self):
        """修复拖拽功能 - 使用Tkinter DnD库"""
        try:
            from tkinterdnd2 import DND_FILES, TkinterDnD
            self.logger.info("TkinterDnD库可用，启用高级拖拽功能")
 
            def handle_drop(event):
                files = self.master.tk.splitlist(event.data)
                for file in files:
                    if file.lower().endswith(('.md', '.txt', '.markdown')):
                        self.load_markdown_file(file)
                        break
 
            self.input_area.drop_target_register(DND_FILES)
            self.input_area.dnd_bind('<<Drop>>', handle_drop)
 
        except ImportError:
            self.logger.warning("TkinterDnD不可用，使用基本文件选择功能")
            if hasattr(self, 'drag_label'):
                self.drag_label.pack_forget()
 
    def setup_ui(self):
        """构建图形用户界面"""
        # 主框架
        main_frame = ttk.Frame(self.master, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
 
        # 标题
        title_label = ttk.Label(main_frame, text="Markdown转Word转换器",
                                font=("微软雅黑", 16, "bold"))
        title_label.pack(pady=(0, 10))
 
        # 输入区域
        input_frame = ttk.LabelFrame(main_frame, text="Markdown输入", padding="5")
        input_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
 
        # 拖拽提示
        self.drag_label = ttk.Label(input_frame,
                                    text="将Markdown文件拖拽到此处或直接粘贴内容",
                                    font=("微软雅黑", 9), foreground="gray")
        self.drag_label.pack(pady=5)
 
        self.input_area = scrolledtext.ScrolledText(
            input_frame,
            wrap=tk.WORD,
            width=80,
            height=15,
            font=('微软雅黑', 10),
            padx=10,
            pady=10
        )
        self.input_area.pack(fill=tk.BOTH, expand=True)
 
        # 配置区域
        config_frame = ttk.LabelFrame(main_frame, text="转换设置", padding="10")
        config_frame.pack(fill=tk.X, pady=(0, 10))
 
        # 字体选择行
        font_row = ttk.Frame(config_frame)
        font_row.pack(fill=tk.X, pady=5)
 
        font_label = ttk.Label(font_row, text="输出字体:", width=10)
        font_label.pack(side=tk.LEFT)
 
        # 扩展字体选项，增加公文字体
        font_options = [
            '微软雅黑', '宋体', '黑体', '楷体', '仿宋',
            '仿宋_GB2312', '方正小标宋简体', '华文楷体', '华文中宋',
            'Arial', 'Times New Roman', 'Calibri'
        ]
 
        self.font_var = tk.StringVar(value="微软雅黑")
        self.font_combo = ttk.Combobox(
            font_row,
            textvariable=self.font_var,
            values=font_options,
            state="readonly",
            width=15
        )
        self.font_combo.pack(side=tk.LEFT, padx=(0, 20))
        self.font_combo.bind('<<ComboboxSelected>>', self._update_font_preview)
 
        # 全角标点选项
        self.full_width_punctuation_var = tk.BooleanVar(value=True)
        full_width_cb = ttk.Checkbutton(
            font_row,
            text="转换标点为全角",
            variable=self.full_width_punctuation_var
        )
        full_width_cb.pack(side=tk.LEFT, padx=(0, 20))
 
        # 首行缩进选项
        self.indent_var = tk.BooleanVar(value=True)
        indent_cb = ttk.Checkbutton(
            font_row,
            text="首行缩进两个字符",
            variable=self.indent_var
        )
        indent_cb.pack(side=tk.LEFT, padx=(0, 20))
 
        # 预览区域
        preview_frame = ttk.Frame(font_row)
        preview_frame.pack(side=tk.RIGHT, fill=tk.X, expand=True)
 
        preview_title = ttk.Label(preview_frame, text="预览:")
        preview_title.pack(side=tk.LEFT)
 
        self.preview_label = ttk.Label(
            preview_frame,
            text="我们秉持的宗旨",
            font=('微软雅黑', 10)
        )
        self.preview_label.pack(side=tk.LEFT, padx=(10, 0))
 
        # 按钮区域
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill=tk.X, pady=(0, 10))
 
        self.convert_btn = ttk.Button(
            button_frame,
            text="转换为Word文档",
            command=self.start_conversion,
            width=20
        )
        self.convert_btn.pack(side=tk.LEFT, padx=(0, 10))
 
        self.clear_btn = ttk.Button(
            button_frame,
            text="清空内容",
            command=self.clear_input,
            width=10
        )
        self.clear_btn.pack(side=tk.LEFT, padx=(0, 10))
 
        load_file_btn = ttk.Button(
            button_frame,
            text="加载Markdown文件",
            command=self.browse_markdown_file,
            width=15
        )
        load_file_btn.pack(side=tk.LEFT, padx=(0, 10))
 
        # 进度条
        self.progress = ttk.Progressbar(
            button_frame,
            mode='indeterminate',
            length=100
        )
        self.progress.pack(side=tk.RIGHT, fill=tk.X, expand=True)
 
        # 状态栏
        self.status = tk.StringVar(value="就绪 - 请输入Markdown内容或拖拽文件")
        status_bar = ttk.Label(
            main_frame,
            textvariable=self.status,
            relief=tk.SUNKEN,
            anchor=tk.W,
            padding=(5, 2)
        )
        status_bar.pack(fill=tk.X)
 
        # 初始化预览
        self._update_font_preview()
 
    def _update_font_preview(self, event=None):
        """更新字体预览"""
        selected_font = self.font_var.get()
        self.preview_label.config(font=(selected_font, 10))
 
    def browse_markdown_file(self):
        """浏览并加载Markdown文件"""
        file_path = filedialog.askopenfilename(
            title="选择Markdown文件",
            filetypes=[
                ("Markdown文件", "*.md"),
                ("文本文件", "*.txt"),
                ("所有文件", "*.*")
            ]
        )
 
        if file_path:
            self.load_markdown_file(file_path)
 
    def load_markdown_file(self, file_path):
        """加载Markdown文件内容"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
 
            self.input_area.delete('1.0', tk.END)
            self.input_area.insert('1.0', content)
 
            # 自动滚动到底部
            self.input_area.see(tk.END)
 
            self.status.set(f"已加载文件: {Path(file_path).name}")
            self.logger.info(f"成功加载Markdown文件: {file_path}")
 
        except Exception as e:
            error_msg = f"读取文件失败: {str(e)}"
            self.logger.error(error_msg)
            messagebox.showerror("错误", error_msg)
 
    def start_conversion(self):
        """在单独线程中开始转换过程"""
        if self.is_converting:
            return
 
        self.is_converting = True
        self.convert_btn.config(state='disabled')
        self.progress.start(10)
        self.status.set("转换中...")
 
        # 在新线程中运行转换
        thread = threading.Thread(target=self.convert)
        thread.daemon = True
        thread.start()
 
    def convert(self):
        """执行Markdown到Word的转换，支持1-6级标题"""
        markdown_text = self.input_area.get("1.0", tk.END).strip()
 
        if not markdown_text:
            self.master.after(0, lambda: messagebox.showwarning("警告", "请输入Markdown内容或选择Markdown文件！"))
            self._conversion_finished()
            return
 
        try:
            # 选择输出路径
            output_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[
                    ("Word文档", "*.docx"),
                    ("所有文件", "*.*")
                ],
                title="保存Word文档"
            )
 
            if not output_path:
                self._conversion_finished()
                return
 
            # 创建临时目录和文件
            with tempfile.TemporaryDirectory() as temp_dir:
                temp_md = Path(temp_dir) / "temp.md"
 
                # 处理标点符号转换
                processed_text = self._preprocess_markdown(markdown_text)
 
                # 保存处理后的Markdown内容到临时文件
                with open(temp_md, 'w', encoding='utf-8') as f:
                    f.write(processed_text)
 
                # 构建pandoc参数 - 支持1-6级标题
                extra_args = [
                    '--standalone',
                    f'--reference-doc={self.template_path}' if self.template_path else '',
                    '--wrap=auto',
                    # 设置标题级别为6级
                    '--top-level-division=section',
                ]
 
                # 过滤空参数
                extra_args = [arg for arg in extra_args if arg]
 
                # 执行转换
                self.logger.info("开始Pandoc转换...")
                pypandoc.convert_file(
                    str(temp_md),
                    'docx',
                    outputfile=output_path,
                    extra_args=extra_args
                )
 
                # 应用字体和段落设置
                self._apply_document_settings(output_path)
 
            # 在主线程中更新UI
            self.master.after(0, lambda: self._conversion_success(output_path))
 
        except Exception as e:
            error_msg = f"转换失败: {str(e)}"
            self.logger.error(error_msg, exc_info=True)
            self.master.after(0, lambda: self._conversion_error(error_msg))
 
    def _preprocess_markdown(self, text):
        """预处理Markdown文本：转换标点符号为全角"""
        if not self.full_width_punctuation_var.get():
            return text
 
        # 半角到全角的标点映射
        half_to_full = {
            ',': '，',
            '.': '。',
            '!': '！',
            '?': '？',
            ':': '：',
            ';': '；',
            '(': '（',
            ')': '）',
            '[': '【',
            ']': '】',
            '<': '《',
            '>': '》',
            '"': '＂',
            "'": '＇',
            '`': '｀',
            '~': '～',
            '@': '＠',
            '#': '＃',
            '$': '＄',
            '%': '％',
            '^': '＾',
            '&': '＆',
            '*': '＊',
            '-': '－',
            '=': '＝',
            '+': '＋',
            '\\': '＼',
            '|': '｜'
        }
 
        # 保护Markdown语法符号，只转换文本内容中的标点
        lines = text.split('\n')
        processed_lines = []
 
        for line in lines:
            # 如果是Markdown语法行（标题、列表、代码块等），跳过转换
            if (line.strip().startswith('#') or
                    line.strip().startswith('-') or
                    line.strip().startswith('*') or
                    line.strip().startswith('>') or
                    line.strip().startswith('`') or
                    '```' in line or
                    line.strip().startswith('|') or  # 表格
                    line.strip().startswith('+') or  # 列表
                    re.match(r'^\s*\d+\.', line.strip()) or  # 有序列表
                    re.match(r'^\s*\d+\)', line.strip())):  # 有序列表
                processed_lines.append(line)
                continue
 
            # 转换普通文本行中的标点
            processed_line = ''
            for char in line:
                if char in half_to_full and self.full_width_punctuation_var.get():
                    processed_line += half_to_full[char]
                else:
                    processed_line += char
            processed_lines.append(processed_line)
 
        return '\n'.join(processed_lines)
 
    def _is_paragraph_text(self, paragraph):
        """判断段落是否为正文文本段落"""
        # 空段落不是正文
        if not paragraph.text.strip():
            return False
 
        # 获取段落样式名称
        style_name = paragraph.style.name.lower() if paragraph.style and paragraph.style.name else ""
 
        # 排除标题段落（支持1-6级标题）
        if ('heading' in style_name or
                any(paragraph.text.strip().startswith('#' * i) for i in range(1, 7))):
            return False
 
        # 排除代码块
        if paragraph.style.name and 'code' in paragraph.style.name.lower():
            return False
 
        # 检查段落是否包含代码格式
        for run in paragraph.runs:
            if run.font.name and 'consolas' in run.font.name.lower():
                return False
            if run.style and run.style.name and 'code' in run.style.name.lower():
                return False
 
        return True
 
    def _apply_document_settings(self, doc_path):
        """应用字体、段落和标点设置到Word文档"""
        try:
            doc = Document(doc_path)
            selected_font = self.font_var.get()
            use_indent = self.indent_var.get()
 
            # 处理所有段落
            for paragraph in doc.paragraphs:
                # 更准确地判断正文段落
                if self._is_paragraph_text(paragraph):
                    if use_indent:
                        # 设置首行缩进两个字符（0.3英寸）
                        paragraph.paragraph_format.first_line_indent = Inches(0.3)
                        self.logger.debug(f"应用首行缩进到段落: {paragraph.text[:50]}...")
 
                # 设置字体
                for run in paragraph.runs:
                    run.font.name = selected_font
                    if hasattr(run.font, '_element'):
                        run.font._element.rPr.rFonts.set(qn('w:eastAsia'), selected_font)
 
            # 处理表格中的字体和段落
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if use_indent and self._is_paragraph_text(paragraph):
                                paragraph.paragraph_format.first_line_indent = Inches(0.3)
                            for run in paragraph.runs:
                                run.font.name = selected_font
                                if hasattr(run.font, '_element'):
                                    run.font._element.rPr.rFonts.set(qn('w:eastAsia'), selected_font)
 
            doc.save(doc_path)
            self.logger.info("文档设置应用成功")
 
        except Exception as e:
            self.logger.warning(f"文档设置应用失败: {str(e)}")
 
    def _conversion_success(self, output_path):
        """转换成功后的处理"""
        self._conversion_finished()
        self.status.set("转换完成！")
        self.logger.info(f"转换成功: {output_path}")
 
        if messagebox.askyesno("完成", "文件转换成功！是否打开所在文件夹？"):
            try:
                os.startfile(str(Path(output_path).parent))
            except:
                try:
                    import subprocess
                    subprocess.run(['open', str(Path(output_path).parent)])
                except:
                    pass
 
    def _conversion_error(self, error_msg):
        """转换失败后的处理"""
        self._conversion_finished()
        messagebox.showerror("错误", error_msg)
        self.status.set("转换失败")
 
    def _conversion_finished(self):
        """转换完成后的清理工作"""
        self.is_converting = False
        self.convert_btn.config(state='normal')
        self.progress.stop()
 
    def clear_input(self):
        """清空输入区域"""
        self.input_area.delete('1.0', tk.END)
        self.status.set("已清空")
        self.logger.info("输入区域已清空")
 
 
def check_dependencies():
    """检查必要的依赖"""
    try:
        pypandoc.get_pandoc_version()
        return True
    except OSError:
        messagebox.showerror(
            "依赖错误"
        )
        return False
    except Exception as e:
        messagebox.showerror("错误", f"初始化失败: {str(e)}")
        return False
 
 
if __name__ == "__main__":
    if not check_dependencies():
        sys.exit(1)
 
    root = tk.Tk()
    try:
        app = MarkdownConverterApp(root)
        root.mainloop()
    except Exception as e:
        messagebox.showerror("致命错误", f"应用程序启动失败: {str(e)}")
        sys.exit(1)
