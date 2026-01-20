import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import os
import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import sys
from pathlib import Path

class SKVALGenerator:
    def __init__(self, root):
        self.root = root
        self.root.title("Генератор обозначения системы ШКВАЛ-Д")
        self.root.geometry("1400x900")
        
        # Стили
        self.setup_styles()
        
        # Данные
        self.elements = []
        self.element_id = 0
        self.additional_text_fields = {}
        
        # Создание интерфейса
        self.create_widgets()
        
    def setup_styles(self):
        # Настройка стилей для tkinter
        style = ttk.Style()
        style.configure('TFrame', background='#d5a6ff')
        style.configure('TLabel', background='#f9f9f9', font=('Segoe UI', 10))
        style.configure('TButton', font=('Segoe UI', 10))
        
    def create_widgets(self):
        # Главный контейнер с прокруткой
        main_frame = ttk.Frame(self.root)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        
        # Canvas для прокрутки
        canvas = tk.Canvas(main_frame, bg='#d5a6ff')
        scrollbar = ttk.Scrollbar(main_frame, orient="vertical", command=canvas.yview)
        self.scrollable_frame = ttk.Frame(canvas)
        
        self.scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # Контейнер приложения
        container = tk.Frame(self.scrollable_frame, bg='white', bd=0, highlightthickness=0)
        container.pack(fill=tk.BOTH, expand=True, padx=30, pady=30)
        
        # Заголовок
        title_frame = tk.Frame(container, bg='white')
        title_frame.pack(fill=tk.X, pady=(0, 20))
        
        title_label = tk.Label(title_frame, text="Генератор обозначения системы ШКВАЛ-Д", 
                              font=('Segoe UI', 24, 'bold'), bg='white', fg='#2c3e50')
        title_label.pack(side=tk.LEFT)
        
        # Кнопка инструкции
        instruction_btn = tk.Button(title_frame, text="Открыть инструкцию", 
                                   bg='#2ecc71', fg='white', font=('Segoe UI', 10, 'bold'),
                                   command=self.open_instruction)
        instruction_btn.pack(side=tk.RIGHT, padx=10)
        
        # Раздел: Данные для заполнения
        self.create_data_section(container)
        
        # Раздел: Основные параметры
        self.create_main_params_section(container)
        
        # Раздел: Элементы системы
        self.create_elements_section(container)
        
        # Раздел: Результат
        self.create_result_section(container)
        
        # Добавляем первый элемент
        self.add_element()
        
    def create_data_section(self, parent):
        frame = tk.LabelFrame(parent, text="Данные для заполнения", font=('Segoe UI', 12, 'bold'),
                             bg='#f9f9f9', fg='#2c3e50', padx=20, pady=20)
        frame.pack(fill=tk.X, pady=(0, 20))
        
        # Создаем сетку для полей
        for i in range(3):  # 3 ряда
            row_frame = tk.Frame(frame, bg='#f9f9f9')
            row_frame.pack(fill=tk.X, pady=5)
            
            for j in range(4):  # 4 колонки
                if i == 2 and j >= 3:  # В последнем ряду только 3 колонки
                    break
                    
                cell_frame = tk.Frame(row_frame, bg='#f9f9f9')
                cell_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=5)
                
                label_text, var_name, hint = self.get_data_field_info(i, j)
                
                label = tk.Label(cell_frame, text=label_text, bg='#f9f9f9', fg='#2c3e50',
                                font=('Segoe UI', 10, 'bold'), anchor='w')
                label.pack(fill=tk.X)
                
                if var_name == 'controllerType':
                    var = tk.StringVar(value='FireVent')
                    combo = ttk.Combobox(cell_frame, textvariable=var, 
                                        values=['FireVent', 'ПР200'], state='readonly')
                    combo.pack(fill=tk.X, pady=(5, 0))
                    setattr(self, var_name, var)
                else:
                    var = tk.StringVar()
                    entry = tk.Entry(cell_frame, textvariable=var, font=('Segoe UI', 10))
                    entry.pack(fill=tk.X, pady=(5, 0))
                    setattr(self, var_name, var)
                
                # Подсказка
                if hint:
                    hint_label = tk.Label(cell_frame, text=hint, bg='#f9f9f9', fg='#7f8c8d',
                                         font=('Segoe UI', 9, 'italic'), anchor='w', justify=tk.LEFT)
                    hint_label.pack(fill=tk.X, pady=(2, 0))
    
    def get_data_field_info(self, row, col):
        fields = [
            [("Номер Б/З", "kaNumber", "Для тега {KA} в документе"),
             ("Номер проекта", "projectNumber", "Для тега {indata} в документе"),
             ("Объект", "object", "Для тега {object} в документе"),
             ("Название системы", "systemName", "Для тега {name_sys} в документе")],
            [("Заказчик", "customer", "Для тега {org} в документе"),
             ("E-mail", "email", "Для тега {org_mail} в документе"),
             ("Телефон/факс", "phone", "Для тега {org_number} в документе"),
             ("Для представителя", "representative", "Для тега {org_name} в документе")],
            [("Разработал", "developer", "Для тега {v_inginer} в документе"),
             ("Менеджер", "manager", "Для тега {v_manager} в документе"),
             ("Тип контроллера", "controllerType", "Для тега {plc} в документе"),
             ("", "", "")]
        ]
        return fields[row][col]
    
    def create_main_params_section(self, parent):
        frame = tk.LabelFrame(parent, text="Основные параметры", font=('Segoe UI', 12, 'bold'),
                             bg='#f9f9f9', fg='#2c3e50', padx=20, pady=20)
        frame.pack(fill=tk.X, pady=(0, 20))
        
        row_frame = tk.Frame(frame, bg='#f9f9f9')
        row_frame.pack(fill=tk.X)
        
        # Напряжение электропитания
        left_frame = tk.Frame(row_frame, bg='#f9f9f9')
        left_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=5)
        
        tk.Label(left_frame, text="Напряжение электропитания (Ф1)", bg='#f9f9f9', 
                fg='#2c3e50', font=('Segoe UI', 10, 'bold'), anchor='w').pack(fill=tk.X)
        
        self.voltage = tk.StringVar(value='О1')
        voltage_combo = ttk.Combobox(left_frame, textvariable=self.voltage, 
                                    values=['О1 – 220 В, 2 ввода, АВР по питанию',
                                            'Ф – 380 В, 2 ввода, АВР по питанию'], 
                                    state='readonly', width=40)
        voltage_combo.pack(fill=tk.X, pady=(5, 0))
        
        # Основная пожарная зона
        right_frame = tk.Frame(row_frame, bg='#f9f9f9')
        right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=5)
        
        tk.Label(right_frame, text="Основная пожарная зона (Ф2)", bg='#f9f9f9', 
                fg='#2c3e50', font=('Segoe UI', 10, 'bold'), anchor='w').pack(fill=tk.X)
        
        self.mainFireZone = tk.StringVar(value='А')
        zone_combo = ttk.Combobox(right_frame, textvariable=self.mainFireZone, 
                                 values=['А – общие элементы для всех зон',
                                         'I – зона 1',
                                         'II – зона 2',
                                         'III – зона 3',
                                         'IV – зона 4'], 
                                 state='readonly', width=40)
        zone_combo.pack(fill=tk.X, pady=(5, 0))
        
        tk.Label(right_frame, text="Основная зона системы. Для элементов можно указывать другие зоны.",
                bg='#f9f9f9', fg='#7f8c8d', font=('Segoe UI', 9, 'italic'), 
                anchor='w', justify=tk.LEFT).pack(fill=tk.X, pady=(2, 0))
    
    def create_elements_section(self, parent):
        self.elements_frame = tk.LabelFrame(parent, text="Элементы системы дымоудаления (Ф3)", 
                                           font=('Segoe UI', 12, 'bold'),
                                           bg='#f9f9f9', fg='#2c3e50', padx=20, pady=20)
        self.elements_frame.pack(fill=tk.X, pady=(0, 20))
        
        # Примечание
        note_frame = tk.Frame(self.elements_frame, bg='#fff3cd', bd=1, relief=tk.SUNKEN)
        note_frame.pack(fill=tk.X, pady=(0, 20))
        
        note_text = """Примечание:
• Если ШКВАЛ-Д управляет только клапанами (без вентиляторов), установите количество вентиляторов = 0.
• Если в зоне нет клапанов, установите количество клапанов = 0.
• Максимальное количество вентиляторов (ВД + ВПД) в системе: 6 шт.
• Максимальное количество клапанов в одном элементе: 16 шт."""
        
        tk.Label(note_frame, text=note_text, bg='#fff3cd', fg='#856404',
                font=('Segoe UI', 9), justify=tk.LEFT, anchor='w').pack(padx=10, pady=10)
        
        # Контейнер для элементов
        self.elements_container = tk.Frame(self.elements_frame, bg='#f9f9f9')
        self.elements_container.pack(fill=tk.X)
        
        # Кнопка добавления элемента
        add_btn = tk.Button(self.elements_frame, text="+ Добавить элемент системы",
                           bg='#2ecc71', fg='white', font=('Segoe UI', 11, 'bold'),
                           command=self.add_element)
        add_btn.pack(pady=(20, 5))
        
        tk.Label(self.elements_frame, text="Добавьте несколько элементов для разных зон или разных типов оборудования",
                bg='#f9f9f9', fg='#7f8c8d', font=('Segoe UI', 9, 'italic')).pack()
    
    def create_result_section(self, parent):
        frame = tk.LabelFrame(parent, text="Результат", font=('Segoe UI', 12, 'bold'),
                             bg='#2c3e50', fg='white', padx=20, pady=20)
        frame.pack(fill=tk.X, pady=(0, 20))
        
        # Результат
        self.result_var = tk.StringVar(value="ШКВАЛ-Д-")
        result_label = tk.Label(frame, textvariable=self.result_var, 
                               font=('Courier', 22, 'bold'), bg='#34495e', 
                               fg='white', padx=15, pady=15, anchor='w', justify=tk.LEFT)
        result_label.pack(fill=tk.X, pady=(0, 20))
        
        # Кнопки
        btn_frame = tk.Frame(frame, bg='#2c3e50')
        btn_frame.pack()
        
        copy_btn = tk.Button(btn_frame, text="Копировать в буфер", 
                            bg='#9b59b6', fg='white', font=('Segoe UI', 11, 'bold'),
                            width=20, height=2, command=self.copy_to_clipboard)
        copy_btn.pack(side=tk.LEFT, padx=5)
        
        generate_btn = tk.Button(btn_frame, text="Сгенерировать бланк", 
                                bg='#e67e22', fg='white', font=('Segoe UI', 11, 'bold'),
                                width=20, height=2, command=self.generate_document)
        generate_btn.pack(side=tk.LEFT, padx=5)
    
    def add_element(self):
        self.element_id += 1
        element_id = self.element_id
        
        # Создаем фрейм для элемента
        element_frame = tk.Frame(self.elements_container, bg='#e8f4fc', 
                                bd=1, relief=tk.RIDGE)
        element_frame.pack(fill=tk.X, pady=(0, 10))
        
        # Заголовок элемента
        header_frame = tk.Frame(element_frame, bg='#e8f4fc')
        header_frame.pack(fill=tk.X, padx=20, pady=(10, 0))
        
        title_label = tk.Label(header_frame, text=f"Элемент системы #{element_id}", 
                              font=('Segoe UI', 11, 'bold'), bg='#e8f4fc', fg='#2c3e50')
        title_label.pack(side=tk.LEFT)
        
        if element_id > 1:
            remove_btn = tk.Button(header_frame, text="Удалить", 
                                  bg='#e74c3c', fg='white', font=('Segoe UI', 9),
                                  command=lambda: self.remove_element(element_id))
            remove_btn.pack(side=tk.RIGHT)
        
        # Зона элемента
        zone_frame = tk.Frame(element_frame, bg='#d4edda', bd=1, relief=tk.SUNKEN)
        zone_frame.pack(fill=tk.X, padx=20, pady=10)
        
        tk.Label(zone_frame, text="Зона элемента", bg='#d4edda', fg='#155724',
                font=('Segoe UI', 10, 'bold'), anchor='w').pack(fill=tk.X, padx=10, pady=(10, 0))
        
        element_zone_var = tk.StringVar(value='А')
        zone_combo = ttk.Combobox(zone_frame, textvariable=element_zone_var, 
                                 values=['А – общие элементы',
                                         'I – зона 1',
                                         'II – зона 2',
                                         'III – зона 3',
                                         'IV – зона 4'], 
                                 state='readonly')
        zone_combo.pack(fill=tk.X, padx=10, pady=(5, 10))
        
        # Основные параметры элемента
        content_frame = tk.Frame(element_frame, bg='#e8f4fc')
        content_frame.pack(fill=tk.X, padx=20, pady=10)
        
        # Создаем сетку для полей элемента
        self.create_element_fields(content_frame, element_id, element_zone_var)
        
        # Сохраняем элемент
        element_data = {
            'id': element_id,
            'frame': element_frame,
            'zone_var': element_zone_var
        }
        self.elements.append(element_data)
        
        # Обновляем результат
        self.update_result()
    
    def create_element_fields(self, parent, element_id, zone_var):
        # Создаем все переменные для элемента
        vars_dict = {}
        
        # Вентиляторы
        ventilator_count_var = tk.StringVar(value='1')
        ventilator_type_var = tk.StringVar(value='ВД')
        motor_power_var = tk.StringVar(value='0.18')
        start_type_var = tk.StringVar(value='П')
        
        # Клапаны
        valve_count_var = tk.StringVar(value='1')
        actuator_count_var = tk.StringVar(value='1')
        valve_type_var = tk.StringVar(value='КЗ')
        valve_voltage_var = tk.StringVar(value='Р')
        heating_type_var = tk.StringVar(value='Х')
        heating_power_var = tk.StringVar(value='0')
        
        # Нагреватель
        heater_presence_var = tk.StringVar(value='X')
        heater_power_var = tk.StringVar(value='0.1')
        
        # Сохраняем переменные
        vars_dict['ventilator_count'] = ventilator_count_var
        vars_dict['ventilator_type'] = ventilator_type_var
        vars_dict['motor_power'] = motor_power_var
        vars_dict['start_type'] = start_type_var
        vars_dict['valve_count'] = valve_count_var
        vars_dict['actuator_count'] = actuator_count_var
        vars_dict['valve_type'] = valve_type_var
        vars_dict['valve_voltage'] = valve_voltage_var
        vars_dict['heating_type'] = heating_type_var
        vars_dict['heating_power'] = heating_power_var
        vars_dict['heater_presence'] = heater_presence_var
        vars_dict['heater_power'] = heater_power_var
        
        # Связываем обновления
        for var in [ventilator_count_var, ventilator_type_var, motor_power_var, 
                   start_type_var, valve_count_var, actuator_count_var, 
                   valve_type_var, valve_voltage_var, heating_type_var, 
                   heating_power_var, heater_presence_var, heater_power_var, zone_var]:
            var.trace_add('write', lambda *args: self.update_result())
        
        # Первый ряд: Вентиляторы
        row1 = tk.Frame(parent, bg='#e8f4fc')
        row1.pack(fill=tk.X, pady=5)
        
        # Количество вентиляторов
        self.create_labeled_entry(row1, "Количество вентиляторов", ventilator_count_var, 
                                 "0 = только клапаны, без вентиляторов", width=10)
        
        # Тип вентилятора
        self.create_labeled_combobox(row1, "Тип вентилятора (Ф31)", ventilator_type_var,
                                    ['ВД – вентилятор дымоудаления', 'ВПД – вентилятор подпора'])
        
        # Мощность двигателя
        self.create_labeled_entry(row1, "Мощность двигателя, кВт (Ф32)", motor_power_var,
                                 "Введите мощность двигателя в кВт (0.18…110.00)", width=10)
        
        # Тип запуска
        self.create_labeled_combobox(row1, "Тип запуска (Ф33)", start_type_var,
                                    ['П – прямой пуск', 'С – софт-стартер', 
                                     'Т – пуск звезда-треугольник', 'Ч – частотный преобразователь'])
        
        # Второй ряд: Клапаны
        row2 = tk.Frame(parent, bg='#e8f4fc')
        row2.pack(fill=tk.X, pady=5)
        
        # Количество клапанов
        self.create_labeled_entry(row2, "Количество клапанов (Ф35)", valve_count_var,
                                 "0 = нет клапанов в этой зоне", width=10)
        
        # Количество приводов
        self.create_labeled_entry(row2, "приводов на клапане", actuator_count_var,
                                 "Укажите, если более 1 привода на одном клапане", width=10)
        
        # Тип клапана
        self.create_labeled_combobox(row2, "Тип клапана (Ф34)", valve_type_var,
                                    ['КЗ – нормально-закрытый', 'КО – нормально-открытый', 'КЭ – электромагнитный'])
        
        # Напряжение клапана
        self.create_labeled_combobox(row2, "Напряжение клапана (Ф36)", valve_voltage_var,
                                    ['Р – 220 В', 'М – 24 В'])
        
        # Обогрев клапана
        self.create_labeled_heating(row2, "Обогрев клапана (Ф37)", heating_type_var, heating_power_var,
                                   "Укажите мощность, если ≥ 0.3 кВт (в кВт)")
        
        # Третий ряд: Нагреватель
        row3 = tk.Frame(parent, bg='#e8f4fc')
        row3.pack(fill=tk.X, pady=5)
        
        # Наличие нагревателя
        self.create_labeled_combobox(row3, "Электронагреватель (Ф38)", heater_presence_var,
                                    ['X – отсутствует', 'Э – имеется'], width=15)
        
        # Мощность нагревателя
        self.create_labeled_entry(row3, "Мощность нагревателя, кВт (Ф39)", heater_power_var,
                                 "Введите мощность нагревателя в кВт (0.1…30.0)", width=10)
        
        # Сохраняем переменные для элемента
        setattr(self, f'element_{element_id}_vars', vars_dict)
    
    def create_labeled_entry(self, parent, label_text, variable, hint=None, width=20):
        frame = tk.Frame(parent, bg='#e8f4fc')
        frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=2)
        
        tk.Label(frame, text=label_text, bg='#e8f4fc', fg='#2c3e50',
                font=('Segoe UI', 9, 'bold'), anchor='w').pack(fill=tk.X)
        
        entry = tk.Entry(frame, textvariable=variable, font=('Segoe UI', 10), width=width)
        entry.pack(fill=tk.X, pady=(2, 0))
        
        if hint:
            tk.Label(frame, text=hint, bg='#e8f4fc', fg='#7f8c8d',
                    font=('Segoe UI', 8, 'italic'), anchor='w').pack(fill=tk.X)
    
    def create_labeled_combobox(self, parent, label_text, variable, values, width=20):
        frame = tk.Frame(parent, bg='#e8f4fc')
        frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=2)
        
        tk.Label(frame, text=label_text, bg='#e8f4fc', fg='#2c3e50',
                font=('Segoe UI', 9, 'bold'), anchor='w').pack(fill=tk.X)
        
        # Берем только значение до дефиса для сохранения
        display_values = values
        actual_values = [v.split(' – ')[0] for v in values]
        
        combo = ttk.Combobox(frame, textvariable=variable, values=display_values,
                            state='readonly', width=width)
        combo.pack(fill=tk.X, pady=(2, 0))
        
        # Сохраняем фактическое значение
        def on_select(event):
            index = combo.current()
            if index >= 0:
                variable.set(actual_values[index])
        
        combo.bind('<<ComboboxSelected>>', on_select)
    
    def create_labeled_heating(self, parent, label_text, type_var, power_var, hint=None):
        frame = tk.Frame(parent, bg='#e8f4fc')
        frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=2)
        
        tk.Label(frame, text=label_text, bg='#e8f4fc', fg='#2c3e50',
                font=('Segoe UI', 9, 'bold'), anchor='w').pack(fill=tk.X)
        
        # Фрейм для выбора типа и мощности
        heating_frame = tk.Frame(frame, bg='#e8f4fc')
        heating_frame.pack(fill=tk.X, pady=(2, 0))
        
        # Комбобокс для типа
        type_combo = ttk.Combobox(heating_frame, textvariable=type_var,
                                 values=['Х – отсутствует', 'К – имеется'],
                                 state='readonly', width=8)
        type_combo.pack(side=tk.LEFT)
        
        # Поле для мощности
        power_entry = tk.Entry(heating_frame, textvariable=power_var, 
                              font=('Segoe UI', 10), width=12)
        power_entry.pack(side=tk.LEFT, padx=(5, 0))
        
        if hint:
            tk.Label(frame, text=hint, bg='#e8f4fc', fg='#7f8c8d',
                    font=('Segoe UI', 8, 'italic'), anchor='w').pack(fill=tk.X)
    
    def remove_element(self, element_id):
        for element in self.elements:
            if element['id'] == element_id:
                element['frame'].destroy()
                self.elements.remove(element)
                break
        
        # Обновляем результат
        self.update_result()
    
    def power_to_index(self, power):
        try:
            power_num = float(str(power).replace(',', '.'))
            index = round(power_num * 100)
            return str(index).zfill(5)
        except:
            return "00000"
    
    def heater_power_to_index(self, power):
        try:
            power_num = float(str(power).replace(',', '.'))
            index = round(power_num * 10)
            return str(index).zfill(3)
        except:
            return "000"
    
    def update_result(self):
        # Получаем основные параметры
        voltage = self.voltage.get().split(' – ')[0] if ' – ' in self.voltage.get() else self.voltage.get()
        main_zone = self.mainFireZone.get().split(' – ')[0] if ' – ' in self.mainFireZone.get() else self.mainFireZone.get()
        
        result = f"ШКВАЛ-Д-{voltage}-{main_zone}"
        
        # Группируем элементы по зонам
        zone_elements = {}
        
        for element in self.elements:
            element_id = element['id']
            zone = element['zone_var'].get().split(' – ')[0] if ' – ' in element['zone_var'].get() else element['zone_var'].get()
            
            # Получаем переменные элемента
            vars_dict = getattr(self, f'element_{element_id}_vars', {})
            
            if not vars_dict:
                continue
            
            # Получаем значения
            ventilator_count = int(vars_dict['ventilator_count'].get() or 0)
            ventilator_type = vars_dict['ventilator_type'].get()
            motor_power = vars_dict['motor_power'].get()
            start_type = vars_dict['start_type'].get()
            valve_count = int(vars_dict['valve_count'].get() or 0)
            actuator_count = int(vars_dict['actuator_count'].get() or 1)
            valve_type = vars_dict['valve_type'].get()
            valve_voltage = vars_dict['valve_voltage'].get()
            heating_type = vars_dict['heating_type'].get()
            heating_power = float(vars_dict['heating_power'].get() or 0)
            heater_presence = vars_dict['heater_presence'].get()
            heater_power_val = vars_dict['heater_power'].get()
            
            # Формируем строку элемента
            element_str = ""
            
            if ventilator_count == 0 and valve_count == 0:
                continue
            elif ventilator_count == 0 and valve_count > 0:
                # Только клапаны
                element_str = f"{valve_type}{valve_count}"
                if actuator_count > 1:
                    element_str += f"х{actuator_count}"
                element_str += f"-{valve_voltage}-"
                if heating_type == 'К' and heating_power >= 0.3:
                    element_str += f"К({heating_power:.1f})".replace('.', ',')
                else:
                    element_str += "Х"
            elif ventilator_count > 0 and valve_count == 0:
                # Только вентиляторы
                if ventilator_count > 1:
                    element_str += str(ventilator_count)
                power_index = self.power_to_index(motor_power)
                element_str += f"{ventilator_type}{power_index}{start_type}-Х"
                if ventilator_type == 'ВПД':
                    element_str += f"-{heater_presence}"
                    if heater_presence == 'Э':
                        heater_index = self.heater_power_to_index(heater_power_val)
                        element_str += heater_index
            else:
                # И вентиляторы, и клапаны
                if ventilator_count > 1:
                    element_str += str(ventilator_count)
                power_index = self.power_to_index(motor_power)
                element_str += f"{ventilator_type}{power_index}{start_type}-{valve_type}{valve_count}"
                if actuator_count > 1:
                    element_str += f"х{actuator_count}"
                element_str += f"-{valve_voltage}-"
                if heating_type == 'К' and heating_power >= 0.3:
                    element_str += f"К({heating_power:.1f})".replace('.', ',')
                else:
                    element_str += "Х"
                if ventilator_type == 'ВПД':
                    element_str += f"-{heater_presence}"
                    if heater_presence == 'Э':
                        heater_index = self.heater_power_to_index(heater_power_val)
                        element_str += heater_index
            
            # Добавляем элемент в зону
            if zone not in zone_elements:
                zone_elements[zone] = []
            zone_elements[zone].append(element_str)
        
        # Формируем итоговую строку
        # Сначала основная зона
        if main_zone in zone_elements and zone_elements[main_zone]:
            if len(zone_elements[main_zone]) == 1:
                result += f"({zone_elements[main_zone][0]})"
            else:
                result += f"({' + '.join(zone_elements[main_zone])})"
            del zone_elements[main_zone]
        
        # Затем остальные зоны в порядке
        for zone in ['А', 'I', 'II', 'III', 'IV']:
            if zone in zone_elements and zone_elements[zone]:
                if result.endswith(')'):
                    result += f"+{zone}({' + '.join(zone_elements[zone])})"
                else:
                    result += f"{zone}({' + '.join(zone_elements[zone])})"
        
        self.result_var.set(result)
    
    def copy_to_clipboard(self):
        self.root.clipboard_clear()
        self.root.clipboard_append(self.result_var.get())
        messagebox.showinfo("Скопировано", "Текст скопирован в буфер обмена")
    
    def generate_document(self):
        # Проверяем обязательные поля
        if not self.kaNumber.get() or not self.systemName.get():
            messagebox.showerror("Ошибка", "Пожалуйста, заполните обязательные поля: 'Номер Б/З' и 'Название системы'")
            return
        
        try:
            # Создаем документ
            doc = Document()
            
            # Настройка стилей
            self.setup_document_styles(doc)
            
            # Добавляем содержимое
            self.add_document_content(doc)
            
            # Сохраняем файл на рабочий стол
            desktop = Path.home() / "Desktop"
            ka_number = self.kaNumber.get().strip()
            system_name = self.systemName.get().strip()
            filename = f"КА_{ka_number}-СПБ {system_name}.docx"
            filepath = desktop / filename
            
            doc.save(filepath)
            
            messagebox.showinfo("Успешно", f"Документ сохранен на рабочем столе:\n{filename}")
            
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось создать документ:\n{str(e)}")
    
    def setup_document_styles(self, doc):
        # Основной стиль документа
        style = doc.styles['Normal']
        style.font.name = 'Cambria'
        style.font.size = Pt(11)
    
    def add_document_content(self, doc):
        # Получаем данные
        ka_number = self.kaNumber.get()
        project_number = self.projectNumber.get()
        object_name = self.object.get()
        system_name = self.systemName.get()
        customer = self.customer.get()
        email = self.email.get()
        phone = self.phone.get()
        representative = self.representative.get()
        developer = self.developer.get()
        manager = self.manager.get()
        controller_type = self.controllerType.get()
        current_date = datetime.datetime.now().strftime("%d.%m.%Y")
        result_text = self.result_var.get()
        
        # Заголовок
        title = doc.add_heading(f'«КА» Комплект Автоматики № {ka_number}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Информационная таблица
        table = doc.add_table(rows=5, cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Настройка ширины колонок
        widths = [Inches(1.5), Inches(3.0), Inches(1.2), Inches(2.5)]
        for i, width in enumerate(widths):
            table.columns[i].width = width
        
        # Заполняем таблицу
        cells_data = [
            ["Объект", object_name, "Тип шкафа ШКВАЛ", result_text],
            ["Название системы", system_name, "", ""],
            ["Заказчик", customer, "Установочная N, кВт", ""],
            ["E-mail", email, "Разработчик от «ВЕЗА»", developer],
            ["Телефон/Факс", phone, "Подпись разработчика", ""],
            ["Для специалиста по автоматике", representative, "Инженер", manager]
        ]
        
        for i, row_data in enumerate(cells_data):
            if i >= len(table.rows):
                table.add_row()
            for j, cell_data in enumerate(row_data):
                cell = table.cell(i, j)
                cell.text = cell_data
                paragraph = cell.paragraphs[0]
                paragraph.style.font.size = Pt(10 if j in [0, 2] else 12)
                paragraph.style.font.bold = (j in [0, 2])
        
        # Разделитель
        doc.add_paragraph()
        
        # Основные положения
        doc.add_heading('1. Шкаф управления для систем приточно-вытяжной противодымной вентиляции ШКВАЛ', level=1)
        
        # Контроллер и габариты
        controller_table = doc.add_table(rows=1, cols=3)
        controller_table.style = 'Table Grid'
        
        controller_cells = controller_table.rows[0].cells
        controller_cells[0].text = "1. Шкаф управления для систем приточно-вытяжной противодымной вентиляции ШКВАЛ"
        controller_cells[1].text = f"Контроллер:\n{controller_type}"
        controller_cells[2].text = "Габариты (ВхШхГ):\nОпределяет завод"
        
        for cell in controller_cells:
            for paragraph in cell.paragraphs:
                paragraph.style.font.bold = True
        
        doc.add_paragraph()
        
        # Перечень приборов
        doc.add_heading('2. Перечень приборов автоматики, входящих в комплект поставки:', level=1)
        
        devices_table = doc.add_table(rows=3, cols=4)
        devices_table.style = 'Table Grid'
        
        # Заголовки
        headers = ["№", "Наименование", "Марка", "Кол-во"]
        for i, header in enumerate(headers):
            devices_table.cell(0, i).text = header
            devices_table.cell(0, i).paragraphs[0].style.font.bold = True
        
        # Данные
        devices_table.cell(1, 0).text = "2.1"
        devices_table.cell(1, 1).text = "Прибор управления пожарный"
        devices_table.cell(1, 2).text = "ШКВАЛ"
        devices_table.cell(1, 3).text = "1 шт."
        
        devices_table.cell(2, 0).text = ""
        devices_table.cell(2, 1).text = ""
        devices_table.cell(2, 2).text = "Итого:"
        devices_table.cell(2, 3).text = "1 шт."
        devices_table.cell(2, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        devices_table.cell(2, 2).paragraphs[0].style.font.bold = True
        
        doc.add_paragraph()
        
        # Основные положения
        doc.add_heading('3. Основные положения комплектов автоматики ШКВАЛ-ВЕЗА:', level=1)
        
        positions = [
            "Шкафы управления для систем приточно-вытяжной противодымной вентиляции, серии «ШКВАЛ» соответствуют техническим условиям ТУ 4371-172-40149153-2014.",
            "Стандартно шкаф имеет степень защиты IP54 по ГОСТ 14254. Температура эксплуатации от 0°С до 40°С по ГОСТ 30631. Относительная влажность 93% при температуре плюс 40°С.",
            "Шкаф ШКВАЛ изготавливается в виде настенного шкафа, совмещающего автоматику и силовую часть. Сетевой фидер, силовые выходы на управляемые устройства и внешние связи вводятся в шкаф через кабельные вводы, расположенные на нижней стенке шкафа.",
            "Шкаф оснащен запираемой дверцей, на которой установлены органы управления и индикации согласно ГОСТ Р 53325-2012.",
            "Питание шкафа ШКВАЛ осуществляется от сети переменного тока 380В частотой 50 Гц, либо 220В в зависимости от исполнения, с глухозаземлённой нейтралью.",
            "Шкаф ШКВАЛ имеет два вода питания и АВР по питанию согласно ГОСТ Р 53325-2012. Установочная мощность шкафа определяется суммарной мощностью коммутируемых элементов.",
            "Шкаф стандартно имеет вход для подключения сигнала от приборов пожарной сигнализации «Пожар» (тип «сухой контакт» Н.О.) для каждой из предусмотренных пожарных зон;",
            "Шкаф управления для систем приточно-вытяжной противодымной вентиляции ШКВАЛ осуществляет контроль целостности линий связи между шкафом и исполнительными устройствами систем противопожарной защиты, техническими средствами, формирующими сигнал «Пожар» и техническими средствами, регистрирующими срабатывание средств противопожарной защиты.",
            "ШКВАЛ имеет функцию тестирования работоспособности устройств звуковой сигнализации световой индикации, расположенных на лицевой панели шкафа.",
            "ШКВАЛ обеспечивает световую индикацию и звуковую сигнализацию (не менее 60 дБ на расстоянии 1-го метра от шкафа) в соответствии с ГОСТ Р 53325-2012.",
            "Шкаф ШКВАЛ обеспечивает возможность автоматического и ручного, местного и дистанционного управления исполнительными устройствами. Выбор способа управления защищен от несанкционированного доступа.",
            "Разработчик оставляет за собой право вносить изменения, не влияющие на основные функции системы без предварительного уведомления с сохранением технических характеристик.",
            "Внимание: в соответствии с п. 7.20 СП 7 7.13130.2013, предусмотрена задержка включения приточной противодымной вентиляции на 25с относительно включения вытяжной противодымной вентиляции.",
            "Согласно п. 7.22, СП 7 7.13130.2013, применение устройств автоматического отключения в цепях электроснабжения исполнительных элементов оборудования систем противодымной вентиляции не допускается и не реализуется в шкафу ШКВАЛ.",
            "При наличии в КА дополнительных требований, противоречащих основным положениям комплектов автоматики и не противоречащих ГОСТ Р 53325-2012, приоритетными являются дополнительные требования."
        ]
        
        for i, position in enumerate(positions, 1):
            p = doc.add_paragraph(f"{i}. {position}")
            p.style.font.size = Pt(9)
        
        doc.add_paragraph()
        
        # Дополнительные требования
        doc.add_heading('4. Дополнительные требования:', level=1)
        
        doc.add_paragraph(f"4.1 Шкаф предназначен для управления системами {system_name}:", style='Heading 2')
        doc.add_paragraph()
        
        # Добавляем зоны (можно расширить для заполнения дополнительных данных)
        zones = ['Общие элементы противодымной системы для всех зон:',
                'Элементы противодымной системы зоны I:',
                'Элементы противодымной системы зоны II:',
                'Элементы противодымной системы зоны III:',
                'Элементы противодымной системы зоны IV:']
        
        for zone in zones:
            p = doc.add_paragraph(zone)
            p.style.font.bold = True
        
        doc.add_paragraph()
        
        # Дискретные выходы
        doc.add_paragraph("4.2 Дискретные выходы (стандартно):")
        doc.add_paragraph("- Н.О. 'сухой' контакт 'ПУСК';")
        doc.add_paragraph("- Н.О. 'сухой' контакт 'НЕИСПРАВНОСТЬ'.")
        
        doc.add_paragraph()
        
        # Внимание
        warning_p = doc.add_paragraph("4.3 ВНИМАНИЮ ПРОЕКТНОЙ ОРГАНИЗАЦИИ согласно ГОСТ Р 53325 п.7.6.1.15 и п.7.6.2.4 сброс сигнала «ПОЖАР» в ППУ Шквал возможен только вручную, с передней панели шкафа.")
        warning_p.runs[0].font.color.rgb = RGBColor(0, 112, 192)
        
        doc.add_paragraph()
        
        # Подпись
        sign_p = doc.add_paragraph("«КА» СОГЛАСОВАНО, Заказчик: ____________________ДАТА_______________2026")
        sign_p.runs[0].underline = True
        sign_p.runs[0].bold = True
        
        warning_p2 = doc.add_paragraph("ВНИМАНИЕ! оплата СЧЕТА одновременно является СОГЛАСОВАНИЕМ «КА»")
        warning_p2.runs[0].underline = True
        warning_p2.runs[0].bold = True
    
    def open_instruction(self):
        import webbrowser
        webbrowser.open("https://example.com/instruction")

def main():
    root = tk.Tk()
    app = SKVALGenerator(root)
    root.mainloop()

if __name__ == "__main__":
    main()