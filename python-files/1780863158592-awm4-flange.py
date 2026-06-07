import os
import winsound
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def play_start_sound():
    try:
        winsound.Beep(880, 300)
        winsound.Beep(1175, 200)
        winsound.Beep(880, 300)
        winsound.Beep(1175, 200)
        winsound.Beep(1397, 400)
    except:
        pass

def get_flange_dimensions(DN):
    db = {
        15: {"D": 95, "D1": 65, "d_bolt": 14, "n": 4, "b": 12, "weight": 0.68},
        20: {"D": 105, "D1": 75, "d_bolt": 14, "n": 4, "b": 12, "weight": 0.87},
        25: {"D": 115, "D1": 85, "d_bolt": 14, "n": 4, "b": 12, "weight": 1.05},
        32: {"D": 135, "D1": 100, "d_bolt": 18, "n": 4, "b": 13, "weight": 1.54},
        40: {"D": 145, "D1": 110, "d_bolt": 18, "n": 4, "b": 13, "weight": 1.85},
        50: {"D": 160, "D1": 125, "d_bolt": 18, "n": 4, "b": 13, "weight": 2.28},
        65: {"D": 180, "D1": 145, "d_bolt": 18, "n": 4, "b": 15, "weight": 3.19},
        80: {"D": 195, "D1": 160, "d_bolt": 18, "n": 4, "b": 17, "weight": 4.21},
        100: {"D": 215, "D1": 180, "d_bolt": 18, "n": 8, "b": 17, "weight": 4.9},
        125: {"D": 245, "D1": 210, "d_bolt": 18, "n": 8, "b": 19, "weight": 6.75},
        150: {"D": 280, "D1": 240, "d_bolt": 22, "n": 8, "b": 19, "weight": 8.3},
        200: {"D": 335, "D1": 295, "d_bolt": 22, "n": 12, "b": 21, "weight": 11.79},
        250: {"D": 405, "D1": 355, "d_bolt": 26, "n": 12, "b": 23, "weight": 17.36},
        300: {"D": 460, "D1": 410, "d_bolt": 26, "n": 12, "b": 24, "weight": 22.76},
        350: {"D": 520, "D1": 470, "d_bolt": 26, "n": 16, "b": 28, "weight": 32.04},
        400: {"D": 580, "D1": 525, "d_bolt": 30, "n": 16, "b": 32, "weight": 43.00},
        500: {"D": 710, "D1": 650, "d_bolt": 33, "n": 20, "b": 38, "weight": 70.97},
        600: {"D": 840, "D1": 770, "d_bolt": 39, "n": 20, "b": 41, "weight": 99.30},
        800: {"D": 1020, "D1": 950, "d_bolt": 39, "n": 24, "b": 45, "weight": 130.57},
        1000: {"D": 1255, "D1": 1170, "d_bolt": 45, "n": 28, "b": 49, "weight": 203.39},
        1200: {"D": 1485, "D1": 1390, "d_bolt": 52, "n": 32, "b": 51, "weight": 284.94},
    }
    return db.get(DN)

def create_word_table(DN, PN, flange_name, flange_code, face_type, dims, save_folder):
    doc = Document()
    title = doc.add_heading(f'ФЛАНЕЦ {flange_name.upper()}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(f'Тип {flange_code} | Исполнение {face_type} | ГОСТ 33259-2015')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    table = doc.add_table(rows=10, cols=2)
    table.style = 'Light Grid Accent 1'
    table.columns[0].width = Cm(5.5)
    table.columns[1].width = Cm(8.5)
    face_names = {'B': 'B - Выступ', 'E': 'E - Шип', 'F': 'F - Впадина'}
    data = [
        ("Тип фланца", f"{flange_name} (тип {flange_code})"),
        ("Исполнение уплотнительной поверхности", face_names.get(face_type, face_type)),
        ("Условный проход DN, мм", str(DN)),
        ("Номинальное давление PN, кгс/см²", str(PN)),
        ("Наружный диаметр D, мм", str(dims['D'])),
        ("Диаметр окружности болтов D1, мм", str(dims['D1'])),
        ("Толщина фланца b, мм", str(dims['b'])),
        ("Диаметр отверстия под болт d, мм", str(dims['d_bolt'])),
        ("Количество отверстий n", str(dims['n'])),
        ("Масса, кг (справочно)", str(dims['weight']))
    ]
    for i, (param, value) in enumerate(data):
        table.cell(i, 0).text = param
        table.cell(i, 1).text = value
        table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
    filename = f"Фланец_{flange_code}_DN{DN}_PN{PN}_{face_type}.docx"
    filepath = os.path.join(save_folder, filename)
    doc.save(filepath)
    return filepath

def play_success_sound():
    try:
        winsound.Beep(1047, 300)
        winsound.Beep(1319, 300)
        winsound.Beep(1568, 500)
    except:
        pass

def main():
    play_start_sound()
    print("\n" + "="*60)
    print("   ФЛАНЦЫ ПО ГОСТ 33259-2015")
    print("   Автор: SuperVardan")
    print("="*60)
    program_folder = os.path.dirname(os.path.abspath(__file__))
    if not program_folder:
        program_folder = os.getcwd()
    print(f"\n📁 Файлы будут сохранены в: {program_folder}")
    print("\nТИПЫ ФЛАНЦЕВ:")
    print("  01 - Плоский приварной")
    print("  02 - Плоский приварной (с отверстиями на D1)")
    print("  11 - Воротниковый приварной")
    print("  12 - Воротниковый приварной (с отверстиями на D1)")
    print("  21 - Свободный на приварном кольце")
    print("  31 - Накидной на приварном кольце")
    flange_code = input("\nВаш выбор (01/02/11/12/21/31): ").strip()
    flange_types = {
        "01": "Плоский приварной",
        "02": "Плоский приварной (с отверстиями на D1)",
        "11": "Воротниковый приварной",
        "12": "Воротниковый приварной (с отверстиями на D1)",
        "21": "Свободный на приварном кольце",
        "31": "Накидной на приварном кольце"
    }
    if flange_code not in flange_types:
        print("❌ Неверный выбор")
        input("Нажмите Enter...")
        return
    flange_name = flange_types[flange_code]
    print("\nИСПОЛНЕНИЕ ПОВЕРХНОСТИ:")
    print("  B - Выступ")
    print("  E - Шип")
    print("  F - Впадина")
    face = input("Выбор (B/E/F): ").strip().upper()
    if face not in ['B', 'E', 'F']:
        print("❌ Неверный выбор")
        input("Нажмите Enter...")
        return
    try:
        dn = int(input("\nВведите DN (15-1200): ").strip())
        pn = int(input("Введите PN: ").strip())
    except:
        print("❌ Ошибка ввода")
        input("Нажмите Enter...")
        return
    dims = get_flange_dimensions(dn)
    if not dims:
        print(f"\n❌ Размеры для DN {dn} не найдены")
        input("Нажмите Enter...")
        return
    print("\n" + "-"*50)
    print(f"DN: {dn} | PN: {pn}")
    print(f"D = {dims['D']} мм | D1 = {dims['D1']} мм")
    print(f"b = {dims['b']} мм | Отверстия: {dims['n']} × {dims['d_bolt']} мм")
    print(f"Масса: {dims['weight']} кг")
    print("-"*50)
    try:
        filepath = create_word_table(dn, pn, flange_name, flange_code, face, dims, program_folder)
        print(f"\n✅ Word-файл создан: {filepath}")
        play_success_sound()
    except Exception as e:
        print(f"\n❌ Ошибка: {e}")
    input("\nНажмите Enter для выхода...")

if __name__ == "__main__":
    main()