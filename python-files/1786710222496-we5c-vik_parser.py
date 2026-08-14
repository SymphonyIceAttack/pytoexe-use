import os
import re
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from datetime import datetime
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.utils import get_column_letter
from docx import Document

def parse_vik_docx(file_path):
    doc = Document(file_path)
    data = {
        "joints": [],
        "category": "",
        "method": "",
        "tools": "",
        "drawing": "",
        "conclusion_no": "",
        "conclusion_date": "",
        "inspected_by": "",
        "supervisor": "",
        "issued_by": ""
    }
    for table in doc.tables:
        for row in table.rows:
            text = " ".join(cell.text.strip() for cell in row.cells)
            if "Категория сварного соединения" in text or "Category of the weld joint" in text:
                for cell in row.cells:
                    if "II" in cell.text or "I" in cell.text:
                        data["category"] = cell.text.strip()
            if "Методика контроля" in text or "Test technique" in text:
                for cell in row.cells:
                    if "РБ" in cell.text or "RB" in cell.text:
                        data["method"] = cell.text.strip()
            if "Контроль проводился с применением" in text or "Inspection was carried out using" in text:
                for cell in row.cells:
                    if "Штангенциркуль" in cell.text or "Caliper" in cell.text:
                        data["tools"] = cell.text.strip()
            if "Чертеж" in text or "Drawing" in text:
                for cell in row.cells:
                    if "RPR" in cell.text or "СФ" in cell.text:
                        data["drawing"] = cell.text.strip()
    for para in doc.paragraphs:
        if "Заключение №" in para.text or "Conclusion №" in para.text:
            match = re.search(r'[№№]?\s*(\d+-\d+)', para.text)
            if match:
                data["conclusion_no"] = match.group(1)
        if "от" in para.text or "as of" in para.text:
            match = re.search(r'(\d{2}\.\d{2}\.\d{4})', para.text)
            if match:
                data["conclusion_date"] = match.group(1)
    for table in doc.tables:
        for i, row in enumerate(table.rows):
            cells = row.cells
            if len(cells) >= 3:
                first_text = cells[0].text.strip()
                if re.match(r'^\d+\.$', first_text):
                    joint_no = cells[1].text.strip()
                    if joint_no and "Welded" not in joint_no and "№" not in joint_no:
                        size = cells[2].text.strip() if len(cells) > 2 else ""
                        scope = cells[3].text.strip() if len(cells) > 3 else "100%"
                        defects = cells[4].text.strip() if len(cells) > 4 else "Дефектов не обнаружено"
                        quality = cells[5].text.strip() if len(cells) > 5 else "А"
                        entry = cells[6].text.strip() if len(cells) > 6 else ""
                        data["joints"].append({
                            "no": joint_no,
                            "size": size,
                            "scope": scope,
                            "defects": defects,
                            "quality": quality,
                            "entry": entry,
                            "conclusion_no": data["conclusion_no"],
                            "date": data["conclusion_date"],
                            "category": data["category"],
                            "method": data["method"],
                            "tools": data["tools"],
                            "drawing": data["drawing"]
                        })
    return data

def process_folder(folder_path, progress_callback=None):
    all_joints = []
    files = [f for f in os.listdir(folder_path) if f.endswith(".docx") and "ВИК" in f]
    for i, file in enumerate(files):
        if progress_callback:
            progress_callback(i + 1, len(files), file)
        file_path = os.path.join(folder_path, file)
        try:
            data = parse_vik_docx(file_path)
            all_joints.extend(data["joints"])
        except Exception as e:
            print(f"Ошибка в файле {file}: {e}")
    return all_joints

def save_to_excel(data, output_path, column_names=None, date_from=None, date_to=None, group_by=None):
    if not data:
        return False, "Нет данных для сохранения"
    if date_from or date_to:
        filtered = []
        for row in data:
            if "date" in row and row["date"]:
                try:
                    d = datetime.strptime(row["date"], "%d.%m.%Y")
                    if date_from and d < date_from:
                        continue
                    if date_to and d > date_to:
                        continue
                except:
                    pass
            filtered.append(row)
        data = filtered
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Сводная ВИК"
    default_headers = ["Номер стыка", "Результат ВИК", "Дата контроля ВИК", "Категория ВИК",
                       "Методика контроля ВИК", "Средства контроля", "Объем контроля ВИК",
                       "Описание обнаруженных дефектов ВИК", "Оценка качества ВИК",
                       "№ заключения ВИК", "№ журнала ВИК", "№ записи в журнале ВИК",
                       "Контроль выполнил ВИК", "Заключение выдал ВИК",
                       "Руководитель работ по контролю ВИК", "Комментарии по ВИК",
                       "Дата Комиссии ВИК", "Чертеж (сварочный формуляр) №"]
    headers = column_names if column_names else default_headers
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal="center", wrap_text=True)
    for row_idx, item in enumerate(data, 2):
        row_data = [
            item.get("no", ""),
            "Сварные соединения признаны годными (А — годен)",
            item.get("date", ""),
            item.get("category", ""),
            item.get("method", ""),
            item.get("tools", ""),
            item.get("scope", "100%"),
            item.get("defects", "Дефектов не обнаружено"),
            item.get("quality", "А (годен)"),
            item.get("conclusion_no", ""),
            "",
            item.get("entry", ""),
            "",
            "",
            "",
            "",
            "",
            item.get("drawing", "")
        ]
        for col, value in enumerate(row_data, 1):
            ws.cell(row=row_idx, column=col, value=value)
    for col in range(1, len(headers) + 1):
        max_length = 15
        for row in range(1, min(len(data) + 2, 100)):
            cell_value = ws.cell(row=row, column=col).value
            if cell_value:
                max_length = max(max_length, len(str(cell_value)))
        ws.column_dimensions[get_column_letter(col)].width = min(max_length + 3, 50)
    if group_by:
        summary_ws = wb.create_sheet("Статистика")
        summary_ws.cell(row=1, column=1, value="Группировка по")
        summary_ws.cell(row=1, column=2, value=group_by)
        summary_ws.cell(row=2, column=1, value="Значение")
        summary_ws.cell(row=2, column=2, value="Количество")
        groups = {}
        for item in data:
            key = item.get(group_by, "Не указано")
            groups[key] = groups.get(key, 0) + 1
        row = 4
        for key, count in groups.items():
            summary_ws.cell(row=row, column=1, value=key)
            summary_ws.cell(row=row, column=2, value=count)
            row += 1
        summary_ws.cell(row=row + 2, column=1, value="Всего стыков")
        summary_ws.cell(row=row + 2, column=2, value=len(data))
        summary_ws.cell(row=row + 3, column=1, value="Годных (А)")
        summary_ws.cell(row=row + 3, column=2, value=sum(1 for d in data if d.get("quality") == "А"))
    wb.save(output_path)
    return True, f"Сохранено {len(data)} записей"

class VikParserApp:
    def __init__(self, root):
        self.root = root
        self.root.title("ВИК Сводная таблица - Парсер")
        self.root.geometry("750x650")
        self.root.resizable(True, True)
        self.folder_path = tk.StringVar()
        self.output_path = tk.StringVar(value="Сводная_ВИК.xlsx")
        self.date_from = tk.StringVar()
        self.date_to = tk.StringVar()
        self.group_by = tk.StringVar(value="№ заключения ВИК")
        self.create_widgets()
    def create_widgets(self):
        main_frame = ttk.Frame(self.root, padding=10)
        main_frame.pack(fill=tk.BOTH, expand=True)
        ttk.Label(main_frame, text="Парсер документов ВИК", font=("Arial", 16, "bold")).pack(pady=5)
        ttk.Label(main_frame, text="Извлечение данных из файлов ВИК-551.docx в Excel", font=("Arial", 10)).pack(pady=5)
        ttk.Separator(main_frame, orient='horizontal').pack(fill=tk.X, pady=10)
        folder_frame = ttk.LabelFrame(main_frame, text="1. Выберите папку с файлами ВИК", padding=10)
        folder_frame.pack(fill=tk.X, pady=5)
        ttk.Entry(folder_frame, textvariable=self.folder_path, width=60).pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        ttk.Button(folder_frame, text="Обзор...", command=self.select_folder).pack(side=tk.RIGHT, padx=5)
        settings_frame = ttk.LabelFrame(main_frame, text="2. Настройки экспорта", padding=10)
        settings_frame.pack(fill=tk.X, pady=5)
        file_frame = ttk.Frame(settings_frame)
        file_frame.pack(fill=tk.X, pady=2)
        ttk.Label(file_frame, text="Имя файла результата:").pack(side=tk.LEFT, padx=5)
        ttk.Entry(file_frame, textvariable=self.output_path, width=40).pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        date_frame = ttk.Frame(settings_frame)
        date_frame.pack(fill=tk.X, pady=2)
        ttk.Label(date_frame, text="Фильтр по дате (ДД.ММ.ГГГГ):").pack(side=tk.LEFT, padx=5)
        ttk.Label(date_frame, text="от").pack(side=tk.LEFT, padx=2)
        ttk.Entry(date_frame, textvariable=self.date_from, width=15).pack(side=tk.LEFT, padx=2)
        ttk.Label(date_frame, text="до").pack(side=tk.LEFT, padx=2)
        ttk.Entry(date_frame, textvariable=self.date_to, width=15).pack(side=tk.LEFT, padx=2)
        group_frame = ttk.Frame(settings_frame)
        group_frame.pack(fill=tk.X, pady=2)
        ttk.Label(group_frame, text="Группировка по:").pack(side=tk.LEFT, padx=5)
        group_options = ["Без группировки", "№ заключения ВИК", "Дата контроля ВИК", "Категория ВИК"]
        ttk.Combobox(group_frame, textvariable=self.group_by, values=group_options, state="readonly", width=25).pack(side=tk.LEFT, padx=5)
        rename_frame = ttk.LabelFrame(settings_frame, text="Переименование столбцов (через запятую, в порядке следования)", padding=5)
        rename_frame.pack(fill=tk.X, pady=5)
        self.rename_entry = ttk.Entry(rename_frame, width=80)
        self.rename_entry.pack(fill=tk.X, padx=5)
        self.rename_entry.insert(0, "Номер стыка,Результат ВИК,Дата контроля,Категория,Методика,Средства,Объем,Дефекты,Оценка,№ заключения,№ журнала,№ записи,Контроль выполнил,Заключение выдал,Руководитель,Комментарии,Дата комиссии,Чертеж")
        action_frame = ttk.Frame(main_frame)
        action_frame.pack(fill=tk.X, pady=10)
        self.process_btn = ttk.Button(action_frame, text="▶ ОБРАБОТАТЬ ВСЕ ФАЙЛЫ", command=self.process_files, width=30)
        self.process_btn.pack(pady=5)
        self.progress_var = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(main_frame, variable=self.progress_var, maximum=100)
        self.progress_bar.pack(fill=tk.X, pady=5)
        self.status_label = ttk.Label(main_frame, text="Готов к работе. Выберите папку с файлами.", foreground="blue")
        self.status_label.pack(pady=5)
        log_frame = ttk.LabelFrame(main_frame, text="Лог операций", padding=5)
        log_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        self.log_text = tk.Text(log_frame, height=8, wrap=tk.WORD)
        self.log_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scroll = ttk.Scrollbar(log_frame, command=self.log_text.yview)
        scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self.log_text.config(yscrollcommand=scroll.set)
    def log(self, message):
        from datetime import datetime
        timestamp = datetime.now().strftime("%H:%M:%S")
        self.log_text.insert(tk.END, f"[{timestamp}] {message}\n")
        self.log_text.see(tk.END)
        self.root.update()
    def select_folder(self):
        folder = filedialog.askdirectory(title="Выберите папку с файлами ВИК")
        if folder:
            self.folder_path.set(folder)
            self.log(f"Выбрана папка: {folder}")
    def process_files(self):
        folder = self.folder_path.get()
        if not folder or not os.path.exists(folder):
            messagebox.showerror("Ошибка", "Выберите существующую папку с файлами")
            return
        self.process_btn.config(state=tk.DISABLED)
        self.progress_var.set(0)
        self.log("Начинаю обработку файлов...")
        try:
            def progress_callback(current, total, filename):
                percent = (current / total) * 100
                self.progress_var.set(percent)
                self.status_label.config(text=f"Обработка: {filename} ({current}/{total})")
                self.root.update()
            all_data = process_folder(folder, progress_callback)
            self.log(f"Найдено {len(all_data)} записей о стыках")
            if not all_data:
                messagebox.showwarning("Предупреждение", "Не найдено данных в файлах")
                self.process_btn.config(state=tk.NORMAL)
                return
            date_from = None
            date_to = None
            if self.date_from.get():
                try:
                    date_from = datetime.strptime(self.date_from.get(), "%d.%m.%Y")
                except:
                    pass
            if self.date_to.get():
                try:
                    date_to = datetime.strptime(self.date_to.get(), "%d.%m.%Y")
                except:
                    pass
            column_names = None
            rename_text = self.rename_entry.get().strip()
            if rename_text:
                column_names = [col.strip() for col in rename_text.split(",") if col.strip()]
                if len(column_names) < 5:
                    column_names = None
            group_by = None
            if self.group_by.get() and self.group_by.get() != "Без группировки":
                group_map = { "№ заключения ВИК": "conclusion_no", "Дата контроля ВИК": "date", "Категория ВИК": "category" }
                group_by = group_map.get(self.group_by.get())
            output = self.output_path.get()
            if not output.endswith(".xlsx"):
                output += ".xlsx"
            success, message = save_to_excel( all_data, output, column_names=column_names, date_from=date_from, date_to=date_to, group_by=group_by )
            if success:
                self.log(f"✅ {message}")
                self.status_label.config(text=f"Готово! Файл сохранён: {output}", foreground="green")
                messagebox.showinfo("Готово", f"Файл сохранён:\n{os.path.abspath(output)}\n\n{message}")
            else:
                self.log(f"❌ {message}")
                messagebox.showerror("Ошибка", message)
        except Exception as e:
            self.log(f"❌ Ошибка: {str(e)}")
            messagebox.showerror("Ошибка", str(e))
        finally:
            self.progress_var.set(100)
            self.process_btn.config(state=tk.NORMAL)
            self.status_label.config(text="Готов")

if __name__ == "__main__":
    root = tk.Tk()
    app = VikParserApp(root)
    root.mainloop()