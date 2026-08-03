# -*- coding: utf-8 -*-
"""
模板文档生成器 v1.0
功能：
1. 管理多个文本模板（增删改）
2. 模板中使用 {字段名} 作为占位符
3. 选择模板后自动生成字段输入框
4. 填写字段值后生成最终文本
5. 支持导出为 .txt 和 .docx 格式

使用说明：
- 运行前请安装依赖：pip install PyQt5 python-docx
- 运行：python main.py
- 打包成EXE：pyinstaller -F -w -i icon.ico main.py
"""

import sys
import os
import json
import re
from datetime import datetime
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QListWidget, QListWidgetItem, QLabel, QLineEdit, QTextEdit,
    QPushButton, QScrollArea, QFrame, QSplitter, QMessageBox,
    QDialog, QDialogButtonBox, QFormLayout, QGroupBox, QFileDialog,
    QComboBox, QSpinBox, QCheckBox, QToolBar, QAction, QStatusBar
)
from PyQt5.QtCore import Qt, pyqtSignal, QSize
from PyQt5.QtGui import QFont, QIcon, QColor, QPalette

# 尝试导入python-docx，如果未安装则提示
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("警告: python-docx 未安装，无法导出 .docx 文件")
    print("请运行: pip install python-docx")


class TemplateManager:
    """模板数据管理类"""
    
    def __init__(self, data_file="templates.json"):
        self.data_file = data_file
        self.templates = []  # 每个模板: {"name": "模板名", "content": "模板内容"}
        self.load()
    
    def load(self):
        """从JSON文件加载模板数据"""
        if os.path.exists(self.data_file):
            try:
                with open(self.data_file, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    self.templates = data.get("templates", [])
            except Exception as e:
                print(f"加载模板失败: {e}")
                self.templates = []
        else:
            # 如果文件不存在，创建默认模板
            self.templates = [
                {
                    "name": "会议通知",
                    "content": "会议通知\n\n各位同事：\n\n兹定于 {date} 在 {location} 召开 {topic} 会议。\n\n请 {attendees} 准时参加。\n\n会议议程：\n{agenda}\n\n注意事项：\n{notes}\n\n特此通知。\n\n{organizer}\n{date}"
                },
                {
                    "name": "合同模板",
                    "content": "合同协议书\n\n甲方（委托方）：{party_a}\n乙方（受托方）：{party_b}\n\n一、项目名称：{project_name}\n二、项目内容：{project_content}\n三、合同金额：{amount} 元（大写：{amount_cn}）\n四、付款方式：{payment_method}\n五、履行期限：{deadline}\n六、违约责任：{liability}\n\n甲方签字：________    乙方签字：________\n日期：{sign_date}"
                },
                {
                    "name": "工作报告",
                    "content": "工作报告\n\n报告人：{reporter}\n报告日期：{report_date}\n\n一、工作概述\n{summary}\n\n二、完成情况\n{completed}\n\n三、存在问题\n{problems}\n\n四、下一步计划\n{next_plan}\n\n五、需协调事项\n{coordination}\n\n{reporter}\n{report_date}"
                }
            ]
            self.save()
    
    def save(self):
        """保存模板数据到JSON文件"""
        try:
            with open(self.data_file, "w", encoding="utf-8") as f:
                json.dump({"templates": self.templates}, f, ensure_ascii=False, indent=2)
            return True
        except Exception as e:
            print(f"保存模板失败: {e}")
            return False
    
    def add_template(self, name, content):
        """添加新模板"""
        # 检查名称是否已存在
        for t in self.templates:
            if t["name"] == name:
                return False, "模板名称已存在"
        self.templates.append({"name": name, "content": content})
        self.save()
        return True, "添加成功"
    
    def delete_template(self, index):
        """删除模板"""
        if 0 <= index < len(self.templates):
            del self.templates[index]
            self.save()
            return True
        return False
    
    def update_template(self, index, name, content):
        """更新模板"""
        if 0 <= index < len(self.templates):
            # 检查名称是否与其他模板冲突
            for i, t in enumerate(self.templates):
                if t["name"] == name and i != index:
                    return False, "模板名称已存在"
            self.templates[index] = {"name": name, "content": content}
            self.save()
            return True, "更新成功"
        return False, "模板不存在"
    
    def get_template(self, index):
        """获取模板"""
        if 0 <= index < len(self.templates):
            return self.templates[index]
        return None
    
    def get_template_names(self):
        """获取所有模板名称"""
        return [t["name"] for t in self.templates]
    
    def extract_fields(self, content):
        """从模板内容中提取所有占位符字段"""
        # 匹配 {字段名} 格式，字段名可以包含中文、字母、数字、下划线
        pattern = r'\{([^{}]+)\}'
        fields = re.findall(pattern, content)
        # 去重并保持顺序
        seen = set()
        unique_fields = []
        for f in fields:
            if f not in seen:
                seen.add(f)
                unique_fields.append(f)
        return unique_fields
    
    def generate_text(self, content, field_values):
        """根据字段值生成最终文本"""
        result = content
        for field, value in field_values.items():
            placeholder = "{" + field + "}"
            result = result.replace(placeholder, value)
        return result


class TemplateEditDialog(QDialog):
    """模板编辑对话框"""
    
    def __init__(self, template_name="", template_content="", parent=None):
        super().__init__(parent)
        self.template_name = template_name
        self.template_content = template_content
        self.init_ui()
    
    def init_ui(self):
        self.setWindowTitle("编辑模板")
        self.setMinimumSize(600, 500)
        
        layout = QVBoxLayout(self)
        
        # 名称输入
        name_layout = QHBoxLayout()
        name_layout.addWidget(QLabel("模板名称:"))
        self.name_edit = QLineEdit()
        self.name_edit.setText(self.template_name)
        self.name_edit.setPlaceholderText("请输入模板名称")
        name_layout.addWidget(self.name_edit)
        layout.addLayout(name_layout)
        
        # 内容输入
        layout.addWidget(QLabel("模板内容 (使用 {字段名} 作为占位符):"))
        self.content_edit = QTextEdit()
        self.content_edit.setPlainText(self.template_content)
        self.content_edit.setPlaceholderText("请输入模板内容，使用 {字段名} 作为占位符...")
        layout.addWidget(self.content_edit)
        
        # 提示信息
        tip_label = QLabel("💡 提示：在内容中使用 {字段名} 格式定义占位符，生成时将自动替换")
        tip_label.setStyleSheet("color: #666; font-size: 12px;")
        layout.addWidget(tip_label)
        
        # 按钮
        button_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        button_box.accepted.connect(self.accept)
        button_box.rejected.connect(self.reject)
        layout.addWidget(button_box)
    
    def get_data(self):
        """获取编辑后的数据"""
        return self.name_edit.text().strip(), self.content_edit.toPlainText()


class MainWindow(QMainWindow):
    """主窗口"""
    
    def __init__(self):
        super().__init__()
        self.manager = TemplateManager()
        self.current_index = -1
        self.field_widgets = {}  # 存储字段名和对应的输入控件
        self.init_ui()
        self.load_template_list()
        self.statusBar().showMessage("就绪")
    
    def init_ui(self):
        self.setWindowTitle("📄 模板文档生成器")
        self.setMinimumSize(1000, 700)
        
        # 创建中央部件
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QHBoxLayout(central_widget)
        main_layout.setContentsMargins(10, 10, 10, 10)
        main_layout.setSpacing(10)
        
        # ===== 左侧面板：模板列表 =====
        left_panel = QWidget()
        left_panel.setMaximumWidth(300)
        left_panel.setMinimumWidth(250)
        left_layout = QVBoxLayout(left_panel)
        left_layout.setContentsMargins(0, 0, 5, 0)
        
        # 标题
        title_label = QLabel("📋 模板列表")
        title_label.setFont(QFont("Microsoft YaHei", 12, QFont.Bold))
        left_layout.addWidget(title_label)
        
        # 模板列表
        self.template_list = QListWidget()
        self.template_list.itemClicked.connect(self.on_template_selected)
        self.template_list.setFont(QFont("Microsoft YaHei", 10))
        left_layout.addWidget(self.template_list)
        
        # 操作按钮
        btn_layout = QHBoxLayout()
        self.add_btn = QPushButton("➕ 添加")
        self.add_btn.clicked.connect(self.add_template)
        self.edit_btn = QPushButton("✏️ 编辑")
        self.edit_btn.clicked.connect(self.edit_template)
        self.del_btn = QPushButton("🗑️ 删除")
        self.del_btn.clicked.connect(self.delete_template)
        btn_layout.addWidget(self.add_btn)
        btn_layout.addWidget(self.edit_btn)
        btn_layout.addWidget(self.del_btn)
        left_layout.addLayout(btn_layout)
        
        # 左侧面板添加到主布局
        main_layout.addWidget(left_panel)
        
        # ===== 分割线 =====
        splitter = QSplitter(Qt.Vertical)
        
        # ===== 右侧上方面板：字段输入 =====
        self.input_panel = QWidget()
        input_layout = QVBoxLayout(self.input_panel)
        input_layout.setContentsMargins(5, 5, 5, 5)
        
        # 当前模板信息
        self.template_info_label = QLabel("请从左侧选择一个模板")
        self.template_info_label.setFont(QFont("Microsoft YaHei", 11, QFont.Bold))
        self.template_info_label.setStyleSheet("color: #2c3e50; padding: 5px; background-color: #ecf0f1; border-radius: 4px;")
        input_layout.addWidget(self.template_info_label)
        
        # 字段输入区域（滚动）
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setFrameShape(QFrame.NoFrame)
        self.field_container = QWidget()
        self.field_layout = QVBoxLayout(self.field_container)
        self.field_layout.setSpacing(8)
        self.field_layout.addStretch()
        scroll_area.setWidget(self.field_container)
        input_layout.addWidget(scroll_area)
        
        # 操作按钮
        action_layout = QHBoxLayout()
        self.generate_btn = QPushButton("🚀 生成文档")
        self.generate_btn.clicked.connect(self.generate_document)
        self.generate_btn.setFont(QFont("Microsoft YaHei", 10, QFont.Bold))
        self.generate_btn.setStyleSheet("""
            QPushButton {
                background-color: #3498db;
                color: white;
                border: none;
                padding: 8px 20px;
                border-radius: 4px;
            }
            QPushButton:hover {
                background-color: #2980b9;
            }
            QPushButton:pressed {
                background-color: #1a6d9b;
            }
        """)
        self.clear_btn = QPushButton("🗑️ 清空")
        self.clear_btn.clicked.connect(self.clear_fields)
        action_layout.addStretch()
        action_layout.addWidget(self.clear_btn)
        action_layout.addWidget(self.generate_btn)
        input_layout.addLayout(action_layout)
        
        # ===== 右侧下方面板：预览 =====
        self.preview_panel = QWidget()
        preview_layout = QVBoxLayout(self.preview_panel)
        preview_layout.setContentsMargins(5, 5, 5, 5)
        
        preview_header = QHBoxLayout()
        preview_header.addWidget(QLabel("📄 预览"))
        preview_header.addStretch()
        
        # 导出按钮
        self.export_txt_btn = QPushButton("📝 导出TXT")
        self.export_txt_btn.clicked.connect(lambda: self.export_document("txt"))
        self.export_txt_btn.setEnabled(False)
        self.export_doc_btn = QPushButton("📘 导出DOCX")
        self.export_doc_btn.clicked.connect(lambda: self.export_document("docx"))
        self.export_doc_btn.setEnabled(False)
        preview_header.addWidget(self.export_txt_btn)
        preview_header.addWidget(self.export_doc_btn)
        preview_layout.addLayout(preview_header)
        
        self.preview_text = QTextEdit()
        self.preview_text.setReadOnly(True)
        self.preview_text.setFont(QFont("微软雅黑", 10))
        self.preview_text.setPlaceholderText("生成后的文档将在此显示...")
        preview_layout.addWidget(self.preview_text)
        
        # 将两个面板添加到分割器
        splitter.addWidget(self.input_panel)
        splitter.addWidget(self.preview_panel)
        splitter.setSizes([400, 300])
        
        # 右侧面板添加到主布局
        main_layout.addWidget(splitter)
        main_layout.setStretchFactor(splitter, 1)
        
        # 状态栏
        self.statusBar = QStatusBar()
        self.setStatusBar(self.statusBar)
        self.statusBar.showMessage("就绪 | 共 0 个模板")
    
    def load_template_list(self):
        """加载模板列表"""
        self.template_list.clear()
        for template in self.manager.templates:
            item = QListWidgetItem(template["name"])
            # 设置字体
            item.setFont(QFont("Microsoft YaHei", 10))
            self.template_list.addItem(item)
        
        # 更新状态
        count = len(self.manager.templates)
        self.statusBar.showMessage(f"就绪 | 共 {count} 个模板")
        
        # 如果当前选中的索引有效，重新选中
        if self.current_index >= 0 and self.current_index < count:
            self.template_list.setCurrentRow(self.current_index)
        else:
            self.current_index = -1
            self.clear_input_fields()
            self.template_info_label.setText("请从左侧选择一个模板")
            self.preview_text.clear()
            self.export_txt_btn.setEnabled(False)
            self.export_doc_btn.setEnabled(False)
    
    def on_template_selected(self, item):
        """模板被选中"""
        index = self.template_list.currentRow()
        if index < 0 or index >= len(self.manager.templates):
            return
        
        self.current_index = index
        template = self.manager.templates[index]
        self.template_info_label.setText(f"📌 {template['name']}")
        
        # 提取字段
        fields = self.manager.extract_fields(template["content"])
        
        # 清空旧控件
        self.clear_input_fields()
        
        # 创建新的输入控件
        if fields:
            for field in fields:
                self.add_field_input(field)
            self.statusBar.showMessage(f"模板: {template['name']} | 共 {len(fields)} 个字段")
        else:
            # 没有字段，显示提示
            no_field_label = QLabel("此模板没有占位符，可直接生成")
            no_field_label.setStyleSheet("color: #999; padding: 10px;")
            self.field_layout.insertWidget(0, no_field_label)
            self.statusBar.showMessage(f"模板: {template['name']} | 无占位符")
        
        # 清空预览
        self.preview_text.clear()
        self.export_txt_btn.setEnabled(False)
        self.export_doc_btn.setEnabled(False)
    
    def add_field_input(self, field_name):
        """添加一个字段输入控件"""
        # 判断是否使用多行输入
        is_multiline = any(keyword in field_name for keyword in ["内容", "备注", "说明", "描述", "详情", "议程", "计划", "问题"])
        
        # 创建水平布局
        row_widget = QWidget()
        row_layout = QHBoxLayout(row_widget)
        row_layout.setContentsMargins(0, 2, 0, 2)
        
        # 标签
        label = QLabel(f"{field_name}:")
        label.setFont(QFont("Microsoft YaHei", 9))
        label.setMinimumWidth(80)
        label.setAlignment(Qt.AlignRight | Qt.AlignCenter)
        row_layout.addWidget(label)
        
        # 输入控件
        if is_multiline:
            input_widget = QTextEdit()
            input_widget.setPlaceholderText(f"请输入 {field_name}...")
            input_widget.setMaximumHeight(80)
            input_widget.setFont(QFont("微软雅黑", 9))
        else:
            input_widget = QLineEdit()
            input_widget.setPlaceholderText(f"请输入 {field_name}...")
            input_widget.setFont(QFont("微软雅黑", 9))
        
        # 存储控件引用
        self.field_widgets[field_name] = input_widget
        row_layout.addWidget(input_widget)
        
        # 插入到布局中（在stretch之前）
        self.field_layout.insertWidget(self.field_layout.count() - 1, row_widget)
    
    def clear_input_fields(self):
        """清空字段输入区域"""
        # 清空所有子控件
        while self.field_layout.count() > 1:
            item = self.field_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()
        self.field_widgets.clear()
    
    def get_field_values(self):
        """获取所有字段的值"""
        values = {}
        for field, widget in self.field_widgets.items():
            if isinstance(widget, QTextEdit):
                values[field] = widget.toPlainText()
            elif isinstance(widget, QLineEdit):
                values[field] = widget.text()
            else:
                values[field] = ""
        return values
    
    def clear_fields(self):
        """清空所有字段输入"""
        for widget in self.field_widgets.values():
            if isinstance(widget, QTextEdit):
                widget.clear()
            elif isinstance(widget, QLineEdit):
                widget.clear()
        self.preview_text.clear()
        self.export_txt_btn.setEnabled(False)
        self.export_doc_btn.setEnabled(False)
    
    def generate_document(self):
        """生成文档"""
        if self.current_index < 0 or self.current_index >= len(self.manager.templates):
            QMessageBox.warning(self, "警告", "请先选择一个模板")
            return
        
        template = self.manager.templates[self.current_index]
        field_values = self.get_field_values()
        
        # 检查是否有字段为空（可选）
        # 这里不强制要求所有字段都填写
        
        # 生成文本
        result = self.manager.generate_text(template["content"], field_values)
        
        # 显示预览
        self.preview_text.setPlainText(result)
        self.export_txt_btn.setEnabled(True)
        self.export_doc_btn.setEnabled(DOCX_AVAILABLE)
        
        if not DOCX_AVAILABLE:
            self.export_doc_btn.setToolTip("请安装 python-docx 以支持导出DOCX格式")
        
        self.statusBar.showMessage("文档生成成功")
    
    def export_document(self, format_type):
        """导出文档"""
        if self.current_index < 0:
            return
        
        template = self.manager.templates[self.current_index]
        content = self.preview_text.toPlainText()
        
        if not content.strip():
            QMessageBox.warning(self, "警告", "没有内容可导出，请先生成文档")
            return
        
        # 生成文件名
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        default_filename = f"{template['name']}_{timestamp}"
        
        if format_type == "txt":
            file_path, _ = QFileDialog.getSaveFileName(
                self, "导出为TXT", f"{default_filename}.txt", "文本文件 (*.txt)"
            )
            if file_path:
                try:
                    with open(file_path, "w", encoding="utf-8") as f:
                        f.write(content)
                    QMessageBox.information(self, "成功", f"已导出到:\n{file_path}")
                except Exception as e:
                    QMessageBox.critical(self, "错误", f"导出失败: {e}")
        
        elif format_type == "docx":
            if not DOCX_AVAILABLE:
                QMessageBox.warning(self, "警告", "请先安装 python-docx:\npip install python-docx")
                return
            
            file_path, _ = QFileDialog.getSaveFileName(
                self, "导出为DOCX", f"{default_filename}.docx", "Word文档 (*.docx)"
            )
            if file_path:
                try:
                    self.export_to_docx(content, file_path)
                    QMessageBox.information(self, "成功", f"已导出到:\n{file_path}")
                except Exception as e:
                    QMessageBox.critical(self, "错误", f"导出失败: {e}")
    
    def export_to_docx(self, text, file_path):
        """导出为DOCX格式"""
        doc = Document()
        
        # 设置默认字体
        style = doc.styles['Normal']
        style.font.name = '微软雅黑'
        style.font.size = Pt(11)
        
        # 按行分割文本
        lines = text.split('\n')
        for line in lines:
            # 检测标题（以 # 开头或全大写）
            if line.strip().startswith('#') or (line.strip() and line.strip().isupper() and len(line.strip()) > 3):
                # 标题
                p = doc.add_heading(line.strip('#').strip(), level=1)
            elif line.strip() and len(line.strip()) < 30 and '：' in line:
                # 可能是带冒号的标题行
                p = doc.add_paragraph(line)
                p.runs[0].bold = True
            else:
                doc.add_paragraph(line)
        
        # 保存
        doc.save(file_path)
    
    def add_template(self):
        """添加新模板"""
        dialog = TemplateEditDialog("", "", self)
        if dialog.exec_() == QDialog.Accepted:
            name, content = dialog.get_data()
            if not name:
                QMessageBox.warning(self, "警告", "模板名称不能为空")
                return
            if not content:
                QMessageBox.warning(self, "警告", "模板内容不能为空")
                return
            
            success, msg = self.manager.add_template(name, content)
            if success:
                self.load_template_list()
                # 选中新添加的模板
                names = self.manager.get_template_names()
                if name in names:
                    index = names.index(name)
                    self.template_list.setCurrentRow(index)
                    self.on_template_selected(self.template_list.currentItem())
                self.statusBar.showMessage(f"已添加模板: {name}")
            else:
                QMessageBox.warning(self, "警告", msg)
    
    def edit_template(self):
        """编辑当前选中的模板"""
        if self.current_index < 0 or self.current_index >= len(self.manager.templates):
            QMessageBox.warning(self, "警告", "请先选择一个模板")
            return
        
        template = self.manager.templates[self.current_index]
        dialog = TemplateEditDialog(template["name"], template["content"], self)
        if dialog.exec_() == QDialog.Accepted:
            name, content = dialog.get_data()
            if not name:
                QMessageBox.warning(self, "警告", "模板名称不能为空")
                return
            if not content:
                QMessageBox.warning(self, "警告", "模板内容不能为空")
                return
            
            success, msg = self.manager.update_template(self.current_index, name, content)
            if success:
                self.load_template_list()
                # 选中编辑后的模板
                names = self.manager.get_template_names()
                if name in names:
                    index = names.index(name)
                    self.template_list.setCurrentRow(index)
                    self.on_template_selected(self.template_list.currentItem())
                self.statusBar.showMessage(f"已更新模板: {name}")
            else:
                QMessageBox.warning(self, "警告", msg)
    
    def delete_template(self):
        """删除当前选中的模板"""
        if self.current_index < 0 or self.current_index >= len(self.manager.templates):
            QMessageBox.warning(self, "警告", "请先选择一个模板")
            return
        
        template = self.manager.templates[self.current_index]
        reply = QMessageBox.question(
            self, "确认删除",
            f"确定要删除模板「{template['name']}」吗？",
            QMessageBox.Yes | QMessageBox.No
        )
        if reply == QMessageBox.Yes:
            self.manager.delete_template(self.current_index)
            self.load_template_list()
            self.clear_input_fields()
            self.template_info_label.setText("请从左侧选择一个模板")
            self.preview_text.clear()
            self.export_txt_btn.setEnabled(False)
            self.export_doc_btn.setEnabled(False)
            self.statusBar.showMessage("模板已删除")


def main():
    app = QApplication(sys.argv)
    app.setStyle("Fusion")
    
    # 设置应用程序图标（如果有）
    # app.setWindowIcon(QIcon("icon.ico"))
    
    window = MainWindow()
    window.show()
    sys.exit(app.exec_())


if __name__ == "__main__":
    main()