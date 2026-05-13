"""
ECN Generator - Engineering Change Note Tool
Generates 3-page PDF, DOCX, and updates Excel log from ECN data
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import os
import re
import json
from datetime import datetime
from reportlab.platypus import Image as RLImage
from PIL import Image, ImageTk

import threading

# ── PDF ──────────────────────────────────────────────────────────────────────
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import mm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    PageBreak, HRFlowable, Image as RLImage
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

# ── DOCX ─────────────────────────────────────────────────────────────────────
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── EXCEL ─────────────────────────────────────────────────────────────────────
import openpyxl
from openpyxl.styles import Font, Fill, PatternFill, Border, Side, Alignment
from openpyxl.utils import get_column_letter

# ─────────────────────────────────────────────────────────────────────────────
# COLOUR PALETTE
# ─────────────────────────────────────────────────────────────────────────────
BRAND_DARK   = "#1A3A5C"   # navy header
BRAND_MID    = "#2E6DA4"   # mid-blue
BRAND_LIGHT  = "#D6E8F7"   # light fill
BRAND_ACCENT = "#F07F3C"   # orange accent
WHITE        = "#FFFFFF"
GREY         = "#F4F6F9"
DARK_TEXT    = "#1C1C1C"

# ─────────────────────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────────────────────
def safe_filename(ecn_number: str) -> str:
    """Remove all non-alphanumeric chars from ECN number for use as filename."""
    return re.sub(r'[^A-Za-z0-9]', '', ecn_number)


def hex_to_rgb(h: str):
    h = h.lstrip('#')
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))


# ─────────────────────────────────────────────────────────────────────────────
# PDF GENERATOR
# ─────────────────────────────────────────────────────────────────────────────
def generate_pdf(data: dict, output_path: str):
    doc = SimpleDocTemplate(
        output_path, pagesize=A4,
        leftMargin=18*mm, rightMargin=18*mm,
        topMargin=14*mm, bottomMargin=14*mm
    )
    styles = getSampleStyleSheet()
    W = A4[0] - 36*mm  # usable width

    # Custom styles
    def ps(name, parent='Normal', **kw):
        return ParagraphStyle(name, parent=styles[parent], **kw)

    title_s  = ps('T', fontSize=16, textColor=colors.HexColor(WHITE),
                  alignment=TA_CENTER, fontName='Helvetica-Bold', leading=20)
    sub_s    = ps('S', fontSize=9,  textColor=colors.HexColor(BRAND_MID),
                  alignment=TA_CENTER, fontName='Helvetica', leading=12)
    hdr_s    = ps('H', fontSize=11, textColor=colors.HexColor(WHITE),
                  fontName='Helvetica-Bold', leading=14, leftIndent=6)
    label_s  = ps('L', fontSize=8,  textColor=colors.HexColor(BRAND_DARK),
                  fontName='Helvetica-Bold', leading=11)
    value_s  = ps('V', fontSize=9,  textColor=colors.HexColor(DARK_TEXT),
                  fontName='Helvetica', leading=12)
    body_s   = ps('B', fontSize=9,  textColor=colors.HexColor(DARK_TEXT),
                  fontName='Helvetica', leading=13, leftIndent=4)
    page_s   = ps('P', fontSize=8,  textColor=colors.HexColor('#888888'),
                  alignment=TA_RIGHT, fontName='Helvetica')
    note_s   = ps('N', fontSize=8,  textColor=colors.HexColor('#555555'),
                  fontName='Helvetica-Oblique', leading=11)

    def band(text, style=hdr_s, bg=BRAND_DARK):
        return Table([[Paragraph(text, style)]],
                     colWidths=[W],
                     style=TableStyle([
                         ('BACKGROUND', (0,0), (-1,-1), colors.HexColor(bg)),
                         ('TOPPADDING', (0,0), (-1,-1), 5),
                         ('BOTTOMPADDING', (0,0), (-1,-1), 5),
                     ]))

    def kv_table(rows, col_ratio=(0.35, 0.65)):
        c1 = W * col_ratio[0]; c2 = W * col_ratio[1]
        tdata = []
        for lbl, val in rows:
            tdata.append([Paragraph(lbl, label_s), Paragraph(str(val), value_s)])
        ts = TableStyle([
            ('BACKGROUND', (0,0), (0,-1), colors.HexColor(BRAND_LIGHT)),
            ('BACKGROUND', (1,0), (1,-1), colors.HexColor(WHITE)),
            ('GRID', (0,0), (-1,-1), 0.4, colors.HexColor('#BBCFE0')),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('TOPPADDING', (0,0), (-1,-1), 4),
            ('BOTTOMPADDING', (0,0), (-1,-1), 4),
            ('LEFTPADDING', (0,0), (-1,-1), 6),
            ('RIGHTPADDING', (0,0), (-1,-1), 6),
        ])
        return Table(tdata, colWidths=[c1, c2], style=ts)

    story = []

    # ── PAGE 1: Basic Information ────────────────────────────────────────────
    # Header banner
    logo = RLImage("logo.jpg", width=100, height=45)
    
    hdr_tbl = Table(
        [[logo, Paragraph("ENGINEERING CHANGE NOTE", title_s)]],
        #colWidths=[W*0.6, W*0.4],
        colWidths=[60, W-60],
        style=TableStyle([
            ('BACKGROUND', (0,0), (-1,-1), colors.HexColor(BRAND_DARK)),
            ('TOPPADDING', (0,0), (-1,-1), 10),
            ('BOTTOMPADDING', (0,0), (-1,-1), 10),
        ])
    )
    story.append(hdr_tbl)
    #story.append(Spacer(1, 4*mm))
    #story.append(Paragraph(f"Company: {data.get('company','LECS Electronics')}  |  "
    #                        f"Date: {data['ecn_date']}  |  "
    #                        f"Prepared by: {data['prepared_by']}  |  "
    #                        f"Page 1 of 3", sub_s))
    #story.append(Spacer(1, 5*mm))

    story.append(band("SECTION 1 — BASIC INFORMATION"))
    story.append(Spacer(1, 2*mm))

    rows1 = [
        ("ECN Number",         data['ecn_number']),
        ("ECN Date",           data['ecn_date']),
        ("Prepared By",        data['prepared_by']),
        ("Main Product Code",  data['main_product_code']),
        ("Product Description",data['main_product_desc']),
        ("ECN Part Code",      data['ecn_part_code']),
        ("Part Description",   data['ecn_part_desc']),
        ("Implementation Date",data['implementation_date']),
        ("Department Involved",data['departments']),
    ]
    story.append(kv_table(rows1))
    story.append(Spacer(1, 5*mm))

    story.append(band("COST IMPACT", bg=BRAND_MID))
    story.append(Spacer(1, 2*mm))
    cost_rows = [
        ("Current Cost", data.get('current_cost', 'Rs. 7.42')),
        ("ECN Change Cost", data.get('ecn_cost', 'Rs. —')),
    ]
    story.append(kv_table(cost_rows))
    story.append(Spacer(1, 5*mm))

    story.append(band("SIGN-OFF", bg=BRAND_MID))
    story.append(Spacer(1, 2*mm))
    sign_data = [
        [Paragraph('Role', label_s), Paragraph('Department', label_s),
         Paragraph('Name', label_s), Paragraph('Signature', label_s),
         Paragraph('Date', label_s)],
        [Paragraph('Verified By', value_s), Paragraph(data.get('verify_dept','R&D'), value_s),
         Paragraph(data.get('verify_name','Muthukrishnan K'), value_s),
         Paragraph('', value_s), Paragraph('', value_s)],
        [Paragraph('Approved By', value_s), Paragraph(data.get('approve_dept','R&D'), value_s),
         Paragraph(data.get('approve_name','Vishnu Prasad (VP)'), value_s),
         Paragraph('', value_s), Paragraph('', value_s)],
    ]
    cw = [W*v for v in (0.15, 0.15, 0.25, 0.25, 0.20)]
    sign_tbl = Table(sign_data, colWidths=cw,
                     style=TableStyle([
                         ('BACKGROUND', (0,0), (-1,0), colors.HexColor(BRAND_LIGHT)),
                         ('GRID', (0,0), (-1,-1), 0.4, colors.HexColor('#BBCFE0')),
                         ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                         ('TOPPADDING', (0,0), (-1,-1), 5),
                         ('BOTTOMPADDING', (0,0), (-1,-1), 5),
                         ('LEFTPADDING', (0,0), (-1,-1), 5),
                         ('ROW', (2,1), (-1,-1), 18),  # taller rows for signature
                     ]))
    story.append(sign_tbl)
    story.append(Spacer(1, 5*mm))
    #story.append(PageBreak())

    # ── PAGE 2: Main Objective ──────────────────────────────────────────────
    #story.append(band("ENGINEERING CHANGE NOTE — PAGE 2 OF 3"))
    #story.append(Spacer(1, 2*mm))
    #story.append(Paragraph(f"ECN #: {data['ecn_number']}  |  Date: {data['ecn_date']}  |  "
    #                        f"Prepared by: {data['prepared_by']}", sub_s))
    #story.append(Spacer(1, 5*mm))

    story.append(band("SECTION 2 — MAIN OBJECTIVE & CHANGE DETAILS"))
    story.append(Spacer(1, 2*mm))

    obj_rows = [
        ("Details of Change",    data['details_of_change']),
        ("Reason for Change",    data['reason_for_change']),
        ("ECN Request Reference",data['ecn_request_ref']),
        ("Addition",             data['addition']),
        ("Deletion",             data['deletion']),
        ("List of Change Documents", data.get('change_docs','BOM')),
    ]
    story.append(kv_table(obj_rows, col_ratio=(0.30, 0.70)))
    story.append(Spacer(1, 5*mm))

    story.append(band("ACTION PLAN", bg=BRAND_MID))
    story.append(Spacer(1, 2*mm))

    actions = [
        ("Action for Change",       data.get('action_change','Updated the MPN in SAP'),
         data.get('action_resp','Suryaprakash - R&D')),
        ("Existing Stock Disposition", data.get('stock_disp','Return existing stock to supplier'),
         data.get('stock_resp','Hemalatha – Production')),
        ("Inventory",               data.get('inventory_action','Buy only updated MPN as per BOM'),
         data.get('inventory_resp','Kalaiselvan – Purchase')),
        ("Change Note",             data.get('change_note','SMDV000042 – 01'),
         data.get('change_note_resp','Suryaprakash - R&D')),
    ]
    act_hdr = [Paragraph(h, label_s) for h in ['Action', 'Description', 'Responsibility']]
    act_data = [act_hdr]
    for a, d, r in actions:
        act_data.append([Paragraph(a, label_s), Paragraph(d, body_s), Paragraph(r, value_s)])

    act_tbl = Table(act_data, colWidths=[W*0.22, W*0.52, W*0.26],
                    style=TableStyle([
                        ('BACKGROUND', (0,0), (-1,0), colors.HexColor(BRAND_LIGHT)),
                        ('GRID', (0,0), (-1,-1), 0.4, colors.HexColor('#BBCFE0')),
                        ('VALIGN', (0,0), (-1,-1), 'TOP'),
                        ('TOPPADDING', (0,0), (-1,-1), 4),
                        ('BOTTOMPADDING', (0,0), (-1,-1), 4),
                        ('LEFTPADDING', (0,0), (-1,-1), 5),
                    ]))
    story.append(act_tbl)
    story.append(Spacer(1, 5*mm))

    story.append(band("IMPLEMENTATION SCHEDULE", bg=BRAND_MID))
    story.append(Spacer(1, 2*mm))
    story.append(kv_table([("Effective Date", data.get('implementation_date','Immediate'))]))
    #story.append(PageBreak())

    # ── PAGE 3: Attachments ─────────────────────────────────────────────────
    #story.append(band("ENGINEERING CHANGE NOTE — PAGE 3 OF 3"))
    #story.append(Spacer(1, 2*mm))
    #story.append(Paragraph(f"ECN #: {data['ecn_number']}  |  Date: {data['ecn_date']}  |  "
    #                        f"Prepared by: {data['prepared_by']}", sub_s))
    #story.append(Spacer(1, 5*mm))

    story.append(band("SECTION 3 — ATTACHMENTS & SUPPORTING DOCUMENTS"))
    story.append(Spacer(1, 3*mm))

    attach_list = data.get('attachments', [])
    if attach_list:
        att_hdr = [Paragraph(h, label_s) for h in ['#', 'Document Name', 'Description', 'Status']]
        att_data = [att_hdr]
        for i, att in enumerate(attach_list, 1):
            att_data.append([
                Paragraph(str(i), value_s),
                Paragraph(att.get('name',''), value_s),
                Paragraph(att.get('desc',''), body_s),
                Paragraph(att.get('status','Attached'), value_s),
            ])
        att_tbl = Table(att_data, colWidths=[W*0.06, W*0.28, W*0.46, W*0.20],
                        style=TableStyle([
                            ('BACKGROUND', (0,0), (-1,0), colors.HexColor(BRAND_LIGHT)),
                            ('GRID', (0,0), (-1,-1), 0.4, colors.HexColor('#BBCFE0')),
                            ('VALIGN', (0,0), (-1,-1), 'TOP'),
                            ('TOPPADDING', (0,0), (-1,-1), 4),
                            ('BOTTOMPADDING', (0,0), (-1,-1), 4),
                            ('LEFTPADDING', (0,0), (-1,-1), 5),
                        ]))
        story.append(att_tbl)
    else:
        story.append(Paragraph("No attachments added.", note_s))

    story.append(Spacer(1, 6*mm))
    story.append(band("NOTES & REMARKS", bg=BRAND_MID))
    story.append(Spacer(1, 2*mm))
    notes = data.get('notes', '')
    story.append(Paragraph(notes if notes else 'No additional notes.', body_s))
    story.append(Spacer(1, 8*mm))

    # Footer line
    story.append(HRFlowable(width=W, thickness=1, color=colors.HexColor(BRAND_MID)))
    story.append(Spacer(1, 2*mm))
    story.append(Paragraph(
        f"Generated on {datetime.now().strftime('%d %b %Y %H:%M')} | "
        f"ECN: {data['ecn_number']} | CONFIDENTIAL",
        page_s))

    doc.build(story)


# ─────────────────────────────────────────────────────────────────────────────
# DOCX GENERATOR
# ─────────────────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), hex_color.lstrip('#'))
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for side in ['top','left','bottom','right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '2E6DA4')
        tcPr.append(b)

def docx_band(doc, text, bg=BRAND_DARK):
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    cell = t.cell(0, 0)
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return t

def docx_kv(doc, rows):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = 'Table Grid'
    t.columns[0].width = Cm(5.5)
    t.columns[1].width = Cm(12)
    for i, (lbl, val) in enumerate(rows):
        lc = t.cell(i, 0); vc = t.cell(i, 1)
        set_cell_bg(lc, BRAND_LIGHT)
        set_cell_bg(vc, WHITE)
        lp = lc.paragraphs[0]
        lr = lp.add_run(lbl)
        lr.bold = True; lr.font.size = Pt(9)
        lr.font.color.rgb = RGBColor(*hex_to_rgb(BRAND_DARK))
        vp = vc.paragraphs[0]
        vr = vp.add_run(str(val))
        vr.font.size = Pt(9)
        for c in (lc, vc):
            set_cell_border(c)
    return t

def generate_docx(data: dict, output_path: str):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(2)

    def sp(n=1):
        for _ in range(n):
            doc.add_paragraph()

    # ── PAGE 1 ───────────────────────────────────────────────────────────────
    ht = doc.add_table(rows=1, cols=2)
    ht.style = 'Table Grid'
    for i, txt in enumerate(["ENGINEERING CHANGE NOTE",
                              f"ECN #: {data['ecn_number']}"]):
        cell = ht.cell(0, i)
        set_cell_bg(cell, BRAND_DARK)
        p = cell.paragraphs[0]
        r = p.add_run(txt)
        r.bold = True; r.font.color.rgb = RGBColor(255,255,255); r.font.size = Pt(13)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub_p.add_run(f"Date: {data['ecn_date']}  |  Prepared by: {data['prepared_by']}  |  Page 1 of 3")
    sr.font.size = Pt(9); sr.font.color.rgb = RGBColor(*hex_to_rgb(BRAND_MID))

    docx_band(doc, "SECTION 1 — BASIC INFORMATION")
    docx_kv(doc, [
        ("ECN Number",          data['ecn_number']),
        ("ECN Date",            data['ecn_date']),
        ("Prepared By",         data['prepared_by']),
        ("Main Product Code",   data['main_product_code']),
        ("Product Description", data['main_product_desc']),
        ("ECN Part Code",       data['ecn_part_code']),
        ("Part Description",    data['ecn_part_desc']),
        ("Implementation Date", data['implementation_date']),
        ("Departments Involved",data['departments']),
    ])
    sp()
    docx_band(doc, "COST IMPACT", bg=BRAND_MID)
    docx_kv(doc, [
        ("Current Cost",    data.get('current_cost','Rs. 7.42')),
        ("ECN Change Cost", data.get('ecn_cost','Rs. —')),
    ])
    sp()
    docx_band(doc, "SIGN-OFF", bg=BRAND_MID)
    sign_t = doc.add_table(rows=3, cols=5)
    sign_t.style = 'Table Grid'
    hdrs = ['Role','Department','Name','Signature','Date']
    for j, h in enumerate(hdrs):
        c = sign_t.cell(0, j)
        set_cell_bg(c, BRAND_LIGHT)
        r = c.paragraphs[0].add_run(h)
        r.bold = True; r.font.size = Pt(9)
    sign_rows = [
        ('Verified By',  data.get('verify_dept','R&D'),   data.get('verify_name','Muthukrishnan K'), '', ''),
        ('Approved By',  data.get('approve_dept','R&D'),  data.get('approve_name','Vishnu Prasad (VP)'), '', ''),
    ]
    for ri, row in enumerate(sign_rows, 1):
        for ci, val in enumerate(row):
            c = sign_t.cell(ri, ci)
            set_cell_bg(c, WHITE)
            r = c.paragraphs[0].add_run(val)
            r.font.size = Pt(9)
            set_cell_border(c)

    doc.add_page_break()

    # ── PAGE 2 ───────────────────────────────────────────────────────────────
    ht2 = doc.add_table(rows=1, cols=1)
    ht2.style = 'Table Grid'
    c2 = ht2.cell(0, 0)
    set_cell_bg(c2, BRAND_DARK)
    p2 = c2.paragraphs[0]
    r2 = p2.add_run("ENGINEERING CHANGE NOTE — PAGE 2 OF 3")
    r2.bold = True; r2.font.color.rgb = RGBColor(255,255,255); r2.font.size = Pt(13)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s2r = sub2.add_run(f"ECN #: {data['ecn_number']}  |  Date: {data['ecn_date']}  |  Prepared by: {data['prepared_by']}")
    s2r.font.size = Pt(9); s2r.font.color.rgb = RGBColor(*hex_to_rgb(BRAND_MID))

    docx_band(doc, "SECTION 2 — MAIN OBJECTIVE & CHANGE DETAILS")
    docx_kv(doc, [
        ("Details of Change",     data['details_of_change']),
        ("Reason for Change",     data['reason_for_change']),
        ("ECN Request Reference", data['ecn_request_ref']),
        ("Addition",              data['addition']),
        ("Deletion",              data['deletion']),
        ("Change Documents",      data.get('change_docs','BOM')),
    ])
    sp()
    docx_band(doc, "ACTION PLAN", bg=BRAND_MID)
    act_t = doc.add_table(rows=5, cols=3)
    act_t.style = 'Table Grid'
    for j, h in enumerate(['Action','Description','Responsibility']):
        c = act_t.cell(0, j)
        set_cell_bg(c, BRAND_LIGHT)
        r = c.paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(9)
    actions = [
        ("Action for Change",       data.get('action_change','Updated the MPN in SAP'), data.get('action_resp','Suryaprakash - R&D')),
        ("Stock Disposition",       data.get('stock_disp','Return existing stock to supplier'), data.get('stock_resp','Hemalatha – Production')),
        ("Inventory",               data.get('inventory_action','Buy only updated MPN as per BOM'), data.get('inventory_resp','Kalaiselvan – Purchase')),
        ("Change Note",             data.get('change_note','SMDV000042 – 01'), data.get('change_note_resp','Suryaprakash - R&D')),
    ]
    for ri, (a, d, r_) in enumerate(actions, 1):
        for ci, val in enumerate([a, d, r_]):
            c = act_t.cell(ri, ci)
            set_cell_bg(c, WHITE)
            rv = c.paragraphs[0].add_run(val); rv.font.size = Pt(9)
            set_cell_border(c)

    doc.add_page_break()

    # ── PAGE 3 ───────────────────────────────────────────────────────────────
    ht3 = doc.add_table(rows=1, cols=1)
    ht3.style = 'Table Grid'
    c3 = ht3.cell(0, 0)
    set_cell_bg(c3, BRAND_DARK)
    p3 = c3.paragraphs[0]
    r3 = p3.add_run("ENGINEERING CHANGE NOTE — PAGE 3 OF 3")
    r3.bold = True; r3.font.color.rgb = RGBColor(255,255,255); r3.font.size = Pt(13)
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub3 = doc.add_paragraph()
    sub3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s3r = sub3.add_run(f"ECN #: {data['ecn_number']}  |  Date: {data['ecn_date']}  |  Prepared by: {data['prepared_by']}")
    s3r.font.size = Pt(9); s3r.font.color.rgb = RGBColor(*hex_to_rgb(BRAND_MID))

    docx_band(doc, "SECTION 3 — ATTACHMENTS & SUPPORTING DOCUMENTS")
    attach_list = data.get('attachments', [])
    if attach_list:
        att_t = doc.add_table(rows=len(attach_list)+1, cols=4)
        att_t.style = 'Table Grid'
        for j, h in enumerate(['#','Document Name','Description','Status']):
            c = att_t.cell(0, j)
            set_cell_bg(c, BRAND_LIGHT)
            r = c.paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(9)
        for ri, att in enumerate(attach_list, 1):
            vals = [str(ri), att.get('name',''), att.get('desc',''), att.get('status','Attached')]
            for ci, val in enumerate(vals):
                c = att_t.cell(ri, ci)
                set_cell_bg(c, WHITE)
                rv = c.paragraphs[0].add_run(val); rv.font.size = Pt(9)
                set_cell_border(c)
    else:
        np = doc.add_paragraph("No attachments added.")
        np.runs[0].font.size = Pt(9)

    sp()
    docx_band(doc, "NOTES & REMARKS", bg=BRAND_MID)
    notes_p = doc.add_paragraph(data.get('notes','No additional notes.'))
    notes_p.runs[0].font.size = Pt(9)

    doc.save(output_path)


# ─────────────────────────────────────────────────────────────────────────────
# EXCEL LOG UPDATER
# ─────────────────────────────────────────────────────────────────────────────
EXCEL_HEADERS = [
    "ECN Number", "ECN Date", "Prepared By", "Main Product Code",
    "Product Description", "ECN Part Code", "Part Description",
    "Details of Change", "Reason for Change", "Implementation Date",
    "Current Cost", "ECN Change Cost", "Departments", "Generated On"
]

def update_excel_log(data: dict, excel_path: str):
    if os.path.exists(excel_path):
        wb = openpyxl.load_workbook(excel_path)
        ws = wb.active
    else:
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "ECN Log"

        # Write header
        hdr_fill  = PatternFill("solid", fgColor=BRAND_DARK.lstrip('#'))
        hdr_font  = Font(bold=True, color="FFFFFF", size=10)
        hdr_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
        thin      = Side(style='thin', color="2E6DA4")
        border    = Border(left=thin, right=thin, top=thin, bottom=thin)

        for col, h in enumerate(EXCEL_HEADERS, 1):
            cell = ws.cell(row=1, column=col, value=h)
            cell.fill  = hdr_fill
            cell.font  = hdr_font
            cell.alignment = hdr_align
            cell.border = border
            ws.column_dimensions[get_column_letter(col)].width = 20

        ws.row_dimensions[1].height = 30
        ws.freeze_panes = 'A2'

    # Find next empty row
    next_row = ws.max_row + 1 if ws.max_row > 1 else 2

    row_fill  = PatternFill("solid", fgColor="EAF2FB")
    row_font  = Font(size=9)
    row_align = Alignment(vertical='center', wrap_text=True)
    thin      = Side(style='thin', color="BBCFE0")
    border    = Border(left=thin, right=thin, top=thin, bottom=thin)

    values = [
        data['ecn_number'], data['ecn_date'], data['prepared_by'],
        data['main_product_code'], data['main_product_desc'],
        data['ecn_part_code'], data['ecn_part_desc'],
        data['details_of_change'], data['reason_for_change'],
        data['implementation_date'], data.get('current_cost',''),
        data.get('ecn_cost',''), data['departments'],
        datetime.now().strftime('%d-%m-%Y %H:%M')
    ]

    for col, val in enumerate(values, 1):
        cell = ws.cell(row=next_row, column=col, value=val)
        cell.fill      = row_fill
        cell.font      = row_font
        cell.alignment = row_align
        cell.border    = border

    ws.row_dimensions[next_row].height = 25
    wb.save(excel_path)


# ─────────────────────────────────────────────────────────────────────────────
# GUI APPLICATION
# ─────────────────────────────────────────────────────────────────────────────
class ECNApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("LECS ECN Generator  ·  Engineering Change Note Tool")
        self.geometry("980x820")
        self.configure(bg=GREY)
        self.resizable(True, True)

        self.attachments = []   # list of dicts
        self._build_ui()
        self._prefill_sample()

    # ── UI CONSTRUCTION ──────────────────────────────────────────────────────
    def _build_ui(self):
        # ─ Top bar
        top = tk.Frame(self, bg=BRAND_DARK, height=56)
        top.pack(fill='x')
        top.pack_propagate(False)
        # Load logo
        img = Image.open("logo.jpg")
        img = img.resize((100, 40))   # optional resize

        self.logo = ImageTk.PhotoImage(img)

        # Logo label
        tk.Label(top, image=self.logo, bg=BRAND_DARK).pack(side='left', padx=(12, 6), pady=6)
        
        tk.Label(top, text="LECS ECN Generator", font=('Segoe UI', 16, 'bold'),
                 bg=BRAND_DARK, fg=WHITE).pack(side='left', padx=20, pady=12)
        tk.Label(top, text="Engineering Change Note Tool",
                 font=('Segoe UI', 9), bg=BRAND_DARK, fg='#8BB8D8').pack(side='left', padx=4, pady=12)

        # ─ Notebook tabs
        nb = ttk.Notebook(self)
        nb.pack(fill='both', expand=True, padx=12, pady=8)

        # Style
        style = ttk.Style()
        style.configure('TNotebook.Tab', font=('Segoe UI', 10, 'bold'), padding=[14, 6])

        self.tab1 = tk.Frame(nb, bg=GREY)
        self.tab2 = tk.Frame(nb, bg=GREY)
        self.tab3 = tk.Frame(nb, bg=GREY)
        self.tab4 = tk.Frame(nb, bg=GREY)

        nb.add(self.tab1, text="  📋 Basic Info  ")
        nb.add(self.tab2, text="  🔧 Change Details  ")
        nb.add(self.tab3, text="  📎 Attachments  ")
        nb.add(self.tab4, text="  ⚙ Output Settings  ")

        self._build_tab1()
        self._build_tab2()
        self._build_tab3()
        self._build_tab4()

        # ─ Status + Generate bar
        bot = tk.Frame(self, bg=BRAND_DARK, height=44)
        bot.pack(fill='x', side='bottom')
        bot.pack_propagate(False)

        self.status_var = tk.StringVar(value="Ready to generate ECN documents.")
        tk.Label(bot, textvariable=self.status_var, bg=BRAND_DARK, fg='#A8C8E8',
                 font=('Segoe UI', 9)).pack(side='left', padx=16, pady=10)

        tk.Button(bot, text="⚡  GENERATE ECN", font=('Segoe UI', 11, 'bold'),
                  bg=BRAND_ACCENT, fg=WHITE, relief='flat', padx=18, pady=6,
                  activebackground='#D06020', cursor='hand2',
                  command=self._generate).pack(side='right', padx=16, pady=6)

    # ── TAB 1: Basic Info ─────────────────────────────────────────────────────
    def _build_tab1(self):
        canvas = tk.Canvas(self.tab1, bg=GREY, highlightthickness=0)
        scrollbar = ttk.Scrollbar(self.tab1, orient='vertical', command=canvas.yview)
        frame = tk.Frame(canvas, bg=GREY)
        frame.bind('<Configure>', lambda e: canvas.configure(scrollregion=canvas.bbox('all')))
        canvas.create_window((0,0), window=frame, anchor='nw')
        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.pack(side='left', fill='both', expand=True)
        scrollbar.pack(side='right', fill='y')

        self.fields1 = {}
        sections = [
            ("ECN Identification", [
                ("ecn_number",        "ECN Number *",          "LECS/DOC/ECN0009"),
                ("ecn_date",          "ECN Date *",            "11.05.2026"),
                ("prepared_by",       "Prepared By *",         "S.Guruprasath"),
                ("company",           "Company Name",          "LECS Electronics"),
            ]),
            ("Product Information", [
                ("main_product_code", "Main Product Code *",   "RDDSMR20001, RDDSMR20004"),
                ("main_product_desc", "Product Description *", "1P 5-30A, 1P 10-60A, 3P 10-60A"),
                ("ecn_part_code",     "ECN Part Code *",       "SMDV000042"),
                ("ecn_part_desc",     "Part Description *",    "Through Hole - 32.768 kHz Crystal"),
            ]),
            ("Cost & Schedule", [
                ("current_cost",         "Current Cost",          "Rs. 7.42"),
                ("ecn_cost",             "ECN Change Cost",       "Rs. —"),
                ("implementation_date",  "Implementation Date",   "Immediate from 10.04.2026"),
                ("departments",          "Departments Involved",  "Purchase, Production, Stores, Quality"),
            ]),
            ("Sign-off", [
                ("verify_dept",    "Verified By – Department", "R&D"),
                ("verify_name",    "Verified By – Name",       "Muthukrishnan K"),
                ("approve_dept",   "Approved By – Department", "R&D"),
                ("approve_name",   "Approved By – Name",       "Vishnu Prasad (VP)"),
            ]),
        ]
        self._render_sections(frame, sections, self.fields1)

    # ── TAB 2: Change Details ──────────────────────────────────────────────────
    def _build_tab2(self):
        canvas = tk.Canvas(self.tab2, bg=GREY, highlightthickness=0)
        scrollbar = ttk.Scrollbar(self.tab2, orient='vertical', command=canvas.yview)
        frame = tk.Frame(canvas, bg=GREY)
        frame.bind('<Configure>', lambda e: canvas.configure(scrollregion=canvas.bbox('all')))
        canvas.create_window((0,0), window=frame, anchor='nw')
        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.pack(side='left', fill='both', expand=True)
        scrollbar.pack(side='right', fill='y')

        self.fields2 = {}
        sections = [
            ("Change Information", [
                ("details_of_change", "Details of Change *",     "Change of make due to quality improvement."),
                ("reason_for_change", "Reason for Change *",     "RTC drift in Raltron makes crystal."),
                ("ecn_request_ref",   "ECN Request Reference *", "RTC drift Mail received from production on 18 March 2026"),
                ("addition",          "Addition",                "Adding Abracon LLC & Wurth Electronics make"),
                ("deletion",          "Deletion",                "Removed the existing Raltron Make"),
                ("change_docs",       "Change Documents",        "BOM"),
            ]),
            ("Action Plan", [
                ("action_change",     "Action for Change",       "Updated the MPN in SAP"),
                ("action_resp",       "Responsibility",          "Suryaprakash - R&D"),
                ("stock_disp",        "Stock Disposition",       "Return existing stock to supplier"),
                ("stock_resp",        "Responsibility",          "Hemalatha – Production"),
                ("inventory_action",  "Inventory Action",        "Buy only updated MPN as per BOM"),
                ("inventory_resp",    "Responsibility",          "Kalaiselvan – Purchase"),
                ("change_note",       "Change Note",             "SMDV000042 – 01"),
                ("change_note_resp",  "Responsibility",          "Suryaprakash - R&D"),
            ]),
        ]
        self._render_sections(frame, sections, self.fields2, tall=['addition','deletion','details_of_change','reason_for_change'])

        # Notes
        self._section_label(frame, "Additional Notes")
        f = tk.Frame(frame, bg=GREY); f.pack(fill='x', padx=20, pady=4)
        tk.Label(f, text="Notes", bg=GREY, fg=BRAND_DARK,
                 font=('Segoe UI', 9, 'bold'), width=22, anchor='w').pack(side='left')
        self.notes_text = scrolledtext.ScrolledText(f, height=4, font=('Segoe UI', 9), wrap='word')
        self.notes_text.pack(side='left', fill='x', expand=True)

    # ── TAB 3: Attachments ────────────────────────────────────────────────────
    def _build_tab3(self):
        self._section_label(self.tab3, "Attachments / Supporting Documents")

        frm = tk.Frame(self.tab3, bg=GREY); frm.pack(fill='x', padx=20, pady=6)
        labels = ['Document Name', 'Description', 'Status']
        self.att_entries = []
        for i, lbl in enumerate(labels):
            tk.Label(frm, text=lbl, bg=GREY, fg=BRAND_DARK,
                     font=('Segoe UI', 9, 'bold')).grid(row=0, column=i, padx=4, sticky='w')
        self._att_vars = []

        self.att_frame = tk.Frame(self.tab3, bg=GREY); self.att_frame.pack(fill='both', padx=20, expand=True)

        btn_frame = tk.Frame(self.tab3, bg=GREY); btn_frame.pack(fill='x', padx=20, pady=4)
        tk.Button(btn_frame, text="➕  Add Attachment", font=('Segoe UI', 9, 'bold'),
                  bg=BRAND_MID, fg=WHITE, relief='flat', padx=10, pady=4,
                  command=self._add_attachment_row).pack(side='left', padx=4)
        tk.Button(btn_frame, text="🗑  Clear All", font=('Segoe UI', 9),
                  bg='#C0392B', fg=WHITE, relief='flat', padx=10, pady=4,
                  command=self._clear_attachments).pack(side='left', padx=4)

        # Default attachment rows
        for att in [
            {"name": "BOM - Updated", "desc": "Updated Bill of Materials with new MPN", "status": "Attached"},
            {"name": "Stock as on 11.05.2026", "desc": "Current stock status report", "status": "Attached"},
            {"name": "Mail From Production", "desc": "RTC drift issue mail from production team", "status": "Attached"},
        ]:
            self._add_attachment_row(att)

    def _add_attachment_row(self, defaults=None):
        row = len(self._att_vars)
        rf = tk.Frame(self.att_frame, bg=GREY); rf.pack(fill='x', pady=2)
        name_var  = tk.StringVar(value=defaults.get('name','')   if defaults else '')
        desc_var  = tk.StringVar(value=defaults.get('desc','')   if defaults else '')
        stat_var  = tk.StringVar(value=defaults.get('status','Attached') if defaults else 'Attached')
        tk.Label(rf, text=f"{row+1}.", bg=GREY, width=3).pack(side='left')
        tk.Entry(rf, textvariable=name_var, font=('Segoe UI',9), width=22).pack(side='left', padx=3)
        tk.Entry(rf, textvariable=desc_var, font=('Segoe UI',9), width=36).pack(side='left', padx=3)
        ttk.Combobox(rf, textvariable=stat_var, values=['Attached','Pending','Reference'],
                     width=12, font=('Segoe UI',9)).pack(side='left', padx=3)
        tk.Button(rf, text='✕', bg='#E74C3C', fg=WHITE, relief='flat', font=('Segoe UI',8),
                  command=lambda f=rf, v=(name_var,desc_var,stat_var): self._del_att_row(f,v)
                  ).pack(side='left', padx=2)
        self._att_vars.append((name_var, desc_var, stat_var))

    def _del_att_row(self, frame, var_tuple):
        frame.destroy()
        if var_tuple in self._att_vars:
            self._att_vars.remove(var_tuple)

    def _clear_attachments(self):
        for w in self.att_frame.winfo_children():
            w.destroy()
        self._att_vars.clear()

    # ── TAB 4: Output Settings ────────────────────────────────────────────────
    def _build_tab4(self):
        self._section_label(self.tab4, "Output Settings")

        frm = tk.Frame(self.tab4, bg=GREY, padx=20, pady=10)
        frm.pack(fill='both', expand=True)

        self.out_dir = tk.StringVar(value="C:/ECN")
        self.excel_path = tk.StringVar(value="C:/ECN/ECN.xlsx")
        self.gen_pdf  = tk.BooleanVar(value=True)
        self.gen_docx = tk.BooleanVar(value=True)
        self.gen_xlsx = tk.BooleanVar(value=True)

        def row(parent, label, var, browse_fn=None, row_n=0):
            tk.Label(parent, text=label, bg=GREY, fg=BRAND_DARK,
                     font=('Segoe UI', 10, 'bold')).grid(row=row_n, column=0, sticky='w', pady=6)
            e = tk.Entry(parent, textvariable=var, font=('Segoe UI', 10), width=50)
            e.grid(row=row_n, column=1, padx=8, sticky='ew')
            if browse_fn:
                tk.Button(parent, text="Browse…", font=('Segoe UI', 9),
                          bg=BRAND_MID, fg=WHITE, relief='flat', padx=8,
                          command=browse_fn).grid(row=row_n, column=2, padx=4)

        row(frm, "Output Folder:", self.out_dir,
            lambda: self.out_dir.set(filedialog.askdirectory(initialdir='C:\\')), 0)
        row(frm, "Excel Log File:", self.excel_path,
            lambda: self.excel_path.set(filedialog.asksaveasfilename(
                initialdir='C:\\', defaultextension='.xlsx',
                filetypes=[('Excel', '*.xlsx')])), 1)

        frm.columnconfigure(1, weight=1)

        tk.Label(frm, text="Generate:", bg=GREY, fg=BRAND_DARK,
                 font=('Segoe UI', 10, 'bold')).grid(row=2, column=0, sticky='w', pady=12)
        cbf = tk.Frame(frm, bg=GREY); cbf.grid(row=2, column=1, sticky='w', columnspan=2)
        for var, lbl in [(self.gen_pdf,'PDF (3 pages)'), (self.gen_xlsx,'Update Excel Log')]:
            tk.Checkbutton(cbf, text=lbl, variable=var, bg=GREY,
                           font=('Segoe UI', 10), fg=BRAND_DARK,
                           activebackground=GREY).pack(side='left', padx=12)

        tk.Label(frm, text="ℹ  File names are derived from ECN Number (alphanumeric only).\n"
                            "  e.g.  LECS/DOC/ECN0009  →  LECSDOCECN0009.pdf",
                 bg=GREY, fg='#666', font=('Segoe UI', 9, 'italic'),
                 justify='left').grid(row=3, column=0, columnspan=3, sticky='w', pady=10)

        # Log output
        self._section_label(self.tab4, "Generation Log")
        self.log_text = scrolledtext.ScrolledText(self.tab4, height=10, font=('Consolas', 9),
                                                   bg='#0D1117', fg='#7FC97F', state='disabled')
        self.log_text.pack(fill='both', expand=True, padx=12, pady=6)

    # ── HELPERS ───────────────────────────────────────────────────────────────
    def _section_label(self, parent, text):
        frm = tk.Frame(parent, bg=BRAND_DARK, height=32)
        frm.pack(fill='x', padx=0, pady=(10, 2))
        frm.pack_propagate(False)
        tk.Label(frm, text=f"  {text}", bg=BRAND_DARK, fg=WHITE,
                 font=('Segoe UI', 10, 'bold')).pack(side='left', padx=10, pady=4)

    def _render_sections(self, parent, sections, store, tall=None):
        tall = tall or []
        for sec_name, fields in sections:
            self._section_label(parent, sec_name)
            for key, label, placeholder in fields:
                frm = tk.Frame(parent, bg=GREY); frm.pack(fill='x', padx=20, pady=3)
                tk.Label(frm, text=label, bg=GREY, fg=BRAND_DARK,
                         font=('Segoe UI', 9, 'bold'), width=26, anchor='w').pack(side='left')
                if key in tall:
                    var = tk.Text(frm, height=3, font=('Segoe UI', 9), wrap='word',
                                  relief='flat', bd=1, highlightthickness=1,
                                  highlightbackground='#BBCFE0')
                    var.insert('1.0', placeholder)
                    var.pack(side='left', fill='x', expand=True)
                    store[key] = ('text', var)
                else:
                    sv = tk.StringVar(value=placeholder)
                    e = tk.Entry(frm, textvariable=sv, font=('Segoe UI', 9),
                                 relief='flat', bd=1, highlightthickness=1,
                                 highlightbackground='#BBCFE0')
                    e.pack(side='left', fill='x', expand=True)
                    store[key] = ('var', sv)

    def _get_field(self, stores, key):
        for store in stores:
            if key in store:
                typ, widget = store[key]
                if typ == 'text':
                    return widget.get('1.0', 'end').strip()
                else:
                    return widget.get().strip()
        return ''

    def _prefill_sample(self):
        pass  # already prefilled via placeholder values

    def _collect_data(self):
        stores = [self.fields1, self.fields2]
        keys = [
            'ecn_number','ecn_date','prepared_by','company',
            'main_product_code','main_product_desc','ecn_part_code','ecn_part_desc',
            'current_cost','ecn_cost','implementation_date','departments',
            'verify_dept','verify_name','approve_dept','approve_name',
            'details_of_change','reason_for_change','ecn_request_ref',
            'addition','deletion','change_docs',
            'action_change','action_resp','stock_disp','stock_resp',
            'inventory_action','inventory_resp','change_note','change_note_resp',
        ]
        data = {k: self._get_field(stores, k) for k in keys}
        data['notes'] = self.notes_text.get('1.0','end').strip()
        data['attachments'] = [
            {'name': n.get(), 'desc': d.get(), 'status': s.get()}
            for n, d, s in self._att_vars
            if n.get().strip()
        ]
        return data

    def _log(self, msg, colour=None):
        self.log_text.config(state='normal')
        self.log_text.insert('end', msg + '\n')
        self.log_text.see('end')
        self.log_text.config(state='disabled')
        self.status_var.set(msg)
        self.update_idletasks()

    # ── GENERATE ──────────────────────────────────────────────────────────────
    def _generate(self):
        threading.Thread(target=self._do_generate, daemon=True).start()

    def _do_generate(self):
        data = self._collect_data()

        # Validate required fields
        required = ['ecn_number', 'ecn_date', 'prepared_by',
                    'main_product_code', 'ecn_part_code',
                    'details_of_change', 'reason_for_change']
        missing = [k for k in required if not data.get(k)]
        if missing:
            messagebox.showerror("Validation Error",
                                 f"Please fill in required fields:\n{', '.join(missing)}")
            return

        fname    = safe_filename(data['ecn_number'])
        out_dir  = self.out_dir.get() or 'C:\\'

        # Ensure output dir exists (if on Linux / sandbox, use /tmp)
        if not os.path.exists(out_dir):
            out_dir = '/tmp'
            self._log(f"⚠  Output folder not found, using /tmp instead")

        pdf_path  = os.path.join(out_dir, fname + '.pdf')
#         docx_path = os.path.join(out_dir, fname + '.docx')
        xlsx_path = self.excel_path.get() or os.path.join(out_dir, 'ECN_Log.xlsx')
        if not os.path.exists(os.path.dirname(xlsx_path) or '.'):
            xlsx_path = os.path.join('/tmp', 'ECN_Log.xlsx')

        self._log("─" * 50)
        self._log(f"⚙  Starting ECN generation for: {data['ecn_number']}")
        self._log(f"📁  Output folder: {out_dir}")
        self._log(f"📄  Filename base: {fname}")

        try:
            if self.gen_pdf.get():
                self._log("🔄  Generating PDF...")
                generate_pdf(data, pdf_path)
                self._log(f"✅  PDF saved → {pdf_path}")
        except Exception as e:
            self._log(f"❌  PDF error: {e}")

       # try:
         #   if self.gen_docx.get():
        #        self._log("🔄  Generating DOCX...")
        #        generate_docx(data, docx_path)
        #        self._log(f"✅  DOCX saved → {docx_path}")
       # except Exception as e:
       #     self._log(f"❌  DOCX error: {e}")

        try:
            if self.gen_xlsx.get():
                self._log("🔄  Updating Excel log...")
                update_excel_log(data, xlsx_path)
                self._log(f"✅  Excel log updated → {xlsx_path}")
        except Exception as e:
            self._log(f"❌  Excel error: {e}")

        self._log("─" * 50)
        self._log("🎉  All done!")
        messagebox.showinfo("ECN Generated",
                            f"ECN documents generated successfully!\n\n"
                            f"PDF:  {pdf_path if self.gen_pdf.get() else 'skipped'}\n"
                           # f"DOCX: {docx_path if self.gen_docx.get() else 'skipped'}\n"
                            f"XLSX: {xlsx_path if self.gen_xlsx.get() else 'skipped'}")


# ─────────────────────────────────────────────────────────────────────────────
if __name__ == '__main__':
    app = ECNApp()
    app.mainloop()