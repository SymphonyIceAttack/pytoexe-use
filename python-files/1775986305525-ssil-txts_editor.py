import tkinter as tk
from tkinter import ttk, filedialog, messagebox, simpledialog
import json
import os

class TXTSEditor:
    def __init__(self, root):
        self.root = root
        self.root.title("Редектор TXTS")
        self.root.geometry("800x600")
        
        # Данные
        self.sections = []          # [{"name": "...", "text": "..."}, ...]
        self.current_file = None
        self.unsaved = False
        
        # Меню
        menubar = tk.Menu(root)
        root.config(menu=menubar)
        
        file_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Файл", menu=file_menu)
        file_menu.add_command(label="Новый", command=self.new_file, accelerator="Ctrl+N")
        file_menu.add_command(label="Открыть .txts...", command=self.open_file, accelerator="Ctrl+O")
        file_menu.add_command(label="Сохранить", command=self.save_file, accelerator="Ctrl+S")
        file_menu.add_command(label="Сохранить как...", command=self.save_as_file, accelerator="Ctrl+Shift+S")
        file_menu.add_separator()
        file_menu.add_command(label="Экспорт в .txt...", command=self.export_txt)
        file_menu.add_command(label="Экспорт в .docx...", command=self.export_docx)
        file_menu.add_separator()
        file_menu.add_command(label="Выход", command=self.quit_app)
        
        section_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Раздел", menu=section_menu)
        section_menu.add_command(label="Новый раздел", command=self.new_section, accelerator="Ctrl+T")
        section_menu.add_command(label="Переименовать", command=self.rename_section, accelerator="F2")
        section_menu.add_command(label="Удалить раздел", command=self.delete_section, accelerator="Ctrl+W")
        
        # Горячие клавиши
        root.bind("<Control-n>", lambda e: self.new_file())
        root.bind("<Control-o>", lambda e: self.open_file())
        root.bind("<Control-s>", lambda e: self.save_file())
        root.bind("<Control-S>", lambda e: self.save_as_file())
        root.bind("<Control-t>", lambda e: self.new_section())
        root.bind("<Control-w>", lambda e: self.delete_section())
        root.bind("<F2>", lambda e: self.rename_section())
        
        # ===== ВКЛАДКИ (как в Chrome) =====
        self.notebook = ttk.Notebook(root)
        self.notebook.pack(fill=tk.BOTH, expand=True, padx=0, pady=0)
        self.notebook.bind("<<NotebookTabChanged>>", self.on_tab_changed)
        
        # Привязка Ctrl+Tab для переключения вкладок
        root.bind("<Control-Tab>", lambda e: self.next_tab())
        root.bind("<Control-Shift-Tab>", lambda e: self.prev_tab())
        
        # Текстовое поле будет создаваться для каждой вкладки динамически
        self.text_widgets = {}  # tab_id -> Text widget
        
        # Кнопка "+" справа от вкладок (нестандартно, добавим отдельно)
        self.add_tab_btn = ttk.Button(root, text="+", width=3, command=self.new_section)
        self.add_tab_btn.place(relx=1.0, x=-5, y=2, anchor="ne")
        
        # Привязка изменения размера для корректировки позиции кнопки "+"
        root.bind("<Configure>", self.reposition_add_button)
        
        # Начинаем с одного пустого раздела
        self.new_file()
    
    def reposition_add_button(self, event=None):
        self.add_tab_btn.place(relx=1.0, x=-5, y=2, anchor="ne")
    
    def next_tab(self):
        idx = self.notebook.index("current")
        if idx < len(self.notebook.tabs()) - 1:
            self.notebook.select(idx + 1)
        else:
            self.notebook.select(0)
        return "break"
    
    def prev_tab(self):
        idx = self.notebook.index("current")
        if idx > 0:
            self.notebook.select(idx - 1)
        else:
            self.notebook.select(len(self.notebook.tabs()) - 1)
        return "break"
    
    def new_section(self, name=None, text=""):
        if name is None:
            name = f"Раздел {len(self.sections) + 1}"
        
        # Создаем фрейм для вкладки
        tab_frame = ttk.Frame(self.notebook)
        
        # Текстовое поле (как в Блокноте)
        text_widget = tk.Text(tab_frame, wrap=tk.WORD, undo=True, font=("Consolas", 11))
        text_widget.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # Скроллбар
        scrollbar = ttk.Scrollbar(text_widget, orient="vertical", command=text_widget.yview)
        text_widget.configure(yscrollcommand=scrollbar.set)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Вставляем текст, если есть
        if text:
            text_widget.insert("1.0", text)
            text_widget.edit_modified(False)
        
        # Привязываем событие изменения текста
        text_widget.bind("<<Modified>>", self.on_text_modified)
        
        # Добавляем вкладку
        self.notebook.add(tab_frame, text=name)
        tab_id = self.notebook.tabs()[-1]
        self.text_widgets[tab_id] = text_widget
        
        # Сохраняем в структуру данных
        self.sections.append({"name": name, "text": text})
        
        # Выбираем новую вкладку
        self.notebook.select(tab_id)
        text_widget.focus_set()
        
        self.mark_unsaved()
    
    def on_tab_changed(self, event=None):
        # Синхронизируем текст из виджета в self.sections при переключении
        current_tab = self.notebook.select()
        if current_tab and current_tab in self.text_widgets:
            # Ничего не делаем, просто обновляем UI при необходимости
            pass
    
    def on_text_modified(self, event=None):
        widget = event.widget
        if widget.edit_modified():
            self.mark_unsaved()
            widget.edit_modified(False)
    
    def get_current_tab_info(self):
        tab_id = self.notebook.select()
        if not tab_id:
            return None, -1
        idx = self.notebook.index(tab_id)
        return tab_id, idx
    
    def sync_current_text_to_data(self):
        """Сохраняет текст из активного виджета в self.sections"""
        tab_id, idx = self.get_current_tab_info()
        if tab_id and idx < len(self.sections):
            widget = self.text_widgets[tab_id]
            self.sections[idx]["text"] = widget.get("1.0", "end-1c")
    
    def sync_all_texts_to_data(self):
        """Сохраняет текст из всех виджетов в self.sections"""
        for i, tab_id in enumerate(self.notebook.tabs()):
            if i < len(self.sections) and tab_id in self.text_widgets:
                widget = self.text_widgets[tab_id]
                self.sections[i]["text"] = widget.get("1.0", "end-1c")
    
    def rename_section(self):
        tab_id, idx = self.get_current_tab_info()
        if idx < 0:
            return
        new_name = simpledialog.askstring("Переименовать", "Новое название раздела:",
                                          initialvalue=self.sections[idx]["name"])
        if new_name:
            self.sections[idx]["name"] = new_name
            self.notebook.tab(tab_id, text=new_name)
            self.mark_unsaved()
    
    def delete_section(self):
        if len(self.sections) <= 1:
            messagebox.showwarning("Предупреждение", "Должен остаться хотя бы один раздел.")
            return
        
        tab_id, idx = self.get_current_tab_info()
        if idx < 0:
            return
        
        if messagebox.askyesno("Удалить", f"Удалить раздел «{self.sections[idx]['name']}»?"):
            self.notebook.forget(tab_id)
            del self.text_widgets[tab_id]
            del self.sections[idx]
            self.mark_unsaved()
    
    def mark_unsaved(self):
        self.unsaved = True
        title = "Редектор TXTS"
        if self.current_file:
            title += f" - {os.path.basename(self.current_file)}"
        title += " *"
        self.root.title(title)
    
    def clear_unsaved(self):
        self.unsaved = False
        title = "Редектор TXTS"
        if self.current_file:
            title += f" - {os.path.basename(self.current_file)}"
        self.root.title(title)
    
    def new_file(self):
        if self.unsaved:
            if not self.ask_save_changes():
                return
        
        # Очищаем все вкладки
        for tab_id in self.notebook.tabs():
            self.notebook.forget(tab_id)
        self.text_widgets.clear()
        self.sections.clear()
        self.current_file = None
        
        # Создаем один пустой раздел
        self.new_section("Раздел 1")
        self.clear_unsaved()
    
    def open_file(self):
        if self.unsaved:
            if not self.ask_save_changes():
                return
        
        filepath = filedialog.askopenfilename(
            title="Открыть файл TXTS",
            filetypes=[("TXTS файлы", "*.txts"), ("Все файлы", "*.*")]
        )
        if not filepath:
            return
        
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                data = json.load(f)
            
            # Проверяем структуру
            if not isinstance(data, list):
                raise ValueError("Неверный формат файла")
            
            # Очищаем текущее
            for tab_id in self.notebook.tabs():
                self.notebook.forget(tab_id)
            self.text_widgets.clear()
            self.sections.clear()
            
            # Загружаем разделы
            for item in data:
                name = item.get("name", "Без названия")
                text = item.get("text", "")
                self.new_section(name, text)
            
            self.current_file = filepath
            self.clear_unsaved()
            
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось открыть файл:\n{e}")
    
    def save_file(self):
        if self.current_file:
            self._do_save(self.current_file)
        else:
            self.save_as_file()
    
    def save_as_file(self):
        filepath = filedialog.asksaveasfilename(
            title="Сохранить как TXTS",
            defaultextension=".txts",
            filetypes=[("TXTS файлы", "*.txts"), ("Все файлы", "*.*")]
        )
        if filepath:
            self._do_save(filepath)
    
    def _do_save(self, filepath):
        self.sync_all_texts_to_data()
        
        try:
            with open(filepath, "w", encoding="utf-8") as f:
                json.dump(self.sections, f, ensure_ascii=False, indent=2)
            self.current_file = filepath
            self.clear_unsaved()
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось сохранить файл:\n{e}")
    
    def ask_save_changes(self):
        answer = messagebox.askyesnocancel("Сохранить изменения?", 
                                           "Сохранить изменения перед закрытием?")
        if answer is None:  # Cancel
            return False
        if answer:  # Yes
            self.save_file()
        return True
    
    def export_txt(self):
        self.sync_all_texts_to_data()
        
        filepath = filedialog.asksaveasfilename(
            title="Экспорт в TXT",
            defaultextension=".txt",
            filetypes=[("Текстовые файлы", "*.txt"), ("Все файлы", "*.*")]
        )
        if not filepath:
            return
        
        try:
            with open(filepath, "w", encoding="utf-8") as f:
                for i, section in enumerate(self.sections):
                    f.write(section["name"] + "\n\n")
                    f.write(section["text"] + "\n")
                    if i < len(self.sections) - 1:
                        f.write("\n\n")
            messagebox.showinfo("Экспорт", f"Успешно экспортировано в {filepath}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка экспорта:\n{e}")
    
    def export_docx(self):
        try:
            from docx import Document
        except ImportError:
            messagebox.showerror("Ошибка", 
                "Для экспорта в DOCX требуется библиотека python-docx.\n"
                "Установите её командой: pip install python-docx")
            return
        
        self.sync_all_texts_to_data()
        
        filepath = filedialog.asksaveasfilename(
            title="Экспорт в DOCX",
            defaultextension=".docx",
            filetypes=[("Документы Word", "*.docx"), ("Все файлы", "*.*")]
        )
        if not filepath:
            return
        
        try:
            doc = Document()
            for section in self.sections:
                doc.add_heading(section["name"], level=1)
                doc.add_paragraph(section["text"])
                doc.add_paragraph()  # пустая строка между разделами
            doc.save(filepath)
            messagebox.showinfo("Экспорт", f"Успешно экспортировано в {filepath}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка экспорта:\n{e}")
    
    def quit_app(self):
        if self.unsaved:
            if not self.ask_save_changes():
                return
        self.root.quit()


if __name__ == "__main__":
    root = tk.Tk()
    app = TXTSEditor(root)
    root.protocol("WM_DELETE_WINDOW", app.quit_app)
    root.mainloop()