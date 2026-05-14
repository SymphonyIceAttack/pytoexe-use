# 安装依赖：pip install pandas openpyxl python-docx
import pandas as pd
import re
import logging
import tkinter as tk
from tkinter import filedialog, messagebox
from datetime import datetime, timedelta
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# -------------------------- 日志配置 --------------------------
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S',
    handlers=[logging.StreamHandler()]
)
logger = logging.getLogger(__name__)

# -------------------------- 全局配置 --------------------------
FONT_NAME = '宋体'
FONT_SIZE_2 = Pt(22)
FONT_SIZE_4 = Pt(12)
INDENT_2_CHAR = Cm(0.74)
SUGGEST_TEXT = "建议辖区重点开展矛盾调解跟进，联合社区开展矛调工作。"

# -------------------------------------------------------------------
# 1. 弹窗选择 Excel 文件
# -------------------------------------------------------------------
def get_excel_file_path():
    root = tk.Tk()
    root.withdraw()
    file_path = filedialog.askopenfilename(
        title="请选择要处理的Excel文件",
        filetypes=[("Excel文件", "*.xlsx;*.xls")],
        initialdir="."
    )
    root.destroy()
    return file_path

# -------------------------------------------------------------------
# 2. 弹窗选择执行日期
# -------------------------------------------------------------------
def select_execution_date():
    root = tk.Tk()
    root.title("选择执行日期")
    root.geometry("300x150")
    root.resizable(False, False)
    selected_date = None

    def confirm():
        nonlocal selected_date
        try:
            date_str = entry.get()
            selected_date = datetime.strptime(date_str, "%Y-%m-%d")
            root.quit()
            root.destroy()
        except:
            messagebox.showerror("错误", "格式错误，请输入 YYYY-MM-DD")

    tk.Label(root, text="请输入执行日期（格式：2025-05-12）", pady=10).pack()
    entry = tk.Entry(root, width=30)
    entry.insert(0, datetime.now().strftime("%Y-%m-%d"))
    entry.pack(pady=5)
    tk.Button(root, text="确认", command=confirm, width=10).pack(pady=10)
    root.mainloop()
    return selected_date

# -------------------------------------------------------------------
# 3. 自动不重复文件名
# -------------------------------------------------------------------
def get_unique_save_path(excel_dir, target_date):
    month = target_date.month
    day = target_date.day
    base_name = f"{month}月{day}日矛盾纠纷调解线索研判预警工作网安汇报"
    save_path = os.path.join(excel_dir, f"{base_name}.docx")
    idx = 1
    while os.path.exists(save_path):
        save_path = os.path.join(excel_dir, f"{base_name}({idx}).docx")
        idx += 1
    return save_path

# -------------------------------------------------------------------
# 主程序
# -------------------------------------------------------------------
def main():
    logger.info("="*60)
    logger.info("程序启动 - 辖区统计修复版")
    logger.info("="*60)

    # 1. 选择 Excel
    excel_path = get_excel_file_path()
    if not excel_path:
        logger.warning("未选择文件，退出")
        return

    # 2. 弹窗选择执行日期
    exec_date = select_execution_date()
    if not exec_date:
        logger.warning("未选择日期，退出")
        return

    # 3. 计算昨天 / 今天
    yesterday = exec_date - timedelta(days=1)
    today = exec_date

    # 4. 读取 Excel
    df = pd.read_excel(excel_path, header=2)
    df.columns = [str(c).strip() for c in df.columns]
    df = df.dropna(how="all")

    # ===================== 只取 1 天：执行日期的昨天 =====================
    df["下发日期"] = pd.to_datetime(df["下发日期"], errors="coerce")
    filtered_df = df[df["下发日期"].dt.date == yesterday.date()].copy()

    # 提取警情编号
    pattern = re.compile(r'【警情编号】\s*(\d+)')
    filtered_df["警情编号"] = filtered_df["线索内容"].astype(str).str.extract(pattern)
    filtered_df = filtered_df.dropna(subset=["警情编号"])
    groups = filtered_df.groupby("警情编号")

    # 统计总数
    total_person = len(filtered_df)
    total_alert = len(groups)

    # ===================== 分辖区统计（恢复你原来的完整格式） =====================
    county_stats = filtered_df.groupby('所属区县').agg(
        警情数=('警情编号', 'nunique'), 人数=('姓名', 'count')
    ).reset_index()
    county_text_parts = [f"{row['所属区县']}{row['警情数']}条{row['人数']}人" for _, row in county_stats.iterrows()]
    county_text = '，'.join(county_text_parts)

    # ===================== 第一段完整文本（100% 恢复你原来的样子） =====================
    y = yesterday.year
    m = yesterday.month
    d1 = yesterday.day
    d2 = today.day

    report_text = (
        f"核查梳理{y}年{m}月{d1}日8时至{m}月{d2}日8时警情 条，"
        f"涉矛盾纠纷警情{total_alert}条，矛盾纠纷排查对象{total_person}人。"
        f"其中：{county_text}。核查情况如下："
    )

    # 保存路径
    save_path = get_unique_save_path(os.path.dirname(excel_path), yesterday)

    # ===================== 生成 Word =====================
    doc = Document()

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run(f"网安支队侦查取证大队{m}月{d1}日矛盾纠纷调解线索研判预警工作汇报")
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_2
    run.font.bold = True

    # 正文样式
    doc.styles["Normal"].font.name = FONT_NAME
    doc.styles["Normal"].font.size = FONT_SIZE_4

    # 第一段
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = INDENT_2_CHAR
    p.add_run(report_text)

    # 小标题
    p2 = doc.add_paragraph()
    r2 = p2.add_run("矛盾纠纷警情排查：")
    r2.font.bold = True

    # 遍历警情
    for idx, (alert_id, rows) in enumerate(groups, 1):
        # 警情编号
        para = doc.add_paragraph()
        para.paragraph_format.first_line_indent = INDENT_2_CHAR
        para.add_run(f"{idx}.【警情编号】{alert_id}")

        # D列 + O列
        first = rows.iloc[0]
        clue = str(first["线索内容"]).strip()
        handle = str(first["处置情况"]).strip()

        for line in clue.split("\n"):
            line = line.strip()
            if line and not line.startswith("【警情编号"):
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = INDENT_2_CHAR
                p.add_run(line)

        for line in handle.split("\n"):
            line = line.strip()
            if line:
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = INDENT_2_CHAR
                p.add_run(line)

        # 网安核查
        net_para = doc.add_paragraph()
        net_para.paragraph_format.first_line_indent = INDENT_2_CHAR
        r_net = net_para.add_run("【网安核查情况】")
        r_net.font.bold = True

        # 拼接格式：姓名，G：值，H：值。J：值K：值L：值
        for i, (_, row) in enumerate(rows.iterrows(), 1):
            Fn = str(row.get("姓名", "")).strip()

            G3 = str(df.columns[6]) if len(df.columns) > 6 else ""
            Gn = str(row.iloc[6]) if len(row) > 6 else ""

            H3 = str(df.columns[7]) if len(df.columns) > 7 else ""
            Hn = str(row.iloc[7]) if len(row) > 7 else ""

            J3 = str(df.columns[9]) if len(df.columns) > 9 else ""
            Jn = str(row.iloc[9]) if len(row) > 9 else ""

            K3 = str(df.columns[10]) if len(df.columns) > 10 else ""
            Kn = str(row.iloc[10]) if len(row) > 10 else ""

            L3 = str(df.columns[11]) if len(df.columns) > 11 else ""
            Ln = str(row.iloc[11]) if len(row) > 11 else ""

            partG = f"{G3}：{Gn}" if Gn else ""
            partH = f"{H3}：{Hn}" if Hn else ""
            partJ = f"{J3}：{Jn}" if Jn else ""
            partK = f"{K3}：{Kn}" if Kn else ""
            partL = f"{L3}：{Ln}" if Ln else ""

            gh = f"{partG}，{partH}" if (partG and partH) else (partG or partH)
            jkl = f"{partJ}{partK}{partL}"
            full = f"{gh}。{jkl}".strip("。")
            content = f"{Fn}，{full}".strip()

            line_p = doc.add_paragraph()
            line_p.paragraph_format.first_line_indent = INDENT_2_CHAR
            line_p.add_run(f"{i}、{content}")

        # 建议
        sug_p = doc.add_paragraph()
        sug_p.paragraph_format.first_line_indent = INDENT_2_CHAR
        sug_p.add_run(SUGGEST_TEXT)

        # 空行
        doc.add_paragraph()

    doc.save(save_path)
    logger.info(f"✅ 生成完成：{save_path}")

if __name__ == "__main__":
    main()