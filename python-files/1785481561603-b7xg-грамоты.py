import pandas as pd
import re
import os
from docx import Document
from tkinter import Tk, Label, Button, Listbox, MULTIPLE, messagebox, filedialog, Frame, Scrollbar, VERTICAL, RIGHT, Y, BOTH, END, Toplevel, Entry
from tkinter import ttk
import tkinter as tk

class CertificateApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Генератор грамот из Excel")
        self.root.geometry("1100x800")
        self.root.resizable(True, True)
        
        self.excel_file = None
        self.template_file = None
        self.sheets_data = {}
        self.all_participants = []
        
        # Настройки столбцов (по умолчанию для вашего файла)
        self.name_col = 1      # столбец B - ФИО
        self.city_col = 0      # столбец A - город/группа
        self.phone_col = 8     # столбец I - телефон
        self.parent_col = 7    # столбец H - родитель
        self.age_col = 3       # столбец D - возраст
        
        self.create_widgets()
    
    def create_widgets(self):
        main_frame = Frame(self.root)
        main_frame.pack(fill=BOTH, expand=True, padx=10, pady=10)
        
        info_frame = Frame(main_frame, relief="groove", bd=2)
        info_frame.pack(fill="x", pady=5)
        
        Label(info_frame, text="📁 Генератор грамот для участников", font=("Arial", 14, "bold")).pack(pady=5)
        
        load_frame = Frame(main_frame)
        load_frame.pack(fill="x", pady=5)
        
        Button(load_frame, text="📂 Загрузить Excel файл", command=self.load_excel, 
               width=25, height=2, bg="#2196F3", fg="white").pack(side="left", padx=5)
        Button(load_frame, text="📄 Загрузить шаблон грамоты", command=self.load_template,
               width=25, height=2, bg="#FF9800", fg="white").pack(side="left", padx=5)
        Button(load_frame, text="⚙️ Настройка столбцов", command=self.column_settings,
               width=20, height=2, bg="#9C27B0", fg="white").pack(side="left", padx=5)
        
        self.info_label = Label(main_frame, text="⚠️ Файлы не загружены", fg="red", font=("Arial", 10))
        self.info_label.pack(pady=5)
        
        ttk.Separator(main_frame, orient="horizontal").pack(fill="x", pady=5)
        
        list_frame = Frame(main_frame)
        list_frame.pack(fill=BOTH, expand=True, pady=10)
        
        list_header_frame = Frame(list_frame)
        list_header_frame.pack(fill="x")
        
        Label(list_header_frame, text="👥 Участники из Excel", font=("Arial", 12, "bold")).pack(side="left")
        Label(list_header_frame, text="(Нажмите Ctrl для множественного выбора)", font=("Arial", 9)).pack(side="left", padx=10)
        
        search_frame = Frame(list_frame)
        search_frame.pack(fill="x", pady=5)
        
        Label(search_frame, text="🔍 Поиск:").pack(side="left", padx=5)
        self.search_entry = Entry(search_frame, width=30)
        self.search_entry.pack(side="left", padx=5)
        self.search_entry.bind('<KeyRelease>', self.search_participants)
        Button(search_frame, text="Очистить", command=self.clear_search, width=10).pack(side="left", padx=5)
        
        listbox_frame = Frame(list_frame)
        listbox_frame.pack(fill=BOTH, expand=True, pady=5)
        
        scrollbar = Scrollbar(listbox_frame, orient=VERTICAL)
        self.listbox = Listbox(listbox_frame, selectmode=MULTIPLE, yscrollcommand=scrollbar.set, 
                              font=("Arial", 10))
        scrollbar.config(command=self.listbox.yview)
        
        self.listbox.pack(side="left", fill=BOTH, expand=True)
        scrollbar.pack(side=RIGHT, fill=Y)
        
        btn_frame = Frame(main_frame)
        btn_frame.pack(fill="x", pady=5)
        
        Button(btn_frame, text="✅ Выбрать всех", command=self.select_all, 
               bg="#4CAF50", fg="white", width=15).pack(side="left", padx=5)
        Button(btn_frame, text="❌ Снять выделение", command=self.deselect_all,
               bg="#f44336", fg="white", width=15).pack(side="left", padx=5)
        Button(btn_frame, text="📋 Выбрать по группе", command=self.select_by_group,
               bg="#9C27B0", fg="white", width=15).pack(side="left", padx=5)
        Button(btn_frame, text="📊 Выбрать по городу", command=self.select_by_city,
               bg="#FF6B00", fg="white", width=15).pack(side="left", padx=5)
        
        self.count_label = Label(main_frame, text="Всего участников: 0 | Выбрано: 0", font=("Arial", 10))
        self.count_label.pack(pady=5)
        
        self.generate_btn = Button(main_frame, text="🎯 Сгенерировать грамоты", 
                                   command=self.generate_certificates,
                                   width=30, height=2, bg="#4CAF50", fg="white", font=("Arial", 12, "bold"))
        self.generate_btn.pack(pady=10)
        
        self.status_label = Label(main_frame, text="✅ Готов к работе", fg="blue", font=("Arial", 10))
        self.status_label.pack()
    
    def load_excel(self):
        file_path = filedialog.askopenfilename(
            title="Выберите Excel файл",
            filetypes=[("Excel files", "*.xlsx *.xls"), ("All files", "*.*")]
        )
        
        if not file_path:
            return
        
        try:
            self.excel_file = file_path
            self.sheets_data = {}
            self.all_participants = []
            self.listbox.delete(0, END)
            self.search_entry.delete(0, END)
            
            excel_file = pd.ExcelFile(file_path)
            sheet_names = excel_file.sheet_names
            
            for sheet_name in sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name, header=None)
                
                name_col = self.name_col
                city_col = self.city_col
                phone_col = self.phone_col
                parent_col = self.parent_col
                age_col = self.age_col
                
                start_row = self.find_start_row(df, name_col)
                
                participants = []
                for idx in range(start_row, len(df)):
                    try:
                        row = df.iloc[idx]
                        
                        if name_col >= len(row) or pd.isna(row[name_col]):
                            continue
                        
                        name = str(row[name_col]).strip()
                        if not name or len(name) < 2:
                            continue
                        
                        if name.startswith(("г.", "Московск", "Кемеров", "Амурская", "Иркутская", "Красноярск", "Таиланд")):
                            continue
                        
                        name_parts = name.split()
                        if len(name_parts) < 2:
                            continue
                        
                        city = ""
                        try:
                            if city_col < len(row) and not pd.isna(row[city_col]) and isinstance(row[city_col], str):
                                city = str(row[city_col]).strip()
                        except:
                            pass
                        
                        phone = ""
                        try:
                            if phone_col < len(row) and not pd.isna(row[phone_col]):
                                phone = str(row[phone_col]).strip()
                        except:
                            pass
                        
                        parent = ""
                        try:
                            if parent_col < len(row) and not pd.isna(row[parent_col]) and isinstance(row[parent_col], str):
                                parent = str(row[parent_col]).strip()
                        except:
                            pass
                        
                        age = ""
                        try:
                            if age_col < len(row) and not pd.isna(row[age_col]):
                                age = str(row[age_col]).strip()
                        except:
                            pass
                        
                        participants.append({
                            "name": name,
                            "city": city,
                            "phone": phone,
                            "parent": parent,
                            "age": age,
                            "sheet": sheet_name,
                            "row": idx
                        })
                    except Exception as e:
                        continue
                
                if participants:
                    self.sheets_data[sheet_name] = participants
                    self.all_participants.extend(participants)
            
            self.update_listbox()
            
            total = len(self.all_participants)
            sheets = len(self.sheets_data)
            
            if total > 0:
                self.info_label.config(text=f"✅ Загружен Excel: {os.path.basename(file_path)}. Найдено {total} участников в {sheets} листах.", fg="green")
                self.status_label.config(text="✅ Excel загружен успешно", fg="green")
                self.count_label.config(text=f"Всего участников: {total} | Выбрано: 0")
            else:
                self.info_label.config(text="⚠️ Не найдено участников. Проверьте структуру файла.", fg="orange")
                self.status_label.config(text="⚠️ Участники не найдены", fg="orange")
            
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось загрузить Excel файл:\n{str(e)}")
            self.status_label.config(text="❌ Ошибка загрузки Excel", fg="red")
    
    def find_start_row(self, df, name_col):
        for i in range(min(20, len(df))):
            try:
                val = df.iloc[i, name_col]
                if pd.notna(val) and isinstance(val, str) and val.strip():
                    if len(val.strip()) > 2 and any(c.isalpha() for c in val):
                        if not val.strip().startswith(("г.", "Московск", "Кемеров", "Амурская", "Иркутская", "Красноярск")):
                            return i
            except:
                continue
        
        for i in range(min(30, len(df))):
            try:
                val = df.iloc[i, name_col]
                if pd.notna(val) and isinstance(val, str) and val.strip():
                    if len(val.strip().split()) >= 2 and len(val.strip()) > 5:
                        return i
            except:
                continue
        
        return 0
    
    def update_listbox(self):
        self.listbox.delete(0, END)
        for p in self.all_participants:
            city_short = p["city"][:30] if p["city"] else ""
            display_text = f"{p['name']} ({city_short})" if city_short else p['name']
            self.listbox.insert(END, display_text)
    
    def load_template(self):
        file_path = filedialog.askopenfilename(
            title="Выберите шаблон грамоты (DOCX)",
            filetypes=[("Word files", "*.docx"), ("All files", "*.*")]
        )
        
        if not file_path:
            return
        
        try:
            doc = Document(file_path)
            self.template_file = file_path
            self.info_label.config(text=f"✅ Шаблон загружен: {os.path.basename(file_path)}", fg="green")
            self.status_label.config(text="✅ Шаблон загружен успешно", fg="green")
            
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось загрузить шаблон:\n{str(e)}")
            self.status_label.config(text="❌ Ошибка загрузки шаблона", fg="red")
    
    def column_settings(self):
        if not self.excel_file:
            messagebox.showinfo("Информация", "Сначала загрузите Excel файл")
            return
        
        settings_window = Toplevel(self.root)
        settings_window.title("Настройка столбцов Excel")
        settings_window.geometry("500x400")
        settings_window.transient(self.root)
        settings_window.grab_set()
        
        Label(settings_window, text="Настройка соответствия столбцов", font=("Arial", 12, "bold")).pack(pady=10)
        Label(settings_window, text="Введите номера столбцов (1, 2, 3...):", font=("Arial", 10)).pack(pady=5)
        
        settings_frame = Frame(settings_window)
        settings_frame.pack(fill=BOTH, expand=True, padx=20, pady=10)
        
        fields = [
            ("Имя участника (столбец):", "name_col"),
            ("Город/Группа (столбец):", "city_col"),
            ("Телефон (столбец):", "phone_col"),
            ("Родитель (столбец):", "parent_col"),
            ("Возраст (столбец):", "age_col")
        ]
        
        self.col_entries = {}
        
        for label, key in fields:
            frame = Frame(settings_frame)
            frame.pack(fill="x", pady=5)
            
            Label(frame, text=label, width=25, anchor="w").pack(side="left")
            
            entry = Entry(frame, width=5)
            entry.insert(0, str(getattr(self, key) + 1))
            entry.pack(side="left", padx=5)
            self.col_entries[key] = entry
        
        def save_settings():
            try:
                for key, entry in self.col_entries.items():
                    value = int(entry.get().strip()) - 1
                    if value >= 0:
                        setattr(self, key, value)
                
                messagebox.showinfo("Успех", "Настройки сохранены!\nПерезагрузите Excel файл для применения изменений.")
                settings_window.destroy()
                
            except ValueError:
                messagebox.showerror("Ошибка", "Введите корректные номера столбцов (целые числа)")
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось сохранить настройки:\n{str(e)}")
        
        Button(settings_window, text="💾 Сохранить", command=save_settings,
               bg="#4CAF50", fg="white", width=15).pack(pady=10)
        Button(settings_window, text="❌ Отмена", command=settings_window.destroy,
               bg="#f44336", fg="white", width=15).pack(pady=5)
    
    def select_all(self):
        self.listbox.selection_set(0, END)
        selected_count = len(self.listbox.curselection())
        total_count = self.listbox.size()
        self.count_label.config(text=f"Всего участников: {total_count} | Выбрано: {selected_count}")
        self.status_label.config(text=f"✅ Выбрано {selected_count} участников", fg="blue")
    
    def deselect_all(self):
        self.listbox.selection_clear(0, END)
        total_count = self.listbox.size()
        self.count_label.config(text=f"Всего участников: {total_count} | Выбрано: 0")
        self.status_label.config(text="✅ Выделение снято", fg="blue")
    
    def select_by_group(self):
        if not self.sheets_data:
            messagebox.showinfo("Информация", "Сначала загрузите Excel файл")
            return
        
        group_window = Toplevel(self.root)
        group_window.title("Выбор группы")
        group_window.geometry("400x350")
        group_window.transient(self.root)
        group_window.grab_set()
        
        Label(group_window, text="Выберите группы для выделения:", font=("Arial", 12)).pack(pady=10)
        
        group_listbox = Listbox(group_window, selectmode=MULTIPLE, height=12, font=("Arial", 10))
        group_listbox.pack(fill=BOTH, expand=True, padx=10, pady=5)
        
        sheet_names = list(self.sheets_data.keys())
        for name in sheet_names:
            count = len(self.sheets_data[name])
            group_listbox.insert(END, f"{name} ({count} чел.)")
        
        def apply_selection():
            selected_indices = group_listbox.curselection()
            if not selected_indices:
                messagebox.showinfo("Информация", "Выберите хотя бы одну группу")
                return
            
            names_to_select = []
            for idx in selected_indices:
                sheet_name = sheet_names[idx]
                for p in self.sheets_data[sheet_name]:
                    names_to_select.append(p["name"])
            
            self.listbox.selection_clear(0, END)
            selected_count = 0
            for i in range(self.listbox.size()):
                item_text = self.listbox.get(i)
                for participant_name in names_to_select:
                    if item_text.startswith(participant_name):
                        self.listbox.selection_set(i)
                        selected_count += 1
                        break
            
            total_count = self.listbox.size()
            self.count_label.config(text=f"Всего участников: {total_count} | Выбрано: {selected_count}")
            self.status_label.config(text=f"✅ Выбрано {selected_count} участников из {len(selected_indices)} групп", fg="blue")
            group_window.destroy()
        
        Button(group_window, text="✅ Применить", command=apply_selection, 
               bg="#4CAF50", fg="white", width=15).pack(pady=10)
        Button(group_window, text="❌ Отмена", command=group_window.destroy,
               bg="#f44336", fg="white", width=15).pack(pady=5)
    
    def select_by_city(self):
        if not self.all_participants:
            messagebox.showinfo("Информация", "Сначала загрузите Excel файл")
            return
        
        cities = set()
        for p in self.all_participants:
            if p["city"]:
                cities.add(p["city"])
        
        if not cities:
            messagebox.showinfo("Информация", "Не найдено городов для фильтрации")
            return
        
        city_window = Toplevel(self.root)
        city_window.title("Выбор города")
        city_window.geometry("400x350")
        city_window.transient(self.root)
        city_window.grab_set()
        
        Label(city_window, text="Выберите города для выделения:", font=("Arial", 12)).pack(pady=10)
        
        city_listbox = Listbox(city_window, selectmode=MULTIPLE, height=12, font=("Arial", 10))
        city_listbox.pack(fill=BOTH, expand=True, padx=10, pady=5)
        
        sorted_cities = sorted(list(cities))
        for city in sorted_cities:
            count = sum(1 for p in self.all_participants if p["city"] == city)
            city_listbox.insert(END, f"{city} ({count} чел.)")
        
        def apply_selection():
            selected_indices = city_listbox.curselection()
            if not selected_indices:
                messagebox.showinfo("Информация", "Выберите хотя бы один город")
                return
            
            names_to_select = []
            for idx in selected_indices:
                city = sorted_cities[idx]
                for p in self.all_participants:
                    if p["city"] == city:
                        names_to_select.append(p["name"])
            
            self.listbox.selection_clear(0, END)
            selected_count = 0
            for i in range(self.listbox.size()):
                item_text = self.listbox.get(i)
                for participant_name in names_to_select:
                    if item_text.startswith(participant_name):
                        self.listbox.selection_set(i)
                        selected_count += 1
                        break
            
            total_count = self.listbox.size()
            self.count_label.config(text=f"Всего участников: {total_count} | Выбрано: {selected_count}")
            self.status_label.config(text=f"✅ Выбрано {selected_count} участников из {len(selected_indices)} городов", fg="blue")
            city_window.destroy()
        
        Button(city_window, text="✅ Применить", command=apply_selection, 
               bg="#4CAF50", fg="white", width=15).pack(pady=10)
        Button(city_window, text="❌ Отмена", command=city_window.destroy,
               bg="#f44336", fg="white", width=15).pack(pady=5)
    
    def search_participants(self, event=None):
        search_text = self.search_entry.get().strip().lower()
        self.listbox.delete(0, END)
        
        if not search_text:
            self.update_listbox()
            return
        
        for p in self.all_participants:
            if search_text in p["name"].lower() or search_text in p["city"].lower():
                city_short = p["city"][:30] if p["city"] else ""
                display_text = f"{p['name']} ({city_short})" if city_short else p['name']
                self.listbox.insert(END, display_text)
    
    def clear_search(self):
        self.search_entry.delete(0, END)
        self.update_listbox()
    
    def replace_text_in_run(self, run, search_text, replace_text):
        if search_text in run.text:
            run.text = run.text.replace(search_text, replace_text)
            return True
        return False
    
    def replace_text_in_paragraph(self, paragraph, search_text, replace_text):
        if not replace_text or search_text not in paragraph.text:
            return False
        
        replaced = False
        for run in paragraph.runs:
            if self.replace_text_in_run(run, search_text, replace_text):
                replaced = True
        
        if not replaced and search_text in paragraph.text:
            paragraph.text = paragraph.text.replace(search_text, replace_text)
            replaced = True
        
        return replaced
    
    def replace_text_in_document(self, doc, search_text, replace_text):
        if not replace_text:
            return False
        
        replaced = False
        
        for paragraph in doc.paragraphs:
            if self.replace_text_in_paragraph(paragraph, search_text, replace_text):
                replaced = True
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if self.replace_text_in_paragraph(paragraph, search_text, replace_text):
                            replaced = True
        
        return replaced
    
    def generate_certificates(self):
        if not self.template_file:
            messagebox.showerror("Ошибка", "Сначала загрузите шаблон грамоты")
            return
        
        if not self.excel_file:
            messagebox.showerror("Ошибка", "Сначала загрузите Excel файл")
            return
        
        selected_indices = self.listbox.curselection()
        if not selected_indices:
            messagebox.showinfo("Информация", "Выберите хотя бы одного участника")
            return
        
        selected_names = []
        for idx in selected_indices:
            item_text = self.listbox.get(idx)
            participant_name = item_text.split(" (")[0] if " (" in item_text else item_text
            selected_names.append(participant_name)
        
        participants_data = []
        for participant_name in selected_names:
            for p in self.all_participants:
                if p["name"] == participant_name:
                    participants_data.append(p)
                    break
        
        if not participants_data:
            messagebox.showerror("Ошибка", "Не удалось найти данные для выбранных участников")
            return
        
        try:
            output_folder = filedialog.askdirectory(title="Выберите папку для сохранения грамот")
            if not output_folder:
                return
            
            certificates_folder = os.path.join(output_folder, "Грамоты")
            os.makedirs(certificates_folder, exist_ok=True)
            
            successful = 0
            failed = 0
            
            for participant in participants_data:
                try:
                    doc = Document(self.template_file)
                    
                    name_placeholders = [
                        "Аликина Вероника",
                        "Астахова Екатерина",
                        "________",
                        "{имя}",
                        "{Имя}",
                        "[Имя]",
                        "<Имя>",
                        "ИМЯ"
                    ]
                    
                    name_replaced = False
                    for placeholder in name_placeholders:
                        if self.replace_text_in_document(doc, placeholder, participant["name"]):
                            name_replaced = True
                            break
                    
                    if not name_replaced:
                        for para in doc.paragraphs:
                            if para.text.strip() and len(para.text.strip()) > 3:
                                para.text = participant["name"]
                                name_replaced = True
                                break
                    
                    if participant["city"]:
                        city_placeholders = ["{город}", "{Город}", "[Город]", "<Город>", "ГОРОД"]
                        for placeholder in city_placeholders:
                            self.replace_text_in_document(doc, placeholder, participant["city"])
                    
                    if participant["phone"]:
                        phone_placeholders = ["{телефон}", "{Телефон}", "[Телефон]", "<Телефон>", "ТЕЛЕФОН"]
                        for placeholder in phone_placeholders:
                            self.replace_text_in_document(doc, placeholder, participant["phone"])
                    
                    if participant["parent"]:
                        parent_placeholders = ["{родитель}", "{Родитель}", "[Родитель]", "<Родитель>", "РОДИТЕЛЬ"]
                        for placeholder in parent_placeholders:
                            self.replace_text_in_document(doc, placeholder, participant["parent"])
                    
                    if participant["age"]:
                        age_placeholders = ["{возраст}", "{Возраст}", "[Возраст]", "<Возраст>", "ВОЗРАСТ"]
                        for placeholder in age_placeholders:
                            self.replace_text_in_document(doc, placeholder, participant["age"])
                    
                    safe_name = re.sub(r'[<>:"/\\|?*]', '', participant["name"])
                    filename = f"Грамота_{safe_name}.docx"
                    filepath = os.path.join(certificates_folder, filename)
                    
                    counter = 1
                    while os.path.exists(filepath):
                        filename = f"Грамота_{safe_name}_{counter}.docx"
                        filepath = os.path.join(certificates_folder, filename)
                        counter += 1
                    
                    doc.save(filepath)
                    successful += 1
                    
                except Exception as e:
                    failed += 1
                    print(f"Ошибка для {participant['name']}: {str(e)}")
            
            result_msg = f"✅ Грамоты созданы!\n\nУспешно: {successful}\nОшибок: {failed}\n\nПапка: {certificates_folder}"
            messagebox.showinfo("Готово", result_msg)
            self.status_label.config(text=f"✅ Создано {successful} грамот", fg="green")
            
            try:
                os.startfile(certificates_folder)
            except:
                pass
            
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при генерации грамот:\n{str(e)}")
            self.status_label.config(text="❌ Ошибка генерации", fg="red")

if __name__ == "__main__":
    root = Tk()
    app = CertificateApp(root)
    root.mainloop()