"""
ONVIF 安防摄像头检测工具 v2.7（优化报告大小 + 汇总页）
- 单设备/批量检测，可停止、清空
- 截图采用帧预热丢弃技术，避免马赛克
- 智能选取 H.264 最高分辨率主码流
- JPEG 质量 85，文件更小且足够清晰
- 批量结果严格按导入顺序排列
- Word 报告：封面 → 汇总表（所有IP，红色状态） → 每设备详细页（仅成功）
"""

import sys
import os
import csv
import time
import concurrent.futures
import threading
from datetime import datetime
from urllib.parse import urlparse

from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QLineEdit, QPushButton, QGroupBox, QFormLayout,
    QTableWidget, QTableWidgetItem, QHeaderView, QMessageBox,
    QProgressBar, QFileDialog, QComboBox, QCheckBox,
    QAbstractItemView, QTextEdit, QDialog, QProgressDialog
)
from PyQt5.QtCore import Qt, QThread, pyqtSignal, QObject

import cv2
from onvif import ONVIFCamera
from zeep.exceptions import Fault, TransportError

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# -------------------- 工具：选取最佳主码流 --------------------
def select_best_profile(profiles):
    h264 = [p for p in profiles if p.VideoEncoderConfiguration.Encoding == 'H264']
    if not h264:
        h264 = profiles
    if not h264:
        return None
    def area(p):
        r = p.VideoEncoderConfiguration.Resolution
        return r.Width * r.Height
    return max(h264, key=area)


# -------------------- 单设备检测 --------------------
class DetectWorker(QObject):
    finished = pyqtSignal(dict)
    error = pyqtSignal(str)
    progress = pyqtSignal(str)

    def __init__(self, ip, port, username, password, brand='', stop_event=None):
        super().__init__()
        self.ip = ip
        self.port = port
        self.username = username
        self.password = password
        self.brand = brand
        self.stop_event = stop_event

    def run(self):
        result = {
            'ip': self.ip, 'port': self.port, 'username': self.username,
            'password': self.password, 'brand': self.brand,
            'model': '', 'firmware': '', 'resolution': '',
            'screenshot': '', 'time': '', 'success': False
        }
        try:
            if self.stop_event and self.stop_event.is_set():
                raise InterruptedError("用户取消")
            self.progress.emit(f"正在连接 {self.ip}:{self.port} ...")
            cam = ONVIFCamera(self.ip, self.port, self.username, self.password)

            if self.stop_event and self.stop_event.is_set():
                raise InterruptedError("用户取消")

            dev_info = cam.devicemgmt.GetDeviceInformation()
            result['model'] = dev_info.Model.strip() if dev_info.Model else "未知"
            result['firmware'] = dev_info.FirmwareVersion.strip() if dev_info.FirmwareVersion else "未知"
            if not self.brand:
                result['brand'] = dev_info.Manufacturer.strip() if dev_info.Manufacturer else "未识别"

            if self.stop_event and self.stop_event.is_set():
                raise InterruptedError("用户取消")

            media_service = cam.create_media_service()
            profiles = media_service.GetProfiles()
            if not profiles:
                raise Exception("未找到媒体配置文件")
            profile = select_best_profile(profiles)
            if not profile:
                raise Exception("未找到可用的视频编码配置")
            video_enc = profile.VideoEncoderConfiguration
            result['resolution'] = f"{video_enc.Resolution.Width}×{video_enc.Resolution.Height}"

            if self.stop_event and self.stop_event.is_set():
                raise InterruptedError("用户取消")

            stream_uri = media_service.GetStreamUri({
                'StreamSetup': {
                    'Stream': 'RTP-Unicast',
                    'Transport': {'Protocol': 'RTSP'}
                },
                'ProfileToken': profile.token
            })
            parsed = urlparse(stream_uri.Uri)
            auth_rtsp = f"rtsp://{self.username}:{self.password}@{parsed.hostname}:{parsed.port}{parsed.path}"

            self.progress.emit("正在获取截图...")
            screenshot_path = self.capture_frame(auth_rtsp, self.ip, self.stop_event)
            result['screenshot'] = screenshot_path

            result['time'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            result['success'] = True
            self.finished.emit(result)

        except InterruptedError:
            self.error.emit("用户取消检测")
        except (Fault, TransportError) as e:
            self.error.emit(f"ONVIF连接错误: {str(e)}")
        except Exception as e:
            self.error.emit(f"检测失败: {str(e)}")

    @staticmethod
    def capture_frame(rtsp_url, ip, stop_event=None, retries=3):
        """截取一帧（含帧预热，消除马赛克），保存为 JPEG 质量 85，减小文件体积"""
        for attempt in range(retries):
            if stop_event and stop_event.is_set():
                raise InterruptedError("用户取消")
            cap = cv2.VideoCapture(rtsp_url, cv2.CAP_FFMPEG)
            cap.set(cv2.CAP_PROP_OPEN_TIMEOUT_MSEC, 10000)
            cap.set(cv2.CAP_PROP_READ_TIMEOUT_MSEC, 10000)

            ret, frame = cap.read()
            if ret and frame is not None and frame.size > 0:
                # 丢弃后续10帧，让解码器充分预热
                for _ in range(10):
                    cap.read()
                ret2, final_frame = cap.read()
                if ret2 and final_frame is not None and final_frame.size > 0:
                    screenshot_dir = "screenshots"
                    os.makedirs(screenshot_dir, exist_ok=True)
                    filename = f"{ip.replace('.', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.jpg"
                    filepath = os.path.join(screenshot_dir, filename)
                    cv2.imwrite(filepath, final_frame, [cv2.IMWRITE_JPEG_QUALITY, 85])
                    cap.release()
                    return filepath
            cap.release()
            time.sleep(1)
        raise Exception("多次尝试仍无法获取有效视频帧，请检查摄像头编码格式或RTSP地址")


# -------------------- 批量检测（结果顺序修复） --------------------
class BatchWorker(QObject):
    progress = pyqtSignal(str, int)
    finished_all = pyqtSignal(list)
    error_item = pyqtSignal(str, str)

    def __init__(self, tasks, max_workers=10, stop_event=None):
        super().__init__()
        self.tasks = tasks
        self.max_workers = max_workers
        self.stop_event = stop_event
        self._executor = None

    def run(self):
        total = len(self.tasks)
        results_list = [None] * total
        completed = 0
        lock = threading.Lock()

        if self.stop_event and self.stop_event.is_set():
            self.finished_all.emit([])
            return

        self._executor = concurrent.futures.ThreadPoolExecutor(max_workers=self.max_workers)
        try:
            future_to_idx = {}
            for idx, task in enumerate(self.tasks):
                future = self._executor.submit(self._detect_single, task)
                future_to_idx[future] = idx

            for future in concurrent.futures.as_completed(future_to_idx):
                if self.stop_event and self.stop_event.is_set():
                    for f in future_to_idx:
                        f.cancel()
                    self._executor.shutdown(wait=False, cancel_futures=True)
                    break
                idx = future_to_idx[future]
                task = self.tasks[idx]
                try:
                    result = future.result()
                except Exception as e:
                    result = {
                        'ip': task['ip'], 'port': task['port'],
                        'username': task['username'], 'password': task['password'],
                        'brand': task.get('brand', ''),
                        'model': '', 'firmware': '', 'resolution': '',
                        'screenshot': '', 'time': '', 'success': False
                    }
                    self.error_item.emit(task['ip'], str(e))
                results_list[idx] = result
                with lock:
                    completed += 1
                self.progress.emit(f"完成 {task['ip']} ({completed}/{total})", completed)
        finally:
            if self._executor:
                self._executor.shutdown(wait=False)
                self._executor = None

        ordered_results = [r for r in results_list if r is not None]
        self.finished_all.emit(ordered_results)

    def stop_executor(self):
        if self._executor:
            self._executor.shutdown(wait=False, cancel_futures=True)

    def _detect_single(self, task):
        if self.stop_event and self.stop_event.is_set():
            raise InterruptedError("用户取消")
        result = {
            'ip': task['ip'], 'port': task['port'],
            'username': task['username'], 'password': task['password'],
            'brand': task.get('brand', ''),
            'model': '', 'firmware': '', 'resolution': '',
            'screenshot': '', 'time': '', 'success': False
        }
        try:
            cam = ONVIFCamera(task['ip'], task['port'], task['username'], task['password'])
            if self.stop_event and self.stop_event.is_set():
                raise InterruptedError("用户取消")

            dev_info = cam.devicemgmt.GetDeviceInformation()
            result['model'] = dev_info.Model.strip() if dev_info.Model else "未知"
            result['firmware'] = dev_info.FirmwareVersion.strip() if dev_info.FirmwareVersion else "未知"
            if not result['brand']:
                result['brand'] = dev_info.Manufacturer.strip() if dev_info.Manufacturer else "未识别"

            if self.stop_event and self.stop_event.is_set():
                raise InterruptedError("用户取消")

            media_service = cam.create_media_service()
            profiles = media_service.GetProfiles()
            if not profiles:
                raise Exception("无媒体配置文件")
            profile = select_best_profile(profiles)
            if not profile:
                raise Exception("未找到可用的视频编码配置")
            video_enc = profile.VideoEncoderConfiguration
            result['resolution'] = f"{video_enc.Resolution.Width}×{video_enc.Resolution.Height}"

            if self.stop_event and self.stop_event.is_set():
                raise InterruptedError("用户取消")

            stream_uri = media_service.GetStreamUri({
                'StreamSetup': {
                    'Stream': 'RTP-Unicast',
                    'Transport': {'Protocol': 'RTSP'}
                },
                'ProfileToken': profile.token
            })
            parsed = urlparse(stream_uri.Uri)
            auth_rtsp = f"rtsp://{task['username']}:{task['password']}@{parsed.hostname}:{parsed.port}{parsed.path}"

            screenshot_path = DetectWorker.capture_frame(auth_rtsp, task['ip'], self.stop_event)
            result['screenshot'] = screenshot_path
            result['time'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            result['success'] = True
        except InterruptedError:
            result['success'] = False
        except Exception as e:
            result['success'] = False
        return result


# -------------------- 主界面 --------------------
class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("ONVIF 安防摄像头检测工具 v2.7")
        self.setMinimumSize(980, 750)
        self.detect_results = []
        self.tasks_batch = []
        self.stop_single_event = threading.Event()
        self.stop_batch_event = threading.Event()
        self.init_ui()
        self.check_dependencies()

    def init_ui(self):
        central = QWidget()
        self.setCentralWidget(central)
        main_layout = QVBoxLayout(central)

        # ---------- 单设备检测 ----------
        single_grp = QGroupBox("单设备检测")
        form = QFormLayout()
        self.ip_edit = QLineEdit()
        self.ip_edit.setPlaceholderText("例如: 192.168.1.10")
        self.port_edit = QLineEdit("80")
        self.port_edit.setMaximumWidth(80)
        self.user_edit = QLineEdit("admin")
        self.pwd_edit = QLineEdit("admin")
        self.pwd_edit.setEchoMode(QLineEdit.Password)

        brand_layout = QHBoxLayout()
        self.brand_combo = QComboBox()
        self.brand_combo.setEditable(True)
        self.brand_combo.addItems(["", "海康威视", "大华", "水星", "宇视", "天地伟业", "其他"])
        self.auto_brand = QCheckBox("自动获取品牌")
        self.auto_brand.setChecked(True)
        brand_layout.addWidget(self.brand_combo)
        brand_layout.addWidget(self.auto_brand)
        brand_widget = QWidget()
        brand_widget.setLayout(brand_layout)

        form.addRow("IP地址:", self.ip_edit)
        form.addRow("端口:", self.port_edit)
        form.addRow("用户名:", self.user_edit)
        form.addRow("密码:", self.pwd_edit)
        form.addRow("品牌:", brand_widget)

        btn_layout = QHBoxLayout()
        self.detect_btn = QPushButton("开始检测")
        self.detect_btn.clicked.connect(self.start_single_detect)
        self.stop_single_btn = QPushButton("停止")
        self.stop_single_btn.setEnabled(False)
        self.stop_single_btn.clicked.connect(self.stop_single_detect)
        self.clear_btn = QPushButton("清空")
        self.clear_btn.clicked.connect(self.clear_single)
        btn_layout.addWidget(self.detect_btn)
        btn_layout.addWidget(self.stop_single_btn)
        btn_layout.addWidget(self.clear_btn)
        btn_layout.addStretch()

        self.single_status = QLabel("就绪")
        self.single_status.setStyleSheet("color: gray")

        single_layout = QVBoxLayout()
        single_layout.addLayout(form)
        single_layout.addLayout(btn_layout)
        single_layout.addWidget(self.single_status)
        single_grp.setLayout(single_layout)
        main_layout.addWidget(single_grp)

        # ---------- 批量检测 ----------
        batch_grp = QGroupBox("批量检测")
        tool_layout = QHBoxLayout()
        self.import_btn = QPushButton("导入CSV")
        self.import_btn.clicked.connect(self.import_csv)
        self.export_btn = QPushButton("导出结果CSV")
        self.export_btn.clicked.connect(self.export_results_csv)
        self.batch_btn = QPushButton("批量检测")
        self.batch_btn.clicked.connect(self.start_batch_detect)
        self.stop_batch_btn = QPushButton("停止")
        self.stop_batch_btn.setEnabled(False)
        self.stop_batch_btn.clicked.connect(self.stop_batch_detect)
        self.report_btn = QPushButton("生成Word报告")
        self.report_btn.clicked.connect(self.generate_report)
        self.clear_results_btn = QPushButton("清除结果")
        self.clear_results_btn.clicked.connect(self.clear_results_list)
        tool_layout.addWidget(self.import_btn)
        tool_layout.addWidget(self.export_btn)
        tool_layout.addWidget(self.batch_btn)
        tool_layout.addWidget(self.stop_batch_btn)
        tool_layout.addWidget(self.report_btn)
        tool_layout.addWidget(self.clear_results_btn)
        tool_layout.addStretch()

        self.table = QTableWidget(0, 6)
        self.table.setHorizontalHeaderLabels(["IP地址", "品牌", "型号", "固件版本", "分辨率", "检测时间"])
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.table.setSelectionBehavior(QAbstractItemView.SelectRows)

        self.progress_bar = QProgressBar()
        self.batch_status = QLabel("批量检测就绪")

        batch_layout = QVBoxLayout()
        batch_layout.addLayout(tool_layout)
        batch_layout.addWidget(self.table)
        batch_layout.addWidget(self.progress_bar)
        batch_layout.addWidget(self.batch_status)
        batch_grp.setLayout(batch_layout)
        main_layout.addWidget(batch_grp)

        # ---------- 日志 ----------
        log_grp = QGroupBox("运行日志")
        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        self.log_text.setMaximumHeight(100)
        log_layout = QVBoxLayout()
        log_layout.addWidget(self.log_text)
        log_grp.setLayout(log_layout)
        main_layout.addWidget(log_grp)

    def check_dependencies(self):
        missing = []
        try:
            import onvif
        except ImportError:
            missing.append("onvif-zeep")
        try:
            import cv2
        except ImportError:
            missing.append("opencv-python")
        try:
            from docx import Document
        except ImportError:
            missing.append("python-docx")
        if missing:
            QMessageBox.warning(self, "缺少依赖",
                f"请先安装以下Python库:\n{', '.join(missing)}\n\n"
                "命令: pip install onvif-zeep opencv-python python-docx")

    # ---- 单设备操作 ----
    def start_single_detect(self):
        ip = self.ip_edit.text().strip()
        if not ip:
            QMessageBox.warning(self, "输入错误", "请输入IP地址")
            return
        try:
            port = int(self.port_edit.text().strip() or 80)
        except ValueError:
            QMessageBox.warning(self, "输入错误", "端口必须为数字")
            return
        username = self.user_edit.text().strip()
        password = self.pwd_edit.text().strip()
        brand = self.brand_combo.currentText().strip() if not self.auto_brand.isChecked() else ""

        self.stop_single_event.clear()
        self.detect_btn.setEnabled(False)
        self.stop_single_btn.setEnabled(True)
        self.single_status.setText("正在检测...")
        self.log(f"开始单设备检测: {ip}:{port}")

        self.worker = DetectWorker(ip, port, username, password, brand, self.stop_single_event)
        self.thread = QThread()
        self.worker.moveToThread(self.thread)
        self.worker.finished.connect(self.on_single_finished)
        self.worker.error.connect(self.on_single_error)
        self.worker.progress.connect(lambda msg: self.single_status.setText(msg))
        self.thread.started.connect(self.worker.run)
        self.thread.start()

    def stop_single_detect(self):
        self.stop_single_event.set()
        self.log("正在尝试停止单设备检测...")

    def on_single_finished(self, result):
        self.detect_btn.setEnabled(True)
        self.stop_single_btn.setEnabled(False)
        self.single_status.setText("检测完成")
        self.thread.quit()
        self.thread.wait()
        self.add_result_to_table(result)
        self.log(f"单设备检测成功: {result['ip']}")
        QMessageBox.information(self, "检测完成", f"设备 {result['ip']} 检测成功！")

    def on_single_error(self, error_msg):
        self.detect_btn.setEnabled(True)
        self.stop_single_btn.setEnabled(False)
        self.single_status.setText("已停止" if "取消" in error_msg else "检测失败")
        self.thread.quit()
        self.thread.wait()
        self.log(f"单设备检测: {error_msg}")
        if "取消" not in error_msg:
            QMessageBox.critical(self, "检测失败", error_msg)

    def clear_single(self):
        self.ip_edit.clear()
        self.port_edit.setText("80")
        self.user_edit.setText("admin")
        self.pwd_edit.setText("admin")
        self.brand_combo.setCurrentText("")
        self.auto_brand.setChecked(True)
        self.single_status.setText("就绪")

    # ---- 批量操作 ----
    def import_csv(self):
        filepath, _ = QFileDialog.getOpenFileName(self, "导入CSV", "", "CSV Files (*.csv)")
        if not filepath:
            return
        try:
            encodings = ['utf-8-sig', 'gbk', 'utf-8']
            data_rows = []
            used_encoding = None
            for enc in encodings:
                try:
                    with open(filepath, 'r', encoding=enc) as f:
                        reader = csv.DictReader(f)
                        if 'ip' not in (reader.fieldnames or []):
                            raise Exception("CSV文件必须包含 'ip' 列")
                        for row in reader:
                            data_rows.append(row)
                        used_encoding = enc
                        break
                except UnicodeDecodeError:
                    continue
            if used_encoding is None:
                raise Exception("无法识别文件编码，请将文件保存为 UTF-8 或 GBK 编码")

            self.tasks_batch = []
            for row in data_rows:
                ip = row.get('ip', '').strip()
                if not ip:
                    continue
                port_str = row.get('port', '80').strip()
                port = int(port_str) if port_str.isdigit() else 80
                username = row.get('username', 'admin').strip()
                password = row.get('password', 'admin').strip()
                brand = row.get('brand', '').strip()
                self.tasks_batch.append({
                    'ip': ip,
                    'port': port,
                    'username': username,
                    'password': password,
                    'brand': brand
                })
            if not self.tasks_batch:
                raise Exception("未找到有效的IP数据")

            self.table.setRowCount(0)
            for t in self.tasks_batch:
                r = self.table.rowCount()
                self.table.insertRow(r)
                self.table.setItem(r, 0, QTableWidgetItem(t['ip']))
                self.table.setItem(r, 1, QTableWidgetItem(t['brand']))
            self.batch_status.setText(f"已导入 {len(self.tasks_batch)} 条任务 (编码: {used_encoding})")
            self.log(f"导入 {len(self.tasks_batch)} 条任务 (编码: {used_encoding})")
        except Exception as e:
            QMessageBox.critical(self, "导入错误", str(e))

    def export_results_csv(self):
        if not self.detect_results:
            QMessageBox.information(self, "提示", "没有可导出的结果")
            return
        filepath, _ = QFileDialog.getSaveFileName(self, "导出结果CSV(可再次导入)", "", "CSV Files (*.csv)")
        if not filepath:
            return
        try:
            with open(filepath, 'w', newline='', encoding='utf-8-sig') as f:
                writer = csv.writer(f)
                writer.writerow(['ip', 'port', 'username', 'password', 'brand',
                                 'model', 'firmware', 'resolution', 'screenshot', 'time', 'success'])
                for r in self.detect_results:
                    writer.writerow([
                        r['ip'], r['port'], r['username'], r['password'], r['brand'],
                        r['model'], r['firmware'], r['resolution'], r['screenshot'], r['time'], r['success']
                    ])
            QMessageBox.information(self, "导出成功", f"结果已保存至 {filepath}\n包含连接参数，可直接用于下次导入。")
        except Exception as e:
            QMessageBox.critical(self, "导出错误", str(e))

    def start_batch_detect(self):
        if not self.tasks_batch:
            QMessageBox.warning(self, "无任务", "请先导入CSV批量任务")
            return
        self.stop_batch_event.clear()
        self.batch_btn.setEnabled(False)
        self.stop_batch_btn.setEnabled(True)
        self.import_btn.setEnabled(False)
        self.report_btn.setEnabled(False)
        self.table.setRowCount(0)
        self.progress_bar.setMaximum(len(self.tasks_batch))
        self.progress_bar.setValue(0)
        self.batch_status.setText("批量检测进行中...")
        self.log("开始批量检测...")

        self.batch_worker = BatchWorker(self.tasks_batch, stop_event=self.stop_batch_event)
        self.batch_thread = QThread()
        self.batch_worker.moveToThread(self.batch_thread)
        self.batch_worker.progress.connect(self.update_batch_progress)
        self.batch_worker.error_item.connect(lambda ip, e: self.log(f"错误 [{ip}]: {e}"))
        self.batch_worker.finished_all.connect(self.on_batch_finished)
        self.batch_thread.started.connect(self.batch_worker.run)
        self.batch_thread.start()

    def stop_batch_detect(self):
        self.stop_batch_event.set()
        if hasattr(self, 'batch_worker'):
            self.batch_worker.stop_executor()
        self.log("正在尝试停止批量检测...")

    def update_batch_progress(self, msg, current):
        self.progress_bar.setValue(current)
        self.batch_status.setText(msg)
        self.log(msg)

    def on_batch_finished(self, results):
        self.batch_btn.setEnabled(True)
        self.stop_batch_btn.setEnabled(False)
        self.import_btn.setEnabled(True)
        self.report_btn.setEnabled(True)
        success_count = sum(1 for r in results if r['success'])
        self.batch_status.setText(f"批量检测结束，成功 {success_count}/{len(results)}")
        self.progress_bar.setValue(len(results))
        self.log("批量检测结束")
        self.detect_results = results
        self.table.setRowCount(0)
        for r in results:
            row = self.table.rowCount()
            self.table.insertRow(row)
            self.table.setItem(row, 0, QTableWidgetItem(r['ip']))
            self.table.setItem(row, 1, QTableWidgetItem(r['brand']))
            self.table.setItem(row, 2, QTableWidgetItem(r['model']))
            self.table.setItem(row, 3, QTableWidgetItem(r['firmware']))
            self.table.setItem(row, 4, QTableWidgetItem(r['resolution']))
            self.table.setItem(row, 5, QTableWidgetItem(r['time']))
            if not r['success']:
                for col in range(6):
                    self.table.item(row, col).setBackground(Qt.red)
        self.batch_thread.quit()
        self.batch_thread.wait()

    def clear_results_list(self):
        self.table.setRowCount(0)
        self.detect_results.clear()
        self.batch_status.setText("结果已清空")
        self.progress_bar.setValue(0)
        self.log("已清除所有检测结果")

    def add_result_to_table(self, result):
        row = self.table.rowCount()
        self.table.insertRow(row)
        self.table.setItem(row, 0, QTableWidgetItem(result['ip']))
        self.table.setItem(row, 1, QTableWidgetItem(result['brand']))
        self.table.setItem(row, 2, QTableWidgetItem(result['model']))
        self.table.setItem(row, 3, QTableWidgetItem(result['firmware']))
        self.table.setItem(row, 4, QTableWidgetItem(result['resolution']))
        self.table.setItem(row, 5, QTableWidgetItem(result['time']))
        self.detect_results.append(result)

    # ---- Word 报告（新增汇总页） ----
    def generate_report(self):
        if not self.detect_results:
            QMessageBox.warning(self, "无数据", "请先执行检测")
            return
        # 所有结果（包括失败）用于汇总
        all_results = self.detect_results
        success = [r for r in all_results if r['success']]

        info_dialog = QDialog(self)
        info_dialog.setWindowTitle("报告封面信息")
        info_dialog.setMinimumWidth(350)
        layout = QFormLayout(info_dialog)
        unit_edit = QLineEdit("安防检测中心")
        person_edit = QLineEdit("管理员")
        layout.addRow("检测单位:", unit_edit)
        layout.addRow("检测人员:", person_edit)
        btn_box = QHBoxLayout()
        ok_btn = QPushButton("确定")
        cancel_btn = QPushButton("取消")
        btn_box.addWidget(ok_btn)
        btn_box.addWidget(cancel_btn)
        layout.addRow(btn_box)
        ok_btn.clicked.connect(info_dialog.accept)
        cancel_btn.clicked.connect(info_dialog.reject)
        if info_dialog.exec_() != QDialog.Accepted:
            return
        unit = unit_edit.text().strip() or "安防检测中心"
        person = person_edit.text().strip() or "管理员"

        filepath, _ = QFileDialog.getSaveFileName(self, "保存Word报告",
                                f"摄像头检测报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                                "Word Files (*.docx)")
        if not filepath:
            return

        total_steps = len(success) + 2  # 封面+汇总+每个详细设备
        progress = QProgressDialog("正在生成报告...", "取消", 0, total_steps, self)
        progress.setWindowTitle("生成报告")
        progress.setWindowModality(Qt.WindowModal)
        progress.setMinimumDuration(0)
        progress.setValue(0)
        progress.show()
        QApplication.processEvents()

        try:
            self.save_word_report(filepath, all_results, success, unit, person, progress)
            progress.setValue(total_steps)
            QMessageBox.information(self, "报告生成成功", f"报告已保存至:\n{filepath}")
            self.log(f"报告已生成: {filepath}")
        except Exception as e:
            QMessageBox.critical(self, "生成报告失败", str(e))
            if os.path.exists(filepath):
                os.remove(filepath)
        finally:
            progress.close()

    def save_word_report(self, filepath, all_results, success_results, unit, person, progress):
        doc = Document()
        section = doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # ===== 封面 =====
        for _ in range(6):
            doc.add_paragraph()
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('摄像头检测报告')
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        line = doc.add_paragraph()
        line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = line.add_run('━' * 20)
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 102, 204)

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(
            f'检测单位：{unit}\n检测人员：{person}\n生成日期：{datetime.now().strftime("%Y年%m月%d日")}'
        )
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(80, 80, 80)
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        doc.add_page_break()
        progress.setValue(1)  # 封面完成

        # ===== 第二页：汇总表 =====
        if progress.wasCanceled():
            raise Exception("用户取消生成")
        progress.setLabelText("正在生成汇总表...")
        QApplication.processEvents()

        h2 = doc.add_heading('检测汇总', level=1)
        for run in h2.runs:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0, 70, 127)
        p = doc.add_paragraph()
        run = p.add_run(f"本次共导入 {len(all_results)} 个IP地址，检测结果如下：")
        run.font.size = Pt(12)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 创建汇总表格
        table = doc.add_table(rows=1 + len(all_results), cols=4, style='Table Grid')
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # 表头
        headers = ['IP地址', '品牌', '型号', '状态']
        for i, header_text in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(header_text)
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'D6E4F0')
            shading.set(qn('w:val'), 'clear')
            cell._tc.get_or_add_tcPr().append(shading)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 数据行
        for i, dev in enumerate(all_results):
            row_cells = [
                dev['ip'],
                dev['brand'],
                dev['model'],
                '在线' if dev['success'] else '离线'
            ]
            for j, value in enumerate(row_cells):
                cell = table.cell(i + 1, j)
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(str(value))
                run.font.size = Pt(11)
                run.font.name = '宋体'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                if j == 3:  # 状态列红色
                    run.font.color.rgb = RGBColor(255, 0, 0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()
        progress.setValue(2)  # 汇总完成

        # ===== 详细设备页（仅成功） =====
        total_success = len(success_results)
        for idx, dev in enumerate(success_results, 1):
            if progress.wasCanceled():
                raise Exception("用户取消生成")
            progress.setValue(2 + idx)
            progress.setLabelText(f"正在生成第 {idx}/{total_success} 台设备报告...")
            QApplication.processEvents()

            h2 = doc.add_heading(f'{idx}. 设备 {dev["ip"]}', level=1)
            for run in h2.runs:
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(0, 70, 127)
                run.font.name = '微软雅黑'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

            table = doc.add_table(rows=6, cols=2, style='Table Grid')
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            tbl = table._tbl
            tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
            borders = OxmlElement('w:tblBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '8')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '2F5496')
                borders.append(border)
            tblPr.append(borders)

            data = [
                ('IP 地址', dev['ip']),
                ('品牌', dev['brand']),
                ('型号', dev['model']),
                ('固件版本', dev['firmware']),
                ('分辨率', dev['resolution']),
                ('检测时间', dev['time']),
            ]
            for i, (label, value) in enumerate(data):
                cell_label = table.cell(i, 0)
                cell_value = table.cell(i, 1)
                cell_label.text = ''
                cell_value.text = ''
                p_label = cell_label.paragraphs[0]
                p_value = cell_value.paragraphs[0]
                run_label = p_label.add_run(label)
                run_label.font.size = Pt(11)
                run_label.bold = True
                run_label.font.name = '宋体'
                run_label.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run_value = p_value.add_run(value)
                run_value.font.size = Pt(11)
                run_value.font.name = '宋体'
                run_value.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'D6E4F0')
                shading.set(qn('w:val'), 'clear')
                cell_label._tc.get_or_add_tcPr().append(shading)
                p_label.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p_value.alignment = WD_ALIGN_PARAGRAPH.LEFT
                cell_label.width = Cm(3.8)
                cell_value.width = Cm(10.2)

            doc.add_paragraph()
            if dev['screenshot'] and os.path.exists(dev['screenshot']):
                cap_title = doc.add_paragraph()
                cap_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cap_title.add_run('设备实时截图')
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(80, 80, 80)
                run.italic = True
                img_table = doc.add_table(rows=1, cols=1)
                img_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                cell = img_table.cell(0, 0)
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '6')
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), 'C0C0C0')
                    tcBorders.append(border)
                tcPr.append(tcBorders)
                p_img = cell.paragraphs[0]
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    run_img = p_img.add_run()
                    run_img.add_picture(dev['screenshot'], width=Inches(5.6))
                except Exception as e:
                    p_img.add_run(f"[截图插入失败: {e}]")
            else:
                doc.add_paragraph("（无截图）")

            if idx < total_success:
                doc.add_page_break()

        # 页眉页脚
        header = doc.sections[0].header
        header.is_linked_to_previous = False
        p_header = header.paragraphs[0]
        p_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p_header.add_run('ONVIF 摄像头检测报告')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)

        footer = doc.sections[0].footer
        footer.is_linked_to_previous = False
        p_footer = footer.paragraphs[0]
        p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_footer.add_run('第 ')
        run.font.size = Pt(9)
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = ' PAGE '
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run2 = p_footer.add_run(' 页')
        run2.font.size = Pt(9)

        doc.save(filepath)

    def log(self, msg):
        self.log_text.append(msg)

    def closeEvent(self, event):
        self.stop_single_event.set()
        self.stop_batch_event.set()
        if hasattr(self, 'batch_worker'):
            self.batch_worker.stop_executor()
        if hasattr(self, 'thread') and self.thread.isRunning():
            self.thread.quit()
            self.thread.wait(3000)
        if hasattr(self, 'batch_thread') and self.batch_thread.isRunning():
            self.batch_thread.quit()
            self.batch_thread.wait(3000)
        event.accept()


if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec_())
