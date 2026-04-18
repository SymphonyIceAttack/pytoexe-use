import tkinter as tk
from tkinter import ttk, messagebox
import json
import random
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

class CharacterCreator:
    def __init__(self, root):
        self.root = root
        self.root.title("Создание листа персонажа D&D 5e")
        self.root.geometry("1150x950")
        self.root.resizable(True, True)
        self.root.configure(bg="#1a1a2e")
        
        self.setup_styles()
        
        self.max_points = 27
        
        self.str_var = tk.IntVar(value=10)
        self.dex_var = tk.IntVar(value=10)
        self.con_var = tk.IntVar(value=10)
        self.int_var = tk.IntVar(value=10)
        self.wis_var = tk.IntVar(value=10)
        self.cha_var = tk.IntVar(value=10)
        
        self.abilities = {
            "Сила": self.str_var,
            "Ловкость": self.dex_var,
            "Телосложение": self.con_var,
            "Интеллект": self.int_var,
            "Мудрость": self.wis_var,
            "Харизма": self.cha_var
        }
        
        self.base_abilities = {
            "Сила": 10, "Ловкость": 10, "Телосложение": 10,
            "Интеллект": 10, "Мудрость": 10, "Харизма": 10
        }
        
        self.race_bonuses = {
            "Человек": {"Сила": 1, "Ловкость": 1, "Телосложение": 1, "Интеллект": 1, "Мудрость": 1, "Харизма": 1},
            "Эльф (Высший)": {"Ловкость": 2, "Интеллект": 1},
            "Эльф (Лесной)": {"Ловкость": 2, "Мудрость": 1},
            "Эльф (Темный - Дроу)": {"Ловкость": 2, "Харизма": 1},
            "Дварф (Холмовой)": {"Телосложение": 2, "Мудрость": 1},
            "Дварф (Горный)": {"Телосложение": 2, "Сила": 2},
            "Полурослик (Легконогий)": {"Ловкость": 2, "Харизма": 1},
            "Полурослик (Крепыш)": {"Ловкость": 2, "Телосложение": 1},
            "Драконорожденный": {"Сила": 2, "Харизма": 1},
            "Гном (Лесной)": {"Интеллект": 2, "Ловкость": 1},
            "Гном (Скальный)": {"Интеллект": 2, "Телосложение": 1},
            "Полуэльф": {"Харизма": 2},
            "Полуорк": {"Сила": 2, "Телосложение": 1},
            "Тифлинг": {"Интеллект": 1, "Харизма": 2}
        }
        
        self.class_requirements = {
            "Варвар": {"Сила": 13},
            "Бард": {"Харизма": 13},
            "Жрец": {"Мудрость": 13},
            "Друид": {"Мудрость": 13},
            "Воин": {"Сила": 13, "Ловкость": 13},
            "Монах": {"Ловкость": 13, "Мудрость": 13},
            "Паладин": {"Сила": 13, "Харизма": 13},
            "Следопыт": {"Ловкость": 13, "Мудрость": 13},
            "Плут": {"Ловкость": 13},
            "Чародей": {"Харизма": 13},
            "Колдун": {"Харизма": 13},
            "Волшебник": {"Интеллект": 13}
        }
        
        self.class_primary_ability = {
            "Варвар": "Сила", "Бард": "Харизма", "Жрец": "Мудрость",
            "Друид": "Мудрость", "Воин": "Сила", "Монах": "Ловкость",
            "Паладин": "Сила", "Следопыт": "Ловкость", "Плут": "Ловкость",
            "Чародей": "Харизма", "Колдун": "Харизма", "Волшебник": "Интеллект"
        }
        
        self.class_saving_throws = {
            "Варвар": ["Сила", "Телосложение"],
            "Бард": ["Ловкость", "Харизма"],
            "Жрец": ["Мудрость", "Харизма"],
            "Друид": ["Интеллект", "Мудрость"],
            "Воин": ["Сила", "Телосложение"],
            "Монах": ["Сила", "Ловкость"],
            "Паладин": ["Мудрость", "Харизма"],
            "Следопыт": ["Сила", "Ловкость"],
            "Плут": ["Ловкость", "Интеллект"],
            "Чародей": ["Телосложение", "Харизма"],
            "Колдун": ["Мудрость", "Харизма"],
            "Волшебник": ["Интеллект", "Мудрость"]
        }
        
        self.skills_by_ability = {
            "Сила": ["Атлетика"],
            "Ловкость": ["Акробатика", "Ловкость рук", "Скрытность"],
            "Телосложение": [],
            "Интеллект": ["Анализ", "История", "Магия", "Природа", "Религия"],
            "Мудрость": ["Внимательность", "Выживание", "Медицина", "Проницательность", "Уход за животными"],
            "Харизма": ["Выступление", "Запугивание", "Обман", "Убеждение"]
        }
        
        self.all_skills = []
        for skills in self.skills_by_ability.values():
            self.all_skills.extend(skills)
        
        self.class_skill_choices = {
            "Варвар": ["Атлетика", "Внимательность", "Выживание", "Запугивание", "Природа", "Уход за животными"],
            "Бард": ["Акробатика", "Анализ", "Атлетика", "Внимательность", "Выживание", "Выступление", 
                     "Запугивание", "История", "Ловкость рук", "Магия", "Медицина", "Обман", "Природа", 
                     "Проницательность", "Религия", "Скрытность", "Убеждение", "Уход за животными"],
            "Жрец": ["История", "Медицина", "Проницательность", "Религия", "Убеждение"],
            "Друид": ["Магия", "Внимательность", "Выживание", "Медицина", "Природа", "Проницательность", 
                      "Религия", "Уход за животными"],
            "Воин": ["Акробатика", "Анализ", "Атлетика", "Внимательность", "Выживание", "Запугивание", 
                     "История", "Проницательность", "Уход за животными"],
            "Монах": ["Акробатика", "Атлетика", "История", "Проницательность", "Религия", "Скрытность"],
            "Паладин": ["Атлетика", "Запугивание", "Медицина", "Проницательность", "Религия", "Убеждение"],
            "Следопыт": ["Атлетика", "Внимательность", "Выживание", "Природа", "Проницательность", 
                         "Скрытность", "Уход за животными"],
            "Плут": ["Акробатика", "Анализ", "Атлетика", "Внимательность", "Выступление", "Запугивание", 
                     "История", "Ловкость рук", "Обман", "Проницательность", "Скрытность", "Убеждение"],
            "Чародей": ["Магия", "Запугивание", "Обман", "Проницательность", "Религия", "Убеждение"],
            "Колдун": ["Магия", "Запугивание", "История", "Обман", "Природа", "Религия", "Убеждение"],
            "Волшебник": ["Анализ", "История", "Магия", "Медицина", "Религия"]
        }
        
        self.class_skill_count = {
            "Варвар": 2, "Бард": 3, "Жрец": 2, "Друид": 2, "Воин": 2,
            "Монах": 2, "Паладин": 2, "Следопыт": 3, "Плут": 4,
            "Чародей": 2, "Колдун": 2, "Волшебник": 2
        }
        
        self.backgrounds = {
            "Благородный": {
                "skills": ["История", "Убеждение"],
                "tools": ["Набор для игры в карты"],
                "languages": 1,
                "trait": "Привилегированное положение",
                "equipment": "Парадная одежда, перстень с печаткой, кошель с 25 зм"
            },
            "Моряк": {
                "skills": ["Атлетика", "Внимательность"],
                "tools": ["Инструменты навигатора", "Водный транспорт"],
                "languages": 0,
                "trait": "Крепкие морские ноги",
                "equipment": "Веревка (50 футов), шёлковый платок, кошель с 10 зм"
            },
            "Мудрец": {
                "skills": ["Магия", "История"],
                "tools": [],
                "languages": 2,
                "trait": "Исследователь",
                "equipment": "Чернила, перо, книга, кошель с 10 зм"
            },
            "Народный герой": {
                "skills": ["Уход за животными", "Выживание"],
                "tools": ["Инструменты ремесленника"],
                "languages": 0,
                "trait": "Гостеприимство простолюдинов",
                "equipment": "Лопата, железный котелок, кошель с 10 зм"
            },
            "Отшельник": {
                "skills": ["Медицина", "Религия"],
                "tools": ["Набор травника"],
                "languages": 1,
                "trait": "Открытие",
                "equipment": "Спальник, одеяло, набор трав, кошель с 5 зм"
            },
            "Преступник": {
                "skills": ["Скрытность", "Обман"],
                "tools": ["Воровские инструменты", "Набор для фальсификации"],
                "languages": 0,
                "trait": "Криминальные связи",
                "equipment": "Ломик, тёмная одежда с капюшоном, кошель с 15 зм"
            },
            "Странник": {
                "skills": ["Выступление", "Ловкость рук"],
                "tools": ["Маскировочный набор", "Музыкальный инструмент"],
                "languages": 0,
                "trait": "Фокусник",
                "equipment": "Костюм, колода карт, кошель с 15 зм"
            },
            "Чужеземец": {
                "skills": ["Атлетика", "Выживание"],
                "tools": ["Музыкальный инструмент"],
                "languages": 1,
                "trait": "Странник",
                "equipment": "Посох, карта местности, кошель с 10 зм"
            },
            "Шарлатан": {
                "skills": ["Обман", "Ловкость рук"],
                "tools": ["Набор для фальсификации"],
                "languages": 0,
                "trait": "Фальшивая личность",
                "equipment": "Поддельные документы, набор грима, кошель с 15 зм"
            },
            "Беспризорник": {
                "skills": ["Ловкость рук", "Скрытность"],
                "tools": ["Воровские инструменты", "Набор для грима"],
                "languages": 0,
                "trait": "Городские тайны",
                "equipment": "Маленький нож, карта города, кошель с 10 зм"
            },
            "Гильдейский ремесленник": {
                "skills": ["Анализ", "Убеждение"],
                "tools": ["Инструменты ремесленника"],
                "languages": 1,
                "trait": "Член гильдии",
                "equipment": "Набор инструментов, письмо от гильдии, кошель с 15 зм"
            },
            "Солдат": {
                "skills": ["Атлетика", "Запугивание"],
                "tools": ["Игровой набор", "Транспорт (сухопутный)"],
                "languages": 0,
                "trait": "Воинское звание",
                "equipment": "Знак отличия, трофей с поля боя, кошель с 10 зм"
            }
        }
        
        self.armors = {
            "Без доспеха": ["none", 10, 99],
            "Стёганый": ["light", 11, 99],
            "Кожаный": ["light", 11, 99],
            "Проклёпанный кожаный": ["light", 12, 99],
            "Шкурный": ["medium", 12, 2],
            "Кольчужная рубаха": ["medium", 13, 2],
            "Чешуйчатый": ["medium", 14, 2],
            "Кираса": ["medium", 14, 2],
            "Полулаты": ["medium", 15, 2],
            "Колечный": ["heavy", 14, 0],
            "Кольчуга": ["heavy", 16, 0],
            "Наборный": ["heavy", 17, 0],
            "Латы": ["heavy", 18, 0]
        }
        
        self.races = list(self.race_bonuses.keys())
        self.classes = list(self.class_requirements.keys())
        self.background_names = list(self.backgrounds.keys())
        self.armor_names = list(self.armors.keys())
        
        self.weapons = [
            "Боевой посох", "Булава", "Дубинка", "Кинжал", "Копье", 
            "Легкий арбалет", "Метательное копье", "Праща", "Ручной топор", "Серп",
            "Алебарда", "Боевой молот", "Боевой топор", "Двуручный меч", 
            "Длинный лук", "Длинный меч", "Короткий меч", "Рапира", "Тяжелый арбалет", "Ятаган"
        ]
        
        self.name_var = tk.StringVar(value="")
        self.race_var = tk.StringVar()
        self.class_var = tk.StringVar()
        self.background_var = tk.StringVar()
        self.weapon_var = tk.StringVar()
        self.armor_var = tk.StringVar(value="Без доспеха")
        self.shield_var = tk.BooleanVar(value=False)
        
        self.level_var = tk.StringVar(value="1")
        self.alignment_var = tk.StringVar(value="Нейтральный")
        
        self.hp_var = tk.StringVar(value="10")
        self.ac_var = tk.StringVar(value="10")
        self.initiative_var = tk.StringVar(value="+0")
        self.speed_var = tk.StringVar(value="30 футов")
        
        self.equipment_text = tk.StringVar()
        self.proficiencies_text = tk.StringVar()
        self.languages_text = tk.StringVar()
        self.trait_text = tk.StringVar()
        
        self.skill_states = {skill: False for skill in self.all_skills}
        self.background_skills = []
        
        self.race_var.set(self.races[0])
        self.class_var.set(self.classes[0])
        self.background_var.set(self.background_names[0])
        self.weapon_var.set(self.weapons[0])
        
        self.remaining_points = tk.IntVar(value=self.max_points)
        self.ability_mode = tk.StringVar(value="pointbuy")
        self.selected_skills_count = 0
        self.max_skills = self.class_skill_count[self.classes[0]]
        
        self.advantage_var = tk.StringVar(value="normal")
        self.dice_result_var = tk.StringVar(value="Результат: -")
        self.dice_history = []
        
        self.race_var.trace_add("write", self.on_race_changed)
        self.class_var.trace_add("write", self.on_class_changed)
        self.background_var.trace_add("write", self.on_background_changed)
        self.level_var.trace_add("write", self.on_level_changed)
        self.armor_var.trace_add("write", self.on_armor_changed)
        self.shield_var.trace_add("write", self.on_armor_changed)
        
        self.create_widgets()
        self.apply_race_bonuses()
        self.update_all_modifiers()
        self.update_skill_list()
        self.update_proficiency_bonus()
        self.update_combat_stats()
        self.on_background_changed()
        self.update_saving_throws_display()
        self.update_race_info()
    
    def setup_styles(self):
        self.colors = {
            "bg_dark": "#1a1a2e",
            "bg_medium": "#16213e",
            "bg_light": "#0f3460",
            "accent": "#e94560",
            "text_light": "#ecf0f1",
            "success": "#27ae60",
            "warning": "#f39c12",
            "info": "#3498db",
            "purple": "#8e44ad"
        }
    
    def on_background_changed(self, *args):
        bg = self.backgrounds[self.background_var.get()]
        self.trait_text.set(f"{bg['trait']}\nСнаряжение: {bg['equipment']}")
        tools_str = ", ".join(bg['tools']) if bg['tools'] else "Нет"
        self.proficiencies_text.set(f"Инструменты: {tools_str}")
        if bg['languages'] > 0:
            self.languages_text.set(f"+{bg['languages']} языка(ов) на выбор")
        else:
            self.languages_text.set("")
        self.reset_background_skills()
        self.apply_background_skills()
    
    def reset_background_skills(self):
        for skill in self.background_skills:
            if skill in self.skill_states:
                self.skill_states[skill] = False
                if skill in self.skill_checkboxes:
                    self.skill_checkboxes[skill].deselect()
                self.selected_skills_count -= 1
        self.background_skills = []
    
    def apply_background_skills(self):
        bg = self.backgrounds[self.background_var.get()]
        bg_skills = bg.get("skills", [])
        available_skills = self.class_skill_choices.get(self.class_var.get(), [])
        
        for skill in bg_skills:
            if skill in self.skill_states:
                if skill in available_skills:
                    self.skill_states[skill] = True
                    if skill in self.skill_checkboxes:
                        self.skill_checkboxes[skill].select()
                    self.selected_skills_count += 1
                    self.background_skills.append(skill)
        
        self.max_skills = self.class_skill_count.get(self.class_var.get(), 2) + len(self.background_skills)
        self.skill_count_label.config(text=f"{self.selected_skills_count}/{self.max_skills}")
        self.update_skill_bonuses()
        self.update_passive_perception()
    
    def toggle_skill(self, skill_name):
        if skill_name in self.background_skills:
            messagebox.showinfo("Навык предыстории", 
                f"Навык '{skill_name}' получен от предыстории и не может быть снят.")
            if skill_name in self.skill_checkboxes:
                self.skill_checkboxes[skill_name].select()
            return
        
        if not self.skill_states[skill_name]:
            if self.selected_skills_count >= self.max_skills:
                messagebox.showinfo("Ограничение", f"Вы уже выбрали максимум навыков ({self.max_skills})")
                return
            self.skill_states[skill_name] = True
            self.selected_skills_count += 1
        else:
            self.skill_states[skill_name] = False
            self.selected_skills_count -= 1
        
        self.skill_count_label.config(text=f"{self.selected_skills_count}/{self.max_skills}")
        self.update_skill_bonuses()
        self.update_passive_perception()
    
    def update_skill_list(self):
        class_name = self.class_var.get()
        available_skills = self.class_skill_choices.get(class_name, [])
        
        for skill in self.all_skills:
            if skill not in self.background_skills:
                self.skill_states[skill] = False
        
        self.selected_skills_count = len([s for s, state in self.skill_states.items() if state])
        
        for skill, cb in self.skill_checkboxes.items():
            if skill in available_skills or skill in self.background_skills:
                cb.config(state="normal")
            else:
                cb.config(state="disabled")
            if self.skill_states[skill]:
                cb.select()
            else:
                cb.deselect()
        
        self.max_skills = self.class_skill_count.get(class_name, 2) + len(self.background_skills)
        self.skill_count_label.config(text=f"{self.selected_skills_count}/{self.max_skills}")
        self.update_skill_bonuses()
        self.update_passive_perception()
    
    def on_class_changed(self, *args):
        class_name = self.class_var.get()
        bg = self.backgrounds[self.background_var.get()]
        bg_skills = bg.get("skills", [])
        available_skills = self.class_skill_choices.get(class_name, [])
        
        incompatible_skills = [s for s in bg_skills if s not in available_skills]
        if incompatible_skills:
            messagebox.showwarning(
                "Несовместимые навыки",
                f"Класс '{class_name}' не имеет доступа к навыкам: {', '.join(incompatible_skills)}."
            )
            for skill in incompatible_skills:
                if skill in self.skill_states:
                    self.skill_states[skill] = False
                    if skill in self.skill_checkboxes:
                        self.skill_checkboxes[skill].deselect()
                    if skill in self.background_skills:
                        self.background_skills.remove(skill)
                    self.selected_skills_count -= 1
        
        self.update_skill_list()
        self.update_class_info()
        self.update_combat_stats()
        self.update_saving_throws_display()
    
    def on_race_changed(self, *args):
        self.apply_race_bonuses()
        self.update_combat_stats()
        self.update_armor_class()
        self.update_race_info()
    
    def on_level_changed(self, *args):
        self.update_proficiency_bonus()
        self.update_combat_stats()
        self.update_saving_throws_display()
        self.update_skill_bonuses()
    
    def on_armor_changed(self, *args):
        self.update_armor_class()
    
    def calculate_ac(self):
        armor_name = self.armor_var.get()
        armor_data = self.armors.get(armor_name, ["none", 10, 99])
        armor_type, base_ac, max_dex = armor_data
        
        dex_mod = self.calculate_modifier(self.abilities["Ловкость"].get())
        effective_dex = min(dex_mod, max_dex)
        
        if armor_type == "none":
            ac = 10 + dex_mod
        elif armor_type == "light":
            ac = base_ac + dex_mod
        elif armor_type == "medium":
            ac = base_ac + effective_dex
        elif armor_type == "heavy":
            ac = base_ac
        else:
            ac = 10 + dex_mod
        
        if self.shield_var.get():
            ac += 2
        
        return ac
    
    def update_armor_class(self):
        ac = self.calculate_ac()
        self.ac_var.set(str(ac))
    
    def calculate_proficiency_bonus(self, level):
        if level < 5: return 2
        elif level < 9: return 3
        elif level < 13: return 4
        elif level < 17: return 5
        else: return 6
    
    def update_proficiency_bonus(self, *args):
        try:
            level = int(self.level_var.get())
            prof_bonus = self.calculate_proficiency_bonus(level)
            self.prof_bonus_label.config(text=f"+{prof_bonus}")
            self.update_skill_bonuses()
            self.update_saving_throws_display()
            self.update_passive_perception()
        except ValueError:
            pass
    
    def get_proficiency_bonus(self):
        try:
            return self.calculate_proficiency_bonus(int(self.level_var.get()))
        except:
            return 2
    
    def apply_race_bonuses(self):
        race = self.race_var.get()
        bonuses = self.race_bonuses.get(race, {})
        
        for ability in self.abilities:
            base = self.base_abilities.get(ability, 10)
            bonus = bonuses.get(ability, 0)
            self.abilities[ability].set(base + bonus)
        
        self.update_all_modifiers()
        self.update_skill_bonuses()
        self.update_saving_throws_display()
        self.update_passive_perception()
        self.update_armor_class()
    
    def check_class_requirements(self):
        class_name = self.class_var.get()
        requirements = self.class_requirements.get(class_name, {})
        
        if class_name == "Воин":
            str_value = self.abilities["Сила"].get()
            dex_value = self.abilities["Ловкость"].get()
            if str_value >= 13 or dex_value >= 13:
                return True, []
            else:
                return False, [("Сила или Ловкость", 13)]
        
        unmet = []
        for ability, min_value in requirements.items():
            current = self.abilities[ability].get()
            if current < min_value:
                unmet.append((ability, min_value))
        
        return len(unmet) == 0, unmet
    
    def calculate_modifier(self, value):
        return (value - 10) // 2
    
    def roll_4d6_drop_lowest(self):
        rolls = [random.randint(1, 6) for _ in range(4)]
        rolls.sort(reverse=True)
        return sum(rolls[:3])
    
    def generate_rolled_stats(self):
        stats = []
        for _ in range(6):
            stats.append(self.roll_4d6_drop_lowest())
        stats.sort(reverse=True)
        return stats
    
    def apply_rolled_stats(self):
        stats = self.generate_rolled_stats()
        abilities_list = list(self.abilities.keys())
        for i, ability in enumerate(abilities_list):
            self.base_abilities[ability] = stats[i]
        self.apply_race_bonuses()
        self.update_remaining_points()
        self.update_all_modifiers()
        self.update_skill_bonuses()
        self.update_combat_stats()
        self.update_saving_throws_display()
        self.update_armor_class()
    
    def calculate_point_cost(self, value):
        if value <= 8: return 0
        elif value == 9: return 1
        elif value == 10: return 2
        elif value == 11: return 3
        elif value == 12: return 4
        elif value == 13: return 5
        elif value == 14: return 7
        elif value == 15: return 9
        else: return 9 + (value - 15) * 2
    
    def update_remaining_points(self):
        total_cost = sum(self.calculate_point_cost(self.base_abilities[a]) for a in self.base_abilities)
        self.remaining_points.set(self.max_points - total_cost)
    
    def can_increase_ability(self, ability_name):
        if self.ability_mode.get() != "pointbuy":
            return True
        current_base = self.base_abilities.get(ability_name, 10)
        if current_base >= 15:
            return False
        new_cost = self.calculate_point_cost(current_base + 1)
        current_cost = self.calculate_point_cost(current_base)
        return self.remaining_points.get() >= (new_cost - current_cost)
    
    def can_decrease_ability(self, ability_name):
        if self.ability_mode.get() != "pointbuy":
            return True
        return self.base_abilities.get(ability_name, 10) > 8
    
    def increase_ability(self, ability_name):
        if self.ability_mode.get() == "pointbuy" and not self.can_increase_ability(ability_name):
            return
        self.base_abilities[ability_name] += 1
        self.apply_race_bonuses()
        self.update_remaining_points()
        self.update_all_modifiers()
        self.update_skill_bonuses()
        self.update_combat_stats()
        self.update_saving_throws_display()
        self.update_armor_class()
    
    def decrease_ability(self, ability_name):
        if self.ability_mode.get() == "pointbuy" and not self.can_decrease_ability(ability_name):
            return
        self.base_abilities[ability_name] -= 1
        self.apply_race_bonuses()
        self.update_remaining_points()
        self.update_all_modifiers()
        self.update_skill_bonuses()
        self.update_combat_stats()
        self.update_saving_throws_display()
        self.update_armor_class()
    
    def switch_ability_mode(self):
        if self.ability_mode.get() == "rolled":
            self.roll_button.config(state="normal")
            for btn in self.increase_buttons.values():
                btn.config(state="disabled")
            for btn in self.decrease_buttons.values():
                btn.config(state="disabled")
        else:
            self.roll_button.config(state="disabled")
            for btn in self.increase_buttons.values():
                btn.config(state="normal")
            for btn in self.decrease_buttons.values():
                btn.config(state="normal")
            for ability in self.base_abilities:
                self.base_abilities[ability] = 10
            self.apply_race_bonuses()
            self.update_remaining_points()
            self.update_all_modifiers()
            self.update_skill_bonuses()
            self.update_combat_stats()
            self.update_armor_class()
    
    def update_modifier_label(self, ability_name):
        value = self.abilities[ability_name].get()
        modifier = self.calculate_modifier(value)
        sign = "+" if modifier >= 0 else ""
        self.modifier_labels[ability_name].config(text=f"{sign}{modifier}")
    
    def update_all_modifiers(self):
        for ability in self.abilities:
            self.update_modifier_label(ability)
    
    def update_skill_bonuses(self):
        prof_bonus = self.get_proficiency_bonus()
        for skill in self.all_skills:
            if skill in self.skill_labels:
                for ability, skills in self.skills_by_ability.items():
                    if skill in skills:
                        ability_mod = self.calculate_modifier(self.abilities[ability].get())
                        total = ability_mod
                        if self.skill_states[skill]:
                            total += prof_bonus
                        sign = "+" if total >= 0 else ""
                        self.skill_labels[skill].config(text=f"{sign}{total}")
                        break
    
    def update_race_info(self, *args):
        race = self.race_var.get()
        bonuses = self.race_bonuses.get(race, {})
        if bonuses:
            bonus_text = "Бонусы: " + ", ".join([f"{k} +{v}" for k, v in bonuses.items()])
        else:
            bonus_text = "Нет бонусов"
        self.race_bonus_label.config(text=bonus_text)
    
    def update_class_info(self, *args):
        class_name = self.class_var.get()
        requirements = self.class_requirements.get(class_name, {})
        primary = self.class_primary_ability.get(class_name, "")
        skill_count = self.class_skill_count.get(class_name, 2)
        
        if class_name == "Воин":
            req_text = f"Требования: Сила 13 ИЛИ Ловкость 13 | Навыки: {skill_count}"
        elif requirements:
            req_text = "Требования: " + ", ".join([f"{k} {v}" for k, v in requirements.items()])
            req_text += f" | Навыки: {skill_count}"
        else:
            req_text = f"Навыки: {skill_count}"
        
        if primary:
            req_text += f" | Осн: {primary}"
        
        self.class_req_label.config(text=req_text)
    
    def update_combat_stats(self):
        dex_mod = self.calculate_modifier(self.abilities["Ловкость"].get())
        con_mod = self.calculate_modifier(self.abilities["Телосложение"].get())
        
        self.initiative_var.set(f"+{dex_mod}" if dex_mod >= 0 else str(dex_mod))
        
        hit_dice = 6
        class_name = self.class_var.get()
        if class_name in ["Варвар"]:
            hit_dice = 12
        elif class_name in ["Воин", "Паладин", "Следопыт"]:
            hit_dice = 10
        elif class_name in ["Бард", "Жрец", "Друид", "Монах", "Плут", "Колдун"]:
            hit_dice = 8
        
        try:
            level = int(self.level_var.get())
            hp = hit_dice + con_mod + (hit_dice // 2 + 1 + con_mod) * (level - 1)
            self.hp_var.set(str(max(hp, 1)))
        except:
            self.hp_var.set(str(hit_dice + con_mod))
        
        race = self.race_var.get()
        if race in ["Дварф (Холмовой)", "Дварф (Горный)", "Гном (Лесной)", "Гном (Скальный)", 
                    "Полурослик (Легконогий)", "Полурослик (Крепыш)"]:
            speed = 25
        else:
            speed = 30
        self.speed_var.set(f"{speed} футов")
        
        self.update_armor_class()
    
    def update_saving_throws_display(self):
        class_name = self.class_var.get()
        proficient_throws = self.class_saving_throws.get(class_name, [])
        prof_bonus = self.get_proficiency_bonus()
        
        for ability in self.abilities:
            ability_mod = self.calculate_modifier(self.abilities[ability].get())
            is_proficient = ability in proficient_throws
            total = ability_mod + (prof_bonus if is_proficient else 0)
            sign = "+" if total >= 0 else ""
            
            if ability in self.saving_throw_labels:
                self.saving_throw_labels[ability].config(text=f"{sign}{total}")
            
            if ability in self.saving_throw_checkboxes:
                if is_proficient:
                    self.saving_throw_checkboxes[ability].select()
                else:
                    self.saving_throw_checkboxes[ability].deselect()
    
    def update_passive_perception(self):
        wis_mod = self.calculate_modifier(self.abilities["Мудрость"].get())
        prof_bonus = self.get_proficiency_bonus()
        has_prof = self.skill_states.get("Внимательность", False)
        passive = 10 + wis_mod + (prof_bonus if has_prof else 0)
        self.passive_perception_label.config(text=str(passive))
    
    def roll_dice(self, dice_type, count=1):
        mode = self.advantage_var.get()
        
        if dice_type == 20 and mode != "normal" and count == 1:
            roll1 = random.randint(1, 20)
            roll2 = random.randint(1, 20)
            if mode == "advantage":
                result = max(roll1, roll2)
                detail = f"{roll1}, {roll2} → {result}"
            else:
                result = min(roll1, roll2)
                detail = f"{roll1}, {roll2} → {result}"
        else:
            results = [random.randint(1, dice_type) for _ in range(count)]
            result = sum(results)
            if count > 1:
                detail = " + ".join(map(str, results)) + f" = {result}"
            else:
                detail = str(result)
        
        history_entry = f"d{dice_type}" + (f"x{count}" if count > 1 else "") + f" ({mode}): {detail}"
        self.dice_history.append(history_entry)
        if len(self.dice_history) > 10:
            self.dice_history.pop(0)
        
        self.dice_result_var.set(f"Результат: {result}")
        self.dice_detail_label.config(text=detail)
        self.update_dice_history()
    
    def update_dice_history(self):
        history_text = "\n".join(self.dice_history[-5:]) if self.dice_history else "Нет бросков"
        self.dice_history_text.config(state="normal")
        self.dice_history_text.delete("1.0", tk.END)
        self.dice_history_text.insert("1.0", history_text)
        self.dice_history_text.config(state="disabled")
    
    def clear_dice_history(self):
        self.dice_history = []
        self.update_dice_history()
        self.dice_result_var.set("Результат: -")
        self.dice_detail_label.config(text="Ожидание броска...")
    
    def create_widgets(self):
        title_frame = tk.Frame(self.root, bg=self.colors["bg_dark"])
        title_frame.pack(fill="x", pady=5)
        
        title_label = tk.Label(title_frame, text="⚔️ СОЗДАЙ СВОЕГО ГЕРОЯ ⚔️", 
                               font=("Arial", 18, "bold"), bg=self.colors["bg_dark"], 
                               fg=self.colors["accent"])
        title_label.pack()
        
        main_container = tk.Frame(self.root, bg=self.colors["bg_dark"])
        main_container.pack(fill="both", expand=True, padx=10, pady=5)
        
        left_panel = tk.Frame(main_container, bg=self.colors["bg_medium"], relief="solid", borderwidth=1)
        left_panel.pack(side="left", fill="both", expand=True, padx=(0, 5))
        
        canvas = tk.Canvas(left_panel, bg=self.colors["bg_medium"], highlightthickness=0)
        scrollbar = ttk.Scrollbar(left_panel, orient="vertical", command=canvas.yview)
        scrollable_frame = tk.Frame(canvas, bg=self.colors["bg_medium"])
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        def _on_mousewheel(event):
            canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        canvas.bind_all("<MouseWheel>", _on_mousewheel)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        right_panel = tk.Frame(main_container, bg=self.colors["bg_medium"], relief="solid", borderwidth=1, width=250)
        right_panel.pack(side="right", fill="y", padx=(5, 0))
        right_panel.pack_propagate(False)
        
        form_frame = tk.Frame(scrollable_frame, bg=self.colors["bg_medium"])
        form_frame.pack(pady=10, padx=15, fill="x")
        
        row = 0
        
        info_frame = tk.LabelFrame(form_frame, text="📋 Основная информация", font=("Arial", 11, "bold"),
                                   bg=self.colors["bg_light"], fg=self.colors["text_light"],
                                   padx=10, pady=10)
        info_frame.grid(row=row, column=0, columnspan=2, pady=5, sticky="ew")
        
        tk.Label(info_frame, text="Имя:", bg=self.colors["bg_light"], fg=self.colors["text_light"]).grid(row=0, column=0, padx=5, pady=2, sticky="w")
        tk.Entry(info_frame, textvariable=self.name_var, width=25, font=("Arial", 10)).grid(row=0, column=1, padx=5, pady=2)
        
        tk.Label(info_frame, text="Уровень:", bg=self.colors["bg_light"], fg=self.colors["text_light"]).grid(row=0, column=2, padx=5, pady=2, sticky="w")
        ttk.Spinbox(info_frame, from_=1, to=20, textvariable=self.level_var, width=5).grid(row=0, column=3, padx=5, pady=2)
        
        tk.Label(info_frame, text="Мировоззрение:", bg=self.colors["bg_light"], fg=self.colors["text_light"]).grid(row=1, column=0, padx=5, pady=2, sticky="w")
        alignments = ["Законно-добрый", "Нейтрально-добрый", "Хаотично-добрый", 
                     "Законно-нейтральный", "Нейтральный", "Хаотично-нейтральный",
                     "Законно-злой", "Нейтрально-злой", "Хаотично-злой"]
        ttk.Combobox(info_frame, textvariable=self.alignment_var, values=alignments, state="readonly", width=18).grid(row=1, column=1, padx=5, pady=2)
        
        tk.Label(info_frame, text="Бонус мастерства:", bg=self.colors["bg_light"], fg=self.colors["text_light"]).grid(row=1, column=2, padx=5, pady=2, sticky="w")
        self.prof_bonus_label = tk.Label(info_frame, text="+2", font=("Arial", 12, "bold"), bg=self.colors["bg_light"], fg=self.colors["info"])
        self.prof_bonus_label.grid(row=1, column=3, padx=5, pady=2, sticky="w")
        
        row += 1
        
        combat_frame = tk.LabelFrame(form_frame, text="⚔️ Боевые показатели", font=("Arial", 11, "bold"),
                                     bg=self.colors["bg_light"], fg=self.colors["accent"],
                                     padx=10, pady=10)
        combat_frame.grid(row=row, column=0, columnspan=2, pady=5, sticky="ew")
        
        tk.Label(combat_frame, text="КД:", bg=self.colors["bg_light"], fg=self.colors["text_light"]).grid(row=0, column=0, padx=5, pady=2)
        tk.Entry(combat_frame, textvariable=self.ac_var, width=5, state="readonly", font=("Arial", 11, "bold")).grid(row=0, column=1, padx=5, pady=2)
        
        tk.Label(combat_frame, text="Хиты:", bg=self.colors["bg_light"], fg=self.colors["text_light"]).grid(row=0, column=2, padx=5, pady=2)
        tk.Entry(combat_frame, textvariable=self.hp_var, width=6, state="readonly", font=("Arial", 11, "bold")).grid(row=0, column=3, padx=5, pady=2)
        
        tk.Label(combat_frame, text="Инициатива:", bg=self.colors["bg_light"], fg=self.colors["text_light"]).grid(row=1, column=0, padx=5, pady=2)
        tk.Entry(combat_frame, textvariable=self.initiative_var, width=5, state="readonly").grid(row=1, column=1, padx=5, pady=2)
        
        tk.Label(combat_frame, text="Скорость:", bg=self.colors["bg_light"], fg=self.colors["text_light"]).grid(row=1, column=2, padx=5, pady=2)
        tk.Entry(combat_frame, textvariable=self.speed_var, width=10, state="readonly").grid(row=1, column=3, padx=5, pady=2)
        
        tk.Label(combat_frame, text="Доспех:", bg=self.colors["bg_light"], fg=self.colors["text_light"]).grid(row=2, column=0, padx=5, pady=2)
        ttk.Combobox(combat_frame, textvariable=self.armor_var, values=self.armor_names, 
                     state="readonly", width=18).grid(row=2, column=1, padx=5, pady=2)
        
        tk.Checkbutton(combat_frame, text="Щит (+2 КД)", variable=self.shield_var,
                      bg=self.colors["bg_light"], fg=self.colors["text_light"],
                      selectcolor=self.colors["bg_light"]).grid(row=2, column=2, columnspan=2, padx=5, pady=2)
        
        tk.Label(combat_frame, text="Пассивная Внимательность:", bg=self.colors["bg_light"], fg=self.colors["text_light"]).grid(row=3, column=0, padx=5, pady=2)
        self.passive_perception_label = tk.Label(combat_frame, text="10", font=("Arial", 11, "bold"), 
                                                  bg=self.colors["bg_light"], fg=self.colors["success"])
        self.passive_perception_label.grid(row=3, column=1, padx=5, pady=2, sticky="w")
        
        row += 1
        
        abilities_frame = tk.LabelFrame(form_frame, text="💪 Характеристики", font=("Arial", 11, "bold"),
                                        bg=self.colors["bg_light"], fg=self.colors["text_light"],
                                        padx=10, pady=10)
        abilities_frame.grid(row=row, column=0, columnspan=2, pady=5, sticky="ew")
        
        mode_frame = tk.Frame(abilities_frame, bg=self.colors["bg_light"])
        mode_frame.pack(pady=5)
        
        tk.Radiobutton(mode_frame, text="Закуп (Point Buy)", variable=self.ability_mode, 
                       value="pointbuy", command=self.switch_ability_mode,
                       bg=self.colors["bg_light"], fg=self.colors["text_light"],
                       selectcolor=self.colors["bg_light"]).pack(side="left", padx=10)
        tk.Radiobutton(mode_frame, text="Броски 4d6", variable=self.ability_mode, 
                       value="rolled", command=self.switch_ability_mode,
                       bg=self.colors["bg_light"], fg=self.colors["text_light"],
                       selectcolor=self.colors["bg_light"]).pack(side="left", padx=10)
        
        self.roll_button = tk.Button(mode_frame, text="🎲 Бросить", command=self.apply_rolled_stats, 
                                      bg=self.colors["info"], fg="white", font=("Arial", 9))
        self.roll_button.pack(side="left", padx=20)
        self.roll_button.config(state="disabled")
        
        self.points_label = tk.Label(mode_frame, textvariable=self.remaining_points, 
                                      font=("Arial", 12, "bold"), bg=self.colors["bg_light"], fg=self.colors["warning"])
        self.points_label.pack(side="left", padx=5)
        tk.Label(mode_frame, text="очков", bg=self.colors["bg_light"], fg=self.colors["text_light"]).pack(side="left")
        
        table_frame = tk.Frame(abilities_frame, bg=self.colors["bg_light"])
        table_frame.pack(pady=5)
        
        headers = ["Характеристика", "Значение", "Модификатор", "Управление"]
        for col, header in enumerate(headers):
            tk.Label(table_frame, text=header, font=("Arial", 9, "bold"), 
                    bg=self.colors["bg_medium"], fg=self.colors["text_light"],
                    width=14 if col == 0 else 10).grid(row=0, column=col, padx=1, pady=1)
        
        self.increase_buttons = {}
        self.decrease_buttons = {}
        self.modifier_labels = {}
        
        abilities_order = ["Сила", "Ловкость", "Телосложение", "Интеллект", "Мудрость", "Харизма"]
        
        for i, ability in enumerate(abilities_order, 1):
            var = self.abilities[ability]
            
            tk.Label(table_frame, text=ability, bg=self.colors["bg_light"], fg=self.colors["text_light"]).grid(
                row=i, column=0, pady=2)
            
            tk.Label(table_frame, textvariable=var, font=("Arial", 11, "bold"), 
                    bg=self.colors["bg_light"], fg=self.colors["text_light"], width=8).grid(row=i, column=1, pady=2)
            
            mod_label = tk.Label(table_frame, text="+0", font=("Arial", 11, "bold"), 
                                 bg=self.colors["bg_light"], fg=self.colors["success"], width=8)
            mod_label.grid(row=i, column=2, pady=2)
            self.modifier_labels[ability] = mod_label
            
            btn_frame = tk.Frame(table_frame, bg=self.colors["bg_light"])
            btn_frame.grid(row=i, column=3, pady=2)
            
            dec_btn = tk.Button(btn_frame, text="-", width=2,
                               command=lambda a=ability: self.decrease_ability(a),
                               bg=self.colors["accent"], fg="white")
            dec_btn.pack(side="left", padx=2)
            self.decrease_buttons[ability] = dec_btn
            
            inc_btn = tk.Button(btn_frame, text="+", width=2,
                               command=lambda a=ability: self.increase_ability(a),
                               bg=self.colors["success"], fg="white")
            inc_btn.pack(side="left", padx=2)
            self.increase_buttons[ability] = inc_btn
        
        row += 1
        
        saves_frame = tk.LabelFrame(form_frame, text="🛡️ Спасброски", font=("Arial", 11, "bold"),
                                    bg=self.colors["bg_light"], fg=self.colors["text_light"],
                                    padx=10, pady=10)
        saves_frame.grid(row=row, column=0, columnspan=2, pady=5, sticky="ew")
        
        saves_inner = tk.Frame(saves_frame, bg=self.colors["bg_light"])
        saves_inner.pack()
        
        self.saving_throw_frames = {}
        self.saving_throw_labels = {}
        self.saving_throw_checkboxes = {}
        
        for ability in abilities_order:
            frame = tk.Frame(saves_inner, bg=self.colors["bg_medium"], relief="solid", borderwidth=1)
            frame.pack(side="left", padx=8, pady=5, ipadx=5, ipady=3)
            self.saving_throw_frames[ability] = frame
            
            tk.Label(frame, text=ability, font=("Arial", 9, "bold"), bg=self.colors["bg_medium"], 
                    fg=self.colors["text_light"]).pack()
            
            cb = tk.Checkbutton(frame, text="Владение", state="disabled",
                               bg=self.colors["bg_medium"], fg=self.colors["text_light"],
                               selectcolor=self.colors["bg_medium"])
            cb.pack()
            self.saving_throw_checkboxes[ability] = cb
            
            label = tk.Label(frame, text="+0", font=("Arial", 11, "bold"), 
                            bg=self.colors["bg_medium"], fg=self.colors["purple"])
            label.pack()
            self.saving_throw_labels[ability] = label
        
        row += 1
        
        selection_frame = tk.LabelFrame(form_frame, text="🎭 Раса, класс и предыстория", font=("Arial", 11, "bold"),
                                        bg=self.colors["bg_light"], fg=self.colors["text_light"],
                                        padx=10, pady=10)
        selection_frame.grid(row=row, column=0, columnspan=2, pady=5, sticky="ew")
        
        tk.Label(selection_frame, text="Раса:", bg=self.colors["bg_light"], fg=self.colors["text_light"]).grid(row=0, column=0, padx=5, pady=3, sticky="w")
        ttk.Combobox(selection_frame, textvariable=self.race_var, values=self.races, 
                     state="readonly", width=25).grid(row=0, column=1, padx=5, pady=3)
        
        self.race_bonus_label = tk.Label(selection_frame, text="", font=("Arial", 9, "italic"), 
                                          bg=self.colors["bg_light"], fg=self.colors["info"])
        self.race_bonus_label.grid(row=1, column=1, padx=5, pady=2, sticky="w")
        
        tk.Label(selection_frame, text="Класс:", bg=self.colors["bg_light"], fg=self.colors["text_light"]).grid(row=2, column=0, padx=5, pady=3, sticky="w")
        ttk.Combobox(selection_frame, textvariable=self.class_var, values=self.classes, 
                     state="readonly", width=25).grid(row=2, column=1, padx=5, pady=3)
        
        self.class_req_label = tk.Label(selection_frame, text="", font=("Arial", 9, "italic"), 
                                         bg=self.colors["bg_light"], fg=self.colors["warning"])
        self.class_req_label.grid(row=3, column=1, padx=5, pady=2, sticky="w")
        
        tk.Label(selection_frame, text="Предыстория:", bg=self.colors["bg_light"], fg=self.colors["text_light"]).grid(row=4, column=0, padx=5, pady=3, sticky="w")
        ttk.Combobox(selection_frame, textvariable=self.background_var, values=self.background_names, 
                     state="readonly", width=25).grid(row=4, column=1, padx=5, pady=3)
        
        row += 1
        
        skills_frame = tk.LabelFrame(form_frame, text="🎯 Навыки", font=("Arial", 11, "bold"),
                                     bg=self.colors["bg_light"], fg=self.colors["text_light"],
                                     padx=10, pady=10)
        skills_frame.grid(row=row, column=0, columnspan=2, pady=5, sticky="ew")
        
        skills_header = tk.Frame(skills_frame, bg=self.colors["bg_light"])
        skills_header.pack(pady=5)
        
        self.skill_count_label = tk.Label(skills_header, text="0/2", font=("Arial", 12, "bold"), 
                                           bg=self.colors["bg_light"], fg=self.colors["purple"])
        self.skill_count_label.pack(side="left", padx=10)
        tk.Label(skills_header, text="(отметьте владение)", font=("Arial", 9, "italic"),
                bg=self.colors["bg_light"], fg=self.colors["text_light"]).pack(side="left")
        
        self.skill_checkboxes = {}
        self.skill_labels = {}
        
        skills_container = tk.Frame(skills_frame, bg=self.colors["bg_light"])
        skills_container.pack(pady=5)
        
        col = 0
        for ability, skills in self.skills_by_ability.items():
            if not skills:
                continue
            
            ability_frame = tk.LabelFrame(skills_container, text=ability, font=("Arial", 9, "bold"), 
                                          bg=self.colors["bg_medium"], fg=self.colors["accent"],
                                          padx=5, pady=5)
            ability_frame.grid(row=0, column=col, padx=3, pady=3, sticky="n")
            
            for skill in skills:
                skill_frame = tk.Frame(ability_frame, bg=self.colors["bg_medium"])
                skill_frame.pack(anchor="w", pady=1)
                
                cb = tk.Checkbutton(skill_frame, text=skill,
                                   command=lambda s=skill: self.toggle_skill(s),
                                   bg=self.colors["bg_medium"], fg=self.colors["text_light"],
                                   selectcolor=self.colors["bg_medium"],
                                   font=("Arial", 9), width=18, anchor="w")
                cb.pack(side="left")
                self.skill_checkboxes[skill] = cb
                
                bonus_label = tk.Label(skill_frame, text="+0", font=("Arial", 9, "bold"), 
                                       bg=self.colors["bg_medium"], fg=self.colors["info"], width=4)
                bonus_label.pack(side="right")
                self.skill_labels[skill] = bonus_label
            
            col += 1
        
        row += 1
        
        traits_frame = tk.LabelFrame(form_frame, text="📜 Умения и владения", font=("Arial", 11, "bold"),
                                     bg=self.colors["bg_light"], fg=self.colors["text_light"],
                                     padx=10, pady=10)
        traits_frame.grid(row=row, column=0, columnspan=2, pady=5, sticky="ew")
        
        tk.Label(traits_frame, text="Умение:", bg=self.colors["bg_light"], fg=self.colors["text_light"]).grid(row=0, column=0, padx=5, pady=3, sticky="w")
        tk.Entry(traits_frame, textvariable=self.trait_text, width=50, state="readonly", 
                bg=self.colors["bg_medium"], fg=self.colors["text_light"]).grid(row=0, column=1, padx=5, pady=3)
        
        tk.Label(traits_frame, text="Владения:", bg=self.colors["bg_light"], fg=self.colors["text_light"]).grid(row=1, column=0, padx=5, pady=3, sticky="w")
        tk.Entry(traits_frame, textvariable=self.proficiencies_text, width=50,
                bg=self.colors["bg_medium"], fg=self.colors["text_light"]).grid(row=1, column=1, padx=5, pady=3)
        
        tk.Label(traits_frame, text="Языки:", bg=self.colors["bg_light"], fg=self.colors["text_light"]).grid(row=2, column=0, padx=5, pady=3, sticky="w")
        tk.Entry(traits_frame, textvariable=self.languages_text, width=50,
                bg=self.colors["bg_medium"], fg=self.colors["text_light"]).grid(row=2, column=1, padx=5, pady=3)
        
        tk.Label(traits_frame, text="Оружие:", bg=self.colors["bg_light"], fg=self.colors["text_light"]).grid(row=3, column=0, padx=5, pady=3, sticky="w")
        ttk.Combobox(traits_frame, textvariable=self.weapon_var, values=self.weapons, 
                     state="readonly", width=47).grid(row=3, column=1, padx=5, pady=3)
        
        row += 1
        
        button_frame = tk.Frame(form_frame, bg=self.colors["bg_medium"])
        button_frame.grid(row=row, column=0, columnspan=2, pady=15)
        
        tk.Button(button_frame, text="📜 Создать лист", command=self.show_character_sheet, 
                 bg=self.colors["success"], fg="white", font=("Arial", 11, "bold"), 
                 padx=15, pady=8).pack(side="left", padx=3)
        
        tk.Button(button_frame, text="💾 JSON", command=self.save_to_json, 
                 bg=self.colors["info"], fg="white", font=("Arial", 11, "bold"), 
                 padx=15, pady=8).pack(side="left", padx=3)
        
        tk.Button(button_frame, text="📄 Word", command=self.save_to_word, 
                 bg="#1abc9c", fg="white", font=("Arial", 11, "bold"), 
                 padx=15, pady=8).pack(side="left", padx=3)
        
        tk.Button(button_frame, text="🔄 Сбросить", command=self.reset_form, 
                 bg=self.colors["warning"], fg="white", font=("Arial", 10), 
                 padx=15, pady=8).pack(side="left", padx=3)
        
        dice_frame = tk.LabelFrame(right_panel, text="🎲 Дайс-роллер", font=("Arial", 12, "bold"),
                                   bg=self.colors["bg_light"], fg=self.colors["accent"],
                                   padx=10, pady=10)
        dice_frame.pack(fill="both", expand=True, padx=5, pady=5)
        
        mode_frame = tk.Frame(dice_frame, bg=self.colors["bg_light"])
        mode_frame.pack(pady=10)
        
        tk.Radiobutton(mode_frame, text="Обычный", variable=self.advantage_var, 
                       value="normal", bg=self.colors["bg_light"], fg=self.colors["text_light"],
                       selectcolor=self.colors["bg_light"]).pack(anchor="w")
        tk.Radiobutton(mode_frame, text="Преимущество", variable=self.advantage_var, 
                       value="advantage", bg=self.colors["bg_light"], fg=self.colors["text_light"],
                       selectcolor=self.colors["bg_light"]).pack(anchor="w")
        tk.Radiobutton(mode_frame, text="Помеха", variable=self.advantage_var, 
                       value="disadvantage", bg=self.colors["bg_light"], fg=self.colors["text_light"],
                       selectcolor=self.colors["bg_light"]).pack(anchor="w")
        
        ttk.Separator(dice_frame, orient='horizontal').pack(fill='x', padx=10, pady=5)
        
        tk.Label(dice_frame, text="Кубики:", font=("Arial", 11, "bold"),
                bg=self.colors["bg_light"], fg=self.colors["text_light"]).pack(pady=5)
        
        dice_buttons_frame = tk.Frame(dice_frame, bg=self.colors["bg_light"])
        dice_buttons_frame.pack(pady=5)
        
        dice_types = [("d2", 2), ("d4", 4), ("d6", 6), ("d8", 8), ("d10", 10), ("d12", 12), ("d20", 20), ("d100", 100)]
        
        for i, (label, dtype) in enumerate(dice_types):
            btn = tk.Button(dice_buttons_frame, text=label, width=5, height=2,
                           command=lambda d=dtype: self.roll_dice(d),
                           bg=self.colors["info"], fg="white", font=("Arial", 10, "bold"))
            btn.grid(row=i // 4, column=i % 4, padx=3, pady=3)
        
        ttk.Separator(dice_frame, orient='horizontal').pack(fill='x', padx=10, pady=5)
        
        result_frame = tk.Frame(dice_frame, bg=self.colors["bg_light"])
        result_frame.pack(pady=10)
        
        tk.Label(result_frame, textvariable=self.dice_result_var, font=("Arial", 14, "bold"), 
                bg=self.colors["bg_light"], fg=self.colors["success"]).pack()
        
        self.dice_detail_label = tk.Label(result_frame, text="Ожидание броска...", 
                                           font=("Arial", 9), bg=self.colors["bg_light"], fg="#7f8c8d")
        self.dice_detail_label.pack()
        
        ttk.Separator(dice_frame, orient='horizontal').pack(fill='x', padx=10, pady=5)
        
        tk.Label(dice_frame, text="История:", font=("Arial", 10, "bold"),
                bg=self.colors["bg_light"], fg=self.colors["text_light"]).pack(pady=5)
        
        self.dice_history_text = tk.Text(dice_frame, height=8, width=28, state="disabled", 
                                          font=("Arial", 9), bg="#2c3e50", fg="#ecf0f1")
        self.dice_history_text.pack(padx=10, pady=5)
        
        tk.Button(dice_frame, text="🗑️ Очистить историю", command=self.clear_dice_history,
                 bg=self.colors["warning"], fg="white", font=("Arial", 9)).pack(pady=5)
    
    def show_character_sheet(self):
        if not self.name_var.get().strip():
            messagebox.showwarning("Внимание", "Пожалуйста, введите имя персонажа!")
            return
        
        meets_req, unmet = self.check_class_requirements()
        if not meets_req:
            if self.class_var.get() == "Воин":
                req_text = "Сила 13 ИЛИ Ловкость 13"
            else:
                req_text = ", ".join([f"{ability} {value}" for ability, value in unmet])
            
            response = messagebox.askyesno(
                "Требования класса не выполнены",
                f"Ваш персонаж не соответствует требованиям класса {self.class_var.get()}!\n"
                f"Необходимо: {req_text}\n\nВсё равно продолжить?"
            )
            if not response:
                return
        
        sheet_window = tk.Toplevel(self.root)
        sheet_window.title(f"Лист персонажа: {self.name_var.get()}")
        sheet_window.geometry("550x800")
        sheet_window.configure(bg=self.colors["bg_dark"])
        sheet_window.resizable(False, False)
        
        main_frame = tk.Frame(sheet_window, bg=self.colors["bg_medium"])
        main_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        tk.Label(main_frame, text="ЛИСТ ПЕРСОНАЖА", font=("Arial", 16, "bold"), 
                bg=self.colors["bg_medium"], fg=self.colors["accent"]).pack(pady=15)
        
        info_frame = tk.Frame(main_frame, bg=self.colors["bg_light"])
        info_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        def create_info_row(parent, label_text, value_text, color="#3498db"):
            frame = tk.Frame(parent, bg=self.colors["bg_light"])
            frame.pack(fill="x", pady=3, padx=20)
            tk.Label(frame, text=label_text, font=("Arial", 11, "bold"), 
                    bg=self.colors["bg_light"], fg=color, width=16, anchor="w").pack(side="left")
            tk.Label(frame, text=value_text, font=("Arial", 11), 
                    bg=self.colors["bg_light"], fg=self.colors["text_light"], 
                    anchor="w", wraplength=300).pack(side="left", padx=10)
        
        prof_bonus = self.get_proficiency_bonus()
        create_info_row(info_frame, "Имя:", self.name_var.get())
        create_info_row(info_frame, "Уровень:", f"{self.level_var.get()} (бонус +{prof_bonus})")
        create_info_row(info_frame, "Мировоззрение:", self.alignment_var.get())
        create_info_row(info_frame, "Раса:", self.race_var.get())
        create_info_row(info_frame, "Класс:", self.class_var.get())
        create_info_row(info_frame, "Предыстория:", self.background_var.get())
        create_info_row(info_frame, "КД/Хиты:", f"{self.ac_var.get()} / {self.hp_var.get()}")
        create_info_row(info_frame, "Доспех:", f"{self.armor_var.get()}" + (" + Щит" if self.shield_var.get() else ""))
        create_info_row(info_frame, "Иниц./Скорость:", f"{self.initiative_var.get()} / {self.speed_var.get()}")
        create_info_row(info_frame, "Пасс. Внимательность:", self.passive_perception_label.cget("text"))
        
        ttk.Separator(info_frame, orient='horizontal').pack(fill='x', padx=20, pady=5)
        
        tk.Label(info_frame, text="Характеристики:", font=("Arial", 12, "bold"), 
                bg=self.colors["bg_light"], fg=self.colors["warning"]).pack(pady=5)
        
        for ability, var in self.abilities.items():
            value = var.get()
            modifier = self.calculate_modifier(value)
            sign = "+" if modifier >= 0 else ""
            create_info_row(info_frame, f"{ability}:", f"{value} ({sign}{modifier})", "#e67e22")
        
        ttk.Separator(info_frame, orient='horizontal').pack(fill='x', padx=20, pady=5)
        
        tk.Label(info_frame, text="Навыки (владение):", font=("Arial", 12, "bold"), 
                bg=self.colors["bg_light"], fg=self.colors["warning"]).pack(pady=5)
        
        selected = [s for s, state in self.skill_states.items() if state]
        skills_text = ", ".join(selected) if selected else "нет"
        create_info_row(info_frame, "Владение:", skills_text, self.colors["purple"])
        
        tk.Label(info_frame, text="Умение предыстории:", font=("Arial", 12, "bold"), 
                bg=self.colors["bg_light"], fg=self.colors["warning"]).pack(pady=5)
        create_info_row(info_frame, "Умение:", self.trait_text.get().split("\n")[0] or "нет", "#1abc9c")
        
        btn_frame = tk.Frame(main_frame, bg=self.colors["bg_medium"])
        btn_frame.pack(pady=20)
        tk.Button(btn_frame, text="Закрыть", command=sheet_window.destroy, 
                 bg=self.colors["accent"], fg="white", font=("Arial", 10, "bold"), 
                 padx=20, pady=5).pack()
    
    def save_to_json(self):
        if not self.name_var.get().strip():
            messagebox.showwarning("Внимание", "Пожалуйста, введите имя персонажа!")
            return
            
        ability_data = {}
        for ability, var in self.abilities.items():
            ability_data[ability] = {
                "value": var.get(),
                "modifier": self.calculate_modifier(var.get())
            }
        
        character_data = {
            "name": self.name_var.get(),
            "level": self.level_var.get(),
            "proficiency_bonus": self.get_proficiency_bonus(),
            "alignment": self.alignment_var.get(),
            "race": self.race_var.get(),
            "class": self.class_var.get(),
            "background": self.background_var.get(),
            "armor": self.armor_var.get(),
            "shield": self.shield_var.get(),
            "ac": self.ac_var.get(),
            "hp": self.hp_var.get(),
            "initiative": self.initiative_var.get(),
            "speed": self.speed_var.get(),
            "passive_perception": self.passive_perception_label.cget("text"),
            "abilities": ability_data,
            "skills": [s for s, state in self.skill_states.items() if state],
            "background_skills": self.background_skills,
            "equipment": self.equipment_text.get(),
            "proficiencies": self.proficiencies_text.get(),
            "languages": self.languages_text.get(),
            "trait": self.trait_text.get()
        }
        
        try:
            filename = f"{self.name_var.get().replace(' ', '_')}_character.json"
            with open(filename, 'w', encoding='utf-8') as f:
                json.dump(character_data, f, ensure_ascii=False, indent=4)
            messagebox.showinfo("Успех", f"Персонаж сохранен в файл:\n{filename}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось сохранить файл:\n{str(e)}")
    
    def save_to_word(self):
        """Сохранение персонажа в документ Word"""
        if not self.name_var.get().strip():
            messagebox.showwarning("Внимание", "Пожалуйста, введите имя персонажа!")
            return
        
        try:
            doc = Document()
            
            # Заголовок
            title = doc.add_heading(f'Лист персонажа: {self.name_var.get()}', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Основная информация
            doc.add_heading('Основная информация', level=1)
            
            info_table = doc.add_table(rows=8, cols=2)
            info_table.style = 'Table Grid'
            
            info_data = [
                ("Имя", self.name_var.get()),
                ("Уровень", f"{self.level_var.get()} (бонус мастерства +{self.get_proficiency_bonus()})"),
                ("Мировоззрение", self.alignment_var.get()),
                ("Раса", self.race_var.get()),
                ("Класс", self.class_var.get()),
                ("Предыстория", self.background_var.get()),
                ("КД", f"{self.ac_var.get()} ({self.armor_var.get()}" + (" + Щит)" if self.shield_var.get() else ")")),
                ("Хиты", self.hp_var.get())
            ]
            
            for i, (label, value) in enumerate(info_data):
                info_table.cell(i, 0).text = label
                info_table.cell(i, 1).text = value
                # Жирный шрифт для заголовков
                info_table.cell(i, 0).paragraphs[0].runs[0].bold = True
            
            # Характеристики
            doc.add_heading('Характеристики', level=1)
            
            stats_table = doc.add_table(rows=7, cols=3)
            stats_table.style = 'Table Grid'
            
            headers = ['Характеристика', 'Значение', 'Модификатор']
            for i, header in enumerate(headers):
                cell = stats_table.cell(0, i)
                cell.text = header
                cell.paragraphs[0].runs[0].bold = True
            
            row = 1
            for ability, var in self.abilities.items():
                value = var.get()
                modifier = self.calculate_modifier(value)
                sign = "+" if modifier >= 0 else ""
                
                stats_table.cell(row, 0).text = ability
                stats_table.cell(row, 1).text = str(value)
                stats_table.cell(row, 2).text = f"{sign}{modifier}"
                row += 1
            
            # Спасброски
            doc.add_heading('Спасброски', level=1)
            
            class_name = self.class_var.get()
            proficient_throws = self.class_saving_throws.get(class_name, [])
            prof_bonus = self.get_proficiency_bonus()
            
            saves_table = doc.add_table(rows=7, cols=3)
            saves_table.style = 'Table Grid'
            
            saves_headers = ['Характеристика', 'Владение', 'Бонус']
            for i, header in enumerate(saves_headers):
                cell = saves_table.cell(0, i)
                cell.text = header
                cell.paragraphs[0].runs[0].bold = True
            
            row = 1
            for ability in self.abilities:
                ability_mod = self.calculate_modifier(self.abilities[ability].get())
                is_proficient = ability in proficient_throws
                total = ability_mod + (prof_bonus if is_proficient else 0)
                sign = "+" if total >= 0 else ""
                
                saves_table.cell(row, 0).text = ability
                saves_table.cell(row, 1).text = "✓" if is_proficient else ""
                saves_table.cell(row, 2).text = f"{sign}{total}"
                row += 1
            
            # Навыки
            doc.add_heading('Навыки', level=1)
            
            skills_table = doc.add_table(rows=len(self.all_skills) + 1, cols=3)
            skills_table.style = 'Table Grid'
            
            skills_headers = ['Навык', 'Характеристика', 'Бонус']
            for i, header in enumerate(skills_headers):
                cell = skills_table.cell(0, i)
                cell.text = header
                cell.paragraphs[0].runs[0].bold = True
            
            row = 1
            for skill in self.all_skills:
                # Находим характеристику
                ability_name = ""
                for ability, skills in self.skills_by_ability.items():
                    if skill in skills:
                        ability_name = ability
                        break
                
                ability_mod = self.calculate_modifier(self.abilities[ability_name].get())
                total = ability_mod
                if self.skill_states[skill]:
                    total += prof_bonus
                sign = "+" if total >= 0 else ""
                
                skills_table.cell(row, 0).text = skill
                skills_table.cell(row, 1).text = ability_name[:3]  # Сокращение
                skills_table.cell(row, 2).text = f"{sign}{total}"
                
                # Отмечаем владение жирным
                if self.skill_states[skill]:
                    skills_table.cell(row, 0).paragraphs[0].runs[0].bold = True
                    skills_table.cell(row, 2).paragraphs[0].runs[0].bold = True
                
                row += 1
            
            # Умения и владения
            doc.add_heading('Умения и снаряжение', level=1)
            
            doc.add_paragraph(f"Умение предыстории: {self.trait_text.get().split(chr(10))[0]}")
            doc.add_paragraph(f"Владения: {self.proficiencies_text.get()}")
            doc.add_paragraph(f"Языки: {self.languages_text.get()}")
            doc.add_paragraph(f"Оружие: {self.weapon_var.get()}")
            doc.add_paragraph(f"Снаряжение: {self.trait_text.get().split('Снаряжение: ')[-1] if 'Снаряжение:' in self.trait_text.get() else 'Нет'}")
            
            # Сохранение файла
            filename = f"{self.name_var.get().replace(' ', '_')}_character.docx"
            doc.save(filename)
            messagebox.showinfo("Успех", f"Персонаж сохранен в файл:\n{filename}")
            
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось сохранить Word-документ:\n{str(e)}")
    
    def reset_form(self):
        self.name_var.set("")
        self.level_var.set("1")
        self.alignment_var.set("Нейтральный")
        self.race_var.set(self.races[0])
        self.class_var.set(self.classes[0])
        self.background_var.set(self.background_names[0])
        self.armor_var.set("Без доспеха")
        self.shield_var.set(False)
        
        for ability in self.base_abilities:
            self.base_abilities[ability] = 10
        
        self.background_skills = []
        for skill in self.all_skills:
            self.skill_states[skill] = False
        self.selected_skills_count = 0
        
        self.equipment_text.set("")
        self.languages_text.set("")
        
        self.apply_race_bonuses()
        self.update_remaining_points()
        self.update_all_modifiers()
        self.update_skill_list()
        self.update_proficiency_bonus()
        self.update_combat_stats()
        self.on_background_changed()
        self.update_saving_throws_display()
        self.update_armor_class()
        self.update_race_info()
        
        self.clear_dice_history()

if __name__ == "__main__":
    root = tk.Tk()
    app = CharacterCreator(root)
    root.mainloop()