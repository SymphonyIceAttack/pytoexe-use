# Auto-generated merged file from Crux source
# NOTE: This merged file is for portability/review; duplicate imports/names may exist.


# ===== BEGIN FILE: config.py =====

"""
MonTool - Linux Admin Monitoring Tool
A professional monitoring dashboard for telecom infrastructure
"""

import customtkinter as ctk
import tkinter as tk
from tkinter import filedialog, messagebox
import threading
import os
import time
from datetime import datetime

try:
    import paramiko
    PARAMIKO_AVAILABLE = True
except ImportError:
    PARAMIKO_AVAILABLE = False

# â”€â”€â”€ Theme â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")

BG_DARK      = "#0D0F14"
BG_PANEL     = "#12151C"
BG_CARD      = "#1A1E2A"
BG_CARD2     = "#1F2435"
ACCENT       = "#00C8FF"
ACCENT2      = "#0084AA"
SUCCESS      = "#00E096"
WARNING      = "#FFB830"
DANGER       = "#FF4D6A"
TEXT_PRIMARY = "#E8EAF0"
TEXT_MUTED   = "#7A8099"
BORDER       = "#2A2F40"
SIDEBAR_W    = 200

# â”€â”€â”€ App State â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
# Shared catalina path for all Aircontrol API nodes
AC_CATALINA_LOG = "/opt/Airlinq/runtime/ac-api/apache-tomcat-9.0.106/logs/catalina.out"

LOG_MODULES = [
    "Aircontrol API-1 (10.221.92.116)",
    "Aircontrol API-2 (10.221.92.117)",
    "Aircontrol API-3 (10.221.92.118)",
    "Aircontrol API-4 (10.221.92.119)",
    "Streaming",
    "OL",
    "K8",
    "Metabase",
]

def _default_module_row(host="", log_path="/var/log/app.log", sudo_user="aqadmin"):
    return {
        "host": host,
        "port": "22",
        "user": "",
        "password": "",
        "key": "",
        "log_path": log_path,
        "sudo_user": sudo_user,
        "sudo_password": "",
    }

class AppState:
    project  = None
    env      = None
    env_hosts = {
        "PROD":        "prod.server.local",
        "ISMIL":       "ismil.server.local",
        "Dallas Prod": "dallas.server.local",
        "Miami":       "miami.server.local",
    }
    log_lines = 200
    # User-saved quick queries for Database page (max 10)
    saved_queries = []

    # Set by MonToolApp: callable() to open the floating evidence panel (or None)
    evidence_panel_opener = None

    # Per-module server config
    # user/password = LDAP credentials for SSH login
    # sudo_user     = target after `sudo su -` (default: aqadmin); uses same password as LDAP
    module_configs = {
        "Aircontrol API-1 (10.221.92.116)": _default_module_row("10.221.92.116", AC_CATALINA_LOG),
        "Aircontrol API-2 (10.221.92.117)": _default_module_row("10.221.92.117", AC_CATALINA_LOG),
        "Aircontrol API-3 (10.221.92.118)": _default_module_row("10.221.92.118", AC_CATALINA_LOG),
        "Aircontrol API-4 (10.221.92.119)": _default_module_row("10.221.92.119", AC_CATALINA_LOG),
        "Streaming":  _default_module_row("", "/var/log/streaming/app.log"),
        "OL":         _default_module_row("", "/var/log/ol/app.log"),
        "K8":         _default_module_row("", ""),
        "Metabase":   _default_module_row("", "/var/log/metabase/app.log"),
    }

    def get_module_host(self, module):
        """Return module-specific host, fallback to env host."""
        h = self.module_configs.get(module, {}).get("host", "").strip()
        return h if h else self.env_hosts.get(self.env, "")

    # Separate DB config for Database Queries page
    # MySQL example: mysql -u aqadmin -p... -h HOST -P PORT -A
    db_configs = {
        "MySQL": {
            "host": "10.221.92.116",
            "port": "6446",
            "user": "aqadmin",
            "password": "Aqadmin@123",
            "db_name": "dallas_prod_aircontrol_db",
        },
        "Postgres": {"host": "", "port": "5432", "user": "", "password": "", "db_name": ""},
    }


# Dallas MySQL database names (dropdown in Config + Database Queries)
MYSQL_DALLAS_DATABASES = [
    "dallas_prod_aircontrol_db",
    "dallas_prod_cmp_interface_db",
    "dallas_prod_cmp_interface_thirdparty_db",
    "dallas_prod_pc_db",
    "dallas_prod_release_details_db",
    "dallas_prod_bss_main_db",
]

state = AppState()

# â”€â”€â”€ Config persistence â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
import json as _json, os as _os

CONFIG_FILE = _os.path.join(_os.path.dirname(_os.path.abspath(__file__)), "montool_config.json")

def save_config():
    try:
        data = {
            "log_lines":      state.log_lines,
            "module_configs": state.module_configs,
            "env_hosts":      state.env_hosts,
            "db_configs":     state.db_configs,
            "saved_queries":  state.saved_queries,
        }
        with open(CONFIG_FILE, "w") as f:
            _json.dump(data, f, indent=2)
    except Exception as e:
        print(f"[Config] Save failed: {e}")

def load_config():
    if not _os.path.exists(CONFIG_FILE):
        return
    try:
        with open(CONFIG_FILE) as f:
            data = _json.load(f)
        if "log_lines" in data:
            state.log_lines = int(data["log_lines"])
        if "env_hosts" in data:
            state.env_hosts.update(data["env_hosts"])
        if "module_configs" in data:
            for mod, cfg in data["module_configs"].items():
                if mod not in state.module_configs:
                    state.module_configs[mod] = _default_module_row()
                state.module_configs[mod].update(cfg)
                for key, default in [("sudo_user", "aqadmin"), ("sudo_password", "")]:
                    if key not in state.module_configs[mod]:
                        state.module_configs[mod][key] = default
        if "db_configs" in data:
            for db, cfg in data["db_configs"].items():
                if db in state.db_configs:
                    state.db_configs[db].update(cfg)
        if "saved_queries" in data and isinstance(data["saved_queries"], list):
            normalized = []
            for item in data["saved_queries"]:
                if not isinstance(item, dict):
                    continue
                name = str(item.get("name", "")).strip()
                sql = str(item.get("sql", "")).strip()
                if not name or not sql:
                    continue
                normalized.append({"name": name, "sql": sql})
            state.saved_queries = normalized[:10]
        print(f"[Config] Loaded from {CONFIG_FILE}")
    except Exception as e:
        print(f"[Config] Load failed: {e}")

load_config()

# ===== END FILE: config.py =====


# ===== BEGIN FILE: app_logging.py =====

"""
MonTool application run logging â€” one UTF-8 log file per process launch.
Use for troubleshooting and analysis (attach file or path in Cursor / support).
"""

from __future__ import annotations

import logging
import os
import sys
from datetime import datetime
from typing import Optional

_RUN_LOG_PATH: Optional[str] = None


def get_logs_directory() -> str:
    base = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(base, "application_logs")


def get_run_log_path() -> Optional[str]:
    """Absolute path to this session's log file, or None if setup not called."""
    return _RUN_LOG_PATH


def setup_application_logging() -> str:
    """
    Configure root logging to a single file per run: application_logs/run_YYYYMMDD_HHMMSS.log
    Returns the log file path.
    """
    global _RUN_LOG_PATH
    log_dir = get_logs_directory()
    os.makedirs(log_dir, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    _RUN_LOG_PATH = os.path.join(log_dir, f"run_{ts}.log")

    fmt = logging.Formatter(
        "%(asctime)s | %(levelname)-8s | %(name)s | %(message)s",
        datefmt="%Y-%m-%d %H:%M:%S",
    )
    fh = logging.FileHandler(_RUN_LOG_PATH, encoding="utf-8")
    fh.setLevel(logging.DEBUG)
    fh.setFormatter(fmt)

    root = logging.getLogger()
    root.handlers.clear()
    root.setLevel(logging.DEBUG)
    root.addHandler(fh)

    def _excepthook(exc_type, exc_value, exc_tb):
        logging.getLogger("uncaught").error(
            "Uncaught exception: %s", exc_value,
            exc_info=(exc_type, exc_value, exc_tb),
        )
        sys.__excepthook__(exc_type, exc_value, exc_tb)

    sys.excepthook = _excepthook

    logging.getLogger("app").info("Session log file: %s", _RUN_LOG_PATH)
    return _RUN_LOG_PATH


def get_logger(name: str) -> logging.Logger:
    return logging.getLogger(name)


def analysis_hint_text() -> str:
    """Short instructions for humans / AI when reviewing logs."""
    p = _RUN_LOG_PATH or "(not initialized)"
    return (
        "This session's application log is:\n"
        f"  {p}\n\n"
        "To analyse: open the file in an editor, or attach it in Cursor / email, "
        "or paste the full path above and ask to review errors and warnings."
    )

# ===== END FILE: app_logging.py =====


# ===== BEGIN FILE: ssh.py =====

import shlex
import time

try:
    import paramiko
    PARAMIKO_AVAILABLE = True
except ImportError:
    PARAMIKO_AVAILABLE = False

from config import state

try:
    from app_logging import get_logger as _get_app_logger
    _ssh_log = _get_app_logger("ssh")
except ImportError:
    _ssh_log = None

# â”€â”€â”€ SSH Manager â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
class SSHManager:
    def __init__(self):
        self.connections = {}   # key â†’ {"client": SSHClient, "sudo_password": str|None, "sudo_user": str|None}

    def is_connected(self, key):
        return key in self.connections

    def connect(self, host, user, password=None, key_path=None, port=22,
                actual_host=None, sudo_user=None, sudo_password=None):
        """
        Connect via SSH as `user` (LDAP).
        sudo_user     : if set, commands will be run as this user via `sudo su -`
        sudo_password : password for the sudo prompt (defaults to the LDAP password)
        """
        if not PARAMIKO_AVAILABLE:
            return False, "paramiko not installed â€” Demo Mode only"

        connect_host = actual_host if actual_host else host
        try:
            client = paramiko.SSHClient()
            client.set_missing_host_key_policy(paramiko.AutoAddPolicy())
            if key_path:
                client.connect(connect_host, port=port, username=user,
                               key_filename=key_path, timeout=10)
            else:
                client.connect(connect_host, port=port, username=user,
                               password=password, timeout=10)

            # Verify we can open a channel
            transport = client.get_transport()
            if transport is None or not transport.is_active():
                return False, "Transport not active after connect"

            self.connections[host] = {
                "client":         client,
                "sudo_user":      sudo_user or "",
                "sudo_password":  sudo_password if sudo_password is not None else (password or ""),
                "use_sudo":       bool(sudo_user and str(sudo_user).strip()),
            }
            if _ssh_log:
                _ssh_log.info(
                    "SSH connected: module=%s host=%s user=%s",
                    host, connect_host, user,
                )
            return True, "Connected"
        except Exception as e:
            if _ssh_log:
                _ssh_log.warning("SSH connect failed: module=%s host=%s err=%s", host, connect_host, e)
            return False, str(e)

    def run(self, host, cmd, timeout=30):
        """
        Run a command on the connected host.
        If sudo_user is configured, wraps the command via an interactive
        shell: sudo su - <sudo_user> -c '<cmd>'
        """
        if host not in self.connections:
            return "", "Not connected"

        info       = self.connections[host]
        client     = info["client"]
        sudo_user  = info.get("sudo_user", "").strip()
        sudo_pwd   = info.get("sudo_password", "")

        try:
            if sudo_user:
                return self._run_with_sudo(client, sudo_user, sudo_pwd, cmd, timeout)
            else:
                _, stdout, stderr = client.exec_command(cmd, timeout=timeout)
                return stdout.read().decode(errors="replace"), stderr.read().decode(errors="replace")
        except Exception as e:
            return "", str(e)

    def _run_with_sudo(self, client, sudo_user, sudo_pwd, cmd, timeout=30):
        """
        Use an interactive shell channel to:
          1. sudo su - <sudo_user>
          2. supply password when prompted
          3. run the actual command
          4. capture output
        """
        # Escape single quotes in the command
        safe_cmd = cmd.replace("'", "'\\''")
        full_cmd = f"sudo su - {sudo_user} -c '{safe_cmd}'\n"

        channel = client.invoke_shell(width=220, height=50)
        channel.settimeout(timeout)

        def _recv(wait=0.3):
            time.sleep(wait)
            out = b""
            while channel.recv_ready():
                out += channel.recv(65536)
            return out.decode(errors="replace")

        # Wait for initial shell prompt
        _recv(0.5)

        # Send the sudo command
        channel.send(full_cmd)

        # Wait for password prompt or immediate output
        deadline = time.time() + 10
        buf = ""
        while time.time() < deadline:
            chunk = _recv(0.3)
            buf += chunk
            lower = buf.lower()
            if "password" in lower or "password for" in lower:
                channel.send(sudo_pwd + "\n")
                break
            # If we already see command output (no sudo prompt), no password needed
            if len(buf.strip()) > len(full_cmd.strip()) + 5:
                break

        # Collect remaining output until shell goes quiet
        time.sleep(0.5)
        output = ""
        last_recv = time.time()
        while time.time() - last_recv < 2.0:
            chunk = _recv(0.3)
            if chunk:
                output += chunk
                last_recv = time.time()

        channel.close()

        # Strip shell echoes / prompts from the output
        lines = output.splitlines()
        cleaned = []
        skip_patterns = (sudo_user + "@", "$ ", "# ", full_cmd.strip(),
                         "sudo su", "[sudo]", "Password:")
        for line in lines:
            stripped = line.strip()
            if any(stripped.startswith(p) or stripped == p for p in skip_patterns):
                continue
            if not stripped:
                continue
            cleaned.append(line)

        return "\n".join(cleaned), ""

    def disconnect(self, host):
        if host in self.connections:
            self.connections[host]["client"].close()
            del self.connections[host]
            if _ssh_log:
                _ssh_log.info("SSH disconnected: module=%s", host)

    def open_tail_log_shell_chain(self, conn_key, log_path, su_password=None):
        """
        Interactive shell chain for log tail (matches server ops):
          1. Already connected as LDAP user (e.g. nandank).
          2. ``sudo su -`` and enter password when prompted.
          3. ``su - <sudo_user>`` (e.g. aqadmin) and enter password if prompted.
          4. ``tail -n 10 -F <log_path>``

        Returns (channel, None) on success, or (None, error_message).
        Caller must read from the channel until done; caller closes the channel.
        """
        if conn_key not in self.connections:
            return None, "Not connected"
        info = self.connections[conn_key]
        client = info.get("client")
        target = (info.get("sudo_user") or "").strip()
        pwd = (info.get("sudo_password") or "").strip()
        if not target:
            return None, "No sudo user (e.g. aqadmin) configured for this module"
        if not pwd:
            return None, "No password configured for sudo/su prompts"
        su_pwd = su_password if su_password is not None else pwd

        def drain(ch):
            out = b""
            while ch.recv_ready():
                out += ch.recv(65536)
            return out

        def feed_password_if_prompted(ch, password):
            time.sleep(0.45)
            buf = drain(ch)
            low = buf.decode(errors="replace").lower()
            # "password", "Password:", French "mot de passe", etc.
            if "assword" in low:
                ch.send(password + "\n")
                time.sleep(0.75)
                drain(ch)

        try:
            ch = client.invoke_shell(term="xterm", width=220, height=48)
            time.sleep(0.35)
            drain(ch)

            # 1) sudo su -  (become root)
            ch.send("sudo su -\n")
            feed_password_if_prompted(ch, pwd)

            # 2) su - <aqadmin>  (or configured sudo_user)
            # su - <aqadmin> (target comes from config; quote for safe shell word)
            ch.send(f"su - {shlex.quote(target)}\n")
            feed_password_if_prompted(ch, su_pwd)

            # 3) tail (path may contain spaces â€” quote for the shell)
            safe = shlex.quote(log_path)
            ch.send(f"tail -n 10 -F {safe} 2>&1\n")
            time.sleep(0.2)
            return ch, None
        except Exception as e:
            return None, str(e)

    def open_command_shell_chain(self, conn_key, command, su_password=None):
        """
        Open interactive shell and switch to configured sudo_user, then run command.
        Sequence:
          1) sudo su -
          2) su - <sudo_user>
          3) <command>
        Returns (channel, None) on success, or (None, error_message).
        Caller reads channel output and closes channel.
        """
        if conn_key not in self.connections:
            return None, "Not connected"
        info = self.connections[conn_key]
        client = info.get("client")
        target = (info.get("sudo_user") or "").strip()
        pwd = (info.get("sudo_password") or "").strip()
        if not target:
            return None, "No sudo user configured for this module"
        if not pwd:
            return None, "No password configured for sudo/su prompts"
        su_pwd = su_password if su_password is not None else pwd

        def drain(ch):
            out = b""
            while ch.recv_ready():
                out += ch.recv(65536)
            return out

        def feed_password_if_prompted(ch, password):
            time.sleep(0.45)
            buf = drain(ch)
            low = buf.decode(errors="replace").lower()
            if "assword" in low:
                ch.send(password + "\n")
                time.sleep(0.75)
                drain(ch)

        try:
            ch = client.invoke_shell(term="xterm", width=220, height=48)
            time.sleep(0.35)
            drain(ch)

            ch.send("sudo su -\n")
            feed_password_if_prompted(ch, pwd)

            ch.send(f"su - {shlex.quote(target)}\n")
            feed_password_if_prompted(ch, su_pwd)

            ch.send(command.rstrip() + "\n")
            time.sleep(0.2)
            return ch, None
        except Exception as e:
            return None, str(e)

    def connect_module(self, module):
        """Connect using per-module config from state. Returns (ok, msg)."""
        mcfg = state.module_configs.get(module, {})
        host = mcfg.get("host", "").strip()
        if not host:
            return False, "No host configured for this module"
        return self.connect(
            host=module,            # logical key (module name)
            actual_host=host,
            user=mcfg.get("user", "").strip(),
            password=mcfg.get("password", "").strip() or None,
            key_path=mcfg.get("key", "").strip() or None,
            port=int(mcfg.get("port", 22) or 22),
            sudo_user=mcfg.get("sudo_user", "").strip() or None,
            sudo_password=mcfg.get("sudo_password", "").strip() or mcfg.get("password", "").strip() or None,
        )

ssh_mgr = SSHManager()

# ===== END FILE: ssh.py =====


# ===== BEGIN FILE: db_mysql.py =====

"""
Direct MySQL TCP access (same idea as DBeaver / JDBC).
MonTool connects from this PC to host:port â€” no SSH hop.
"""
from __future__ import annotations

from typing import List, Optional, Tuple

MAX_ROWS_DEFAULT = 500
MAX_COL_WIDTH = 120

# Result from mysql_query_to_grid: (ok, columns, rows, info, error)
# info is set for non-SELECT success (e.g. rows affected); columns/rows are None then.
GridResult = Tuple[
    bool,
    Optional[List[str]],
    Optional[List[Tuple]],
    Optional[str],
    Optional[str],
]


def _normalize_sql(sql: str) -> str:
    s = sql.strip()
    # mysql CLI vertical format (\G) â€” not used by PyMySQL
    if s.endswith(r"\G"):
        s = s[:-2].rstrip()
    return s


def _fmt_cell(x) -> str:
    if x is None:
        return "NULL"
    s = str(x)
    if len(s) > MAX_COL_WIDTH:
        return s[: MAX_COL_WIDTH - 3] + "..."
    return s.replace("\t", " ").replace("\n", "\\n")


def _cell_display_dbeaver(x) -> str:
    """Display string for grid cells (NULL like DBeaver)."""
    if x is None:
        return "[NULL]"
    s = str(x)
    if len(s) > MAX_COL_WIDTH:
        return s[: MAX_COL_WIDTH - 3] + "..."
    return s.replace("\n", " ").replace("\r", " ")


def mysql_query_to_grid(
    host: str,
    port: str,
    user: str,
    password: str,
    database: str,
    sql: str,
    max_rows: int = MAX_ROWS_DEFAULT,
) -> GridResult:
    """
    Run one SQL statement.
    Returns (ok, columns, rows, info_message, error_message).
    - SELECT: columns + rows (each row tuple of display strings); info None.
    - Non-SELECT: columns/rows None; info has success text.
    - Failure: ok False, error set.
    """
    try:
        import pymysql
    except ImportError:
        return (
            False,
            None,
            None,
            None,
            "PyMySQL is not installed. From the Crux folder run:\n"
            "  pip install pymysql\n\n"
            "MonTool uses direct TCP to MySQL (like DBeaver), not SSH.",
        )

    host = (host or "").strip() or "127.0.0.1"
    try:
        port_i = int((port or "3306").strip() or "3306")
    except ValueError:
        return False, None, None, None, f"Invalid port: {port!r}"

    sql = _normalize_sql(sql)
    if not sql:
        return False, None, None, None, "Empty SQL"

    user = (user or "").strip()
    pwd = password or ""

    conn = None
    try:
        conn = pymysql.connect(
            host=host,
            port=port_i,
            user=user,
            password=pwd,
            database=(database or "").strip() or None,
            connect_timeout=25,
            read_timeout=180,
            charset="utf8mb4",
        )
    except Exception as e:
        return False, None, None, None, f"Connection failed: {e}"

    try:
        with conn.cursor() as cur:
            cur.execute(sql)
            if cur.description:
                colnames = [d[0] for d in cur.description]
                raw_rows = cur.fetchmany(max_rows + 1)
                truncated = len(raw_rows) > max_rows
                if truncated:
                    raw_rows = raw_rows[:max_rows]
                rows: List[Tuple] = [
                    tuple(_cell_display_dbeaver(v) for v in row) for row in raw_rows
                ]
                info = f"... truncated to {max_rows} rows" if truncated else None
                return True, colnames, rows, info, None
            return True, None, None, f"OK (rows affected: {cur.rowcount})", None
    except Exception as e:
        return False, None, None, None, f"Query error: {e}"
    finally:
        conn.close()


def mysql_query_to_text(
    host: str,
    port: str,
    user: str,
    password: str,
    database: str,
    sql: str,
    max_rows: int = MAX_ROWS_DEFAULT,
) -> Tuple[bool, Optional[str], Optional[str]]:
    """Plain-text TSV (legacy); prefer mysql_query_to_grid for UI tables."""
    ok, cols, rows, info, err = mysql_query_to_grid(
        host, port, user, password, database, sql, max_rows
    )
    if not ok:
        return False, None, err
    if cols is None or rows is None:
        return True, info or "", None
    lines = ["\t".join(str(c) for c in cols)]
    for row in rows:
        lines.append("\t".join(str(x) for x in row))
    out = "\n".join(lines)
    if info:
        out += f"\n\n{info}"
    return True, out, None

# ===== END FILE: db_mysql.py =====


# ===== BEGIN FILE: db_postgres.py =====

"""
Direct Postgres TCP access (same idea as DBeaver / JDBC).
"""
from __future__ import annotations

from typing import List, Optional, Tuple

MAX_ROWS_DEFAULT = 500
MAX_COL_WIDTH = 120

GridResult = Tuple[
    bool,
    Optional[List[str]],
    Optional[List[Tuple]],
    Optional[str],
    Optional[str],
]


def _cell_display(x) -> str:
    if x is None:
        return "[NULL]"
    s = str(x)
    if len(s) > MAX_COL_WIDTH:
        return s[: MAX_COL_WIDTH - 3] + "..."
    return s.replace("\n", " ").replace("\r", " ")


def postgres_query_to_grid(
    host: str,
    port: str,
    user: str,
    password: str,
    database: str,
    sql: str,
    max_rows: int = MAX_ROWS_DEFAULT,
) -> GridResult:
    try:
        import psycopg2
    except Exception:
        return (
            False,
            None,
            None,
            None,
            "psycopg2 is not installed. Run:\n  pip install psycopg2-binary",
        )

    host = (host or "").strip() or "127.0.0.1"
    try:
        port_i = int((port or "5432").strip() or "5432")
    except ValueError:
        return False, None, None, None, f"Invalid port: {port!r}"

    sql = (sql or "").strip()
    if not sql:
        return False, None, None, None, "Empty SQL"

    conn = None
    try:
        conn = psycopg2.connect(
            host=host,
            port=port_i,
            user=(user or "").strip(),
            password=password or "",
            dbname=(database or "").strip() or None,
            connect_timeout=25,
        )
    except Exception as e:
        return False, None, None, None, f"Connection failed: {e}"

    try:
        with conn.cursor() as cur:
            cur.execute(sql)
            if cur.description:
                cols = [d[0] for d in cur.description]
                raw_rows = cur.fetchmany(max_rows + 1)
                truncated = len(raw_rows) > max_rows
                if truncated:
                    raw_rows = raw_rows[:max_rows]
                rows = [tuple(_cell_display(v) for v in r) for r in raw_rows]
                info = f"... truncated to {max_rows} rows" if truncated else None
                return True, cols, rows, info, None
            conn.commit()
            return True, None, None, f"OK (rows affected: {cur.rowcount})", None
    except Exception as e:
        return False, None, None, None, f"Query error: {e}"
    finally:
        conn.close()


# ===== END FILE: db_postgres.py =====


# ===== BEGIN FILE: widgets.py =====

import customtkinter as ctk
from config import (BG_DARK, BG_PANEL, BG_CARD, BG_CARD2, ACCENT, ACCENT2,
                    SUCCESS, WARNING, DANGER, TEXT_PRIMARY, TEXT_MUTED, BORDER,
                    SIDEBAR_W, state)

SIDEBAR_ITEMS = [
    "Logs",
    "Database Queries",
    "Password Vault",
    "Mock & Sim Activation",
    "Invoice Pregeneration",
    "Consumption",
    "Sim Deletion",
    "Log Analysis",
]

# â”€â”€â”€ Widget helpers â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
def lbl(parent, text, size=13, weight="normal", color=TEXT_PRIMARY, **kw):
    return ctk.CTkLabel(parent, text=text, font=("Segoe UI", size, weight),
                        text_color=color, **kw)

def card(parent, **kw):
    d = dict(fg_color=BG_CARD, corner_radius=12, border_color=BORDER, border_width=1)
    d.update(kw)
    return ctk.CTkFrame(parent, **d)

def btn(parent, text, cmd, fg=ACCENT, hover=ACCENT2, tc="#000000", w=160, h=40, **kw):
    return ctk.CTkButton(parent, text=text, command=cmd,
                         fg_color=fg, hover_color=hover, text_color=tc,
                         font=("Segoe UI", 12, "bold"), corner_radius=8,
                         width=w, height=h, **kw)

def inp(parent, ph="", w=260, show="", **kw):
    return ctk.CTkEntry(parent, placeholder_text=ph, show=show,
                        fg_color=BG_CARD2, border_color=BORDER,
                        text_color=TEXT_PRIMARY, placeholder_text_color=TEXT_MUTED,
                        font=("Segoe UI", 12), width=w, height=36, **kw)

def hsep(parent):
    ctk.CTkFrame(parent, height=1, fg_color=BORDER).pack(fill="x", pady=6)

def badge(parent, text, color=ACCENT):
    f = ctk.CTkFrame(parent, fg_color=BG_CARD2, corner_radius=6,
                     border_color=color, border_width=1)
    ctk.CTkLabel(f, text=text, font=("Segoe UI", 10, "bold"),
                 text_color=color).pack(padx=8, pady=2)
    return f


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  SIDEBAR
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
class Sidebar(ctk.CTkFrame):
    def __init__(self, master, items, active=None, on_select=None, on_back=None, on_config=None):
        super().__init__(master, width=SIDEBAR_W, fg_color=BG_PANEL, corner_radius=0)
        self.pack_propagate(False)
        self.on_select = on_select

        # Header
        hdr = ctk.CTkFrame(self, fg_color=ACCENT, corner_radius=0, height=48)
        hdr.pack(fill="x")
        hdr.pack_propagate(False)
        ctk.CTkLabel(hdr, text="MONTOOL", font=("Courier New", 14, "bold"),
                     text_color="#000").pack(expand=True)

        # State badges
        bf = ctk.CTkFrame(self, fg_color="transparent")
        bf.pack(fill="x", padx=10, pady=6)
        if state.project:
            badge(bf, state.project, color=WARNING).pack(fill="x", pady=2)
        if state.env:
            badge(bf, state.env, color=SUCCESS).pack(fill="x", pady=2)

        hsep(self)

        # Menu buttons
        self.btns = {}
        for item in items:
            b = ctk.CTkButton(
                self, text=item, anchor="w",
                fg_color=BG_CARD2 if item == active else "transparent",
                hover_color=BG_CARD,
                text_color=ACCENT if item == active else TEXT_PRIMARY,
                font=("Segoe UI", 12, "bold" if item == active else "normal"),
                corner_radius=6, height=38,
                command=lambda i=item: self._sel(i)
            )
            b.pack(fill="x", padx=8, pady=2)
            self.btns[item] = b

        # Bottom controls
        bot = ctk.CTkFrame(self, fg_color="transparent")
        bot.pack(side="bottom", fill="x", padx=8, pady=10)
        if getattr(state, "evidence_panel_opener", None):
            ctk.CTkButton(
                bot,
                text="ðŸ“·  Evidence",
                fg_color=BG_CARD2,
                hover_color=BG_CARD,
                text_color=ACCENT,
                font=("Segoe UI", 11, "bold"),
                corner_radius=6,
                height=32,
                command=state.evidence_panel_opener,
            ).pack(fill="x", pady=2)
        if on_config:
            ctk.CTkButton(bot, text="âš™  Config", fg_color=BG_CARD2, hover_color=BG_CARD,
                          text_color=TEXT_MUTED, font=("Segoe UI", 11), corner_radius=6,
                          height=32, command=on_config).pack(fill="x", pady=2)
        if on_back:
            ctk.CTkButton(bot, text="â† Back", fg_color="transparent", hover_color=BG_CARD,
                          text_color=TEXT_MUTED, font=("Segoe UI", 11), corner_radius=6,
                          height=32, command=on_back).pack(fill="x", pady=2)

    def _sel(self, item):
        for k, b in self.btns.items():
            active = k == item
            b.configure(fg_color=BG_CARD2 if active else "transparent",
                        text_color=ACCENT if active else TEXT_PRIMARY,
                        font=("Segoe UI", 12, "bold" if active else "normal"))
        if self.on_select:
            self.on_select(item)


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  CENTERED CONTAINER helper (replaces place())
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
def centered(parent):
    """Returns a centered column frame using pure pack â€” works on Windows."""
    outer = ctk.CTkFrame(parent, fg_color="transparent")
    outer.pack(fill="both", expand=True)
    outer.columnconfigure(0, weight=1)
    outer.rowconfigure(0, weight=1)
    inner = ctk.CTkFrame(outer, fg_color="transparent")
    inner.grid(row=0, column=0)
    return outer, inner


# ===== END FILE: widgets.py =====


# ===== BEGIN FILE: evidence_capture.py =====

"""
Floating evidence-capture panel (screenshots + Word evidence report), draggable,
always-on-top.
"""
from __future__ import annotations

import os
import subprocess
import sys
import time
from datetime import datetime

import customtkinter as ctk
import tkinter as tk
from tkinter import filedialog, messagebox

from app_logging import get_logger
from config import (
    ACCENT,
    ACCENT2,
    BG_CARD,
    BG_CARD2,
    BG_DARK,
    BG_PANEL,
    DANGER,
    TEXT_MUTED,
    TEXT_PRIMARY,
    state,
)

_log = get_logger("evidence")


class _CommentDialog(ctk.CTkToplevel):
    def __init__(self, master, title="Capture Comment"):
        super().__init__(master)
        self.title(title)
        self.geometry("520x280")
        self.minsize(420, 240)
        self.configure(fg_color=BG_DARK)
        self.attributes("-topmost", True)
        self.transient(master)
        self.grab_set()

        self.result = None

        ctk.CTkLabel(
            self,
            text="Add comment for this screenshot:",
            text_color=TEXT_PRIMARY,
            font=("Segoe UI", 12, "bold"),
            anchor="w",
        ).pack(fill="x", padx=14, pady=(12, 6))

        self.text = tk.Text(
            self,
            bg=BG_CARD,
            fg=TEXT_PRIMARY,
            insertbackground=ACCENT,
            relief="flat",
            font=("Segoe UI", 11),
            padx=8,
            pady=6,
            wrap="word",
        )
        self.text.pack(fill="both", expand=True, padx=14, pady=(0, 8))
        self.text.focus_set()
        self.text.bind("<Return>", self._on_enter_save)

        row = ctk.CTkFrame(self, fg_color="transparent")
        row.pack(fill="x", padx=14, pady=(0, 12))
        ctk.CTkButton(
            row,
            text="Cancel",
            fg_color=BG_CARD2,
            hover_color=BG_CARD,
            text_color=TEXT_MUTED,
            width=90,
            command=self._cancel,
        ).pack(side="right", padx=(6, 0))
        ctk.CTkButton(
            row,
            text="Save Comment",
            fg_color=ACCENT,
            hover_color=ACCENT2,
            text_color="#000",
            width=120,
            command=self._ok,
        ).pack(side="right")

        self.protocol("WM_DELETE_WINDOW", self._cancel)

    def _ok(self):
        self.result = self.text.get("1.0", "end").strip()
        self.destroy()

    def _on_enter_save(self, event):
        self._ok()
        return "break"

    def _cancel(self):
        self.result = None
        self.destroy()


class _AnnotateDialog(ctk.CTkToplevel):
    """Draw rectangular highlight boxes on screenshot before saving."""

    def __init__(self, master, image_path: str):
        super().__init__(master)
        self.title("Annotate Screenshot")
        self._set_fullscreen_like()
        self.minsize(900, 640)
        self.configure(fg_color=BG_DARK)
        self.attributes("-topmost", True)
        self.transient(master)
        self.grab_set()

        self.result = None
        self._image_path = image_path
        self._rects = []
        self._start_x = None
        self._start_y = None
        self._cur_rect = None

        try:
            from PIL import Image, ImageDraw, ImageTk  # noqa: PLC0415
        except Exception as e:
            messagebox.showerror(
                "Missing dependency",
                f"Pillow is required for annotation.\n\n{e}",
                parent=self,
            )
            self.destroy()
            return

        self._Image = Image
        self._ImageDraw = ImageDraw
        self._ImageTk = ImageTk
        self._orig_img = Image.open(image_path).convert("RGB")
        self._orig_w, self._orig_h = self._orig_img.size

        max_w, max_h = 1120, 650
        scale = min(max_w / self._orig_w, max_h / self._orig_h, 1.0)
        self._disp_w = int(self._orig_w * scale)
        self._disp_h = int(self._orig_h * scale)
        self._scale_x = self._orig_w / self._disp_w
        self._scale_y = self._orig_h / self._disp_h
        self._disp_img = self._orig_img.resize((self._disp_w, self._disp_h))
        self._tk_img = self._ImageTk.PhotoImage(self._disp_img)

        ctk.CTkLabel(
            self,
            text="Drag to draw square/rectangle highlights. Save to apply.",
            text_color=TEXT_MUTED,
            font=("Segoe UI", 11),
        ).pack(fill="x", padx=12, pady=(10, 4))

        self._canvas = tk.Canvas(
            self,
            width=self._disp_w,
            height=self._disp_h,
            bg=BG_CARD,
            highlightthickness=1,
            highlightbackground=BG_CARD2,
            relief="flat",
        )
        self._canvas.pack(padx=12, pady=6, fill="both", expand=True)
        self._canvas.create_image(0, 0, anchor="nw", image=self._tk_img)
        self._canvas.bind("<ButtonPress-1>", self._start_box)
        self._canvas.bind("<B1-Motion>", self._drag_box)
        self._canvas.bind("<ButtonRelease-1>", self._end_box)

        row = ctk.CTkFrame(self, fg_color="transparent")
        row.pack(fill="x", padx=12, pady=(0, 12))
        ctk.CTkButton(
            row,
            text="Undo Last",
            width=100,
            fg_color=BG_CARD2,
            hover_color=BG_CARD,
            text_color=TEXT_PRIMARY,
            command=self._undo_last,
        ).pack(side="left")
        ctk.CTkButton(
            row,
            text="Skip Annotation",
            width=140,
            fg_color=BG_CARD2,
            hover_color=BG_CARD,
            text_color=TEXT_MUTED,
            command=self._skip,
        ).pack(side="right", padx=(6, 0))
        ctk.CTkButton(
            row,
            text="Save Annotation",
            width=140,
            fg_color=ACCENT,
            hover_color=ACCENT2,
            text_color="#000",
            command=self._save,
        ).pack(side="right")

        self.protocol("WM_DELETE_WINDOW", self._skip)

    def _set_fullscreen_like(self):
        # Windows: maximize to full screen work area.
        # Fallback to exact screen geometry for other platforms.
        try:
            self.state("zoomed")
            return
        except Exception:
            pass
        try:
            w = self.winfo_screenwidth()
            h = self.winfo_screenheight()
            self.geometry(f"{w}x{h}+0+0")
        except Exception:
            self.geometry("1180x780")

    def _start_box(self, event):
        self._start_x, self._start_y = event.x, event.y
        self._cur_rect = self._canvas.create_rectangle(
            event.x, event.y, event.x, event.y,
            outline="#ff3355", width=3
        )

    def _drag_box(self, event):
        if self._cur_rect is None:
            return
        self._canvas.coords(self._cur_rect, self._start_x, self._start_y, event.x, event.y)

    def _end_box(self, event):
        if self._cur_rect is None:
            return
        x1, y1, x2, y2 = self._canvas.coords(self._cur_rect)
        if abs(x2 - x1) < 6 or abs(y2 - y1) < 6:
            self._canvas.delete(self._cur_rect)
        else:
            self._rects.append(self._cur_rect)
        self._cur_rect = None

    def _undo_last(self):
        if not self._rects:
            return
        rid = self._rects.pop()
        self._canvas.delete(rid)

    def _skip(self):
        self.result = self._image_path
        self.destroy()

    def _save(self):
        if not self._rects:
            self.result = self._image_path
            self.destroy()
            return
        draw = self._ImageDraw.Draw(self._orig_img)
        for rid in self._rects:
            x1, y1, x2, y2 = self._canvas.coords(rid)
            ox1 = int(min(x1, x2) * self._scale_x)
            oy1 = int(min(y1, y2) * self._scale_y)
            ox2 = int(max(x1, x2) * self._scale_x)
            oy2 = int(max(y1, y2) * self._scale_y)
            draw.rectangle([ox1, oy1, ox2, oy2], outline="#ff3355", width=6)
        self._orig_img.save(self._image_path, "PNG")
        self.result = self._image_path
        self.destroy()


def get_evidence_root() -> str:
    base = os.path.dirname(os.path.abspath(__file__))
    d = os.path.join(base, "evidence_capture")
    os.makedirs(d, exist_ok=True)
    return d


def _capture_screen_png(dest_path: str) -> tuple[bool, str]:
    try:
        from PIL import ImageGrab

        try:
            img = ImageGrab.grab(all_screens=True)
        except TypeError:
            img = ImageGrab.grab()
        img.save(dest_path, "PNG")
        return True, dest_path
    except Exception as e:
        _log.warning("Screenshot failed: %s", e)
        return False, str(e)


class EvidenceFloatPanel(ctk.CTkToplevel):
    """Small floating window: drag top bar to move; stays above windows."""

    def __init__(self, master):
        super().__init__(master)
        self.title("MonTool â€” Evidence")
        self.geometry("320x420+120+120")
        self.minsize(300, 360)
        self.configure(fg_color=BG_DARK)
        self.attributes("-topmost", True)
        self.resizable(True, True)

        self._drag_start_x = 0
        self._drag_start_y = 0
        self._win_start_x = 0
        self._win_start_y = 0
        self._topmost = True

        self._report_path = ""
        self._capture_dir = ""
        self._capture_count = 0

        self._build_ui()
        self.protocol("WM_DELETE_WINDOW", self._close)

        self._drag_handle.bind("<ButtonPress-1>", self._start_drag)
        self._drag_handle.bind("<B1-Motion>", self._on_drag)

    def _build_ui(self) -> None:
        self._drag_handle = ctk.CTkFrame(self, fg_color=BG_PANEL, corner_radius=0, height=40)
        self._drag_handle.pack(fill="x")
        self._drag_handle.pack_propagate(False)

        ctk.CTkLabel(
            self._drag_handle,
            text="â‹®â‹®  Evidence capture  (drag to move)",
            font=("Segoe UI", 11, "bold"),
            text_color=TEXT_MUTED,
            anchor="w",
        ).pack(side="left", padx=12, pady=10)

        ctk.CTkButton(
            self._drag_handle,
            text="âœ•",
            width=36,
            height=28,
            fg_color=DANGER,
            hover_color="#cc3355",
            text_color="#fff",
            font=("Segoe UI", 12, "bold"),
            corner_radius=6,
            command=self._close,
        ).pack(side="right", padx=8, pady=6)

        pad = ctk.CTkFrame(self, fg_color="transparent")
        pad.pack(fill="both", expand=True, padx=14, pady=12)

        self._topmost_var = ctk.BooleanVar(value=True)

        def toggle_topmost():
            self._topmost = self._topmost_var.get()
            self.attributes("-topmost", self._topmost)

        ctk.CTkCheckBox(
            pad,
            text="Stay on top",
            variable=self._topmost_var,
            command=toggle_topmost,
            font=("Segoe UI", 11),
            text_color=TEXT_MUTED,
            fg_color=ACCENT,
            hover_color=ACCENT2,
        ).pack(anchor="w", pady=(0, 10))

        btn_kw = dict(
            fg_color=ACCENT,
            hover_color=ACCENT2,
            text_color="#000000",
            font=("Segoe UI", 12, "bold"),
            corner_radius=10,
            height=44,
        )
        ctk.CTkButton(
            pad,
            text="ðŸ“¸  Capture",
            command=self._do_capture,
            **btn_kw,
        ).pack(fill="x", pady=6)
        ctk.CTkButton(
            pad,
            text="ðŸ“  New report",
            command=self._new_report,
            **btn_kw,
        ).pack(fill="x", pady=6)
        ctk.CTkButton(
            pad,
            text="ðŸ’¾  Save & Close report",
            command=self._save_close_report,
            fg_color=BG_CARD2,
            hover_color=BG_CARD,
            text_color=TEXT_PRIMARY,
            font=("Segoe UI", 12, "bold"),
            corner_radius=10,
            height=40,
        ).pack(fill="x", pady=6)
        ctk.CTkButton(
            pad,
            text="ðŸ“‚  Open evidence folder",
            command=self._open_folder,
            fg_color=BG_CARD2,
            hover_color=BG_CARD,
            text_color=TEXT_PRIMARY,
            font=("Segoe UI", 12, "bold"),
            corner_radius=10,
            height=40,
        ).pack(fill="x", pady=6)

        self._active_report_lbl = ctk.CTkLabel(
            pad,
            text="Report: (none)",
            font=("Consolas", 10),
            text_color=TEXT_MUTED,
            wraplength=280,
            anchor="w",
            justify="left",
        )
        self._active_report_lbl.pack(fill="x", pady=(8, 0))

        self._status = ctk.CTkLabel(
            pad,
            text="",
            font=("Segoe UI", 10),
            text_color=TEXT_MUTED,
            wraplength=280,
            anchor="w",
            justify="left",
        )
        self._status.pack(fill="x", pady=(6, 0))

    def _start_drag(self, event) -> None:
        self._drag_start_x = event.x_root
        self._drag_start_y = event.y_root
        self._win_start_x = self.winfo_x()
        self._win_start_y = self.winfo_y()

    def _on_drag(self, event) -> None:
        dx = event.x_root - self._drag_start_x
        dy = event.y_root - self._drag_start_y
        self.geometry(f"+{self._win_start_x + dx}+{self._win_start_y + dy}")

    def _set_status(self, text: str) -> None:
        self._status.configure(text=text)
        _log.info("%s", text)

    def _set_report(self, path: str, capture_dir: str, count: int = 0) -> None:
        self._report_path = path
        self._capture_dir = capture_dir
        self._capture_count = count
        if path:
            self._active_report_lbl.configure(text=f"Report: {path}")
        else:
            self._active_report_lbl.configure(text="Report: (none)")

    def _ensure_docx(self):
        try:
            from docx import Document  # noqa: PLC0415
            from docx.shared import Inches  # noqa: PLC0415
            return Document, Inches, None
        except Exception as e:
            msg = (
                "python-docx is not installed.\n"
                "Install once:\n  pip install python-docx\n\n"
                f"Details: {e}"
            )
            return None, None, msg

    def _ask_capture_comment(self) -> str | None:
        dlg = _CommentDialog(self)
        self.wait_window(dlg)
        return dlg.result

    def _ask_annotation(self, image_path: str) -> str | None:
        dlg = _AnnotateDialog(self, image_path=image_path)
        self.wait_window(dlg)
        return dlg.result

    def _new_report(self) -> None:
        Document, _, err = self._ensure_docx()
        if err:
            messagebox.showerror("Missing dependency", err, parent=self)
            return

        default_dir = get_evidence_root()
        default_name = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        path = filedialog.asksaveasfilename(
            parent=self,
            title="Save Evidence Report",
            initialdir=default_dir,
            initialfile=default_name,
            defaultextension=".docx",
            filetypes=[("Word document", "*.docx")],
        )
        if not path:
            return

        path = os.path.abspath(path)
        if not path.lower().endswith(".docx"):
            path += ".docx"

        base_no_ext = os.path.splitext(path)[0]
        capture_dir = f"{base_no_ext}_images"
        os.makedirs(capture_dir, exist_ok=True)

        proj = getattr(state, "project", None) or "(not set)"
        env = getattr(state, "env", None) or "(not set)"
        try:
            doc = Document()
            doc.add_heading("MonTool Evidence Report", level=1)
            doc.add_paragraph(f"Generated: {datetime.now().isoformat(timespec='seconds')}")
            doc.add_paragraph(f"Project: {proj}")
            doc.add_paragraph(f"Environment: {env}")
            doc.add_paragraph(" ")
            doc.add_paragraph("Screenshots and comments")
            doc.save(path)
        except Exception as e:
            messagebox.showerror("Report creation failed", str(e), parent=self)
            return

        self._set_report(path, capture_dir, count=0)
        self._set_status(f"Report created:\n{path}")

    def _do_capture(self) -> None:
        if not self._report_path:
            messagebox.showinfo(
                "No active report",
                "Click 'New report' first and choose where to save the Word document.",
                parent=self,
            )
            return

        Document, Inches, err = self._ensure_docx()
        if err:
            messagebox.showerror("Missing dependency", err, parent=self)
            return

        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        os.makedirs(self._capture_dir, exist_ok=True)
        path = os.path.join(self._capture_dir, f"capture_{ts}.png")

        # Hide the floating panel so it never appears in the screenshot.
        try:
            self.withdraw()
            self.update_idletasks()
            time.sleep(0.2)
        except Exception:
            pass

        ok, msg = _capture_screen_png(path)

        try:
            self.deiconify()
            self.lift()
            if self._topmost:
                self.attributes("-topmost", True)
            self.focus_force()
        except Exception:
            pass

        if not ok:
            hint = ""
            if "PIL" in msg or "No module named" in msg:
                hint = "\nInstall: pip install Pillow"
            messagebox.showerror("Capture failed", f"{msg}{hint}", parent=self)
            self._set_status("Capture failed â€” see dialog")
            return

        annotated_path = self._ask_annotation(path)
        if annotated_path is None:
            try:
                os.remove(path)
            except Exception:
                pass
            self._set_status("Capture cancelled (annotation skipped with cancel)")
            return

        comment = self._ask_capture_comment()
        if comment is None:
            try:
                os.remove(path)
            except Exception:
                pass
            self._set_status("Capture cancelled (comment dialog cancelled)")
            return

        if not comment.strip():
            comment = "(No comment provided)"

        try:
            doc = Document(self._report_path)
            self._capture_count += 1
            doc.add_paragraph(
                f"Capture {self._capture_count} â€” {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            )
            doc.add_picture(annotated_path, width=Inches(6.5))
            doc.add_paragraph(f"Comment: {comment}")
            doc.add_paragraph(" ")
            doc.save(self._report_path)
        except Exception as e:
            messagebox.showerror("Report update failed", str(e), parent=self)
            self._capture_count -= 1
            return

        self._set_status(
            f"Capture #{self._capture_count} added to report.\n"
            f"Image: {path}"
        )

    def _save_close_report(self) -> None:
        if not self._report_path:
            messagebox.showinfo("No active report", "No report is currently open.", parent=self)
            return

        Document, _, err = self._ensure_docx()
        if err:
            messagebox.showerror("Missing dependency", err, parent=self)
            return

        try:
            doc = Document(self._report_path)
            doc.add_paragraph(
                f"Report closed at {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            )
            doc.save(self._report_path)
        except Exception as e:
            messagebox.showerror("Save failed", str(e), parent=self)
            return

        closed_path = self._report_path
        self._set_report("", "", count=0)
        self._set_status(f"Saved and closed report:\n{closed_path}")

    def _open_folder(self) -> None:
        path = get_evidence_root()
        try:
            if sys.platform == "win32":
                os.startfile(path)  # noqa: S606
            elif sys.platform == "darwin":
                subprocess.run(["open", path], check=False)
            else:
                subprocess.run(["xdg-open", path], check=False)
            self._set_status(f"Opened:\n{path}")
        except Exception as e:
            messagebox.showerror("Open folder failed", str(e), parent=self)

    def _close(self) -> None:
        if self._report_path:
            close_now = messagebox.askyesno(
                "Close Evidence Panel",
                "A report is still active. Save & close it before exiting panel?",
                parent=self,
            )
            if close_now:
                self._save_close_report()

        try:
            self.destroy()
        except Exception:
            pass
        app = self.master
        if app is not None and hasattr(app, "_evidence_panel_cleared"):
            app._evidence_panel_cleared()

# ===== END FILE: evidence_capture.py =====


# ===== BEGIN FILE: page_setup.py =====

import customtkinter as ctk
from config import (BG_DARK, BG_CARD, BG_CARD2, ACCENT, ACCENT2,
                    SUCCESS, WARNING, DANGER, TEXT_PRIMARY, TEXT_MUTED, BORDER,
                    BG_PANEL, state)
from widgets import lbl, hsep, badge, centered, btn

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  PAGE: Select Project
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
class PageSelectProject(ctk.CTkFrame):
    def __init__(self, master, on_select):
        super().__init__(master, fg_color=BG_DARK, corner_radius=0)
        self.on_select = on_select
        self._build()

    def _build(self):
        outer, c = centered(self)

        ctk.CTkLabel(c, text="MONTOOL",
                     font=("Courier New", 42, "bold"), text_color=ACCENT).pack(pady=(0, 4))
        ctk.CTkLabel(c, text="Environment Monitoring Tool  â€¢  Developed by Nandan",
                     font=("Segoe UI", 13), text_color=TEXT_MUTED).pack(pady=(0, 40))

        lbl(c, "SELECT PROJECT", size=11, color=TEXT_MUTED).pack()
        hsep(c)

        row = ctk.CTkFrame(c, fg_color="transparent")
        row.pack(pady=20)

        for proj in ["AMX", "STC"]:
            f = ctk.CTkFrame(row, fg_color=BG_CARD, corner_radius=16,
                             border_color=BORDER, border_width=1)
            f.pack(side="left", padx=16, pady=4)

            ctk.CTkLabel(f, text=proj, font=("Courier New", 32, "bold"),
                         text_color=ACCENT).pack(padx=40, pady=(30, 10))
            ctk.CTkButton(f, text="Select â†’",
                          command=lambda p=proj: self.on_select(p),
                          fg_color=ACCENT, hover_color=ACCENT2, text_color="#000",
                          font=("Segoe UI", 12, "bold"), corner_radius=8,
                          width=140, height=36).pack(padx=20, pady=(0, 24))

            f.bind("<Enter>", lambda e, w=f: w.configure(border_color=ACCENT))
            f.bind("<Leave>", lambda e, w=f: w.configure(border_color=BORDER))


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  PAGE: Select Environment
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
class PageSelectEnv(ctk.CTkFrame):
    def __init__(self, master, on_select, on_back):
        super().__init__(master, fg_color=BG_DARK, corner_radius=0)
        self.on_select = on_select
        self.on_back   = on_back
        self._build()

    def _build(self):
        outer, c = centered(self)

        ctk.CTkLabel(c, text=state.project,
                     font=("Courier New", 36, "bold"), text_color=ACCENT).pack(pady=(0, 4))
        lbl(c, "Choose Environment", size=14, color=TEXT_MUTED).pack(pady=(0, 32))

        envs = [("Dallas Production", DANGER), ("15 Million", WARNING), ("PreProd-2", SUCCESS), ("Miami Production", ACCENT)]
        row = ctk.CTkFrame(c, fg_color="transparent")
        row.pack()

        for env, color in envs:
            f = ctk.CTkFrame(row, fg_color=BG_CARD, corner_radius=14,
                             border_color=color, border_width=1)
            f.pack(side="left", padx=10, pady=4)

            ctk.CTkLabel(f, text=env, font=("Segoe UI", 14, "bold"),
                         text_color=color).pack(padx=30, pady=(24, 10))
            ctk.CTkButton(f, text="Enter",
                          command=lambda e=env: self.on_select(e),
                          fg_color=BG_CARD2, hover_color=BG_CARD,
                          text_color=color, font=("Segoe UI", 11, "bold"),
                          corner_radius=8, width=110, height=30,
                          border_color=color, border_width=1).pack(padx=16, pady=(0, 20))

        ctk.CTkButton(c, text="â† Back", command=self.on_back,
                      fg_color="transparent", hover_color=BG_CARD,
                      text_color=TEXT_MUTED, font=("Segoe UI", 12),
                      corner_radius=8, width=100, height=34).pack(pady=(24, 0))


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  PAGE: Dashboard
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
class PageDashboard(ctk.CTkFrame):
    ITEMS = [
        ("Evidence Capture",      "ðŸ“‚", "Floating panel: screenshots & reports"),
        ("Logs",                  "ðŸ“‹", "View application & system logs"),
        ("Database Queries",      "ðŸ—„",  "Run MySQL / Postgres queries"),
        ("Password Vault",        "ðŸ”", "Store and manage credentials"),
        ("Mock & Sim Activation", "ðŸ”§", "Upload & activate mock files"),
        ("Invoice Pregeneration", "ðŸ’³", "Generate invoices in bulk"),
        ("Consumption",           "ðŸ“Š", "Voice / SMS / Data usage lookup"),
        ("Sim Deletion",          "ðŸ—‘", "Remove SIMs from inventory / HLR"),
        ("Log Analysis",          "ðŸ“ˆ", "MonTool session logs & analysis hints"),
    ]

    def __init__(self, master, on_select, on_back):
        super().__init__(master, fg_color=BG_DARK, corner_radius=0)
        self.on_select = on_select
        self.on_back   = on_back
        self._build()

    def _build(self):
        # Top bar
        hdr = ctk.CTkFrame(self, fg_color=BG_PANEL, corner_radius=0, height=52)
        hdr.pack(fill="x")
        hdr.pack_propagate(False)

        ctk.CTkLabel(hdr, text="MONTOOL",
                     font=("Courier New", 16, "bold"), text_color=ACCENT).pack(side="left", padx=16, pady=12)
        badge(hdr, state.project, color=WARNING).pack(side="left", padx=4, pady=16)
        badge(hdr, state.env,     color=SUCCESS).pack(side="left", padx=4, pady=16)
        ctk.CTkButton(hdr, text="â† Back", command=self.on_back,
                      fg_color="transparent", hover_color=BG_CARD,
                      text_color=TEXT_MUTED, font=("Segoe UI", 11),
                      corner_radius=6, width=80, height=28).pack(side="right", padx=16)

        # Grid of cards
        body = ctk.CTkFrame(self, fg_color="transparent")
        body.pack(fill="both", expand=True, padx=32, pady=24)

        cols = 3
        for i, (name, icon, desc) in enumerate(self.ITEMS):
            r, col = divmod(i, cols)
            f = ctk.CTkFrame(body, fg_color=BG_CARD, corner_radius=14,
                             border_color=BORDER, border_width=1)
            f.grid(row=r, column=col, padx=10, pady=10, sticky="nsew")
            body.grid_columnconfigure(col, weight=1)
            body.grid_rowconfigure(r, weight=1)

            ctk.CTkLabel(f, text=icon, font=("Segoe UI", 28)).pack(pady=(20, 4))
            ctk.CTkLabel(f, text=name, font=("Segoe UI", 13, "bold"),
                         text_color=TEXT_PRIMARY).pack()
            ctk.CTkLabel(f, text=desc, font=("Segoe UI", 10),
                         text_color=TEXT_MUTED, wraplength=170).pack(pady=(4, 10))
            ctk.CTkButton(f, text="Open", command=lambda n=name: self.on_select(n),
                          fg_color=ACCENT, hover_color=ACCENT2, text_color="#000",
                          font=("Segoe UI", 12, "bold"), corner_radius=8,
                          width=120, height=34).pack(pady=(0, 18))

            f.bind("<Enter>", lambda e, w=f: w.configure(border_color=ACCENT))
            f.bind("<Leave>", lambda e, w=f: w.configure(border_color=BORDER))


# ===== END FILE: page_setup.py =====


# ===== BEGIN FILE: page_logs.py =====

import customtkinter as ctk
import tkinter as tk
from tkinter import filedialog
from tkinter import messagebox
from tkinter import simpledialog
import threading
import shlex
from datetime import datetime
from config import (BG_DARK, BG_PANEL, BG_CARD, BG_CARD2, ACCENT, ACCENT2,
                    SUCCESS, WARNING, DANGER, TEXT_PRIMARY, TEXT_MUTED, BORDER,
                    state, LOG_MODULES, save_config)
from ssh import ssh_mgr
from widgets import lbl, btn, Sidebar, SIDEBAR_ITEMS


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  PAGE: Logs
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
class PageLogs(ctk.CTkFrame):
    """Live streaming log viewer â€” tail -F per module, fully isolated per tab."""
    MAX_LINES = 2000
    FONT_MIN = 6
    FONT_MAX = 28
    FONT_DEFAULT = 10

    def __init__(self, master, on_nav, on_back, on_config):
        super().__init__(master, fg_color=BG_DARK, corner_radius=0)
        self.on_nav         = on_nav
        extra_tabs = [k for k in state.module_configs.keys() if k not in LOG_MODULES]
        self.TABS           = list(LOG_MODULES) + sorted(extra_tabs)
        self.active_tab     = self.TABS[0]
        self._tab_stream_stop = {t: threading.Event() for t in self.TABS}
        self._tab_stream_threads = {t: None for t in self.TABS}
        self._line_counts   = {t: 0 for t in self.TABS}
        self._auto_scroll   = True
        self._alive         = True          # set False when page is destroyed
        self._grep_term     = ""
        self._grep_filter   = False
        self._log_font_pt   = self.FONT_DEFAULT
        self._log_texts     = {}            # tab â†’ tk.Text (stacked; one visible)
        self._sy = None                   # vertical scrollbar only (log lines wrap)
        self._conn_fields   = {}
        self._conn_config_open = False
        self.k8_namespace_var = ctk.StringVar(value="default")
        self.k8_pod_var = ctk.StringVar(value="")
        self.k8_tail_lines_var = ctk.StringVar(value="200")
        self._build(on_back, on_config)
        self.bind("<Destroy>", self._on_destroy)

    def _on_destroy(self, e):
        self._alive = False
        self._stop_all_streams()

    def _config_log_tags(self, w):
        w.tag_config("ERR",   foreground=DANGER)
        w.tag_config("WARN",  foreground=WARNING)
        w.tag_config("INFO",  foreground=SUCCESS)
        w.tag_config("DEBUG", foreground=TEXT_MUTED)
        w.tag_config("HIT",   background="#5C4400")
        w.tag_config("GREP",  background="#003060", foreground="#FFD700")

    def _make_log_text(self, parent):
        w = tk.Text(parent, bg="#0A0C10", fg=TEXT_PRIMARY,
                    font=("Consolas", self._log_font_pt),
                    wrap="char", relief="flat", padx=10, pady=8,
                    insertbackground=ACCENT, selectbackground=ACCENT2,
                    state="disabled")
        self._config_log_tags(w)
        return w

    def _bind_scrollbars_to(self, w):
        if self._sy is None:
            return
        self._sy.configure(command=w.yview)
        w.configure(yscrollcommand=self._sy.set)

    def _show_log_tab(self, tab):
        """Raise the log Text for this tab and point self.log at it."""
        if tab not in self._log_texts:
            return
        self.active_tab = tab
        self._log_texts[tab].lift()
        self.log = self._log_texts[tab]
        self._bind_scrollbars_to(self.log)

    def _apply_font_zoom(self):
        for w in self._log_texts.values():
            try:
                w.configure(font=("Consolas", self._log_font_pt))
            except Exception:
                pass
        if hasattr(self, "zoom_lbl"):
            self.zoom_lbl.configure(text=f"{self._log_font_pt} pt")

    def _zoom_delta(self, delta):
        self._log_font_pt = max(self.FONT_MIN, min(self.FONT_MAX, self._log_font_pt + delta))
        self._apply_font_zoom()

    def _build(self, on_back, on_config):
        Sidebar(self, SIDEBAR_ITEMS, active="Logs",
                on_select=self._nav_and_stop, on_back=on_back,
                on_config=on_config).pack(side="left", fill="y")

        main = ctk.CTkFrame(self, fg_color=BG_DARK, corner_radius=0)
        main.pack(side="left", fill="both", expand=True)

        # â”€â”€ Tab bar (improved look + add tab action) â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
        tb_outer = ctk.CTkFrame(
            main,
            fg_color=BG_PANEL,
            corner_radius=8,
            height=52,
            border_width=1,
            border_color=BORDER,
        )
        tb_outer.pack(side="top", fill="x")
        tb_outer.pack_propagate(False)
        tb = ctk.CTkScrollableFrame(tb_outer, fg_color=BG_PANEL, height=46,
                                    orientation="horizontal", corner_radius=0)
        tb.pack(side="left", fill="both", expand=True, padx=(4, 2), pady=2)
        self._tab_btn_wrap = tb

        btn(
            tb_outer,
            "+ Add Connection Tab",
            self._add_connection_tab,
            fg=BG_CARD2,
            hover=BG_CARD,
            tc=TEXT_PRIMARY,
            w=170,
            h=32,
        ).pack(side="right", padx=8, pady=9)

        self.tab_btns = {}
        for t in self.TABS:
            self._create_tab_button(t)

        # â”€â”€ Toolbar â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
        toolbar = ctk.CTkFrame(main, fg_color=BG_PANEL, corner_radius=0, height=42)
        toolbar.pack(side="top", fill="x")
        toolbar.pack_propagate(False)

        self.search_var = ctk.StringVar()
        ctk.CTkEntry(toolbar, textvariable=self.search_var,
                     placeholder_text="ðŸ”  Highlight keyword...",
                     fg_color=BG_CARD2, border_color=BORDER, text_color=TEXT_PRIMARY,
                     placeholder_text_color=TEXT_MUTED,
                     font=("Segoe UI", 11), width=180, height=28).pack(side="left", padx=8, pady=7)
        self.search_var.trace_add("write", lambda *a: self._filter())

        self.live_var = ctk.BooleanVar(value=True)
        self.live_btn = ctk.CTkButton(toolbar, text="â¸ Pause",
                                      fg_color=DANGER, hover_color="#AA2040",
                                      text_color="#fff", font=("Segoe UI", 11, "bold"),
                                      corner_radius=6, width=100, height=28,
                                      command=self._toggle_live)
        self.live_btn.pack(side="left", padx=4, pady=7)

        btn(toolbar, "ðŸ—‘ Clear", self._clear_log,
            fg=BG_CARD2, hover=BG_CARD, tc=TEXT_MUTED, w=80, h=28).pack(side="left", padx=4, pady=7)
        btn(toolbar, "ðŸ’¾ Export", self._export,
            fg=BG_CARD2, hover=BG_CARD, tc=TEXT_MUTED, w=90, h=28).pack(side="left", padx=4, pady=7)

        ctk.CTkLabel(toolbar, text="Zoom", font=("Segoe UI", 10),
                     text_color=TEXT_MUTED).pack(side="left", padx=(8, 2), pady=7)
        btn(toolbar, "Aâˆ’", lambda: self._zoom_delta(-1),
            fg=BG_CARD2, hover=BG_CARD, tc=TEXT_MUTED, w=36, h=28).pack(side="left", padx=1, pady=7)
        btn(toolbar, "A+", lambda: self._zoom_delta(1),
            fg=BG_CARD2, hover=BG_CARD, tc=TEXT_MUTED, w=36, h=28).pack(side="left", padx=1, pady=7)
        self.zoom_lbl = lbl(toolbar, f"{self._log_font_pt} pt", size=10, color=TEXT_MUTED)
        self.zoom_lbl.pack(side="left", padx=(4, 8), pady=7)

        ctk.CTkLabel(toolbar, text="Auto-scroll", font=("Segoe UI", 10),
                     text_color=TEXT_MUTED).pack(side="left", padx=(4, 2), pady=7)
        self.auto_scroll_sw = ctk.CTkSwitch(
            toolbar, text="",
            width=42,
            command=self._on_auto_scroll_switch,
            fg_color=BG_CARD2, progress_color=SUCCESS,
            button_color=TEXT_PRIMARY, button_hover_color=ACCENT,
        )
        self.auto_scroll_sw.pack(side="left", padx=0, pady=7)
        self.auto_scroll_sw.select()

        self.status = lbl(toolbar, "Connectingâ€¦", size=10, color=TEXT_MUTED)
        self.status.pack(side="right", padx=12)

        conn_toggle_row = ctk.CTkFrame(main, fg_color="transparent")
        conn_toggle_row.pack(side="top", fill="x", padx=10, pady=(8, 0))
        self.conn_toggle_btn = btn(
            conn_toggle_row,
            "Connection Config â–¼",
            self._toggle_connection_config,
            fg=BG_CARD2,
            hover=BG_PANEL,
            tc=TEXT_PRIMARY,
            w=170,
            h=30,
        )
        self.conn_toggle_btn.pack(side="left")

        # â”€â”€ Connection details (active log tab; collapsible) â”€â”€â”€â”€â”€â”€â”€â”€
        self._conn_wrap = ctk.CTkFrame(
            main,
            fg_color=BG_CARD,
            corner_radius=8,
            border_width=1,
            border_color=DANGER,
        )
        lbl(self._conn_wrap, "CONNECTION DETAILS (ACTIVE TAB)", size=10, weight="bold", color=DANGER).pack(
            anchor="w", padx=10, pady=(8, 4)
        )
        conn_grid = ctk.CTkFrame(self._conn_wrap, fg_color="transparent")
        conn_grid.pack(fill="x", padx=8, pady=(0, 8))
        for c in range(3):
            conn_grid.columnconfigure(c, weight=1)

        def mk_field(parent, row, col, key, title, show=""):
            cell = ctk.CTkFrame(parent, fg_color="transparent")
            cell.grid(row=row, column=col, sticky="ew", padx=6, pady=4)
            lbl(cell, title, size=10, color=TEXT_MUTED).pack(anchor="w")
            ent = ctk.CTkEntry(
                cell,
                fg_color=BG_CARD2,
                border_color=BORDER,
                text_color=TEXT_PRIMARY,
                font=("Segoe UI", 11),
                height=30,
                show=show,
            )
            ent.pack(fill="x")
            self._conn_fields[key] = ent

        mk_field(conn_grid, 0, 0, "host", "Host / IP")
        mk_field(conn_grid, 0, 1, "port", "Port")
        mk_field(conn_grid, 0, 2, "user", "User")
        mk_field(conn_grid, 1, 0, "password", "Password", show="â—")
        mk_field(conn_grid, 1, 1, "log_path", "Log Path")

        btn(
            conn_grid,
            "Save Tab Details",
            self._save_active_tab_connection_details,
            fg=BG_CARD2,
            hover=BG_PANEL,
            tc=TEXT_PRIMARY,
            w=150,
            h=30,
        ).grid(row=1, column=2, sticky="e", padx=6, pady=4)
        self._tab_connect_btn = btn(
            conn_grid,
            "Connect",
            self._toggle_active_tab_connection,
            fg=BG_CARD2,
            hover=BG_PANEL,
            tc=TEXT_PRIMARY,
            w=110,
            h=30,
        )
        self._tab_connect_btn.grid(row=2, column=2, sticky="e", padx=6, pady=(2, 4))
        self._tab_conn_status = lbl(conn_grid, "", size=10, color=TEXT_MUTED)
        self._tab_conn_status.grid(row=2, column=0, columnspan=2, sticky="w", padx=6, pady=(2, 4))

        # â”€â”€ Transaction ID / Grep bar â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
        grep_bar = ctk.CTkFrame(main, fg_color="#0D1520", corner_radius=0, height=44)
        grep_bar.pack(side="top", fill="x")
        grep_bar.pack_propagate(False)

        lbl(grep_bar, "Transaction ID", size=11, weight="bold", color=WARNING).pack(side="left", padx=(12, 6), pady=10)
        self.grep_var = ctk.StringVar()
        self.grep_entry = ctk.CTkEntry(
            grep_bar, textvariable=self.grep_var,
            placeholder_text="Paste or type transaction / correlation / request IDâ€¦",
            fg_color=BG_CARD, border_color=WARNING, border_width=1,
            text_color=TEXT_PRIMARY, placeholder_text_color=TEXT_MUTED,
            font=("Consolas", 11), width=360, height=28,
        )
        self.grep_entry.pack(side="left", padx=(0, 8), pady=8)
        self.grep_entry.bind("<Return>", lambda e: self._apply_grep())

        self.grep_mode = ctk.StringVar(value="highlight")
        ctk.CTkRadioButton(grep_bar, text="Highlight", variable=self.grep_mode, value="highlight",
                           fg_color=WARNING, hover_color="#AA7700", text_color=TEXT_PRIMARY,
                           font=("Segoe UI", 11), command=self._apply_grep).pack(side="left", padx=4)
        ctk.CTkRadioButton(grep_bar, text="Filter only", variable=self.grep_mode, value="filter",
                           fg_color=WARNING, hover_color="#AA7700", text_color=TEXT_PRIMARY,
                           font=("Segoe UI", 11), command=self._apply_grep).pack(side="left", padx=4)
        ctk.CTkButton(grep_bar, text="âš¡ Search", command=self._apply_grep,
                      fg_color=WARNING, hover_color="#AA7700", text_color="#000",
                      font=("Segoe UI", 11, "bold"), corner_radius=6,
                      width=88, height=28).pack(side="left", padx=4, pady=8)
        ctk.CTkButton(grep_bar, text="âœ• Clear", command=self._clear_grep,
                      fg_color=BG_CARD2, hover_color=BG_CARD, text_color=TEXT_MUTED,
                      font=("Segoe UI", 11), corner_radius=6,
                      width=80, height=28).pack(side="left", padx=4, pady=8)
        self.grep_status = lbl(grep_bar, "", size=10, color=WARNING)
        self.grep_status.pack(side="left", padx=8)

        # â”€â”€ K8 shortcuts (shown only for K8 tab) â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
        self.k8_bar = ctk.CTkFrame(main, fg_color="#122022", corner_radius=0, height=44)
        self.k8_bar.pack_propagate(False)
        lbl(self.k8_bar, "K8 Namespace", size=10, color=SUCCESS).pack(side="left", padx=(12, 6), pady=8)
        ctk.CTkEntry(
            self.k8_bar,
            textvariable=self.k8_namespace_var,
            fg_color=BG_CARD,
            border_color=SUCCESS,
            text_color=TEXT_PRIMARY,
            font=("Consolas", 11),
            width=160,
            height=28,
        ).pack(side="left", padx=(0, 8), pady=8)
        lbl(self.k8_bar, "Pod Name", size=10, color=SUCCESS).pack(side="left", padx=(0, 6), pady=8)
        ctk.CTkEntry(
            self.k8_bar,
            textvariable=self.k8_pod_var,
            fg_color=BG_CARD,
            border_color=SUCCESS,
            text_color=TEXT_PRIMARY,
            font=("Consolas", 11),
            width=300,
            height=28,
            placeholder_text="pod name here",
        ).pack(side="left", padx=(0, 8), pady=8)
        lbl(self.k8_bar, "Tail Lines", size=10, color=SUCCESS).pack(side="left", padx=(0, 6), pady=8)
        ctk.CTkEntry(
            self.k8_bar,
            textvariable=self.k8_tail_lines_var,
            fg_color=BG_CARD,
            border_color=SUCCESS,
            text_color=TEXT_PRIMARY,
            font=("Consolas", 11),
            width=70,
            height=28,
        ).pack(side="left", padx=(0, 8), pady=8)
        btn(
            self.k8_bar, "List Pods", self._k8_list_pods,
            fg=SUCCESS, hover="#00B37A", tc="#000", w=100, h=28
        ).pack(side="left", padx=4, pady=8)
        btn(
            self.k8_bar, "Tail Pod Logs", self._k8_tail_pod_logs,
            fg=ACCENT, hover=ACCENT2, tc="#000", w=120, h=28
        ).pack(side="left", padx=4, pady=8)
        btn(
            self.k8_bar, "Stop Tail", self._k8_stop_tail,
            fg=BG_CARD2, hover=BG_CARD, tc=TEXT_MUTED, w=90, h=28
        ).pack(side="left", padx=4, pady=8)

        # â”€â”€ Log text widgets (one per tab, stacked â€” preserves content when switching) â”€â”€
        self._logs_host = ctk.CTkFrame(main, fg_color="#0A0C10", corner_radius=0)
        self._logs_host.pack(side="top", fill="both", expand=True)
        self._logs_host.grid_rowconfigure(0, weight=1)
        self._logs_host.grid_columnconfigure(0, weight=1)

        for t in self.TABS:
            self._init_tab_runtime(t)

        self._sy = ctk.CTkScrollbar(self._logs_host, command=lambda *a: None)
        self._sy.grid(row=0, column=1, sticky="ns")

        self._show_log_tab(self.active_tab)  # sets self.log and scrollbar bindings
        self._load_active_tab_connection_details()
        self._refresh_tab_connection_button_state()
        self._refresh_tab_button_styles()
        self._sync_k8_ui_for_active_tab()

        self.bind("<Control-plus>", lambda e: self._zoom_delta(1))
        self.bind("<Control-equal>", lambda e: self._zoom_delta(1))
        self.bind("<Control-minus>", lambda e: self._zoom_delta(-1))
        self.bind("<Control-0>", lambda e: self._zoom_reset())

        self._start_all_streams()

    def _zoom_reset(self):
        self._log_font_pt = self.FONT_DEFAULT
        self._apply_font_zoom()

    # â”€â”€ Navigation â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    def _nav_and_stop(self, item):
        self._stop_all_streams()
        self.on_nav(item)

    # â”€â”€ Tab switching â€” swap visible widget only; streams keep running in background â”€â”€
    def _switch(self, tab):
        self._show_log_tab(tab)
        self._load_active_tab_connection_details()
        self._refresh_tab_connection_button_state()
        self._refresh_tab_button_styles()
        self._sync_k8_ui_for_active_tab()
        self._filter()
        self._apply_grep()
        self._refresh_active_tab_status()
        # Pick up newly connected modules (e.g. Connect in Config while Logs is open)
        for t in self.TABS:
            self._start_stream_for_tab(t)

    def _tab_width(self, tab_name: str) -> int:
        if "Aircontrol API" in tab_name:
            return 210
        return max(110, min(230, 88 + len(tab_name) * 6))

    def _create_tab_button(self, tab_name: str):
        b = ctk.CTkButton(
            self._tab_btn_wrap,
            text=tab_name,
            fg_color="transparent",
            hover_color=BG_CARD,
            text_color=TEXT_MUTED,
            font=("Segoe UI", 10, "bold"),
            corner_radius=8,
            border_width=1,
            border_color=BORDER,
            height=32,
            width=self._tab_width(tab_name),
            command=lambda x=tab_name: self._switch(x),
        )
        b.pack(side="left", padx=4, pady=8)
        self.tab_btns[tab_name] = b

    def _refresh_tab_button_styles(self):
        for t, b in self.tab_btns.items():
            active = (t == self.active_tab)
            b.configure(
                fg_color=ACCENT if active else "transparent",
                text_color="#000000" if active else TEXT_MUTED,
                border_color=ACCENT if active else BORDER,
            )

    def _init_tab_runtime(self, tab_name: str):
        if tab_name in self._log_texts:
            return
        txt = self._make_log_text(self._logs_host)
        txt.grid(row=0, column=0, sticky="nsew")
        self._log_texts[tab_name] = txt
        txt.bind("<MouseWheel>", self._on_log_wheel)
        txt.bind("<Button-4>", self._on_log_wheel_linux)
        txt.bind("<Button-5>", self._on_log_wheel_linux)
        self._tab_stream_stop[tab_name] = threading.Event()
        self._tab_stream_threads[tab_name] = None
        self._line_counts[tab_name] = 0

    def _add_connection_tab(self):
        name = simpledialog.askstring("Add Connection Tab", "Enter new tab name:", parent=self)
        if name is None:
            return
        name = name.strip()
        if not name:
            messagebox.showwarning("Add Connection Tab", "Tab name cannot be empty.", parent=self)
            return
        if name in self.TABS:
            messagebox.showwarning("Add Connection Tab", f"Tab '{name}' already exists.", parent=self)
            return

        self.TABS.append(name)
        self._create_tab_button(name)
        self._init_tab_runtime(name)
        state.module_configs.setdefault(
            name,
            {
                "host": "",
                "port": "22",
                "user": "",
                "password": "",
                "key": "",
                "log_path": "/var/log/app.log",
                "sudo_user": "aqadmin",
                "sudo_password": "",
            },
        )
        save_config()
        self._switch(name)
        self.status.configure(text=f"Created new connection tab: {name}", text_color=SUCCESS)

    def _toggle_connection_config(self):
        self._conn_config_open = not self._conn_config_open
        if self._conn_config_open:
            self._conn_wrap.pack(side="top", fill="x", padx=10, pady=(6, 6), after=self.conn_toggle_btn.master)
            self.conn_toggle_btn.configure(text="Connection Config â–²")
        else:
            self._conn_wrap.pack_forget()
            self.conn_toggle_btn.configure(text="Connection Config â–¼")

    def _sync_k8_ui_for_active_tab(self):
        if self.active_tab == "K8":
            self.k8_bar.pack(side="top", fill="x")
        else:
            self.k8_bar.pack_forget()

    def _load_active_tab_connection_details(self):
        cfg = state.module_configs.get(self.active_tab, {})
        values = {
            "host": cfg.get("host", ""),
            "port": cfg.get("port", "22"),
            "user": cfg.get("user", ""),
            "password": cfg.get("password", ""),
            "log_path": cfg.get("log_path", ""),
        }
        for key, ent in self._conn_fields.items():
            try:
                ent.delete(0, "end")
                ent.insert(0, values.get(key, ""))
            except Exception:
                pass
        self._refresh_tab_connection_button_state()

    def _save_active_tab_connection_details(self):
        tab = self.active_tab
        if tab not in state.module_configs:
            state.module_configs[tab] = {}
        cfg = state.module_configs[tab]
        cfg["host"] = self._conn_fields["host"].get().strip()
        cfg["port"] = self._conn_fields["port"].get().strip() or "22"
        cfg["user"] = self._conn_fields["user"].get().strip()
        cfg["password"] = self._conn_fields["password"].get().strip()
        cfg["log_path"] = self._conn_fields["log_path"].get().strip()
        save_config()
        self.status.configure(
            text=f"Saved connection details for {tab}",
            text_color=SUCCESS,
        )
        self._refresh_tab_connection_button_state()

    def _refresh_tab_connection_button_state(self):
        connected = ssh_mgr.is_connected(self.active_tab)
        self._tab_connect_btn.configure(text="Disconnect" if connected else "Connect")
        self._tab_conn_status.configure(
            text="â— Connected" if connected else "â—‹ Disconnected",
            text_color=SUCCESS if connected else TEXT_MUTED,
        )

    def _toggle_active_tab_connection(self):
        tab = self.active_tab
        self._save_active_tab_connection_details()
        if ssh_mgr.is_connected(tab):
            self._stop_stream_for_tab(tab)
            if tab == "K8":
                self._k8_stop_tail()
            ssh_mgr.disconnect(tab)
            self.status.configure(text=f"Disconnected: {tab}", text_color=WARNING)
            self._refresh_tab_connection_button_state()
            return

        cfg = state.module_configs.get(tab, {})
        host = (cfg.get("host") or "").strip()
        user = (cfg.get("user") or "").strip()
        if not host or not user:
            self.status.configure(
                text="Set Host and User before connect",
                text_color=DANGER,
            )
            return

        self._tab_connect_btn.configure(state="disabled")
        self._tab_conn_status.configure(text="Connectingâ€¦", text_color=WARNING)

        def work():
            try:
                port = int((cfg.get("port") or "22").strip() or "22")
            except Exception:
                port = 22
            ok, msg = ssh_mgr.connect(
                host=tab,  # logical connection key
                actual_host=host,
                user=user,
                password=(cfg.get("password") or "").strip() or None,
                key_path=(cfg.get("key") or "").strip() or None,
                port=port,
                sudo_user=(cfg.get("sudo_user") or "").strip() or None,
                sudo_password=(cfg.get("sudo_password") or "").strip() or (cfg.get("password") or "").strip() or None,
            )

            def done():
                self._tab_connect_btn.configure(state="normal")
                if ok:
                    self.status.configure(text=f"Connected: {tab}", text_color=SUCCESS)
                    self._start_stream_for_tab(tab)
                else:
                    self.status.configure(text=f"Connect failed: {msg}", text_color=DANGER)
                self._refresh_tab_connection_button_state()

            self.after(0, done)

        threading.Thread(target=work, daemon=True).start()

    def _refresh_active_tab_status(self):
        """Show line count for the visible tab after switching; jump to tail if auto-scroll on."""
        try:
            n = self._line_counts.get(self.active_tab, 0)
            self.status.configure(
                text=f"â¬¤ LIVE  {self.active_tab}  â€¢  {n} lines",
                text_color=SUCCESS,
            )
            if self._auto_scroll:
                self.log.see("end")
        except Exception:
            pass

    # â”€â”€ Stream control â€” one tail thread per tab (runs even when tab not visible) â”€â”€
    def _start_all_streams(self):
        if self._alive:
            self.status.configure(text="â¬¤ Starting log streamsâ€¦", text_color=WARNING)
        for tab in self.TABS:
            self._start_stream_for_tab(tab)

    def _start_stream_for_tab(self, tab):
        if not self._alive:
            return
        # K8 uses shortcut commands (list/tail pod), not static tail -F path.
        if tab == "K8":
            return
        if not ssh_mgr.is_connected(tab):
            return
        th = self._tab_stream_threads.get(tab)
        if th is not None and th.is_alive():
            return
        self._tab_stream_stop[tab].clear()
        t = threading.Thread(target=self._stream_worker, args=(tab,), daemon=True)
        self._tab_stream_threads[tab] = t
        t.start()

    def _stop_stream_for_tab(self, tab):
        self._tab_stream_stop[tab].set()
        th = self._tab_stream_threads.get(tab)
        if th is not None and th.is_alive():
            th.join(timeout=0.5)
        self._tab_stream_threads[tab] = None

    def _stop_all_streams(self):
        for tab in self.TABS:
            self._stop_stream_for_tab(tab)

    def _apply_status_if_active(self, tab, text, color):
        """Avoid status flicker from background tab streams."""
        if tab != self.active_tab:
            return
        try:
            self.status.configure(text=text, text_color=color)
        except Exception:
            pass

    def _stream_worker(self, tab):
        conn_key  = tab
        mcfg      = state.module_configs.get(tab, {})
        host      = mcfg.get("host", "").strip() or state.env_hosts.get(state.env, "")
        log_path  = mcfg.get("log_path", "").strip()

        def safe_after(fn):
            """Only schedule UI update if page is still alive."""
            if self._alive:
                try:
                    self.after(0, fn)
                except Exception:
                    pass

        if not ssh_mgr.is_connected(conn_key):
            safe_after(lambda tb=tab: self._append_line(
                f"[NOT CONNECTED] Configure {tb} in âš™ Config and click Connect.", tag="ERR", for_tab=tb))
            safe_after(lambda t=tab: self._apply_status_if_active(t, f"â¬¤ {t} not connected", DANGER))
            return

        if not log_path:
            safe_after(lambda tb=tab: self._append_line(
                f"[NO LOG PATH] Set log path for {tb} in âš™ Config.", tag="ERR", for_tab=tb))
            safe_after(lambda t=tab: self._apply_status_if_active(t, f"â¬¤ {t} no log path", DANGER))
            return

        tail_cmd = f"tail -n 10 -F {log_path} 2>&1"

        conn_info        = ssh_mgr.connections.get(conn_key, {})
        stored_sudo_user = conn_info.get("sudo_user", "").strip()

        safe_after(lambda t=tab, h=host: self._apply_status_if_active(
            t, f"â¬¤ LIVE  {t} @ {h}", SUCCESS))

        try:
            client = conn_info.get("client") if isinstance(conn_info, dict) else conn_info
            import time as _t

            if stored_sudo_user:
                channel, chain_err = ssh_mgr.open_tail_log_shell_chain(conn_key, log_path)
                if chain_err:
                    safe_after(lambda tb=tab: self._append_line(
                        f"[STREAM ERROR] {chain_err}", tag="ERR", for_tab=tb))
                    safe_after(lambda t=tab: self._apply_status_if_active(t, f"â¬¤ {t} stream failed", DANGER))
                    return
            else:
                transport = client.get_transport()
                channel = transport.open_session()
                channel.set_combine_stderr(True)
                channel.exec_command(tail_cmd)

            buf = b""
            ev = self._tab_stream_stop[tab]
            while not ev.is_set():
                if channel.recv_ready():
                    chunk = channel.recv(4096)
                    if not chunk:
                        break
                    buf += chunk
                    while b"\n" in buf:
                        line_bytes, buf = buf.split(b"\n", 1)
                        line = line_bytes.decode(errors="replace").rstrip()
                        if line and not ev.is_set():
                            captured = line
                            safe_after(lambda l=captured, tb=tab: self._append_line(l, for_tab=tb))
                elif not stored_sudo_user and channel.exit_status_ready():
                    break
                else:
                    _t.sleep(0.05)
            try:
                channel.close()
            except Exception:
                pass
        except Exception as e:
            err = str(e)
            safe_after(lambda tb=tab: self._append_line(f"[STREAM ERROR] {err}", tag="ERR", for_tab=tb))

        if not self._tab_stream_stop[tab].is_set():
            safe_after(lambda t=tab: self._apply_status_if_active(t, f"â¬¤ {t} stream ended", WARNING))

    def _k8_list_pods(self):
        if self.active_tab != "K8":
            return
        if not ssh_mgr.is_connected("K8"):
            self.status.configure(text="Connect K8 first", text_color=DANGER)
            return
        ns = (self.k8_namespace_var.get() or "default").strip()
        cmd = f"kubectl get pods -n {shlex.quote(ns)} -o wide 2>&1"
        self.status.configure(text=f"Listing pods in namespace: {ns}", text_color=WARNING)

        def work():
            ch, chain_err = ssh_mgr.open_command_shell_chain("K8", cmd)
            if chain_err:
                result = f"[K8 ERROR] {chain_err}"
            else:
                import time as _t
                parts = []
                last = _t.time()
                while _t.time() - last < 1.6:
                    if ch.recv_ready():
                        chunk = ch.recv(65536)
                        if not chunk:
                            break
                        parts.append(chunk.decode(errors="replace"))
                        last = _t.time()
                    else:
                        _t.sleep(0.08)
                try:
                    ch.close()
                except Exception:
                    pass
                result = "".join(parts).strip() or "[No output]"
            def done():
                self._append_line(f"----- kubectl get pods -n {ns} -----", tag="INFO", for_tab="K8")
                for line in result.splitlines():
                    self._append_line(line, for_tab="K8")
                self.status.configure(text=f"Pod list fetched for {ns}", text_color=SUCCESS)
            self.after(0, done)

        threading.Thread(target=work, daemon=True).start()

    def _k8_stop_tail(self):
        self._stop_stream_for_tab("K8")
        self.status.configure(text="Stopped K8 pod tail", text_color=WARNING)

    def _k8_tail_pod_logs(self):
        if self.active_tab != "K8":
            return
        if not ssh_mgr.is_connected("K8"):
            self.status.configure(text="Connect K8 first", text_color=DANGER)
            return
        ns = (self.k8_namespace_var.get() or "default").strip()
        pod = (self.k8_pod_var.get() or "").strip()
        lines = (self.k8_tail_lines_var.get() or "200").strip()
        if not pod:
            self.status.configure(text="Enter pod name to tail logs", text_color=DANGER)
            return
        try:
            lines_i = int(lines)
            if lines_i <= 0:
                raise ValueError
        except Exception:
            lines_i = 200

        self._stop_stream_for_tab("K8")
        self._line_counts["K8"] = 0
        self._tab_stream_stop["K8"].clear()
        t = threading.Thread(
            target=self._k8_tail_worker,
            args=(ns, pod, lines_i),
            daemon=True,
        )
        self._tab_stream_threads["K8"] = t
        t.start()

    def _k8_tail_worker(self, namespace: str, pod_name: str, tail_lines: int):
        tab = "K8"
        stop_ev = self._tab_stream_stop[tab]

        cmd = (
            f"kubectl logs -f -n {shlex.quote(namespace)} --tail={tail_lines} "
            f"{shlex.quote(pod_name)} 2>&1"
        )

        def safe_after(fn):
            if self._alive:
                try:
                    self.after(0, fn)
                except Exception:
                    pass

        safe_after(lambda: self._append_line(
            f"----- tailing pod logs: ns={namespace} pod={pod_name} -----",
            tag="INFO",
            for_tab=tab,
        ))
        safe_after(lambda: self.status.configure(
            text=f"Tailing K8 pod logs: {pod_name}",
            text_color=SUCCESS,
        ))

        ch, chain_err = ssh_mgr.open_command_shell_chain(tab, cmd)
        if chain_err:
            safe_after(lambda e=chain_err: self._append_line(f"[K8 STREAM ERROR] {e}", tag="ERR", for_tab=tab))
            safe_after(lambda: self.status.configure(text=f"K8 tail failed: {chain_err}", text_color=DANGER))
            return

        try:
            buf = b""
            import time as _t
            while not stop_ev.is_set():
                if ch.recv_ready():
                    chunk = ch.recv(4096)
                    if not chunk:
                        break
                    buf += chunk
                    while b"\n" in buf:
                        line_bytes, buf = buf.split(b"\n", 1)
                        line = line_bytes.decode(errors="replace").rstrip()
                        if line and not stop_ev.is_set():
                            safe_after(lambda l=line: self._append_line(l, for_tab=tab))
                else:
                    _t.sleep(0.05)

            try:
                ch.close()
            except Exception:
                pass
        except Exception as e:
            safe_after(lambda err=str(e): self._append_line(f"[K8 STREAM ERROR] {err}", tag="ERR", for_tab=tab))

        if not stop_ev.is_set():
            safe_after(lambda: self.status.configure(text=f"K8 pod tail ended: {pod_name}", text_color=WARNING))

    def _append_line(self, line, tag=None, for_tab=None):
        """Append line to the given tab's log widget (stream tab may differ from active during races)."""
        tab = for_tab if for_tab is not None else self.active_tab
        if not self._alive:
            return
        logw = self._log_texts.get(tab)
        if not logw:
            return
        try:
            if not logw.winfo_exists():
                return
        except Exception:
            return

        if self.live_var.get() is False:
            return
        if self._grep_filter and self._grep_term:
            if self._grep_term.lower() not in line.lower():
                return

        if tag is None:
            u = line.upper()
            tag = ("ERR"   if "ERROR" in u or "CRITICAL" in u else
                   "WARN"  if "WARN"  in u else
                   "INFO"  if "INFO"  in u else
                   "DEBUG" if "DEBUG" in u else "")

        try:
            logw.configure(state="normal")
            self._line_counts[tab] = self._line_counts.get(tab, 0) + 1
            if self._line_counts[tab] > self.MAX_LINES:
                logw.delete("1.0", "2.0")
                self._line_counts[tab] -= 1
            logw.insert("end", line + "\n", tag)

            if self._grep_term and not self._grep_filter:
                last = f"{logw.index('end-1c').split('.')[0]}.0"
                pos  = last
                while True:
                    pos = logw.search(self._grep_term, pos, nocase=True, stopindex="end")
                    if not pos: break
                    logw.tag_add("GREP", pos, f"{pos}+{len(self._grep_term)}c")
                    pos = f"{pos}+{len(self._grep_term)}c"

            term = self.search_var.get()
            if term:
                last = f"{logw.index('end-1c').split('.')[0]}.0"
                pos  = last
                while True:
                    pos = logw.search(term, pos, nocase=True, stopindex="end")
                    if not pos: break
                    logw.tag_add("HIT", pos, f"{pos}+{len(term)}c")
                    pos = f"{pos}+{len(term)}c"

            logw.configure(state="disabled")
            # Auto-scroll every tab's buffer when enabled (parallel streams run in background)
            if self._auto_scroll:
                try:
                    logw.see("end")
                except Exception:
                    pass
            if tab == self.active_tab:
                self.status.configure(
                    text=f"â¬¤ LIVE  {self.active_tab}  â€¢  {self._line_counts.get(tab, 0)} lines  â€¢  {datetime.now().strftime('%H:%M:%S')}",
                    text_color=SUCCESS)
        except Exception:
            pass

    # â”€â”€ Controls â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    def _toggle_live(self):
        if self.live_var.get():
            self.live_var.set(False)
            self.live_btn.configure(text="â–¶ Resume", fg_color=SUCCESS,
                                    hover_color="#009960", text_color="#000")
        else:
            self.live_var.set(True)
            self.live_btn.configure(text="â¸ Pause", fg_color=DANGER,
                                    hover_color="#AA2040", text_color="#fff")

    def _on_auto_scroll_switch(self):
        self._auto_scroll = bool(self.auto_scroll_sw.get())
        if self._auto_scroll:
            for w in self._log_texts.values():
                try:
                    w.see("end")
                except Exception:
                    pass

    def _disable_auto_scroll_ui(self):
        self._auto_scroll = False
        try:
            self.auto_scroll_sw.deselect()
        except Exception:
            pass

    def _on_log_wheel(self, event):
        """Ctrl+wheel = zoom; otherwise scroll log; wheel toward older lines may disable auto-scroll."""
        w = event.widget
        try:
            if event.state & 0x4:  # Control
                d = getattr(event, "delta", 0)
                if d:
                    self._zoom_delta(1 if d > 0 else -1)
                return "break"
        except Exception:
            pass
        if event.num == 4 or (hasattr(event, "delta") and event.delta > 0):
            self._disable_auto_scroll_ui()
        try:
            if hasattr(event, "delta") and event.delta:
                w.yview_scroll(int(-1 * (event.delta / 120)), "units")
        except Exception:
            pass
        return "break"

    def _on_log_wheel_linux(self, event):
        w = event.widget
        try:
            if (event.state or 0) & 0x4:
                self._zoom_delta(1 if event.num == 4 else -1)
                return "break"
        except Exception:
            pass
        if event.num == 4:
            self._disable_auto_scroll_ui()
        try:
            w.yview_scroll(-1 if event.num == 4 else 1, "units")
        except Exception:
            pass
        return "break"

    def _clear_log(self):
        try:
            self.log.configure(state="normal")
            self.log.delete("1.0", "end")
            self.log.configure(state="disabled")
            self._line_counts[self.active_tab] = 0
        except Exception:
            pass

    def _filter(self):
        try:
            term = self.search_var.get()
            self.log.tag_remove("HIT", "1.0", "end")
            if not term:
                return
            pos = "1.0"
            while True:
                pos = self.log.search(term, pos, nocase=True, stopindex="end")
                if not pos: break
                end = f"{pos}+{len(term)}c"
                self.log.tag_add("HIT", pos, end)
                pos = end
        except Exception:
            pass

    def _apply_grep(self):
        term = self.grep_var.get().strip()
        self._grep_term   = term
        self._grep_filter = self.grep_mode.get() == "filter"
        try:
            self.log.configure(state="normal")
            self.log.tag_remove("GREP", "1.0", "end")
            if not term:
                self.grep_status.configure(text="")
                self.log.configure(state="disabled")
                return
            if self._grep_filter:
                all_text = self.log.get("1.0", "end")
                matching = [l for l in all_text.splitlines() if term.lower() in l.lower()]
                self.log.delete("1.0", "end")
                count = 0
                for line in matching:
                    u = line.upper()
                    t = ("ERR" if "ERROR" in u or "CRITICAL" in u else
                         "WARN" if "WARN" in u else "INFO" if "INFO" in u else "")
                    self.log.insert("end", line + "\n", t)
                    count += 1
                self._line_counts[self.active_tab] = count
                pos = "1.0"
                while True:
                    pos = self.log.search(term, pos, nocase=True, stopindex="end")
                    if not pos: break
                    self.log.tag_add("GREP", pos, f"{pos}+{len(term)}c")
                    pos = f"{pos}+{len(term)}c"
                self.grep_status.configure(text=f"  {count} matching lines (filter mode)")
            else:
                count = 0
                pos = "1.0"
                while True:
                    pos = self.log.search(term, pos, nocase=True, stopindex="end")
                    if not pos: break
                    self.log.tag_add("GREP", pos, f"{pos}+{len(term)}c")
                    pos = f"{pos}+{len(term)}c"
                    count += 1
                if count:
                    first = self.log.search(term, "1.0", nocase=True, stopindex="end")
                    if first:
                        self.log.see(first)
                    self.grep_status.configure(text=f"  {count} matches")
                else:
                    self.grep_status.configure(text="  No matches")
            self.log.configure(state="disabled")
        except Exception:
            pass

    def _clear_grep(self):
        self.grep_var.set("")
        self._grep_term   = ""
        self._grep_filter = False
        self.grep_status.configure(text="")
        try:
            self.log.configure(state="normal")
            self.log.tag_remove("GREP", "1.0", "end")
            self.log.configure(state="disabled")
        except Exception:
            pass

    def _export(self):
        path = filedialog.asksaveasfilename(defaultextension=".log",
               filetypes=[("Log files", "*.log"), ("Text", "*.txt")])
        if path:
            with open(path, "w", encoding="utf-8") as f:
                f.write(self.log.get("1.0", "end"))
            messagebox.showinfo("Exported", f"Saved to:\n{path}")

# ===== END FILE: page_logs.py =====


# ===== BEGIN FILE: page_database.py =====

import customtkinter as ctk
import tkinter as tk
from tkinter import messagebox, simpledialog, ttk
import threading
from datetime import datetime
from typing import Any, List, Optional, Tuple

from config import (BG_DARK, BG_PANEL, BG_CARD, BG_CARD2, ACCENT, ACCENT2,
                    SUCCESS, WARNING, DANGER, TEXT_PRIMARY, TEXT_MUTED, BORDER,
                    state, MYSQL_DALLAS_DATABASES, save_config)
from db_mysql import mysql_query_to_grid
from db_postgres import postgres_query_to_grid
from widgets import lbl, btn, hsep, Sidebar, SIDEBAR_ITEMS

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  PAGE: Database Queries
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

class PageDatabase(ctk.CTkFrame):
    def __init__(self, master, on_nav, on_back, on_config):
        super().__init__(master, fg_color=BG_DARK, corner_radius=0)
        self._db_conn_fields = {}
        self._db_conn_config_open = False
        self._build(on_nav, on_back, on_config)

    def _build(self, on_nav, on_back, on_config):
        Sidebar(self, SIDEBAR_ITEMS, active="Database Queries",
                on_select=on_nav, on_back=on_back, on_config=on_config).pack(side="left", fill="y")

        main = ctk.CTkFrame(self, fg_color=BG_DARK, corner_radius=0)
        main.pack(side="left", fill="both", expand=True)

        # DB type selector bar
        self._tb = ctk.CTkFrame(main, fg_color=BG_PANEL, corner_radius=0, height=46)
        self._tb.pack(fill="x")
        self._tb.pack_propagate(False)
        lbl(self._tb, "DATABASE TYPE:", size=11, color=TEXT_MUTED).pack(side="left", padx=16, pady=14)
        self.db_type = ctk.StringVar(value="MySQL")
        for db in ["MySQL", "Postgres"]:
            ctk.CTkRadioButton(self._tb, text=db, variable=self.db_type, value=db,
                               fg_color=ACCENT, hover_color=ACCENT2, text_color=TEXT_PRIMARY,
                               font=("Segoe UI", 12, "bold")).pack(side="left", padx=16, pady=14)

        self.mysql_db_row = ctk.CTkFrame(main, fg_color=BG_PANEL, corner_radius=0, height=46)
        self.mysql_db_row.pack_propagate(False)
        lbl(self.mysql_db_row, "MYSQL DATABASE:", size=11, color=TEXT_MUTED).pack(
            side="left", padx=16, pady=14)
        self.mysql_db_combo = ctk.CTkComboBox(
            self.mysql_db_row,
            values=MYSQL_DALLAS_DATABASES,
            width=520,
            height=32,
            fg_color=BG_CARD2,
            button_color=ACCENT2,
            button_hover_color=ACCENT,
            dropdown_fg_color=BG_CARD,
            font=("Segoe UI", 11),
            text_color=TEXT_PRIMARY,
            command=self._on_mysql_db_combo,
        )
        self.mysql_db_combo.pack(side="left", padx=8, pady=8)
        self.db_type.trace_add("write", self._on_db_type_change)

        conn_toggle_row = ctk.CTkFrame(main, fg_color="transparent")
        conn_toggle_row.pack(fill="x", padx=12, pady=(8, 0))
        self.db_conn_toggle_btn = btn(
            conn_toggle_row,
            "Connection Config â–¼",
            self._toggle_db_connection_config,
            fg=BG_CARD2,
            hover=BG_PANEL,
            tc=TEXT_PRIMARY,
            w=170,
            h=30,
        )
        self.db_conn_toggle_btn.pack(side="left")

        # Connection details panel (active DB type; collapsible)
        self._db_conn_wrap = ctk.CTkFrame(
            main,
            fg_color=BG_CARD,
            corner_radius=8,
            border_width=1,
            border_color=DANGER,
        )
        lbl(self._db_conn_wrap, "CONNECTION DETAILS (ACTIVE DATABASE)", size=10, weight="bold", color=DANGER).pack(
            anchor="w", padx=10, pady=(8, 4)
        )

        conn_grid = ctk.CTkFrame(self._db_conn_wrap, fg_color="transparent")
        conn_grid.pack(fill="x", padx=8, pady=(0, 8))
        for c in range(3):
            conn_grid.columnconfigure(c, weight=1)

        def mk_field(row, col, key, title, show=""):
            cell = ctk.CTkFrame(conn_grid, fg_color="transparent")
            cell.grid(row=row, column=col, sticky="ew", padx=6, pady=4)
            lbl(cell, title, size=10, color=TEXT_MUTED).pack(anchor="w")
            ent = ctk.CTkEntry(
                cell,
                fg_color=BG_CARD2,
                border_color=BORDER,
                text_color=TEXT_PRIMARY,
                font=("Segoe UI", 11),
                height=30,
                show=show,
            )
            ent.pack(fill="x")
            self._db_conn_fields[key] = ent

        mk_field(0, 0, "host", "Host / IP")
        mk_field(0, 1, "port", "Port")
        mk_field(0, 2, "user", "User")
        mk_field(1, 0, "password", "Password", show="â—")
        mk_field(1, 1, "db_name", "Database")
        btn(
            conn_grid,
            "Save DB Details",
            self._save_active_db_connection_details,
            fg=BG_CARD2,
            hover=BG_PANEL,
            tc=TEXT_PRIMARY,
            w=140,
            h=30,
        ).grid(row=1, column=2, sticky="e", padx=6, pady=4)

        body = ctk.CTkFrame(main, fg_color="transparent")
        body.pack(fill="both", expand=True, padx=12, pady=10)

        # Left panel: quick queries
        left = ctk.CTkFrame(body, fg_color=BG_CARD, corner_radius=10, width=240)
        left.pack(side="left", fill="y", padx=(0, 8))
        left.pack_propagate(False)
        lbl(left, "QUICK QUERIES", size=10, color=TEXT_MUTED).pack(pady=(14, 4), padx=14, anchor="w")
        hsep(left)
        self._quick_wrap = ctk.CTkFrame(left, fg_color="transparent")
        self._quick_wrap.pack(fill="both", expand=True, padx=2, pady=(0, 6))
        self._render_quick_queries()

        # Right: editor + results
        right = ctk.CTkFrame(body, fg_color="transparent")
        right.pack(side="left", fill="both", expand=True)

        lbl(right, "QUERY EDITOR", size=10, color=TEXT_MUTED).pack(anchor="w", pady=(0, 4))
        self.editor = tk.Text(right, bg=BG_CARD, fg=TEXT_PRIMARY, font=("Consolas", 11),
                              height=7, relief="flat", padx=10, pady=8,
                              insertbackground=ACCENT)
        self.editor.pack(fill="x")

        toolbar = ctk.CTkFrame(right, fg_color="transparent")
        toolbar.pack(fill="x", pady=6)
        btn(toolbar, "â–¶ Run", self._run, w=120, h=34).pack(side="left")
        btn(toolbar, "âœ• Clear", self._clear, fg=BG_CARD2, hover=BG_CARD, tc=TEXT_MUTED, w=90, h=34).pack(side="left", padx=8)
        btn(
            toolbar,
            "ðŸ’¾ Save Query",
            self._save_current_query,
            fg=BG_CARD2,
            hover=BG_CARD,
            tc=TEXT_PRIMARY,
            w=130,
            h=34,
        ).pack(side="left", padx=8)
        self.q_status = lbl(toolbar, "", size=10, color=TEXT_MUTED)
        self.q_status.pack(side="right")

        lbl(right, "RESULTS", size=10, color=TEXT_MUTED).pack(anchor="w", pady=(6, 4))

        self._result_wrap = ctk.CTkFrame(right, fg_color="transparent")
        self._result_wrap.pack(fill="both", expand=True)

        # Plain text (errors, Postgres, non-SELECT MySQL messages)
        self._text_frame = ctk.CTkFrame(self._result_wrap, fg_color=BG_CARD, corner_radius=6)
        self.result = tk.Text(
            self._text_frame,
            bg=BG_CARD,
            fg=TEXT_PRIMARY,
            font=("Consolas", 10),
            relief="flat",
            padx=10,
            pady=8,
            state="disabled",
        )
        self.result.pack(fill="both", expand=True, padx=2, pady=2)

        # DBeaver-style grid (MySQL SELECT)
        self._grid_frame = tk.Frame(self._result_wrap, bg=BG_DARK)
        self._grid_footer = ctk.CTkLabel(
            self._grid_frame,
            text="",
            font=("Segoe UI", 10),
            text_color=TEXT_MUTED,
            anchor="w",
        )
        self._tree: Optional[ttk.Treeview] = None
        self._setup_result_grid()
        self._text_frame.pack(fill="both", expand=True)

        self._sync_mysql_combo_from_state()
        self._on_db_type_change()

    def _toggle_db_connection_config(self):
        self._db_conn_config_open = not self._db_conn_config_open
        if self._db_conn_config_open:
            self._db_conn_wrap.pack(fill="x", padx=12, pady=(6, 6), after=self.db_conn_toggle_btn.master)
            self.db_conn_toggle_btn.configure(text="Connection Config â–²")
            self._load_active_db_connection_details()
        else:
            self._db_conn_wrap.pack_forget()
            self.db_conn_toggle_btn.configure(text="Connection Config â–¼")

    def _setup_result_grid(self):
        style = ttk.Style(self)
        try:
            style.theme_use("clam")
        except tk.TclError:
            pass
        style.configure(
            "Grid.Treeview",
            background=BG_CARD,
            foreground=TEXT_PRIMARY,
            fieldbackground=BG_CARD,
            rowheight=22,
            font=("Segoe UI", 10),
        )
        style.configure(
            "Grid.Treeview.Heading",
            background=BG_PANEL,
            foreground=TEXT_PRIMARY,
            font=("Segoe UI", 10, "bold"),
            relief="flat",
        )
        style.map(
            "Grid.Treeview",
            background=[("selected", ACCENT2)],
            foreground=[("selected", "#ffffff")],
        )
        style.map("Grid.Treeview.Heading", background=[("active", BG_CARD2)])

        self._tree = ttk.Treeview(
            self._grid_frame,
            show="tree headings",
            style="Grid.Treeview",
            height=14,
        )
        vsb = ttk.Scrollbar(self._grid_frame, orient="vertical", command=self._tree.yview)
        hsb = ttk.Scrollbar(self._grid_frame, orient="horizontal", command=self._tree.xview)
        self._tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)

        self._grid_frame.grid_rowconfigure(0, weight=1)
        self._grid_frame.grid_columnconfigure(0, weight=1)
        self._tree.grid(row=0, column=0, sticky="nsew")
        vsb.grid(row=0, column=1, sticky="ns")
        hsb.grid(row=1, column=0, sticky="ew")
        self._grid_footer.grid(row=2, column=0, columnspan=2, sticky="ew", padx=4, pady=(4, 0))

    def _sync_mysql_combo_from_state(self):
        cur = (state.db_configs.get("MySQL", {}).get("db_name") or "").strip()
        if not cur:
            cur = MYSQL_DALLAS_DATABASES[0]
        vals = list(MYSQL_DALLAS_DATABASES)
        if cur not in vals:
            vals = [cur] + vals
        self.mysql_db_combo.configure(values=vals)
        self.mysql_db_combo.set(cur)

    def _on_mysql_db_combo(self, choice):
        state.db_configs.setdefault("MySQL", {})
        state.db_configs["MySQL"]["db_name"] = (choice or "").strip()
        if "db_name" in self._db_conn_fields:
            self._db_conn_fields["db_name"].delete(0, "end")
            self._db_conn_fields["db_name"].insert(0, state.db_configs["MySQL"]["db_name"])
        save_config()

    def _on_db_type_change(self, *args):
        if self.db_type.get() == "MySQL":
            self.mysql_db_row.pack(fill="x", after=self._tb)
            self._sync_mysql_combo_from_state()
        else:
            self.mysql_db_row.pack_forget()
        self._load_active_db_connection_details()

    def _load_active_db_connection_details(self):
        db_key = self.db_type.get()
        cfg = state.db_configs.get(db_key, {})
        defaults = {"host": "", "port": "5432", "user": "", "password": "", "db_name": ""}
        if db_key == "MySQL":
            defaults["port"] = "6446"
            defaults["db_name"] = (self.mysql_db_combo.get() or "").strip()
        for key, ent in self._db_conn_fields.items():
            ent.delete(0, "end")
            ent.insert(0, str(cfg.get(key, defaults.get(key, "")) or ""))

    def _save_active_db_connection_details(self):
        db_key = self.db_type.get()
        state.db_configs.setdefault(db_key, {})
        cfg = state.db_configs[db_key]
        cfg["host"] = self._db_conn_fields["host"].get().strip()
        cfg["port"] = self._db_conn_fields["port"].get().strip() or ("6446" if db_key == "MySQL" else "5432")
        cfg["user"] = self._db_conn_fields["user"].get().strip()
        cfg["password"] = self._db_conn_fields["password"].get().strip()
        cfg["db_name"] = self._db_conn_fields["db_name"].get().strip()
        if db_key == "MySQL":
            self._sync_mysql_combo_from_state()
        save_config()
        self.q_status.configure(
            text=f"Saved {db_key} connection details  â€¢  {datetime.now().strftime('%H:%M:%S')}",
            text_color=SUCCESS,
        )

    def _load(self, sql):
        self.editor.delete("1.0", "end")
        self.editor.insert("1.0", sql)

    def _render_quick_queries(self):
        for child in self._quick_wrap.winfo_children():
            child.destroy()

        saved = list(getattr(state, "saved_queries", []) or [])
        if saved:
            lbl(self._quick_wrap, "Saved", size=10, color=ACCENT).pack(
                anchor="w", padx=8, pady=(2, 4)
            )
            for item in saved:
                name = str(item.get("name", "")).strip() or "Saved Query"
                sql = str(item.get("sql", "")).strip()
                if not sql:
                    continue
                row = ctk.CTkFrame(self._quick_wrap, fg_color="transparent")
                row.pack(fill="x", padx=6, pady=2)
                ctk.CTkButton(
                    row,
                    text=f"â˜… {name}",
                    anchor="w",
                    fg_color=BG_CARD2,
                    hover_color=BG_PANEL,
                    text_color=TEXT_PRIMARY,
                    font=("Segoe UI", 11),
                    corner_radius=6,
                    height=32,
                    command=lambda s=sql: self._load(s),
                ).pack(side="left", fill="x", expand=True)
                ctk.CTkButton(
                    row,
                    text="âœŽ",
                    width=30,
                    height=32,
                    fg_color=BG_CARD2,
                    hover_color=BG_PANEL,
                    text_color=TEXT_MUTED,
                    font=("Segoe UI", 11, "bold"),
                    corner_radius=6,
                    command=lambda n=name: self._rename_saved_query(n),
                ).pack(side="left", padx=(4, 2))
                ctk.CTkButton(
                    row,
                    text="ðŸ—‘",
                    width=30,
                    height=32,
                    fg_color=BG_CARD2,
                    hover_color=BG_PANEL,
                    text_color=DANGER,
                    font=("Segoe UI", 11, "bold"),
                    corner_radius=6,
                    command=lambda n=name: self._delete_saved_query(n),
                ).pack(side="left")

    def _save_current_query(self):
        sql = self.editor.get("1.0", "end").strip()
        if not sql:
            messagebox.showinfo("Save Query", "Query editor is empty.")
            return

        default_name = "Saved Query"
        first_line = sql.splitlines()[0].strip() if sql.splitlines() else ""
        if first_line:
            default_name = first_line[:40]

        name = simpledialog.askstring(
            "Save Query",
            "Enter query name:",
            initialvalue=default_name,
            parent=self,
        )
        if name is None:
            return
        name = name.strip()
        if not name:
            messagebox.showwarning("Save Query", "Query name cannot be empty.")
            return

        saved = list(getattr(state, "saved_queries", []) or [])
        existing_idx = next(
            (i for i, q in enumerate(saved) if str(q.get("name", "")).strip().lower() == name.lower()),
            -1,
        )
        payload = {"name": name, "sql": sql}

        if existing_idx >= 0:
            saved[existing_idx] = payload
        else:
            if len(saved) >= 10:
                messagebox.showwarning(
                    "Save Query",
                    "Maximum 10 saved queries allowed. Delete/rename an existing one first.",
                )
                return
            saved.append(payload)

        state.saved_queries = saved
        save_config()
        self._render_quick_queries()
        self.q_status.configure(
            text=f"Saved query: {name}  â€¢  {datetime.now().strftime('%H:%M:%S')}",
            text_color=TEXT_MUTED,
        )

    def _rename_saved_query(self, old_name: str):
        saved = list(getattr(state, "saved_queries", []) or [])
        idx = next(
            (i for i, q in enumerate(saved) if str(q.get("name", "")).strip() == old_name),
            -1,
        )
        if idx < 0:
            return
        new_name = simpledialog.askstring(
            "Rename Query",
            "Enter new query name:",
            initialvalue=old_name,
            parent=self,
        )
        if new_name is None:
            return
        new_name = new_name.strip()
        if not new_name:
            messagebox.showwarning("Rename Query", "Query name cannot be empty.")
            return
        dup = next(
            (
                i
                for i, q in enumerate(saved)
                if i != idx and str(q.get("name", "")).strip().lower() == new_name.lower()
            ),
            -1,
        )
        if dup >= 0:
            messagebox.showwarning("Rename Query", f"A query named '{new_name}' already exists.")
            return
        saved[idx]["name"] = new_name
        state.saved_queries = saved
        save_config()
        self._render_quick_queries()
        self.q_status.configure(
            text=f"Renamed query to: {new_name}  â€¢  {datetime.now().strftime('%H:%M:%S')}",
            text_color=TEXT_MUTED,
        )

    def _delete_saved_query(self, name: str):
        if not messagebox.askyesno("Delete Query", f"Delete saved query '{name}'?", parent=self):
            return
        saved = list(getattr(state, "saved_queries", []) or [])
        new_saved = [q for q in saved if str(q.get("name", "")).strip() != name]
        if len(new_saved) == len(saved):
            return
        state.saved_queries = new_saved
        save_config()
        self._render_quick_queries()
        self.q_status.configure(
            text=f"Deleted query: {name}  â€¢  {datetime.now().strftime('%H:%M:%S')}",
            text_color=TEXT_MUTED,
        )

    def _run(self):
        sql      = self.editor.get("1.0", "end").strip()
        db_key   = self.db_type.get()        # "MySQL" or "Postgres"
        dcfg     = state.db_configs.get(db_key, {})
        self.q_status.configure(text="Runningâ€¦")
        def work():
            # MySQL: direct TCP from this machine (same as DBeaver / JDBC). No SSH.
            if db_key == "MySQL":
                host = (dcfg.get("host") or "").strip()
                if not host:
                    msg = (
                        "[NOT CONFIGURED] Set MySQL host, port, user, and password in "
                        "âš™ Config â†’ Database (same values as DBeaver)."
                    )
                    self.after(0, lambda m=msg: self._show_result_text(m))
                    return
                ok, cols, rows, info, err = mysql_query_to_grid(
                    host=host,
                    port=str(dcfg.get("port") or "3306"),
                    user=(dcfg.get("user") or "").strip(),
                    password=dcfg.get("password") or "",
                    database=(dcfg.get("db_name") or "").strip(),
                    sql=sql,
                )
                self.after(
                    0,
                    lambda o=ok, c=cols, r=rows, i=info, e=err: self._show_mysql_result(
                        o, c, r, i, e
                    ),
                )
                return
            # Postgres: direct TCP from this machine (same as DBeaver / JDBC). No SSH.
            host = (dcfg.get("host") or "").strip()
            if not host:
                msg = (
                    "[NOT CONFIGURED] Set Postgres host, port, user, password, and database in "
                    "âš™ Config â†’ Database."
                )
                self.after(0, lambda m=msg: self._show_result_text(m))
                return
            ok, cols, rows, info, err = postgres_query_to_grid(
                host=host,
                port=str(dcfg.get("port") or "5432"),
                user=(dcfg.get("user") or "").strip(),
                password=dcfg.get("password") or "",
                database=(dcfg.get("db_name") or "").strip(),
                sql=sql,
            )
            self.after(
                0,
                lambda o=ok, c=cols, r=rows, i=info, e=err: self._show_mysql_result(
                    o, c, r, i, e
                ),
            )
        threading.Thread(target=work, daemon=True).start()

    def _show_mysql_result(
        self,
        ok: bool,
        columns: Optional[List[str]],
        rows: Optional[List[Tuple[Any, ...]]],
        info: Optional[str],
        err: Optional[str],
    ):
        if not ok:
            self._show_result_text(err or "Unknown error")
            self.q_status.configure(
                text=f"Error  â€¢  {datetime.now().strftime('%H:%M:%S')}", text_color=DANGER
            )
            return
        if columns is not None and rows is not None:
            self._show_result_grid(columns, rows, info)
            n = len(rows)
            extra = f"  â€¢  {info}" if info else ""
            self.q_status.configure(
                text=f"{n} row{'s' if n != 1 else ''}{extra}  â€¢  {datetime.now().strftime('%H:%M:%S')}",
                text_color=TEXT_MUTED,
            )
            return
        self._show_result_text(info or "")
        self.q_status.configure(text=f"Done  â€¢  {datetime.now().strftime('%H:%M:%S')}")

    def _show_result_grid(self, columns: List[str], rows: List[Tuple], footer_note: Optional[str]):
        assert self._tree is not None
        self._text_frame.pack_forget()
        self._grid_frame.pack(fill="both", expand=True)

        for item in self._tree.get_children():
            self._tree.delete(item)

        n = len(columns)
        ids = [f"c{i}" for i in range(n)]
        self._tree["columns"] = tuple(ids)

        self._tree.heading("#0", text="#")
        self._tree.column("#0", width=44, minwidth=36, anchor="e", stretch=False)

        for i, col in enumerate(columns):
            cid = ids[i]
            disp = str(col)
            self._tree.heading(cid, text=disp, anchor="w")
            w = min(max(len(disp) * 7 + 28, 72), 300)
            self._tree.column(cid, width=w, minwidth=52, anchor="w", stretch=True)

        for i, row in enumerate(rows):
            tag = "even" if i % 2 == 0 else "odd"
            self._tree.insert("", "end", text=str(i + 1), values=row, tags=(tag,))
        self._tree.tag_configure("even", background=BG_CARD)
        self._tree.tag_configure("odd", background=BG_CARD2)

        self._grid_footer.configure(text=footer_note or "")

    def _show_result_text(self, text: str):
        self._grid_frame.pack_forget()
        self._text_frame.pack(fill="both", expand=True)
        self.result.configure(state="normal")
        self.result.delete("1.0", "end")
        self.result.insert("1.0", text)
        self.result.configure(state="disabled")
        lines = max(1, text.count("\n") + 1)
        self.q_status.configure(
            text=f"{lines} line{'s' if lines != 1 else ''}  â€¢  {datetime.now().strftime('%H:%M:%S')}",
            text_color=TEXT_MUTED,
        )

    def _clear(self):
        self.editor.delete("1.0", "end")
        self.result.configure(state="normal")
        self.result.delete("1.0", "end")
        self.result.configure(state="disabled")
        if self._tree is not None:
            for item in self._tree.get_children():
                self._tree.delete(item)
        self._grid_footer.configure(text="")
        self.q_status.configure(text="")



# ===== END FILE: page_database.py =====


# ===== BEGIN FILE: page_config.py =====

import os
import customtkinter as ctk
import tkinter as tk
from tkinter import messagebox
from config import (state, save_config, MYSQL_DALLAS_DATABASES,
                    BG_DARK, BG_PANEL, BG_CARD, BG_CARD2, ACCENT, ACCENT2,
                    SUCCESS, WARNING, DANGER, TEXT_PRIMARY, TEXT_MUTED, BORDER,
                    LOG_MODULES, SIDEBAR_W)
from ssh import ssh_mgr
from app_logging import get_run_log_path, get_logs_directory, analysis_hint_text
from widgets import lbl, btn, inp, hsep, badge, Sidebar, SIDEBAR_ITEMS

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  PAGE: Config
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
class PageConfig(ctk.CTkFrame):
    def __init__(self, master, on_back):
        super().__init__(master, fg_color=BG_DARK, corner_radius=0)
        self.on_back = on_back
        self._build()

    def _build(self):
        # â”€â”€ Top bar â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
        hdr = ctk.CTkFrame(self, fg_color=BG_PANEL, corner_radius=0, height=48)
        hdr.pack(fill="x")
        hdr.pack_propagate(False)
        ctk.CTkLabel(hdr, text="MONTOOL  â€”  Configuration",
                     font=("Courier New", 14, "bold"), text_color=ACCENT).pack(side="left", padx=16, pady=12)
        ctk.CTkButton(hdr, text="â† Back", command=self.on_back,
                      fg_color="transparent", hover_color=BG_CARD,
                      text_color=TEXT_MUTED, font=("Segoe UI", 11),
                      corner_radius=6, width=80, height=28).pack(side="right", padx=16)

        # â”€â”€ Scrollable body â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
        scroll = ctk.CTkScrollableFrame(self, fg_color="transparent")
        scroll.pack(fill="both", expand=True, padx=0, pady=0)

        body = ctk.CTkFrame(scroll, fg_color="transparent")
        body.pack(fill="both", expand=True, padx=24, pady=16)

        # â”€â”€ Section: Log Lines â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
        self._section_header(body, "GENERAL")
        gen = ctk.CTkFrame(body, fg_color=BG_CARD, corner_radius=12)
        gen.pack(fill="x", pady=(0, 16))
        lbl(gen, "Default log lines to fetch:", size=11, color=TEXT_MUTED).pack(
            side="left", padx=16, pady=14)
        self.log_lines_var = ctk.StringVar(value=str(state.log_lines))
        ctk.CTkEntry(gen, textvariable=self.log_lines_var,
                     fg_color=BG_CARD2, border_color=BORDER,
                     text_color=TEXT_PRIMARY, font=("Segoe UI", 12),
                     width=100, height=32).pack(side="left", padx=8)

        # â”€â”€ Section: Module SSH Configs â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
        self._section_header(body, "MODULE SSH CONNECTIONS")
        self._mod_fields = {}
        for module in LOG_MODULES:
            self._build_module_card(body, module)

        # â”€â”€ Section: Database Configs â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
        self._section_header(body, "DATABASE CONNECTIONS")
        self._db_fields = {}
        for db_key in ["MySQL", "Postgres"]:
            self._build_db_card(body, db_key)

        # â”€â”€ Section: Environment Hosts â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
        self._section_header(body, "ENVIRONMENT HOSTS")
        self._env_fields = {}
        env_card = ctk.CTkFrame(body, fg_color=BG_CARD, corner_radius=12)
        env_card.pack(fill="x", pady=(0, 16))
        for env_name, host in state.env_hosts.items():
            row = ctk.CTkFrame(env_card, fg_color="transparent")
            row.pack(fill="x", padx=16, pady=6)
            lbl(row, f"{env_name}:", size=11, color=TEXT_MUTED, width=160,
                anchor="w").pack(side="left")
            e = inp(row, ph="hostname or IP", w=400)
            e.insert(0, host)
            e.pack(side="left", padx=8)
            self._env_fields[env_name] = e

        # â”€â”€ Section: Application run log (MonTool diagnostics) â”€â”€â”€
        self._section_header(body, "APPLICATION RUN LOG")
        log_card = ctk.CTkFrame(body, fg_color=BG_CARD, corner_radius=12)
        log_card.pack(fill="x", pady=(0, 16))
        self._app_log_path_lbl = lbl(
            log_card,
            "Each app launch writes one UTF-8 file under application_logs\\",
            size=11,
            color=TEXT_MUTED,
        )
        self._app_log_path_lbl.pack(anchor="w", padx=16, pady=(12, 4))
        self._app_log_file_lbl = ctk.CTkLabel(
            log_card,
            text=self._app_log_path_display(),
            font=("Consolas", 10),
            text_color=TEXT_PRIMARY,
            anchor="w",
            justify="left",
            wraplength=920,
        )
        self._app_log_file_lbl.pack(anchor="w", padx=16, pady=(0, 10))
        log_row = ctk.CTkFrame(log_card, fg_color="transparent")
        log_row.pack(fill="x", padx=16, pady=(0, 14))
        btn(log_row, "Open folder", self._app_log_open_folder,
            fg=BG_CARD2, hover=BG_PANEL, tc=TEXT_MUTED, w=120, h=30).pack(side="left", padx=(0, 8))
        btn(log_row, "Open log file", self._app_log_open_file,
            fg=BG_CARD2, hover=BG_PANEL, tc=TEXT_MUTED, w=120, h=30).pack(side="left", padx=(0, 8))
        btn(log_row, "Copy path", self._app_log_copy_path,
            fg=BG_CARD2, hover=BG_PANEL, tc=TEXT_MUTED, w=100, h=30).pack(side="left", padx=(0, 8))
        btn(log_row, "Analyse logs", self._app_log_analyse,
            fg=ACCENT, hover=ACCENT2, tc="#000", w=130, h=30).pack(side="left", padx=(0, 8))

        # â”€â”€ Save button â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
        save_row = ctk.CTkFrame(body, fg_color="transparent")
        save_row.pack(fill="x", pady=(8, 24))
        self._status_lbl = lbl(save_row, "", size=11, color=SUCCESS)
        self._status_lbl.pack(side="right", padx=16)
        btn(save_row, "ðŸ’¾  Save All", self._save_all, w=160, h=40).pack(side="right")

    # â”€â”€ Helpers â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

    def _section_header(self, parent, text):
        lbl(parent, text, size=10, color=TEXT_MUTED).pack(anchor="w", pady=(8, 4))
        hsep(parent)

    def _app_log_path_display(self):
        p = get_run_log_path()
        return p if p else "(Log file is created when you start MonTool from main.py)"

    def _app_log_open_folder(self):
        d = get_logs_directory()
        if not os.path.isdir(d):
            messagebox.showwarning("Logs", f"Folder not found:\n{d}")
            return
        try:
            os.startfile(d)
        except Exception as e:
            messagebox.showerror("Logs", str(e))

    def _app_log_open_file(self):
        p = get_run_log_path()
        if not p or not os.path.isfile(p):
            messagebox.showinfo("Logs", "No session log file yet. Restart the app from main.py.")
            return
        try:
            os.startfile(p)
        except Exception as e:
            messagebox.showerror("Logs", str(e))

    def _app_log_copy_path(self):
        p = get_run_log_path()
        if not p:
            messagebox.showinfo("Logs", "No path yet. Start MonTool from main.py first.")
            return
        try:
            self.clipboard_clear()
            self.clipboard_append(p)
            self.update()
            messagebox.showinfo("Logs", "Full path copied to clipboard.")
        except Exception as e:
            messagebox.showerror("Logs", str(e))

    def _app_log_analyse(self):
        """Show how to analyse logs + optional tail of current session file."""
        top = ctk.CTkToplevel(self)
        top.title("Analyse application logs")
        top.geometry("720x480")
        top.configure(fg_color=BG_DARK)
        tx = ctk.CTkTextbox(top, font=("Consolas", 11), text_color=TEXT_PRIMARY,
                            fg_color=BG_CARD2, border_color=BORDER, border_width=1)
        tx.pack(fill="both", expand=True, padx=12, pady=12)
        body = analysis_hint_text() + "\n\n"
        p = get_run_log_path()
        if p and os.path.isfile(p):
            body += "â€” Last 200 lines from this session â€”\n\n"
            try:
                with open(p, encoding="utf-8", errors="replace") as f:
                    lines = f.readlines()
                body += "".join(lines[-200:])
            except Exception as e:
                body += f"(Could not read file: {e})"
        else:
            body += "(No log file on disk yet â€” run MonTool via main.py.)"
        tx.insert("1.0", body)
        tx.configure(state="disabled")
        btn(top, "Close", top.destroy, fg=BG_CARD2, hover=BG_CARD, tc=TEXT_MUTED, w=100, h=32).pack(pady=(0, 12))

    def _build_module_card(self, parent, module):
        mcfg = state.module_configs[module]

        card = ctk.CTkFrame(parent, fg_color=BG_CARD, corner_radius=12)
        card.pack(fill="x", pady=(0, 10))

        # Header row with module name + connect/disconnect button
        hrow = ctk.CTkFrame(card, fg_color="transparent")
        hrow.pack(fill="x", padx=16, pady=(12, 4))
        lbl(hrow, module, size=12, weight="bold", color=ACCENT).pack(side="left")

        conn_key = module
        status_lbl = lbl(hrow, self._conn_status(conn_key), size=10,
                         color=SUCCESS if ssh_mgr.is_connected(conn_key) else TEXT_MUTED)
        status_lbl.pack(side="right", padx=8)

        conn_btn = ctk.CTkButton(hrow, text=self._conn_btn_text(conn_key),
                                 fg_color=BG_CARD2, hover_color=BG_PANEL,
                                 text_color=TEXT_MUTED, font=("Segoe UI", 11),
                                 corner_radius=6, width=110, height=28,
                                 border_color=BORDER, border_width=1)
        conn_btn.pack(side="right")
        conn_btn.configure(command=lambda m=module, b=conn_btn, s=status_lbl:
                           self._toggle_connection(m, b, s))

        hsep(card)

        fields = {}
        rows = [
            ("host",      "Host / IP",      "",                False),
            ("port",      "Port",           "22",              False),
            ("user",      "Username (LDAP)","",                False),
            ("password",  "Password",       "",                True),
            ("key",       "SSH Key Path",   "/path/to/key",    False),
            ("log_path",  "Log File Path",  "/var/log/app.log",False),
            ("sudo_user", "Sudo User",      "aqadmin",         False),
        ]
        grid = ctk.CTkFrame(card, fg_color="transparent")
        grid.pack(fill="x", padx=16, pady=(4, 12))
        for col in range(4):
            grid.columnconfigure(col, weight=1)

        for i, (key, label, placeholder, secret) in enumerate(rows):
            r, c = divmod(i, 2)
            cell = ctk.CTkFrame(grid, fg_color="transparent")
            cell.grid(row=r, column=c*2, columnspan=2, sticky="ew", padx=4, pady=4)
            lbl(cell, label, size=10, color=TEXT_MUTED).pack(anchor="w")
            e = inp(cell, ph=placeholder, w=300, show="â—" if secret else "")
            e.insert(0, mcfg.get(key, ""))
            e.pack(fill="x")
            fields[key] = e

        self._mod_fields[module] = fields

    def _build_db_card(self, parent, db_key):
        dcfg = state.db_configs[db_key]
        color = ACCENT if db_key == "MySQL" else WARNING

        card = ctk.CTkFrame(parent, fg_color=BG_CARD, corner_radius=12)
        card.pack(fill="x", pady=(0, 10))

        hrow = ctk.CTkFrame(card, fg_color="transparent")
        hrow.pack(fill="x", padx=16, pady=(12, 4))
        lbl(hrow, db_key, size=12, weight="bold", color=color).pack(side="left")

        conn_key = f"DB_{db_key}"
        if db_key in ("MySQL", "Postgres"):
            lbl(
                hrow,
                "Direct TCP from this PC (like DBeaver) â€” no SSH Connect",
                size=10,
                color=TEXT_MUTED,
            ).pack(side="right", padx=8)
        else:
            status_lbl = lbl(hrow, self._conn_status(conn_key), size=10,
                             color=SUCCESS if ssh_mgr.is_connected(conn_key) else TEXT_MUTED)
            status_lbl.pack(side="right", padx=8)

            conn_btn = ctk.CTkButton(hrow, text=self._conn_btn_text(conn_key),
                                     fg_color=BG_CARD2, hover_color=BG_PANEL,
                                     text_color=TEXT_MUTED, font=("Segoe UI", 11),
                                     corner_radius=6, width=110, height=28,
                                     border_color=BORDER, border_width=1)
            conn_btn.pack(side="right")
            conn_btn.configure(command=lambda dk=db_key, ck=conn_key, b=conn_btn, s=status_lbl:
                               self._toggle_db_connection(dk, ck, b, s))

        hsep(card)

        fields = {}
        if db_key == "MySQL":
            rows = [
                ("host",     "MySQL host", ""),
                ("port",     "MySQL port", "6446"),
                ("user",     "MySQL user", ""),
                ("password", "MySQL password", ""),
            ]
        else:
            rows = [
                ("host",     "Postgres host",  ""),
                ("port",     "Postgres port",  "5432"),
                ("user",     "Postgres user",  ""),
                ("password", "Postgres password", ""),
                ("db_name",  "Database",       ""),
            ]
        grid = ctk.CTkFrame(card, fg_color="transparent")
        grid.pack(fill="x", padx=16, pady=(4, 12))

        for i, (key, label, placeholder) in enumerate(rows):
            r, c = divmod(i, 2)
            cell = ctk.CTkFrame(grid, fg_color="transparent")
            cell.grid(row=r, column=c, sticky="ew", padx=8, pady=4)
            grid.columnconfigure(c, weight=1)
            lbl(cell, label, size=10, color=TEXT_MUTED).pack(anchor="w")
            secret = key == "password"
            e = inp(cell, ph=placeholder, w=260, show="â—" if secret else "")
            e.insert(0, dcfg.get(key, ""))
            e.pack(fill="x")
            fields[key] = e

        if db_key == "MySQL":
            db_row = ctk.CTkFrame(card, fg_color="transparent")
            db_row.pack(fill="x", padx=16, pady=(0, 12))
            lbl(db_row, "Database", size=10, color=TEXT_MUTED).pack(anchor="w")
            cur = (dcfg.get("db_name") or "").strip() or MYSQL_DALLAS_DATABASES[0]
            combo_vals = list(MYSQL_DALLAS_DATABASES)
            if cur and cur not in combo_vals:
                combo_vals = [cur] + combo_vals
            db_combo = ctk.CTkComboBox(
                db_row,
                values=combo_vals,
                width=560,
                height=32,
                fg_color=BG_CARD2,
                button_color=ACCENT2,
                button_hover_color=ACCENT,
                dropdown_fg_color=BG_CARD,
                font=("Segoe UI", 11),
                text_color=TEXT_PRIMARY,
            )
            db_combo.set(cur)
            db_combo.pack(anchor="w", fill="x", pady=(4, 0))
            fields["db_name"] = db_combo

        self._db_fields[db_key] = fields

    def _conn_status(self, key):
        return "â— Connected" if ssh_mgr.is_connected(key) else "â—‹ Disconnected"

    def _conn_btn_text(self, key):
        return "Disconnect" if ssh_mgr.is_connected(key) else "Connect"

    def _toggle_connection(self, module, btn_widget, status_lbl):
        conn_key = module
        if ssh_mgr.is_connected(conn_key):
            ssh_mgr.disconnect(conn_key)
            btn_widget.configure(text="Connect")
            status_lbl.configure(text="â—‹ Disconnected", text_color=TEXT_MUTED)
            return

        # Pull current field values (in case user hasn't saved yet)
        flds = self._mod_fields[module]
        host = flds["host"].get().strip()
        if not host:
            status_lbl.configure(text="âœ— No host set", text_color=DANGER)
            return

        status_lbl.configure(text="Connectingâ€¦", text_color=WARNING)
        btn_widget.configure(state="disabled")

        import threading
        def work():
            ldap_pwd   = flds["password"].get().strip() or None
            sudo_user  = flds["sudo_user"].get().strip() or None
            # sudo uses the same LDAP password by default
            ok, msg = ssh_mgr.connect(
                host=conn_key,          # logical key
                user=flds["user"].get().strip(),
                password=ldap_pwd,
                key_path=flds["key"].get().strip() or None,
                port=int(flds["port"].get().strip() or 22),
                actual_host=host,
                sudo_user=sudo_user,
                sudo_password=ldap_pwd,
            )
            def update():
                btn_widget.configure(state="normal",
                                     text="Disconnect" if ok else "Connect")
                status_lbl.configure(
                    text="â— Connected" if ok else f"âœ— {msg}",
                    text_color=SUCCESS if ok else DANGER)
            self.after(0, update)
        threading.Thread(target=work, daemon=True).start()

    def _toggle_db_connection(self, db_key, conn_key, btn_widget, status_lbl):
        if ssh_mgr.is_connected(conn_key):
            ssh_mgr.disconnect(conn_key)
            btn_widget.configure(text="Connect")
            status_lbl.configure(text="â—‹ Disconnected", text_color=TEXT_MUTED)
            return

        flds = self._db_fields[db_key]
        host = flds["host"].get().strip()
        if not host:
            status_lbl.configure(text="âœ— No host set", text_color=DANGER)
            return

        status_lbl.configure(text="Connectingâ€¦", text_color=WARNING)
        btn_widget.configure(state="disabled")

        import threading
        def work():
            ok, msg = ssh_mgr.connect(
                host=conn_key,
                user=flds["user"].get().strip(),
                password=flds["password"].get().strip() or None,
                key_path=None,
                port=22,
                actual_host=host,
            )
            def update():
                btn_widget.configure(state="normal",
                                     text="Disconnect" if ok else "Connect")
                status_lbl.configure(
                    text="â— Connected" if ok else f"âœ— {msg}",
                    text_color=SUCCESS if ok else DANGER)
            self.after(0, update)
        threading.Thread(target=work, daemon=True).start()

    def _save_all(self):
        # General
        try:
            state.log_lines = int(self.log_lines_var.get())
        except ValueError:
            pass

        # Module configs
        for module, fields in self._mod_fields.items():
            for key, widget in fields.items():
                state.module_configs[module][key] = widget.get()

        # DB configs
        for db_key, fields in self._db_fields.items():
            for key, widget in fields.items():
                state.db_configs[db_key][key] = widget.get()

        # Env hosts
        for env_name, widget in self._env_fields.items():
            state.env_hosts[env_name] = widget.get()

        save_config()
        self._status_lbl.configure(text="âœ“ Saved", text_color=SUCCESS)
        self.after(3000, lambda: self._status_lbl.configure(text=""))

# ===== END FILE: page_config.py =====


# ===== BEGIN FILE: page_log_analysis.py =====

import os
import customtkinter as ctk
from tkinter import messagebox
from config import (BG_DARK, BG_PANEL, BG_CARD, BG_CARD2, ACCENT, ACCENT2,
                    SUCCESS, TEXT_PRIMARY, TEXT_MUTED, BORDER)
from app_logging import get_run_log_path, get_logs_directory, analysis_hint_text
from widgets import lbl, btn, hsep, Sidebar, SIDEBAR_ITEMS


class PageLogAnalysis(ctk.CTkFrame):
    """MonTool application run logs â€” open, copy, or review for analysis."""

    def __init__(self, master, on_nav, on_back, on_config):
        super().__init__(master, fg_color=BG_DARK, corner_radius=0)
        self._build(on_nav, on_back, on_config)

    def _build(self, on_nav, on_back, on_config):
        Sidebar(self, SIDEBAR_ITEMS, active="Log Analysis",
                on_select=on_nav, on_back=on_back, on_config=on_config).pack(side="left", fill="y")

        main = ctk.CTkFrame(self, fg_color=BG_DARK, corner_radius=0)
        main.pack(side="left", fill="both", expand=True)

        hdr = ctk.CTkFrame(main, fg_color=BG_PANEL, corner_radius=0, height=48)
        hdr.pack(fill="x")
        hdr.pack_propagate(False)
        ctk.CTkLabel(hdr, text="Log Analysis",
                     font=("Courier New", 16, "bold"), text_color=ACCENT).pack(side="left", padx=16, pady=12)

        body = ctk.CTkFrame(main, fg_color="transparent")
        body.pack(fill="both", expand=True, padx=24, pady=16)

        card = ctk.CTkFrame(body, fg_color=BG_CARD, corner_radius=12)
        card.pack(fill="both", expand=True)

        lbl(card, "APPLICATION RUN LOG (this MonTool session)", size=10, color=TEXT_MUTED).pack(
            anchor="w", padx=16, pady=(16, 4))
        hsep(card)
        self.path_lbl = ctk.CTkLabel(
            card, text=self._path_display(), font=("Consolas", 11),
            text_color=TEXT_PRIMARY, anchor="w", justify="left", wraplength=900,
        )
        self.path_lbl.pack(anchor="w", padx=16, pady=(8, 12))

        row = ctk.CTkFrame(card, fg_color="transparent")
        row.pack(fill="x", padx=16, pady=(0, 12))
        btn(row, "Open folder", self._open_folder, fg=BG_CARD2, hover=BG_CARD, tc=TEXT_MUTED, w=120, h=32).pack(
            side="left", padx=(0, 8))
        btn(row, "Open log file", self._open_file, fg=BG_CARD2, hover=BG_CARD, tc=TEXT_MUTED, w=120, h=32).pack(
            side="left", padx=(0, 8))
        btn(row, "Copy path", self._copy_path, fg=BG_CARD2, hover=BG_CARD, tc=TEXT_MUTED, w=100, h=32).pack(
            side="left", padx=(0, 8))
        btn(row, "Preview in window", self._preview, fg=ACCENT, hover=ACCENT2, tc="#000", w=150, h=32).pack(
            side="left", padx=(0, 8))

        lbl(card, "How to analyse", size=10, color=TEXT_MUTED).pack(anchor="w", padx=16, pady=(12, 4))
        hsep(card)
        hint = ctk.CTkTextbox(card, font=("Segoe UI", 12), text_color=TEXT_PRIMARY,
                              fg_color=BG_CARD2, border_color=BORDER, border_width=1, height=120)
        hint.pack(fill="x", padx=16, pady=(0, 16))
        hint.insert("1.0", analysis_hint_text())
        hint.configure(state="disabled")

    def _path_display(self):
        p = get_run_log_path()
        return p if p else "(Start MonTool via main.py to create a session log file.)"

    def _open_folder(self):
        d = get_logs_directory()
        if not os.path.isdir(d):
            messagebox.showwarning("Logs", f"Folder not found:\n{d}")
            return
        try:
            os.startfile(d)
        except Exception as e:
            messagebox.showerror("Logs", str(e))

    def _open_file(self):
        p = get_run_log_path()
        if not p or not os.path.isfile(p):
            messagebox.showinfo("Logs", "No session log file yet. Restart from main.py.")
            return
        try:
            os.startfile(p)
        except Exception as e:
            messagebox.showerror("Logs", str(e))

    def _copy_path(self):
        p = get_run_log_path()
        if not p:
            messagebox.showinfo("Logs", "No path yet. Start MonTool from main.py first.")
            return
        try:
            self.clipboard_clear()
            self.clipboard_append(p)
            self.update()
            messagebox.showinfo("Logs", "Full path copied to clipboard.")
        except Exception as e:
            messagebox.showerror("Logs", str(e))

    def _preview(self):
        top = ctk.CTkToplevel(self)
        top.title("Log preview â€” analysis")
        top.geometry("720x480")
        top.configure(fg_color=BG_DARK)
        tx = ctk.CTkTextbox(top, font=("Consolas", 11), text_color=TEXT_PRIMARY,
                            fg_color=BG_CARD2, border_color=BORDER, border_width=1)
        tx.pack(fill="both", expand=True, padx=12, pady=12)
        body = analysis_hint_text() + "\n\n"
        p = get_run_log_path()
        if p and os.path.isfile(p):
            body += "â€” Last 200 lines from this session â€”\n\n"
            try:
                with open(p, encoding="utf-8", errors="replace") as f:
                    lines = f.readlines()
                body += "".join(lines[-200:])
            except Exception as e:
                body += f"(Could not read file: {e})"
        else:
            body += "(No log file on disk yet â€” run MonTool via main.py.)"
        tx.insert("1.0", body)
        tx.configure(state="disabled")
        btn(top, "Close", top.destroy, fg=BG_CARD2, hover=BG_CARD, tc=TEXT_MUTED, w=100, h=32).pack(pady=(0, 12))

# ===== END FILE: page_log_analysis.py =====


# ===== BEGIN FILE: page_mocksim.py =====

import customtkinter as ctk
import tkinter as tk
from tkinter import filedialog
import threading
from config import (BG_DARK, BG_PANEL, BG_CARD, BG_CARD2, ACCENT, ACCENT2,
                    SUCCESS, WARNING, DANGER, TEXT_PRIMARY, TEXT_MUTED, BORDER, state)
from ssh import ssh_mgr
from widgets import lbl, btn, inp, hsep, Sidebar, SIDEBAR_ITEMS, centered

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  PAGE: Mock & Sim Activation
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
class PageMockSim(ctk.CTkFrame):
    def __init__(self, master, on_nav, on_back, on_config):
        super().__init__(master, fg_color=BG_DARK, corner_radius=0)
        self.file_path = ctk.StringVar()
        self._build(on_nav, on_back, on_config)

    def _build(self, on_nav, on_back, on_config):
        Sidebar(self, SIDEBAR_ITEMS, active="Mock & Sim Activation",
                on_select=on_nav, on_back=on_back, on_config=on_config).pack(side="left", fill="y")

        main = ctk.CTkFrame(self, fg_color=BG_DARK, corner_radius=0)
        main.pack(side="left", fill="both", expand=True)

        outer, c = centered(main)

        f = ctk.CTkFrame(c, fg_color=BG_CARD, corner_radius=16, width=520)
        f.pack()
        f.pack_propagate(False)

        lbl(f, "MOCK & SIM ACTIVATION", size=14, weight="bold", color=ACCENT).pack(pady=(28, 4))
        lbl(f, "Upload a config file to activate on the target host",
            size=11, color=TEXT_MUTED).pack(pady=(0, 16), padx=24)
        hsep(f)

        lbl(f, "ACTIVATION TYPE", size=10, color=TEXT_MUTED).pack(pady=(14, 6), padx=24, anchor="w")
        self.act_type = ctk.StringVar(value="Mock")
        tr = ctk.CTkFrame(f, fg_color="transparent")
        tr.pack(anchor="w", padx=24, pady=(0, 14))
        for t in ["Mock", "Kim", "Activation"]:
            ctk.CTkRadioButton(tr, text=t, variable=self.act_type, value=t,
                               fg_color=ACCENT, hover_color=ACCENT2, text_color=TEXT_PRIMARY,
                               font=("Segoe UI", 12)).pack(side="left", padx=14)

        lbl(f, "UPLOAD FILE", size=10, color=TEXT_MUTED).pack(pady=(0, 6), padx=24, anchor="w")
        fr = ctk.CTkFrame(f, fg_color="transparent")
        fr.pack(padx=24, fill="x", pady=(0, 10))
        ctk.CTkEntry(fr, textvariable=self.file_path,
                     fg_color=BG_CARD2, border_color=BORDER, text_color=TEXT_PRIMARY,
                     font=("Segoe UI", 11), width=330, height=34,
                     placeholder_text="No file selected",
                     placeholder_text_color=TEXT_MUTED).pack(side="left")
        ctk.CTkButton(fr, text="Browseâ€¦", command=self._browse,
                      fg_color=BG_CARD2, hover_color=BG_PANEL, text_color=TEXT_MUTED,
                      font=("Segoe UI", 11), corner_radius=6, width=110, height=34,
                      border_color=BORDER, border_width=1).pack(side="left", padx=8)

        self.prog = ctk.CTkProgressBar(f, fg_color=BG_CARD2, progress_color=ACCENT, height=6)
        self.prog.pack(fill="x", padx=24, pady=(12, 4))
        self.prog.set(0)
        self.stat = lbl(f, "Ready to upload", size=11, color=TEXT_MUTED)
        self.stat.pack(pady=(0, 14))
        btn(f, "â¬†  Submit & Activate", self._submit, w=460, h=44).pack(padx=24, pady=(0, 28))

    def _browse(self):
        p = filedialog.askopenfilename(
            filetypes=[("Config files", "*.json *.yaml *.yml *.conf *.csv"), ("All", "*.*")])
        if p:
            self.file_path.set(p)

    def _submit(self):
        path = self.file_path.get()
        if not path:
            messagebox.showwarning("No File", "Please select a file first.")
            return
        host = state.env_hosts.get(state.env, "")
        def work():
            self.after(0, lambda: self.stat.configure(text="Uploadingâ€¦", text_color=WARNING))
            for i in range(1, 11):
                time.sleep(0.15)
                self.after(0, lambda v=i/10: self.prog.set(v))
            if host in ssh_mgr.connections:
                sftp = ssh_mgr.connections[host].open_sftp()
                remote = f"/tmp/{os.path.basename(path)}"
                sftp.put(path, remote)
                sftp.close()
                out, _ = ssh_mgr.run(host, f"activate_mock.sh {remote} 2>&1")
                msg = out or "Activation complete"
            else:
                msg = f"[DEMO] '{os.path.basename(path)}' uploaded & activated"
            self.after(0, lambda: self.stat.configure(text=f"âœ“ {msg}", text_color=SUCCESS))
        threading.Thread(target=work, daemon=True).start()



# ===== END FILE: page_mocksim.py =====


# ===== BEGIN FILE: page_invoice.py =====

import customtkinter as ctk
import tkinter as tk
import threading
import time
from config import (BG_DARK, BG_PANEL, BG_CARD, BG_CARD2, ACCENT, ACCENT2,
                    SUCCESS, WARNING, DANGER, TEXT_PRIMARY, TEXT_MUTED, BORDER, state)
from ssh import ssh_mgr
from widgets import lbl, btn, inp, hsep, Sidebar, SIDEBAR_ITEMS, centered

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  PAGE: Invoice Pregeneration
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
class PageInvoice(ctk.CTkFrame):
    def __init__(self, master, on_nav, on_back, on_config):
        super().__init__(master, fg_color=BG_DARK, corner_radius=0)
        self._build(on_nav, on_back, on_config)

    def _build(self, on_nav, on_back, on_config):
        Sidebar(self, SIDEBAR_ITEMS, active="Invoice Pregeneration",
                on_select=on_nav, on_back=on_back, on_config=on_config).pack(side="left", fill="y")

        main = ctk.CTkFrame(self, fg_color=BG_DARK, corner_radius=0)
        main.pack(side="left", fill="both", expand=True)

        outer, c = centered(main)

        f = ctk.CTkFrame(c, fg_color=BG_CARD, corner_radius=16, width=560)
        f.pack()
        f.pack_propagate(False)

        lbl(f, "INVOICE PREGENERATION", size=14, weight="bold", color=ACCENT).pack(pady=(28, 4))
        lbl(f, "Trigger bulk invoice generation for a billing cycle",
            size=11, color=TEXT_MUTED).pack(pady=(0, 16), padx=24)
        hsep(f)

        form = ctk.CTkFrame(f, fg_color="transparent")
        form.pack(padx=24, fill="x", pady=12)
        self.inv_fields = {}
        for lbl_text, ph in [("Billing Cycle", "YYYY-MM"),
                              ("Account Range From", "e.g. ACC-0001"),
                              ("Account Range To",   "e.g. ACC-9999"),
                              ("Batch Size",         "e.g. 500")]:
            lbl(form, lbl_text, size=11, color=TEXT_MUTED).pack(anchor="w", pady=(10, 2))
            e = inp(form, ph=ph, w=508)
            e.pack(fill="x")
            self.inv_fields[lbl_text] = e

        self.dry_run = ctk.BooleanVar(value=True)
        ctk.CTkCheckBox(form, text="Dry Run (preview only)",
                        variable=self.dry_run, fg_color=ACCENT, hover_color=ACCENT2,
                        text_color=TEXT_PRIMARY, font=("Segoe UI", 12)).pack(anchor="w", pady=(14, 0))

        self.inv_log = tk.Text(f, bg=BG_CARD2, fg=TEXT_PRIMARY, font=("Consolas", 10),
                               height=5, relief="flat", padx=10, pady=8, state="disabled")
        self.inv_log.pack(fill="x", padx=24, pady=12)
        btn(f, "â–¶ Generate Invoices", self._generate, w=508, h=44).pack(padx=24, pady=(0, 28))

    def _generate(self):
        vals = {k: v.get() for k, v in self.inv_fields.items()}
        dry  = self.dry_run.get()
        host = state.env_hosts.get(state.env, "")
        def work():
            self.after(0, lambda: self._log("Starting invoice generation...\n", WARNING))
            if host in ssh_mgr.connections:
                cmd = (f"invoice_gen.sh --cycle {vals['Billing Cycle']} "
                       f"--from {vals['Account Range From']} "
                       f"--to {vals['Account Range To']} "
                       f"--batch {vals['Batch Size']} {'--dry-run' if dry else ''} 2>&1")
                out, err = ssh_mgr.run(host, cmd)
                self.after(0, lambda: self._log(out or err, SUCCESS))
            else:
                for line in [f"Cycle: {vals['Billing Cycle']}", f"Batch: {vals['Batch Size']}",
                             f"Mode: {'DRY RUN' if dry else 'LIVE'}", "â”€"*40,
                             "[DEMO] 1200 accounts queued", "[DEMO] 0 errors", "âœ“ Done"]:
                    time.sleep(0.1)
                    self.after(0, lambda l=line: self._log(l+"\n", SUCCESS))
        threading.Thread(target=work, daemon=True).start()

    def _log(self, text, color=TEXT_PRIMARY):
        self.inv_log.configure(state="normal")
        self.inv_log.insert("end", text)
        self.inv_log.see("end")
        self.inv_log.configure(state="disabled")



# ===== END FILE: page_invoice.py =====


# ===== BEGIN FILE: page_consumption.py =====

import customtkinter as ctk
import tkinter as tk
import threading
from config import (BG_DARK, BG_PANEL, BG_CARD, BG_CARD2, ACCENT, ACCENT2,
                    SUCCESS, WARNING, DANGER, TEXT_PRIMARY, TEXT_MUTED, BORDER, state)
from ssh import ssh_mgr
from widgets import lbl, btn, inp, hsep, badge, Sidebar, SIDEBAR_ITEMS

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  PAGE: Consumption
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
class PageConsumption(ctk.CTkFrame):
    def __init__(self, master, on_nav, on_back, on_config):
        super().__init__(master, fg_color=BG_DARK, corner_radius=0)
        self.tab_type = ctk.StringVar(value="Voice")
        self._build(on_nav, on_back, on_config)

    def _build(self, on_nav, on_back, on_config):
        Sidebar(self, SIDEBAR_ITEMS, active="Consumption",
                on_select=on_nav, on_back=on_back, on_config=on_config).pack(side="left", fill="y")

        main = ctk.CTkFrame(self, fg_color=BG_DARK, corner_radius=0)
        main.pack(side="left", fill="both", expand=True)

        # Type tabs
        tb = ctk.CTkFrame(main, fg_color=BG_PANEL, corner_radius=0, height=46)
        tb.pack(fill="x")
        tb.pack_propagate(False)
        self.tab_colors = {"Voice": SUCCESS, "SMS": WARNING, "Data": ACCENT}
        self.tab_btns = {}
        for t in ["Voice", "SMS", "Data"]:
            c = self.tab_colors[t]
            b = ctk.CTkButton(tb, text=t,
                              fg_color=c if t == "Voice" else "transparent",
                              hover_color=BG_CARD,
                              text_color="#000" if t == "Voice" else TEXT_MUTED,
                              font=("Segoe UI", 12, "bold"), corner_radius=6,
                              height=30, width=110, command=lambda x=t: self._switch(x))
            b.pack(side="left", padx=6, pady=8)
            self.tab_btns[t] = b

        body = ctk.CTkFrame(main, fg_color="transparent")
        body.pack(fill="both", expand=True, padx=24, pady=20)

        # Form
        form = ctk.CTkFrame(body, fg_color=BG_CARD, corner_radius=12, width=380)
        form.pack(side="left", fill="y", padx=(0, 16))
        form.pack_propagate(False)

        lbl(form, "QUERY PARAMETERS", size=10, color=TEXT_MUTED).pack(pady=(16, 4), padx=16, anchor="w")
        hsep(form)

        self.fields = {}
        for field, ph in [("IMSI", "e.g. 310150123456789"),
                          ("MSISDN", "e.g. +1234567890"),
                          ("Message Type", "e.g. MT, MO"),
                          ("Amount", "e.g. 100")]:
            lbl(form, field, size=11, color=TEXT_MUTED).pack(pady=(10, 2), padx=16, anchor="w")
            e = inp(form, ph=ph, w=340)
            e.pack(padx=16)
            self.fields[field] = e

        btn(form, "â–¶ Submit Query", self._submit, w=340, h=42).pack(pady=20, padx=16)

        # Results
        res = ctk.CTkFrame(body, fg_color=BG_CARD, corner_radius=12)
        res.pack(side="left", fill="both", expand=True)
        lbl(res, "RESULTS", size=10, color=TEXT_MUTED).pack(pady=(16, 4), padx=16, anchor="w")
        hsep(res)
        self.result_box = tk.Text(res, bg=BG_CARD, fg=TEXT_PRIMARY, font=("Consolas", 10),
                                  relief="flat", padx=12, pady=10, state="disabled")
        self.result_box.pack(fill="both", expand=True, padx=4, pady=4)

    def _switch(self, tab):
        self.tab_type.set(tab)
        for t, b in self.tab_btns.items():
            c = self.tab_colors[t]
            a = t == tab
            b.configure(fg_color=c if a else "transparent",
                        text_color="#000" if a else TEXT_MUTED)

    def _submit(self):
        vals = {k: v.get() for k, v in self.fields.items()}
        tab  = self.tab_type.get()
        host = state.env_hosts.get(state.env, "")
        def work():
            if host in ssh_mgr.connections:
                cmd = (f"mysql -e \"SELECT * FROM consumption WHERE type='{tab}' "
                       f"AND imsi='{vals['IMSI']}' LIMIT 50;\" 2>&1")
                out, _ = ssh_mgr.run(host, cmd)
            else:
                import random
                rows = [f"{'Type':<10} {'IMSI':<20} {'MSISDN':<16} {'Amount':>10}", "â”€"*60]
                for _ in range(12):
                    rows.append(f"{tab:<10} {vals['IMSI'] or '310150XXXXXXX':<20} "
                                f"{vals['MSISDN'] or '+1XXXXXXXXXX':<16} {random.randint(10,9999):>10}")
                rows.append("â”€"*60 + "\n[DEMO MODE]")
                out = "\n".join(rows)
            self.after(0, lambda: self._show(out))
        threading.Thread(target=work, daemon=True).start()

    def _show(self, text):
        self.result_box.configure(state="normal")
        self.result_box.delete("1.0", "end")
        self.result_box.insert("1.0", text)
        self.result_box.configure(state="disabled")



# ===== END FILE: page_consumption.py =====


# ===== BEGIN FILE: page_sim_deletion.py =====

import customtkinter as ctk
from tkinter import messagebox
from config import (BG_DARK, BG_PANEL, BG_CARD, BG_CARD2, ACCENT, ACCENT2,
                    TEXT_PRIMARY, TEXT_MUTED, BORDER)
from widgets import lbl, btn, inp, hsep, Sidebar, SIDEBAR_ITEMS


class PageSimDeletion(ctk.CTkFrame):
    """SIM deletion workflow â€” placeholder for bulk / single SIM removal."""

    def __init__(self, master, on_nav, on_back, on_config):
        super().__init__(master, fg_color=BG_DARK, corner_radius=0)
        self._build(on_nav, on_back, on_config)

    def _build(self, on_nav, on_back, on_config):
        Sidebar(self, SIDEBAR_ITEMS, active="Sim Deletion",
                on_select=on_nav, on_back=on_back, on_config=on_config).pack(side="left", fill="y")

        main = ctk.CTkFrame(self, fg_color=BG_DARK, corner_radius=0)
        main.pack(side="left", fill="both", expand=True)

        hdr = ctk.CTkFrame(main, fg_color=BG_PANEL, corner_radius=0, height=48)
        hdr.pack(fill="x")
        hdr.pack_propagate(False)
        ctk.CTkLabel(hdr, text="SIM Deletion",
                     font=("Courier New", 16, "bold"), text_color=ACCENT).pack(side="left", padx=16, pady=12)

        body = ctk.CTkScrollableFrame(main, fg_color="transparent")
        body.pack(fill="both", expand=True, padx=24, pady=16)

        card = ctk.CTkFrame(body, fg_color=BG_CARD, corner_radius=12)
        card.pack(fill="x", pady=(0, 12))

        self.iccid = inp(card, ph="ICCID (optional)", w=400)
        self.iccid.pack(padx=16, pady=(16, 8), anchor="w")
        self.imsi = inp(card, ph="IMSI (optional)", w=400)
        self.imsi.pack(padx=16, pady=8, anchor="w")
        self.msisdn = inp(card, ph="MSISDN (optional)", w=400)
        self.msisdn.pack(padx=16, pady=8, anchor="w")

        row = ctk.CTkFrame(card, fg_color="transparent")
        row.pack(fill="x", padx=16, pady=16)
        btn(row, "Submit deletion request", self._submit, w=220, h=40).pack(side="left")

    def _submit(self):
        messagebox.showinfo(
            "Sim Deletion",
            "Deletion workflow is not connected yet.\n"
            "Hook this action to your API or job runner.",
        )

# ===== END FILE: page_sim_deletion.py =====


# ===== BEGIN FILE: page_password_vault.py =====

import customtkinter as ctk
from tkinter import messagebox
from datetime import datetime
from pathlib import Path

from config import (
    BG_DARK,
    BG_PANEL,
    BG_CARD,
    BG_CARD2,
    ACCENT,
    ACCENT2,
    TEXT_PRIMARY,
    TEXT_MUTED,
)
from widgets import lbl, btn, inp, hsep, Sidebar, SIDEBAR_ITEMS


class PagePasswordVault(ctk.CTkFrame):
    """Password Vault page with save/list/edit/delete using a text file."""

    def __init__(self, master, on_nav, on_back, on_config):
        super().__init__(master, fg_color=BG_DARK, corner_radius=0)
        self._vault_file = Path(__file__).resolve().parent / "password_vault_details.txt"
        self._editing_index = None
        self._build(on_nav, on_back, on_config)

    def _build(self, on_nav, on_back, on_config):
        Sidebar(
            self,
            SIDEBAR_ITEMS,
            active="Password Vault",
            on_select=on_nav,
            on_back=on_back,
            on_config=on_config,
        ).pack(side="left", fill="y")

        main = ctk.CTkFrame(self, fg_color=BG_DARK, corner_radius=0)
        main.pack(side="left", fill="both", expand=True)

        hdr = ctk.CTkFrame(main, fg_color=BG_PANEL, corner_radius=0, height=48)
        hdr.pack(fill="x")
        hdr.pack_propagate(False)
        ctk.CTkLabel(
            hdr,
            text="Password Vault",
            font=("Courier New", 16, "bold"),
            text_color=ACCENT,
        ).pack(side="left", padx=16, pady=12)

        body = ctk.CTkScrollableFrame(main, fg_color="transparent")
        body.pack(fill="both", expand=True, padx=24, pady=16)

        # Section 1: create/update
        card = ctk.CTkFrame(body, fg_color=BG_CARD, corner_radius=12)
        card.pack(fill="x", pady=(0, 12))

        lbl(card, "Section 1 - Save / Update password", size=11, color=TEXT_MUTED).pack(
            anchor="w", padx=16, pady=(16, 4)
        )
        hsep(card)

        self.service = inp(card, ph="Service / System name", w=380)
        self.service.pack(anchor="w", padx=16, pady=(8, 6))
        self.username = inp(card, ph="Username", w=380)
        self.username.pack(anchor="w", padx=16, pady=6)
        self.password = inp(card, ph="Password", w=380, show="â—")
        self.password.pack(anchor="w", padx=16, pady=6)

        row = ctk.CTkFrame(card, fg_color="transparent")
        row.pack(fill="x", padx=16, pady=16)
        self.save_btn = btn(row, "Save Entry", self._save_entry, w=140, h=36)
        self.save_btn.pack(side="left")
        btn(
            row,
            "Clear",
            self._clear,
            fg=BG_CARD2,
            hover=BG_CARD,
            tc=TEXT_MUTED,
            w=100,
            h=36,
        ).pack(side="left", padx=8)
        btn(
            row,
            "Cancel Edit",
            self._cancel_edit,
            fg=BG_CARD2,
            hover=BG_CARD,
            tc=TEXT_MUTED,
            w=110,
            h=36,
        ).pack(side="left", padx=8)

        self.status = lbl(card, "", size=10, color=TEXT_MUTED)
        self.status.pack(anchor="w", padx=16, pady=(0, 12))

        self.path_lbl = lbl(
            card,
            f"Storage file: {self._vault_file}",
            size=10,
            color=TEXT_MUTED,
        )
        self.path_lbl.pack(anchor="w", padx=16, pady=(0, 14))

        note = ctk.CTkFrame(body, fg_color=BG_CARD, corner_radius=12)
        note.pack(fill="x", pady=(0, 12))
        lbl(
            note,
            "Note: This is a new tab scaffold. Wire secure storage/back-end before production use.",
            size=11,
            color=TEXT_MUTED,
        ).pack(anchor="w", padx=16, pady=14)

        # Section 2: list all
        list_card = ctk.CTkFrame(body, fg_color=BG_CARD, corner_radius=12)
        list_card.pack(fill="both", expand=True, pady=(0, 12))
        lbl(list_card, "Section 2 - Saved passwords", size=11, color=TEXT_MUTED).pack(
            anchor="w", padx=16, pady=(16, 4)
        )
        hsep(list_card)
        self._list_wrap = ctk.CTkFrame(list_card, fg_color="transparent")
        self._list_wrap.pack(fill="both", expand=True, padx=12, pady=(4, 12))

        self._refresh_entries_view()

    def _sanitize(self, value: str) -> str:
        return value.replace("|", "/").replace("\n", " ").strip()

    def _parse_line(self, line: str):
        parts = [p.strip() for p in line.split("|")]
        if len(parts) < 4:
            return None
        ts = parts[0]
        service = parts[1].replace("service=", "", 1).strip()
        username = parts[2].replace("username=", "", 1).strip()
        password = parts[3].replace("password=", "", 1).strip()
        if not service and not username:
            return None
        return {"ts": ts, "service": service, "username": username, "password": password}

    def _read_entries(self):
        if not self._vault_file.exists():
            return []
        rows = []
        try:
            for raw in self._vault_file.read_text(encoding="utf-8").splitlines():
                item = self._parse_line(raw)
                if item:
                    rows.append(item)
        except Exception:
            return []
        return rows

    def _write_entries(self, rows):
        self._vault_file.parent.mkdir(parents=True, exist_ok=True)
        lines = []
        for r in rows:
            line = (
                f"{r['ts']} | service={self._sanitize(r['service'])} | "
                f"username={self._sanitize(r['username'])} | password={self._sanitize(r['password'])}"
            )
            lines.append(line)
        self._vault_file.write_text("\n".join(lines) + ("\n" if lines else ""), encoding="utf-8")

    def _refresh_entries_view(self):
        for child in self._list_wrap.winfo_children():
            child.destroy()

        rows = self._read_entries()
        if not rows:
            lbl(self._list_wrap, "No saved passwords yet.", size=11, color=TEXT_MUTED).pack(
                anchor="w", padx=8, pady=8
            )
            return

        for idx, r in enumerate(rows):
            row = ctk.CTkFrame(self._list_wrap, fg_color=BG_CARD2, corner_radius=8)
            row.pack(fill="x", padx=6, pady=4)
            left = ctk.CTkFrame(row, fg_color="transparent")
            left.pack(side="left", fill="x", expand=True, padx=10, pady=8)
            lbl(
                left,
                f"{idx + 1}. {r['service']}  |  {r['username']}",
                size=12,
                color=TEXT_PRIMARY,
            ).pack(anchor="w")
            lbl(left, f"Password: {r['password']}", size=10, color=TEXT_MUTED).pack(anchor="w")
            lbl(left, f"Saved: {r['ts']}", size=10, color=TEXT_MUTED).pack(anchor="w")

            btn(
                row,
                "Edit",
                lambda i=idx: self._start_edit(i),
                fg=BG_CARD,
                hover=BG_DARK,
                tc=TEXT_PRIMARY,
                w=80,
                h=32,
            ).pack(side="right", padx=(6, 10), pady=8)
            btn(
                row,
                "Delete",
                lambda i=idx: self._delete_entry(i),
                fg=BG_CARD,
                hover=BG_DARK,
                tc="#ff8080",
                w=80,
                h=32,
            ).pack(side="right", pady=8)

    def _save_entry(self):
        svc = self.service.get().strip()
        usr = self.username.get().strip()
        pwd = self.password.get().strip()
        if not svc or not usr or not pwd:
            messagebox.showwarning("Password Vault", "Please fill service, username, and password.")
            return

        rows = self._read_entries()
        ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        item = {"ts": ts, "service": svc, "username": usr, "password": pwd}
        if self._editing_index is None:
            rows.append(item)
            action = f"Saved: {svc}"
        else:
            rows[self._editing_index] = item
            action = f"Updated: {svc}"
        try:
            self._write_entries(rows)
        except Exception as e:
            messagebox.showerror("Password Vault", f"Failed to save details:\n{e}")
            return

        self.status.configure(text=f"{action}  â€¢  File: {self._vault_file.name}")
        self._editing_index = None
        self.save_btn.configure(text="Save Entry")
        self._clear()
        self._refresh_entries_view()

    def _start_edit(self, idx: int):
        rows = self._read_entries()
        if idx < 0 or idx >= len(rows):
            return
        r = rows[idx]
        self._editing_index = idx
        self.service.delete(0, "end")
        self.service.insert(0, r["service"])
        self.username.delete(0, "end")
        self.username.insert(0, r["username"])
        self.password.delete(0, "end")
        self.password.insert(0, r["password"])
        self.save_btn.configure(text="Update Entry")
        self.status.configure(text=f"Editing entry #{idx + 1}")

    def _delete_entry(self, idx: int):
        rows = self._read_entries()
        if idx < 0 or idx >= len(rows):
            return
        target = rows[idx]
        if not messagebox.askyesno(
            "Delete Entry",
            f"Delete password for service '{target['service']}' and user '{target['username']}'?",
        ):
            return
        rows.pop(idx)
        try:
            self._write_entries(rows)
        except Exception as e:
            messagebox.showerror("Password Vault", f"Failed to delete entry:\n{e}")
            return
        if self._editing_index == idx:
            self._cancel_edit()
        self.status.configure(text=f"Deleted entry: {target['service']}")
        self._refresh_entries_view()

    def _clear(self):
        self.service.delete(0, "end")
        self.username.delete(0, "end")
        self.password.delete(0, "end")

    def _cancel_edit(self):
        self._editing_index = None
        self.save_btn.configure(text="Save Entry")
        self._clear()
        self.status.configure(text="Edit cancelled")


# ===== END FILE: page_password_vault.py =====


# ===== BEGIN FILE: main.py =====

import customtkinter as ctk
from tkinter import messagebox
from app_logging import setup_application_logging, get_logger
from config import (state, BG_DARK, BG_PANEL, BG_CARD, BG_CARD2, ACCENT, ACCENT2,
                    SUCCESS, WARNING, DANGER, TEXT_PRIMARY, TEXT_MUTED, BORDER)
from page_setup import PageSelectProject, PageSelectEnv, PageDashboard
from page_logs import PageLogs
from page_database import PageDatabase
from page_consumption import PageConsumption
from page_mocksim import PageMockSim
from page_invoice import PageInvoice
from page_config import PageConfig
from page_sim_deletion import PageSimDeletion
from page_log_analysis import PageLogAnalysis
from page_password_vault import PagePasswordVault
from evidence_capture import EvidenceFloatPanel

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  MAIN APP
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
class MonToolApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        self._log = get_logger("app.ui")
        self.title("MonTool â€” Linux Admin Monitor")
        self.geometry("1280x800")
        self.minsize(1024, 680)
        self.configure(fg_color=BG_DARK)
        self._current = None
        self._evidence_panel = None
        state.evidence_panel_opener = self._open_evidence_panel
        self.protocol("WM_DELETE_WINDOW", self._on_close_request)
        self._show(PageSelectProject(self, self._on_project))

    def _open_evidence_panel(self):
        """Floating always-on-top evidence window (screenshots / reports)."""
        p = self._evidence_panel
        if p is not None:
            try:
                if p.winfo_exists():
                    p.lift()
                    p.focus_force()
                    return
            except Exception:
                self._evidence_panel = None
        self._evidence_panel = EvidenceFloatPanel(self)
        self._log.info("Evidence float panel opened")

    def _evidence_panel_cleared(self):
        self._evidence_panel = None

    def _on_close_request(self):
        self._log.info("Session end (user closed window)")
        if self._evidence_panel is not None:
            try:
                self._evidence_panel.destroy()
            except Exception:
                pass
            self._evidence_panel = None
        self.destroy()

    def _show(self, page):
        if self._current:
            self._current.destroy()
        self._current = page
        page.pack(fill="both", expand=True)

    def _on_project(self, project):
        state.project = project
        get_logger("app").info("Project selected: %s", project)
        self._show(PageSelectEnv(self, self._on_env,
                   lambda: self._show(PageSelectProject(self, self._on_project))))

    def _on_env(self, env):
        state.env = env
        get_logger("app").info("Environment selected: %s", env)
        self._show_dashboard()

    def _show_dashboard(self):
        self._show(PageDashboard(self, self._on_menu, self._back_to_env))

    def _back_to_env(self):
        get_logger("app").info("Back to environment selection")
        self._show(PageSelectEnv(self, self._on_env,
                   lambda: self._show(PageSelectProject(self, self._on_project))))

    def _on_menu(self, item):
        if item == "Evidence Capture":
            get_logger("app").info("Menu: Evidence Capture (float panel)")
            self._open_evidence_panel()
            return
        nav, back, cfg = self._on_menu, self._show_dashboard, self._show_config
        get_logger("app").info("Menu: %s", item)
        pages = {
            "Logs":                  lambda: PageLogs(self, nav, back, cfg),
            "Database Queries":      lambda: PageDatabase(self, nav, back, cfg),
            "Password Vault":        lambda: PagePasswordVault(self, nav, back, cfg),
            "Mock & Sim Activation": lambda: PageMockSim(self, nav, back, cfg),
            "Invoice Pregeneration": lambda: PageInvoice(self, nav, back, cfg),
            "Consumption":           lambda: PageConsumption(self, nav, back, cfg),
            "Sim Deletion":          lambda: PageSimDeletion(self, nav, back, cfg),
            "Log Analysis":          lambda: PageLogAnalysis(self, nav, back, cfg),
        }
        if item in pages:
            try:
                result = pages[item]()
                if result is not None:
                    self._show(result)
            except Exception as e:
                get_logger("app").exception("Menu open failed for '%s': %s", item, e)
                messagebox.showerror("Open Page Failed", f"{item}\n\n{e}")

    def _show_config(self):
        self._show(PageConfig(self, self._show_dashboard))


if __name__ == "__main__":
    log_path = setup_application_logging()
    log = get_logger("main")
    log.info("=== MonTool session start ===")
    log.info("Log file: %s", log_path)
    try:
        app = MonToolApp()
        app.mainloop()
    except Exception:
        get_logger("main").exception("Fatal error in mainloop")
        raise
    finally:
        get_logger("main").info("=== MonTool process exit ===")

# ===== END FILE: main.py =====

